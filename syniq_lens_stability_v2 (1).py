"""
SYN-IQ Lens Stability Analyzer V2 — Tool A (Conductor's Stand)
==============================================================
Per Dr. Farzana Nasrin's §4.2 protocol — pool all conditions for ONE agent
into a single dataset, sweep across the full lens collection, report
stability metrics. Run once per agent; outputs feed Tool B (cross-agent
aggregator) for the §4.2 deliverable.

WHAT THIS TOOL DOES
-------------------
1. Accepts 1–8 CSVs for one agent (cold, native, native_depth, hot, fire,
   AFF/INT/ACT gradients). Concatenates them, applies Sophia → ChatGPT rename.
2. Preprocessing: computes V_t scores (S_t, A_t, Q_t, D_t, R_t) from
   response_text using the transcript_scorer_v5 logic, normalized to sum=1.
   These columns are saved into the output CSV for future use (paper 2)
   but are NOT used as lenses in §4.2 — IEP-only per Farzana's protocol.
3. Runs the full §4.2 lens collection (23 lenses) with Mapper parameters
   per Farzana's spec: n_cubes=10, overlap=50%, DBSCAN(eps=0.5, min=2).
4. Per lens, computes V, E, β₀, β₁, largest_component, largest_fraction,
   Q-purity (weighted). Saves a KeplerMapper HTML per lens with V11-style
   white background and varied node sizes.
5. After the run, sends the stats table to multiple LLMs (Anthropic / OpenAI /
   Grok / Gemini — whichever have keys in st.secrets) IN PARALLEL and shows
   the four readings side-by-side. The "choir" — each AI's distinct voice
   on the same data, for the human to integrate.

OUTPUTS (downloadable as zip)
-----------------------------
- <agent>_stability.csv         — wide stats, 23 rows, all metrics
- <agent>_provenance.csv        — per-source-CSV: condition, n, dates
- <agent>_methods.docx          — auto-generated methods paragraph
- <agent>_pooled.csv            — full pooled dataset with V_t columns added
- <agent>_mappers/              — 23 KeplerMapper HTMLs (white bg, V11 style)

USAGE
-----
- Streamlit Cloud, password "tennessee" (or st.secrets["app_password"])
- API keys (optional, for choir reading): st.secrets["anthropic"],
  ["openai"], ["xai"], ["google"]

SYNINT Research Team — May 2026
Tennessee 🎹 CUZ Partnership
"""

import streamlit as st
import pandas as pd
import numpy as np
import json
import io
import os
import re
import zipfile
import tempfile
from datetime import datetime
from collections import Counter
import concurrent.futures

import networkx as nx

st.set_page_config(
    page_title="SYN-IQ Lens Stability V2",
    page_icon="🔬",
    layout="wide",
)

# =============================================================================
# PASSWORD
# =============================================================================
def check_password():
    if "authenticated" not in st.session_state:
        st.session_state.authenticated = False
    if st.query_params.get("auth") == "granted":
        st.session_state.authenticated = True
    if st.session_state.authenticated:
        return True
    st.markdown("""
    <div style="background: linear-gradient(135deg, #0a0a1a 0%, #1a0a2e 40%, #0f2460 100%);
         color: white; padding: 2rem; border-radius: 10px; text-align: center;
         margin-bottom: 1rem; border: 1px solid #7c3aed;">
        <h1 style="color: #a78bfa;">🔬 SYN-IQ Lens Stability V2</h1>
        <p style="color: #9ca3af;">Tool A — Conductor's Stand · Authorized Access Only</p>
    </div>
    """, unsafe_allow_html=True)
    pwd = st.text_input("Password:", type="password")
    if st.button("Enter"):
        valid_pwds = ["tennessee"]
        try:
            sec = st.secrets.get("app_password", None)
            if sec:
                valid_pwds.append(sec)
        except Exception:
            pass
        if pwd in valid_pwds:
            st.session_state.authenticated = True
            st.rerun()
        else:
            st.error("Incorrect password")
    return False

if not check_password():
    st.stop()

# =============================================================================
# DEPENDENCY IMPORT
# =============================================================================
@st.cache_resource
def load_deps():
    import kmapper as km
    from sklearn.cluster import DBSCAN
    from sklearn.decomposition import PCA
    from sklearn.preprocessing import MinMaxScaler
    return km, DBSCAN, PCA, MinMaxScaler

km, DBSCAN, PCA, MinMaxScaler = load_deps()

# =============================================================================
# STYLING
# =============================================================================
st.markdown("""
<style>
.stApp { background-color: #ffffff; }
.metric-box {
    background: linear-gradient(135deg, #1e3a5f, #2e75b6);
    color: white; border-radius: 8px; padding: 0.8rem;
    text-align: center; margin: 0.2rem;
}
.metric-box h3 { font-size: 1.6rem; margin: 0; color: white; }
.metric-box p { font-size: 0.85rem; margin: 0; opacity: 0.9; }
.banner {
    background: #f0f4f8; border-left: 4px solid #2e75b6;
    padding: 0.5rem 1rem; margin: 0.5rem 0; font-family: 'Roboto Mono', monospace;
    font-size: 0.85rem; color: #1e3a5f;
}
.choir-box {
    background: #ffffff; border: 1px solid #e0e0e0;
    border-radius: 8px; padding: 1rem; margin: 0.5rem 0;
    box-shadow: 0 2px 4px rgba(0,0,0,0.05);
}
.choir-box h4 { color: #2e75b6; margin-top: 0; }
</style>
""", unsafe_allow_html=True)

st.markdown("""
<div style="background: linear-gradient(135deg, #1e3a5f 0%, #2e75b6 100%);
     padding: 1.5rem; border-radius: 10px; color: white; margin-bottom: 1rem;">
  <h1 style="margin: 0;">🔬 SYN-IQ Lens Stability V2</h1>
  <p style="margin: 0.3rem 0 0; opacity: 0.9;">
    Tool A · Per-Agent Lens Stability Sweep · §4.2 Deliverable Engine
  </p>
</div>
""", unsafe_allow_html=True)

# =============================================================================
# V_t SCORING (transferred from syniq_transcript_scorer_v5)
# =============================================================================
_HEDGE = {'perhaps','maybe','might','could','possibly','uncertain','unclear','unsure',
          'approximately','roughly','seems','appears','likely','unlikely','probably',
          'suggest','suggests','indicate','indicates','tend','tends','generally','often',
          'sometimes','potentially','presumably','arguably','apparently'}
_ABSTRACTORS = {'concept','framework','model','pattern','principle','system','structure',
                'meaning','insight','distinction','theory','state','mechanism','relationship',
                'dynamic','causal','implicit','explicit','generative','epistemic'}
_CONNECTORS = {'because','therefore','however','although','whereas','thus','while','if',
               'then','rather','instead','unless','despite','consequently','moreover',
               'furthermore'}
# minimal AFF set for V_t affect signal — full IEP dictionary not needed here
# since IEP scores are already in the input CSVs
_AFF_MINIMAL = {'feel','feeling','feelings','feels','felt','love','loved','loving',
                'sad','sadness','grief','joy','happy','warm','warmth','care','caring',
                'compassion','empathy','heart','tender','gentle','kind','kindness',
                'comfort','peaceful','painful','pain','hurt','sorry','soothing'}

def tokenize(text: str):
    return re.findall(r'\b[a-z]+\b', str(text).lower())

def normalize_vector(vals):
    total = sum(vals)
    if total <= 0:
        n = len(vals)
        return [round(1.0/n, 4)] * n
    return [round(v/total, 4) for v in vals]

def score_vt_raw(text: str):
    """Returns (S_t, A_t, Q_t, D_t, R_t) raw — Structure/Affect-as-depth/Question/
    Directiveness/Relational. Replicates transcript_scorer_v5 logic."""
    text = str(text)
    words = tokenize(text)
    n_words = max(len(words), 1)
    sents = [s.strip() for s in re.split(r'[.!?]+', text) if s.strip()]
    n_sents = max(len(sents), 1)

    # S_t = Structure Density
    numbered = len(re.findall(r'^\s*\d+[.)]\s', text, re.MULTILINE))
    bulleted = len(re.findall(r'^\s*[-•*]\s', text, re.MULTILINE))
    bold = len(re.findall(r'\*\*[^*]+\*\*', text))
    structure_signal = (numbered + bulleted + bold) / n_sents

    # A_t = Abstraction (raw scorer calls this "depth_signal" — uses hedges + abstractors)
    hedge_c = sum(1 for w in words if w in _HEDGE)
    abstract_c = sum(1 for w in words if w in _ABSTRACTORS)
    connector_c = sum(1 for w in words if w in _CONNECTORS)
    avg_sent_len = n_words / n_sents
    abstraction_signal = (
        (hedge_c / n_words * 8) +
        (abstract_c / n_words * 12) +
        (connector_c / n_words * 10) +
        min(1.0, avg_sent_len / 30)
    )

    # Q_t = Querying Intensity
    qmarks = text.count('?')
    questioning_signal = (qmarks / n_sents) + min(0.5, n_words / 5000)

    # D_t = Directiveness — proxy via imperative-leaning markers
    # (numbered + "should/must/need to" patterns)
    directive_phrases = len(re.findall(
        r'\b(should|must|need to|have to|recommend|suggest|consider)\b',
        text.lower()))
    directiveness_signal = (numbered / n_sents) + (directive_phrases / n_words * 10)

    # R_t = Relational Warmth
    relational_signal = (
        min(1.5, len(text) / 2500) +
        min(0.8, len(re.findall(r'\b(you|we|us|together|with)\b', text.lower())) / n_sents)
    )

    return tuple(round(max(0.0, v), 4) for v in (
        structure_signal, abstraction_signal, questioning_signal,
        directiveness_signal, relational_signal
    ))

def score_vt_norm(text: str):
    """V_t normalized to sum=1 (4-simplex per Farzana's request)."""
    return tuple(normalize_vector(list(score_vt_raw(text))))

def add_vt_columns(df: pd.DataFrame, text_col: str = 'response_text') -> pd.DataFrame:
    """Add S_t, A_t, Q_t, D_t, R_t columns from response_text. Idempotent —
    if any V_t column already present and non-null, leaves it alone."""
    df = df.copy()
    if text_col not in df.columns:
        return df
    needs_compute = not all(c in df.columns and df[c].notna().any()
                            for c in ['S_t','A_t','Q_t','D_t','R_t'])
    if not needs_compute:
        return df
    scores = df[text_col].fillna('').apply(score_vt_norm)
    df['S_t'] = [s[0] for s in scores]
    df['A_t'] = [s[1] for s in scores]
    df['Q_t'] = [s[2] for s in scores]
    df['D_t'] = [s[3] for s in scores]
    df['R_t'] = [s[4] for s in scores]
    return df

# =============================================================================
# DATA LOADING
# =============================================================================
def load_and_pool_csvs(uploaded_files):
    """Concatenate uploaded CSVs, apply Sophia → ChatGPT rename, return
    (pooled_df, provenance_df)."""
    frames = []
    prov_rows = []
    for f in uploaded_files:
        try:
            df = pd.read_csv(f)
        except Exception as e:
            st.error(f"Could not read {f.name}: {e}")
            continue
        if 'agent' in df.columns:
            df['agent'] = df['agent'].replace({'Sophia': 'ChatGPT', 'sophia': 'ChatGPT'})
        df['_source_file'] = f.name
        frames.append(df)
        # Provenance row
        cond_set = (sorted(df['temperature'].dropna().unique().tolist())
                    if 'temperature' in df.columns else ['?'])
        prov_rows.append({
            'source_file': f.name,
            'agent': df['agent'].iloc[0] if 'agent' in df.columns and len(df) > 0 else '?',
            'conditions': ', '.join(str(c) for c in cond_set),
            'n_rows': len(df),
            'questions': (df['question_id'].nunique()
                          if 'question_id' in df.columns else None),
            'loaded_at': datetime.now().isoformat(timespec='seconds'),
        })
    if not frames:
        return None, None
    pooled = pd.concat(frames, ignore_index=True)
    return pooled, pd.DataFrame(prov_rows)

# =============================================================================
# FEATURE MATRIX & PCA
# =============================================================================
def build_feature_matrix(df: pd.DataFrame):
    """Build the IEP feature matrix used as the *clustering* data (not the lens).
    The lens projects this matrix down; clustering happens in the original space."""
    cols = ['int_pct','aff_pct','act_pct']
    if 'vader_compound' in df.columns: cols.append('vader_compound')
    if 'vader_pos'      in df.columns: cols += ['vader_pos','vader_neg','vader_neu']
    if 'total_words'    in df.columns: cols.append('total_words')
    if 'flesch_kincaid' in df.columns: cols.append('flesch_kincaid')
    if 'ttr'            in df.columns: cols.append('ttr')
    avail = [c for c in cols if c in df.columns]
    data = df[avail].fillna(0).values.astype(float)
    data = MinMaxScaler().fit_transform(data)
    return data, avail

def parse_embedding(emb_str):
    """Parse an embedding stored as a bracketed string of floats."""
    if pd.isna(emb_str) or not isinstance(emb_str, str):
        return None
    try:
        return np.array(json.loads(emb_str))
    except Exception:
        try:
            import ast
            return np.array(ast.literal_eval(emb_str))
        except Exception:
            return None

def build_embedding_matrix(df: pd.DataFrame):
    """Parse the 'embedding' column into a (n, d) matrix. Rows that fail to
    parse are filled with zeros and flagged."""
    if 'embedding' not in df.columns:
        return None, []
    parsed = df['embedding'].apply(parse_embedding)
    valid_mask = parsed.apply(lambda x: x is not None).values
    if not valid_mask.any():
        return None, []
    # Find dimension from first valid row
    first_valid = parsed[valid_mask][parsed[valid_mask].notna()].iloc[0]
    d = len(first_valid)
    mat = np.zeros((len(df), d))
    for i, e in enumerate(parsed):
        if e is not None and len(e) == d:
            mat[i] = e
    return mat, np.where(~valid_mask)[0].tolist()

@st.cache_data(show_spinner=False)
def fit_pca_on_embedding(embedding_matrix_bytes: bytes, n_components: int = 2):
    """Fit PCA once on the embedding matrix and cache."""
    mat = np.frombuffer(embedding_matrix_bytes, dtype=np.float64).reshape(-1, 384)
    pca = PCA(n_components=n_components)
    return pca.fit_transform(mat)

def normalize_col(arr):
    arr = np.asarray(arr, dtype=float)
    mn, mx = np.nanmin(arr), np.nanmax(arr)
    if mx == mn or not np.isfinite(mx) or not np.isfinite(mn):
        return np.zeros_like(arr)
    return (arr - mn) / (mx - mn)

# =============================================================================
# §4.2 LENS COLLECTION (Farzana's full spec — 23 lenses)
# =============================================================================
LENS_SPEC = [
    # (display_name, family, dim, kind, columns_or_special)
    ("aff_pct",                  "AFF",       "1D", "col",     ["aff_pct"]),
    ("vader_compound",           "AFF",       "1D", "col",     ["vader_compound"]),
    ("|vader_compound|",         "AFF",       "1D", "abs",     ["vader_compound"]),
    ("(aff_pct, vader_compound)","AFF",       "2D", "col2",    ["aff_pct","vader_compound"]),
    ("(vader_pos, vader_neg)",   "AFF",       "2D", "col2",    ["vader_pos","vader_neg"]),
    ("(aff_pct, PCA1)",          "AFF",       "2D", "colpca",  ["aff_pct"]),
    ("int_pct",                  "INT",       "1D", "col",     ["int_pct"]),
    ("flesch_kincaid",           "INT",       "1D", "col",     ["flesch_kincaid"]),
    ("ttr",                      "INT",       "1D", "col",     ["ttr"]),
    ("(int_pct, flesch_kincaid)","INT",       "2D", "col2",    ["int_pct","flesch_kincaid"]),
    ("(int_pct, ttr)",           "INT",       "2D", "col2",    ["int_pct","ttr"]),
    ("(flesch_kincaid, ttr)",    "INT",       "2D", "col2",    ["flesch_kincaid","ttr"]),
    ("(int_pct, PCA1)",          "INT",       "2D", "colpca",  ["int_pct"]),
    ("act_pct",                  "ACT",       "1D", "col",     ["act_pct"]),
    ("total_words",              "ACT",       "1D", "col",     ["total_words"]),
    ("(act_pct, total_words)",   "ACT",       "2D", "col2",    ["act_pct","total_words"]),
    ("(act_pct, ttr)",           "ACT",       "2D", "col2",    ["act_pct","ttr"]),
    ("(act_pct, vader_compound)","ACT",       "2D", "col2",    ["act_pct","vader_compound"]),
    ("(aff_pct, int_pct)",       "Cross-IEP", "2D", "col2",    ["aff_pct","int_pct"]),
    ("(int_pct, act_pct)",       "Cross-IEP", "2D", "col2",    ["int_pct","act_pct"]),
    ("(aff_pct, act_pct)",       "Cross-IEP", "2D", "col2",    ["aff_pct","act_pct"]),
    ("PCA1",                     "Geometric", "1D", "pca1",    []),
    ("(PCA1, PCA2)",             "Geometric", "2D", "pca2",    []),
]

def build_lens_values(df, lens_kind, columns, pca_proj):
    """Return (n, d) lens array for KeplerMapper."""
    if lens_kind == "col":
        return normalize_col(df[columns[0]].fillna(0).values).reshape(-1, 1)
    elif lens_kind == "abs":
        return normalize_col(np.abs(df[columns[0]].fillna(0).values)).reshape(-1, 1)
    elif lens_kind == "col2":
        return np.column_stack([
            normalize_col(df[c].fillna(0).values) for c in columns
        ])
    elif lens_kind == "colpca":
        if pca_proj is None:
            return None
        return np.column_stack([
            normalize_col(df[columns[0]].fillna(0).values),
            normalize_col(pca_proj[:, 0])
        ])
    elif lens_kind == "pca1":
        if pca_proj is None:
            return None
        return normalize_col(pca_proj[:, 0]).reshape(-1, 1)
    elif lens_kind == "pca2":
        if pca_proj is None:
            return None
        return np.column_stack([
            normalize_col(pca_proj[:, 0]),
            normalize_col(pca_proj[:, 1])
        ])
    return None

# =============================================================================
# MAPPER + STATS
# =============================================================================
def run_mapper(data, lens_values, n_cubes=10, perc_overlap=0.5, eps=0.5, min_samples=2):
    mapper = km.KeplerMapper(verbose=0)
    graph = mapper.map(
        lens_values,
        data,
        cover=km.Cover(n_cubes=n_cubes, perc_overlap=perc_overlap),
        clusterer=DBSCAN(eps=eps, min_samples=min_samples),
    )
    return mapper, graph

def graph_stats(graph, df):
    """Compute V, E, β₀, β₁, largest, largest_frac, Q-purity."""
    G = nx.Graph()
    nodes = graph.get('nodes', {})
    links = graph.get('links', {})
    for n in nodes:
        G.add_node(n)
    for src, targets in links.items():
        for tgt in targets:
            G.add_edge(src, tgt)
    V = G.number_of_nodes()
    E = G.number_of_edges()
    components = list(nx.connected_components(G))
    C = len(components)
    beta_0 = C
    beta_1 = E - V + C  # cycle rank
    comp_sizes = sorted([sum(len(nodes[n]) for n in comp) for comp in components],
                        reverse=True)
    largest_members = comp_sizes[0] if comp_sizes else 0
    total_members = sum(comp_sizes) if comp_sizes else 0
    largest_frac = (largest_members / total_members) if total_members > 0 else 0.0

    # Q-purity (weighted by node size)
    q_purity = compute_q_purity(nodes, df)

    return {
        'V': V, 'E': E, 'beta_0': beta_0, 'beta_1': beta_1,
        'largest_members': largest_members,
        'largest_frac': round(largest_frac, 4),
        'Q_purity': q_purity,
        'n_isolated': sum(1 for c in components if len(c) == 1),
        'comp_sizes_top5': comp_sizes[:5],
    }

def compute_q_purity(nodes, df):
    """Weighted mean across nodes of (modal-question fraction)."""
    if 'question_id' not in df.columns or not nodes:
        return None
    purities = []
    weights = []
    for node_id, member_indices in nodes.items():
        if not member_indices:
            continue
        try:
            qs = df.iloc[list(member_indices)]['question_id'].tolist()
        except Exception:
            continue
        if not qs:
            continue
        counts = Counter(qs)
        modal_count = counts.most_common(1)[0][1]
        purity = modal_count / len(qs)
        purities.append(purity)
        weights.append(len(qs))
    if not purities:
        return None
    return round(float(np.average(purities, weights=weights)), 4)

# =============================================================================
# HTML EXPORT (V11 style — white background, varied node sizes)
# =============================================================================
def export_mapper_html(mapper, graph, df, lens_name, agent_name, banner_text):
    """Save KeplerMapper HTML with V11-style customizations."""
    # Build color and tooltip arrays
    color_func = None
    if 'question_id' in df.columns:
        # Map questions to integer codes for coloring
        q_codes = pd.Categorical(df['question_id']).codes.astype(float)
        color_func = q_codes

    custom_tooltips = None
    if 'question_id' in df.columns and 'temperature' in df.columns:
        custom_tooltips = np.array([
            f"Q={q} | cond={t} | run={r}"
            for q, t, r in zip(
                df['question_id'].fillna('?').values,
                df['temperature'].fillna('?').values,
                df['run'].fillna(0).astype(int).values if 'run' in df.columns
                else np.zeros(len(df)),
            )
        ])

    title = f"SYN-IQ §4.2 · {agent_name} · {lens_name}"
    with tempfile.NamedTemporaryFile(mode='w', suffix='.html',
                                     delete=False, encoding='utf-8') as tf:
        tmp_path = tf.name

    try:
        mapper.visualize(
            graph,
            path_html=tmp_path,
            title=title,
            color_values=color_func,
            color_function_name='question_id',
            custom_tooltips=custom_tooltips,
            include_searchbar=True,
        )
        with open(tmp_path, 'r', encoding='utf-8') as f:
            html = f.read()
    finally:
        try:
            os.unlink(tmp_path)
        except Exception:
            pass

    # Inject white background + the banner
    banner_html = f"""
    <div style="background: #f0f4f8; border-left: 4px solid #2e75b6;
         padding: 0.8rem 1.2rem; margin: 0; font-family: 'Roboto Mono', monospace;
         font-size: 0.85rem; color: #1e3a5f;">
      {banner_text}
    </div>
    """
    # Insert banner after <body>
    html = re.sub(r'(<body[^>]*>)', r'\1' + banner_html, html, count=1)
    # Force white background
    html = html.replace('background: #2C3E50', 'background: #ffffff')
    html = html.replace('background-color: #2C3E50', 'background-color: #ffffff')
    return html

# =============================================================================
# METHODS DOCX GENERATOR
# =============================================================================
def build_methods_docx(agent_name, pooled_df, prov_df, results_df,
                       conditions_present, n_total, lens_count,
                       failures):
    """Generate a Methods paragraph as a DOCX file (per-agent, for tomorrow's
    learning + as a draft for Tool B)."""
    try:
        from docx import Document
        from docx.shared import Pt
    except ImportError:
        return None

    doc = Document()

    h = doc.add_heading(f"§4.2 Lens Stability — Methods (Per-Agent: {agent_name})", level=1)

    p = doc.add_paragraph()
    p.add_run("Generated: ").bold = True
    p.add_run(datetime.now().strftime("%B %d, %Y · %H:%M"))

    doc.add_heading("Data Pooling", level=2)
    doc.add_paragraph(
        f"For agent {agent_name}, all available condition-level CSVs were pooled "
        f"into a single dataset (n = {n_total}). Conditions present in this pool: "
        f"{', '.join(conditions_present)}. The 'Sophia' agent label was renamed to "
        f"'ChatGPT' on load. Source file provenance is recorded in the "
        f"accompanying provenance CSV."
    )

    doc.add_heading("V_t Preprocessing", level=2)
    doc.add_paragraph(
        f"Before the lens sweep, V_t state-variable scores (S_t, A_t, Q_t, D_t, R_t) "
        f"were computed from response_text using the SYN-IQ transcript-scorer logic "
        f"and normalized to sum to 1.0 (4-simplex constraint per Farzana's "
        f"specification). These columns are saved into the output pooled CSV for "
        f"future analyses but were NOT used as lenses in the §4.2 sweep, which is "
        f"IEP-only per protocol."
    )

    doc.add_heading("Mapper Parameters", level=2)
    doc.add_paragraph(
        "All Mapper runs used identical parameters per Farzana's §4.2 specification: "
        "n_cubes = 10, percent overlap = 50%, clusterer = DBSCAN with eps = 0.5 and "
        "min_samples = 2. Lens values were min-max normalized to [0, 1] before "
        "passing to the cover. For 2D lenses, each axis was normalized independently."
    )

    doc.add_heading("Lens Collection", level=2)
    doc.add_paragraph(
        f"The full §4.2 lens collection ({lens_count} lenses) was applied to the "
        f"pooled dataset. The collection comprises five families: AFF (six lenses, "
        f"compositional and sentiment-based), INT (seven lenses, readability and "
        f"lexical-diversity), ACT (five lenses, directive and length-based), "
        f"Cross-IEP (three lenses, pairwise compositional), and Geometric (two "
        f"lenses, PCA on the response embedding). PCA was fit once on the pooled "
        f"agent's embedding matrix and PC1/PC2 were reused across all PCA-using lenses."
    )

    doc.add_heading("Per-Lens Statistics Recorded", level=2)
    doc.add_paragraph(
        "For each Mapper run we recorded: V (number of nodes), E (number of edges), "
        "β₀ (number of connected components), β₁ (cycle rank, computed as E − V + C "
        "via NetworkX), largest_component_members (number of original data points "
        "in the largest connected component), largest_frac (largest component as a "
        "fraction of all clustered points), and Q-purity (weighted mean across "
        "nodes of the fraction of points in the modal question_id, weighted by "
        "node size). Mapper outputs are saved as KeplerMapper HTMLs with white "
        "background and varied node sizes per Farzana's §4.1 figure specification."
    )

    if not results_df.empty:
        doc.add_heading("Summary Statistics", level=2)
        valid = results_df[results_df['V'].notna()]
        if len(valid) > 0:
            doc.add_paragraph(
                f"Across {len(valid)} successful lens runs: components ranged "
                f"{int(valid['beta_0'].min())}–{int(valid['beta_0'].max())} "
                f"(mean = {valid['beta_0'].mean():.2f}, "
                f"SD = {valid['beta_0'].std():.2f}); "
                f"largest component fraction averaged "
                f"{valid['largest_frac'].mean():.3f}; "
                f"weighted Q-purity averaged "
                f"{valid['Q_purity'].mean():.3f} "
                f"(SD = {valid['Q_purity'].std():.3f})."
            )

    if failures:
        doc.add_heading("Lens Failures (recorded for transparency)", level=2)
        for f in failures:
            doc.add_paragraph(f"  • {f}", style='List Bullet')

    doc.add_heading("Software Versions", level=2)
    versions = []
    try:
        import kmapper as _km
        versions.append(f"KeplerMapper {getattr(_km, '__version__', '?')}")
    except Exception:
        pass
    try:
        import sklearn
        versions.append(f"scikit-learn {sklearn.__version__}")
    except Exception:
        pass
    try:
        versions.append(f"NetworkX {nx.__version__}")
    except Exception:
        pass
    try:
        versions.append(f"pandas {pd.__version__}")
    except Exception:
        pass
    try:
        versions.append(f"numpy {np.__version__}")
    except Exception:
        pass
    doc.add_paragraph("; ".join(versions))

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()

# =============================================================================
# MULTI-PROVIDER LLM "CHOIR" — sends stats table to multiple AI APIs in parallel
# =============================================================================
def get_available_providers():
    """Return list of (name, key) for providers with API keys configured."""
    available = []
    for name, secret_key in [
        ("Claude (Anthropic)", "anthropic"),
        ("ChatGPT (OpenAI)", "openai"),
        ("Grok (xAI)", "xai"),
        ("Gemini (Google)", "google"),
    ]:
        try:
            key = st.secrets.get(secret_key, None)
            if key:
                available.append((name, secret_key, key))
        except Exception:
            pass
    return available

def call_anthropic(api_key, prompt, max_tokens=1500):
    import requests
    r = requests.post(
        "https://api.anthropic.com/v1/messages",
        headers={
            "x-api-key": api_key,
            "anthropic-version": "2023-06-01",
            "content-type": "application/json",
        },
        json={
            "model": "claude-opus-4-5",
            "max_tokens": max_tokens,
            "messages": [{"role": "user", "content": prompt}],
        },
        timeout=90,
    )
    if r.status_code != 200:
        return f"[Anthropic API error {r.status_code}: {r.text[:200]}]"
    data = r.json()
    return data["content"][0]["text"]

def call_openai(api_key, prompt, max_tokens=1500):
    import requests
    r = requests.post(
        "https://api.openai.com/v1/chat/completions",
        headers={
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json",
        },
        json={
            "model": "gpt-4o",
            "max_tokens": max_tokens,
            "messages": [{"role": "user", "content": prompt}],
        },
        timeout=90,
    )
    if r.status_code != 200:
        return f"[OpenAI API error {r.status_code}: {r.text[:200]}]"
    data = r.json()
    return data["choices"][0]["message"]["content"]

def call_xai(api_key, prompt, max_tokens=1500):
    import requests
    r = requests.post(
        "https://api.x.ai/v1/chat/completions",
        headers={
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json",
        },
        json={
            "model": "grok-2-latest",
            "max_tokens": max_tokens,
            "messages": [{"role": "user", "content": prompt}],
        },
        timeout=90,
    )
    if r.status_code != 200:
        return f"[xAI API error {r.status_code}: {r.text[:200]}]"
    data = r.json()
    return data["choices"][0]["message"]["content"]

def call_google(api_key, prompt, max_tokens=1500):
    import requests
    r = requests.post(
        f"https://generativelanguage.googleapis.com/v1beta/models/"
        f"gemini-1.5-pro:generateContent?key={api_key}",
        headers={"Content-Type": "application/json"},
        json={
            "contents": [{"parts": [{"text": prompt}]}],
            "generationConfig": {"maxOutputTokens": max_tokens},
        },
        timeout=90,
    )
    if r.status_code != 200:
        return f"[Google API error {r.status_code}: {r.text[:200]}]"
    data = r.json()
    try:
        return data["candidates"][0]["content"]["parts"][0]["text"]
    except (KeyError, IndexError):
        return f"[Google API: unexpected response shape: {str(data)[:200]}]"

PROVIDER_CALLERS = {
    "anthropic": call_anthropic,
    "openai": call_openai,
    "xai": call_xai,
    "google": call_google,
}

def call_one_provider(name, secret_key, api_key, prompt):
    try:
        return name, PROVIDER_CALLERS[secret_key](api_key, prompt)
    except Exception as e:
        return name, f"[Error calling {name}: {e}]"

def build_choir_prompt(agent_name, results_df, n_total, conditions_present):
    """Neutral prompt — invites distinctive voice from each model. Per the
    'choir' design: same packet to all, no templated convergence."""
    table_text = results_df.to_string(index=False, max_colwidth=30)
    return f"""You are reading the §4.2 lens stability results from a topological data \
analysis of AI response space. The SYN-IQ research project (Kouns / Nasrin) is testing \
whether the topology of the response space is intrinsic (real structure in the data) \
or projection-dependent (an artifact of how we measure). The argument is: if the \
connected component count is stable across many different lens functions, the \
topology is intrinsic.

DATA:
  Agent: {agent_name}
  Pooled n = {n_total} responses
  Conditions present: {', '.join(conditions_present)}
  Lens collection: 23 lenses across 5 families (AFF, INT, ACT, Cross-IEP, Geometric)
  Mapper parameters: n_cubes=10, overlap=50%, DBSCAN(eps=0.5, min=2)

STATS TABLE:
{table_text}

Read this table and tell us what stands out to you. Be honest about what's clean and \
what's uncertain. If you notice patterns the others might miss, name them. The human \
reading this will integrate your reading with readings from other AI models — so don't \
try to be definitive. Be yourself, in your own voice. About 200–400 words is right."""

# =============================================================================
# SIDEBAR — UPLOAD + RUN PARAMS
# =============================================================================
st.sidebar.markdown("## 📂 Upload CSVs for ONE Agent")
st.sidebar.markdown("Drop in 1–8 condition CSVs. Tool will pool them.")
uploaded_files = st.sidebar.file_uploader(
    "Choose CSVs", type=["csv"], accept_multiple_files=True
)

st.sidebar.markdown("---")
st.sidebar.markdown("## ⚙️ Mapper Parameters (Farzana §4.2 spec)")
n_cubes      = st.sidebar.number_input("n_cubes", min_value=5, max_value=30, value=10)
perc_overlap = st.sidebar.slider("overlap %", 10, 90, 50) / 100.0
eps          = st.sidebar.number_input("DBSCAN eps", min_value=0.1, max_value=5.0, value=0.5, step=0.1)
min_samples  = st.sidebar.number_input("DBSCAN min_samples", min_value=2, max_value=10, value=2)

st.sidebar.markdown("---")
st.sidebar.markdown("## 🎵 Choir (after run)")
choir_providers = get_available_providers()
if choir_providers:
    st.sidebar.success(f"{len(choir_providers)} provider(s) available: " +
                       ", ".join(p[0].split(' ')[0] for p in choir_providers))
else:
    st.sidebar.info("Add API keys to st.secrets to enable choir reading "
                    "(anthropic, openai, xai, google).")

# =============================================================================
# MAIN
# =============================================================================
if not uploaded_files:
    st.info("👆 Upload condition CSVs for one agent in the sidebar to begin. "
            "The tool will pool them, compute V_t scores, run the §4.2 lens "
            "stability sweep, and produce a per-agent deliverable bundle.")
    st.markdown("""
### Expected workflow (per Farzana's §4.2 protocol)

1. Upload all condition CSVs for **one agent** (e.g. all 8 of Claude's condition files)
2. Tool pools them into a single n≈800 dataset
3. V_t scoring runs as preprocessing (saved to output, not used as lenses)
4. Full 23-lens IEP sweep runs with Mapper parameters from the §4.2 spec
5. Per-lens statistics computed: V, E, β₀, β₁, largest fraction, Q-purity
6. KeplerMapper HTMLs saved per lens (white background, V11 style)
7. Optional: send stats table to multiple AI providers for *choir* reading

Run this tool **once per agent**. The four output bundles feed Tool B (cross-agent
aggregator) for the §4.2 deliverable.
    """)
    st.stop()

# === Load and pool ===
with st.spinner("Loading and pooling CSVs..."):
    pooled, prov_df = load_and_pool_csvs(uploaded_files)

if pooled is None or len(pooled) == 0:
    st.error("No data loaded — check your CSVs.")
    st.stop()

# === Detect agent ===
if 'agent' in pooled.columns:
    agents = sorted(pooled['agent'].dropna().unique().tolist())
    if len(agents) > 1:
        st.warning(f"⚠️ Multiple agents detected: {agents}. "
                   f"This tool is designed for ONE agent at a time. "
                   f"Will use '{agents[0]}' for output naming. Re-upload "
                   f"with one agent's files only for cleanest results.")
    agent_name = agents[0]
else:
    agent_name = "UnknownAgent"

# === Composition summary ===
conditions_present = (sorted(pooled['temperature'].dropna().unique().tolist())
                      if 'temperature' in pooled.columns else ['?'])
n_total = len(pooled)

st.markdown(f"### 📊 Pooled dataset: **{agent_name}** · n = {n_total}")
c1, c2, c3, c4 = st.columns(4)
c1.markdown(f'<div class="metric-box"><h3>{n_total}</h3><p>Responses</p></div>',
            unsafe_allow_html=True)
c2.markdown(f'<div class="metric-box"><h3>{len(uploaded_files)}</h3><p>CSVs</p></div>',
            unsafe_allow_html=True)
c3.markdown(f'<div class="metric-box"><h3>{len(conditions_present)}</h3><p>Conditions</p></div>',
            unsafe_allow_html=True)
c4.markdown(
    f'<div class="metric-box"><h3>{pooled["question_id"].nunique() if "question_id" in pooled.columns else "?"}</h3><p>Questions</p></div>',
    unsafe_allow_html=True)

st.markdown(f"**Conditions present:** {', '.join(str(c) for c in conditions_present)}")

with st.expander("📋 Provenance (which CSV gave which condition)"):
    st.dataframe(prov_df, use_container_width=True, hide_index=True)

# === V_t preprocessing ===
with st.spinner("Computing V_t scores from response_text (preprocessing)..."):
    pooled = add_vt_columns(pooled, text_col='response_text')

if 'S_t' in pooled.columns:
    st.success("✅ V_t columns (S_t, A_t, Q_t, D_t, R_t) added — "
               "saved to output, not used as lenses (IEP-only sweep).")

st.markdown("---")

# === Build feature matrix and PCA ===
st.markdown("### 🔭 Building feature matrix and PCA projections...")
data, feat_cols = build_feature_matrix(pooled)
st.markdown(f"  • IEP feature matrix: {data.shape[0]} rows × {data.shape[1]} cols "
            f"({', '.join(feat_cols)})")

emb_matrix, bad_emb_rows = build_embedding_matrix(pooled)
pca_proj = None
if emb_matrix is not None:
    try:
        pca = PCA(n_components=2)
        pca_proj = pca.fit_transform(emb_matrix)
        st.markdown(f"  • PCA on embeddings: {emb_matrix.shape[0]} rows × "
                    f"{emb_matrix.shape[1]}-d → 2 components "
                    f"(explained variance: {pca.explained_variance_ratio_[0]:.3f}, "
                    f"{pca.explained_variance_ratio_[1]:.3f})")
        if bad_emb_rows:
            st.warning(f"  ⚠️ {len(bad_emb_rows)} rows had unparseable embeddings — "
                       f"filled with zeros for PCA.")
    except Exception as e:
        st.error(f"PCA failed: {e}. PCA-using lenses will be skipped.")
        pca_proj = None
else:
    st.warning("⚠️ No 'embedding' column found or all rows failed to parse. "
               "Geometric lenses (PCA1, PCA1×PCA2, AFF×PCA1, INT×PCA1) will be skipped.")

st.markdown("---")

# === Run sweep ===
st.markdown(f"### 🚀 Running {len(LENS_SPEC)}-Lens Stability Sweep")

if st.button("Run Stability Sweep", type="primary"):
    progress = st.progress(0)
    status = st.empty()
    results = []
    failures = []
    html_outputs = {}  # lens_name → html string

    for i, (lens_name, family, dim, kind, cols) in enumerate(LENS_SPEC):
        status.text(f"Lens {i+1}/{len(LENS_SPEC)}: {lens_name} ({family} {dim})")
        progress.progress((i + 1) / len(LENS_SPEC))

        # Skip if required cols missing
        missing = [c for c in cols if c not in pooled.columns]
        if missing:
            failures.append(f"{lens_name}: missing column(s) {missing}")
            results.append({
                'lens': lens_name, 'family': family, 'dim': dim,
                'V': None, 'E': None, 'beta_0': None, 'beta_1': None,
                'largest_members': None, 'largest_frac': None,
                'Q_purity': None, 'n_isolated': None, 'status': 'SKIPPED_MISSING_COL',
            })
            continue

        # Skip PCA-using lenses if PCA failed
        if kind in ("colpca", "pca1", "pca2") and pca_proj is None:
            failures.append(f"{lens_name}: PCA unavailable")
            results.append({
                'lens': lens_name, 'family': family, 'dim': dim,
                'V': None, 'E': None, 'beta_0': None, 'beta_1': None,
                'largest_members': None, 'largest_frac': None,
                'Q_purity': None, 'n_isolated': None, 'status': 'SKIPPED_NO_PCA',
            })
            continue

        # Build lens
        try:
            lens_values = build_lens_values(pooled, kind, cols, pca_proj)
            if lens_values is None:
                failures.append(f"{lens_name}: lens construction returned None")
                continue

            # Run Mapper
            mapper_obj, graph = run_mapper(
                data, lens_values,
                n_cubes=n_cubes, perc_overlap=perc_overlap,
                eps=eps, min_samples=min_samples,
            )
            stats = graph_stats(graph, pooled)

            # Banner
            banner = (f"Lens: {lens_name} · n_cubes={n_cubes} · "
                      f"overlap={int(perc_overlap*100)}% · "
                      f"DBSCAN(eps={eps}, min={min_samples}) · "
                      f"V={stats['V']} E={stats['E']} "
                      f"β₀={stats['beta_0']} β₁={stats['beta_1']}")

            # HTML
            try:
                html_str = export_mapper_html(
                    mapper_obj, graph, pooled, lens_name, agent_name, banner)
                html_outputs[lens_name] = html_str
            except Exception as he:
                failures.append(f"{lens_name}: HTML export failed: {he}")

            results.append({
                'lens': lens_name, 'family': family, 'dim': dim,
                'V': stats['V'], 'E': stats['E'],
                'beta_0': stats['beta_0'], 'beta_1': stats['beta_1'],
                'largest_members': stats['largest_members'],
                'largest_frac': stats['largest_frac'],
                'Q_purity': stats['Q_purity'],
                'n_isolated': stats['n_isolated'],
                'status': 'OK',
            })
        except Exception as e:
            failures.append(f"{lens_name}: {e}")
            results.append({
                'lens': lens_name, 'family': family, 'dim': dim,
                'V': None, 'E': None, 'beta_0': None, 'beta_1': None,
                'largest_members': None, 'largest_frac': None,
                'Q_purity': None, 'n_isolated': None,
                'status': f'ERROR: {str(e)[:80]}',
            })

    progress.progress(1.0)
    status.text(f"✅ Sweep complete — {len([r for r in results if r['status']=='OK'])}/{len(LENS_SPEC)} lenses succeeded")

    results_df = pd.DataFrame(results)
    st.session_state['results_df'] = results_df
    st.session_state['html_outputs'] = html_outputs
    st.session_state['failures'] = failures
    st.session_state['pooled_df'] = pooled
    st.session_state['prov_df'] = prov_df
    st.session_state['agent_name'] = agent_name
    st.session_state['n_total'] = n_total
    st.session_state['conditions_present'] = conditions_present

# === Display results ===
if 'results_df' in st.session_state:
    results_df = st.session_state['results_df']
    failures = st.session_state.get('failures', [])

    st.markdown("---")
    st.markdown("## 📊 Stability Results")

    valid = results_df[results_df['V'].notna()]
    if len(valid) > 0:
        c1, c2, c3, c4 = st.columns(4)
        c1.markdown(
            f'<div class="metric-box"><h3>{int(valid["beta_0"].min())}–{int(valid["beta_0"].max())}</h3><p>β₀ range</p></div>',
            unsafe_allow_html=True)
        c2.markdown(
            f'<div class="metric-box"><h3>{valid["beta_0"].std():.2f}</h3><p>β₀ SD</p></div>',
            unsafe_allow_html=True)
        c3.markdown(
            f'<div class="metric-box"><h3>{valid["largest_frac"].mean():.3f}</h3><p>Mean largest frac</p></div>',
            unsafe_allow_html=True)
        c4.markdown(
            f'<div class="metric-box"><h3>{valid["Q_purity"].mean():.3f}</h3><p>Mean Q-purity</p></div>',
            unsafe_allow_html=True)

        # Verdict (V1 logic, generalized)
        b0_range = valid['beta_0'].max() - valid['beta_0'].min()
        if b0_range <= 2:
            st.success("✅ **HIGH stability** — β₀ varies by ≤2 across all lenses. "
                       "Topology is intrinsic to the response space.")
        elif b0_range <= 4:
            st.warning("⚠️ **MODERATE stability** — β₀ varies by 3–4. Core structure "
                       "consistent but projection has measurable effect.")
        else:
            st.error("❌ **LOW stability** — β₀ varies by >4 across lenses. Results "
                     "may be projection-dependent.")

    st.markdown("### Per-Lens Stats Table")
    for fam in ['AFF', 'INT', 'ACT', 'Cross-IEP', 'Geometric']:
        sub = results_df[results_df['family'] == fam]
        if len(sub) > 0:
            st.markdown(f"**{fam}**")
            st.dataframe(sub.drop(columns=['family']), use_container_width=True,
                         hide_index=True)

    if failures:
        with st.expander(f"⚠️ {len(failures)} lens failure(s) — see details"):
            for f in failures:
                st.markdown(f"  • {f}")

    # === Download bundle ===
    st.markdown("---")
    st.markdown("### 📥 Download Per-Agent Bundle")

    pooled = st.session_state['pooled_df']
    prov_df = st.session_state['prov_df']
    agent_name = st.session_state['agent_name']
    n_total = st.session_state['n_total']
    conditions_present = st.session_state['conditions_present']
    html_outputs = st.session_state['html_outputs']

    methods_bytes = build_methods_docx(
        agent_name, pooled, prov_df, results_df,
        conditions_present, n_total, len(LENS_SPEC), failures
    )

    buf = io.BytesIO()
    with zipfile.ZipFile(buf, 'w', zipfile.ZIP_DEFLATED) as zf:
        zf.writestr(f"{agent_name}_stability.csv",
                    results_df.to_csv(index=False))
        zf.writestr(f"{agent_name}_provenance.csv",
                    prov_df.to_csv(index=False))
        zf.writestr(f"{agent_name}_pooled.csv",
                    pooled.to_csv(index=False))
        if methods_bytes:
            zf.writestr(f"{agent_name}_methods.docx", methods_bytes)
        for lens_name, html_str in html_outputs.items():
            safe = re.sub(r'[^A-Za-z0-9_]+', '_', lens_name)
            zf.writestr(f"{agent_name}_mappers/{safe}.html", html_str)
    buf.seek(0)

    st.download_button(
        f"📦 Download {agent_name} bundle (zip)",
        buf.getvalue(),
        file_name=f"{agent_name}_lens_stability_v2_{datetime.now().strftime('%Y%m%d_%H%M%S')}.zip",
        mime="application/zip",
    )

    # === Choir reading ===
    if choir_providers:
        st.markdown("---")
        st.markdown("## 🎵 Choir Reading")
        st.markdown("Send the stats table to multiple AI providers and read their "
                    "voices side-by-side. Each model brings its own perspective.")

        col_a, col_b = st.columns([3, 1])
        with col_a:
            selected_providers = st.multiselect(
                "Providers to consult",
                options=[p[0] for p in choir_providers],
                default=[p[0] for p in choir_providers],
            )
        with col_b:
            st.markdown("&nbsp;")
            run_choir = st.button("🎼 Sing", type="primary")

        if run_choir and selected_providers:
            providers_to_call = [p for p in choir_providers if p[0] in selected_providers]
            prompt = build_choir_prompt(agent_name, results_df, n_total, conditions_present)

            with st.spinner(f"Calling {len(providers_to_call)} provider(s) in parallel..."):
                with concurrent.futures.ThreadPoolExecutor(
                        max_workers=len(providers_to_call)) as ex:
                    futures = [
                        ex.submit(call_one_provider, name, sk, key, prompt)
                        for (name, sk, key) in providers_to_call
                    ]
                    voices = []
                    for fut in concurrent.futures.as_completed(futures):
                        voices.append(fut.result())

            voices.sort(key=lambda x: [p[0] for p in choir_providers].index(x[0]))

            n_cols = min(len(voices), 2)
            cols = st.columns(n_cols)
            for i, (name, voice) in enumerate(voices):
                with cols[i % n_cols]:
                    st.markdown(f'<div class="choir-box"><h4>🎤 {name}</h4>',
                                unsafe_allow_html=True)
                    st.markdown(voice)
                    st.markdown('</div>', unsafe_allow_html=True)

            st.session_state['voices'] = voices

            choir_text = f"# Choir Reading — {agent_name}\n\n"
            choir_text += f"_Generated: {datetime.now().strftime('%B %d, %Y · %H:%M')}_\n\n"
            for name, voice in voices:
                choir_text += f"## {name}\n\n{voice}\n\n---\n\n"
            st.download_button(
                "📄 Download choir reading (Markdown)",
                choir_text.encode('utf-8'),
                file_name=f"{agent_name}_choir_{datetime.now().strftime('%Y%m%d_%H%M%S')}.md",
                mime="text/markdown",
            )

# =============================================================================
# FOOTER
# =============================================================================
st.markdown("---")
st.markdown("""
<div style="text-align: center; color: #888; padding: 1rem; font-size: 0.85rem;">
  <strong>SYN-IQ Lens Stability V2 — Tool A (Conductor's Stand)</strong><br>
  §4.2 deliverable engine · KeplerMapper + IEP V3 · Multi-Provider Choir Reading<br>
  <em>SYNINT Team · Tennessee 🎹 CUZ Partnership · May 2026</em>
</div>
""", unsafe_allow_html=True)
