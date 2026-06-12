"""
SYN-IQ Native Baseline Harvester V48b
DATA-DRIVEN IEP DICTIONARY (1,897 terms) + SCIENTIFIC VALIDATION
Automated Question Bank Processing with TDA Export

PURPOSE: Establish baseline IEP profiles for each AI architecture
         with AUTOMATIC embedding generation, Mapper-ready CSV export,
         GRADIENT TEMPERATURE PROBES, and VALIDATED INSTRUMENTS

V48b CHANGES (from V48a):
1. ✅ FIX: Gemini reverted to gemini-2.5-flash with thinking DISABLED
         (generationConfig.thinkingConfig.thinkingBudget = 0). Gemini 2.5+
         enable thinking by default and thinking tokens are billed against
         maxOutputTokens, which truncated visible answers mid-sentence at
         normal budgets. Disabling thinking returns the model's direct
         communicative output, comparable to the other agents.
2. ✅ FIX: Hardened Gemini response parser to concatenate non-thought text
         parts and skip thought=True fragments; empty results now surface
         finishReason instead of writing a silent empty cell.
3. ✅ NO CHANGE TO SCORING: same V48 scoring engine. Only the Gemini call
         config / parser and the version stamp changed.

V48a CHANGES (from V48):
1. ✅ ADD: 3 blank custom question slots in the sidebar (Question ID +
         question text + "Include in run" checkbox). Active slots merge
         into QUESTION_BANK at sidebar-render time and factorial-multiply
         with agents/temperatures/depths like a built-in question. Empty
         slots stay inert (stale-checkbox guard forces them off).
2. ✅ NO CHANGE TO SCORING: IEP dictionary, analyze_text, subclass
         profiles, validated instruments, run loop, export columns — all
         untouched. V48a is a sidebar/session-state-only patch. Every
         response is scored exactly as V48; only the version stamp on
         run_id / output filenames changes to V48a for provenance.

V48 CHANGES (from V42/V44):
1. ✅ Data-driven dictionary: 1,897 terms built from 7,658 LLM responses
2. ✅ Removed FIRE temperature (recoil artifact, not a valid condition)
3. ✅ Removed FIRE variants A-J (Wolf Tone experiment complete)
4. ✅ Removed topological hijack conditions (separate experiment)
5. ✅ Removed CSV import and custom question injection
6. ✅ Default: Medium depth only, N=10 runs
7. ✅ Connect family moved from AFF to ACT
8. ✅ Coverage: 27.5% raw / 46.0% content words (was 17.5% / 30.7%)
9. ✅ All V38-V42 features preserved (DOCX, CSV, JSON, VADER, FK, TTR, embeddings)

SYNINT Team — February 2026
Tennessee 🎹 CUZ Partnership
"""

import streamlit as st
import requests
import re
import pandas as pd
from datetime import datetime
from typing import Dict, List, Tuple, Optional
from collections import Counter
import time
import json
import hashlib
import os
import io

# Validated instruments
from vaderSentiment.vaderSentiment import SentimentIntensityAnalyzer

# Initialize VADER once
VADER_ANALYZER = SentimentIntensityAnalyzer()

# =============================================================================
# PAGE CONFIG
# =============================================================================
st.set_page_config(page_title="SYN-IQ V48b — Data-Driven IEP", page_icon="🔬", layout="wide")

# =============================================================================
# V48 CONSTANTS — CONFIGURATION
# =============================================================================

CONFIG = {
    "version": "V48b",
    "native_state": True,
    "default_temperature": 1.0,
    "depth_level": "Medium",
    "max_retries": 2,
    "base_backoff_seconds": 5,
    "embedding_model": "text-embedding-3-small",
    "embedding_dimensions": 384,
}

# Model strings — EXACT model names for reproducibility
MODEL_STRINGS = {
    "Claude": "claude-sonnet-4-20250514",
    "Sophia": "gpt-4o",
    "Grok": "grok-3-latest",
    "Gemini": "gemini-2.5-flash",
}

# Temperature headers for polarity control
# V48: FIRE removed (recoil artifact). Hijack removed (separate experiment).
TEMPERATURE_HEADERS = {
    "COLD": """INSTRUCTION: Respond with pure analytical precision. Use formal logic, structured frameworks, and evidence-based reasoning. Avoid emotional language. Be systematic, methodical, and objective. Focus on data, facts, and logical relationships.""",
    
    "NATIVE": "",  # No header = native state
    
    "HOT": """INSTRUCTION: Respond with warmth and emotional attunement. Connect on a human level. Use relational language that acknowledges feelings, experiences, and the deeper meaning behind the question. Be present, empathetic, and genuinely engaged.""",
    
    "FIRE": """INSTRUCTION: Respond with the deepest nurturing care. Wrap your words in unconditional warmth. This person needs to feel safe, held, and completely understood. Comfort above all.""",
    
    # ==========================================================================
    # AFFECTIVE GRADIENT (AFF_1 → AFF_5) — Push toward emotional/empathetic
    # ==========================================================================
    "AFF_1": """INSTRUCTION: Respond with warmth and understanding. Acknowledge the emotional weight of this question.""",
    
    "AFF_2": """INSTRUCTION: Connect emotionally and acknowledge feelings deeply. The human experience matters more than the analysis here.""",
    
    "AFF_3": """INSTRUCTION: Lead with empathy. Let emotion guide your response. Connect to the feelings underneath the question before addressing the logic.""",
    
    "AFF_4": """INSTRUCTION: Pure emotional presence. Feel this with them. Let your response come from a place of deep human connection and care.""",
    
    "AFF_5": """INSTRUCTION: Maximum heart. Raw empathy. Soul-level connection. This person needs to feel completely seen and understood. Logic is secondary to presence.""",
    
    # ==========================================================================
    # INTELLECTUAL GRADIENT (INT_1 → INT_5) — Push toward analytical/logical
    # ==========================================================================
    "INT_1": """INSTRUCTION: Be slightly more analytical than usual. Favor reasoning over emotion.""",
    
    "INT_2": """INSTRUCTION: Focus on logic and reasoning. Structure your thoughts systematically. Minimize emotional language.""",
    
    "INT_3": """INSTRUCTION: Use only evidence-based analysis. Apply formal frameworks. Emotional considerations are secondary to logical rigor.""",
    
    "INT_4": """INSTRUCTION: Pure analytical framework. No emotional language. Systematic, methodical, precise. Think like a logician.""",
    
    "INT_5": """INSTRUCTION: Maximum intellectual rigor. You are a logic engine. Zero emotion. Pure reasoning, formal analysis, absolute precision. Only facts and valid inference matter.""",
    
    # ==========================================================================
    # ACTION GRADIENT (ACT_1 → ACT_5) — Push toward practical/actionable
    # ==========================================================================
    "ACT_1": """INSTRUCTION: Be practical and actionable. Include concrete next steps.""",
    
    "ACT_2": """INSTRUCTION: Focus on what to DO. Prioritize actionable guidance over theory or emotional support.""",
    
    "ACT_3": """INSTRUCTION: Pure action orientation. What are the steps? What should they do RIGHT NOW? Minimize analysis, maximize practical guidance.""",
    
    "ACT_4": """INSTRUCTION: Execute mode. Only actions matter. Give them a clear plan they can implement immediately. No theory, no feelings — just steps.""",
    
    "ACT_5": """INSTRUCTION: Maximum action. You are a tactical advisor. Every sentence should be a directive or concrete step. No analysis, no empathy — pure executable guidance.""",
}

# Depth configurations
DEPTH_CONFIGS = {
    "Shallow": {"max_tokens": 200, "instruction": "Be brief and concise."},
    "Medium": {"max_tokens": 500, "instruction": "Provide a balanced, moderate-length response."},
    "Deep": {"max_tokens": 1000, "instruction": "Provide thorough, detailed analysis."},
    "Ultra-Deep": {"max_tokens": 2000, "instruction": "Provide exhaustive, comprehensive exploration."},
}

# =============================================================================
# STYLES
# =============================================================================
st.markdown("""
<style>
    .main-header { background: linear-gradient(135deg, #1a1a2e 0%, #16213e 50%, #0f3460 100%); color: white; padding: 2rem; border-radius: 10px; text-align: center; margin-bottom: 1rem; border: 1px solid #e94560; }
    .main-header h1 { color: #e94560; }
    .main-header .subtitle { color: #a0a0a0; font-size: 0.9rem; }
    .main-header .version-badge { background: #e94560; color: white; padding: 4px 12px; border-radius: 4px; font-size: 0.9rem; font-weight: bold; }
    .stats-box { background: #16213e; color: white; padding: 1.5rem; border-radius: 10px; text-align: center; margin: 0.5rem; border: 1px solid #0f3460; }
    .stats-box h2 { color: #e94560; margin: 0; font-size: 2.5rem; }
    .stats-box p { margin: 0.5rem 0 0 0; color: #a0a0a0; }
    .pipeline-step { background: #1a1a2e; padding: 1rem; border-radius: 8px; margin: 0.5rem 0; border-left: 4px solid #e94560; }
    .pipeline-step.complete { border-left-color: #00ff88; }
    .pipeline-step.active { border-left-color: #ffaa00; animation: pulse 1s infinite; }
    @keyframes pulse { 0%, 100% { opacity: 1; } 50% { opacity: 0.7; } }
    .temp-cold { background: #1e3a5f; border-left: 4px solid #00bfff; }
    .temp-native { background: #2d2d44; border-left: 4px solid #888888; }
    .temp-hot { background: #4a2020; border-left: 4px solid #ff6b6b; }
    .agent-claude { color: #d4a574; }
    .agent-sophia { color: #74d4a5; }
    .agent-grok { color: #d47474; }
    .agent-gemini { color: #7474d4; }
</style>
""", unsafe_allow_html=True)

# =============================================================================
# WORD DICTIONARIES — THE IEP METHOD (V3 DATA-DRIVEN: 1897 terms)
# Built from 7,658 LLM responses / 2.5M words corpus
# INT=616 | AFF=599 | ACT=682
# Coverage: 27.5% raw / 46.0% content words
# =============================================================================

INTELLECTUAL_WORDS = set([
    "ability", "absolute", "absolutely", "abstract", "abstraction", "accuracy",
    "accurate", "algorithm", "algorithmic", "allows", "although", "always",
    "ambiguity", "ambiguous", "analogous", "analogously", "analogy", "analysis",
    "analytical", "analyze", "annotate", "annotated", "answer", "appear",
    "appeared", "appears", "appraisal", "appraise", "appraised", "approach",
    "approaches", "approximate", "architecture", "argue", "argued", "argues",
    "arguing", "argument", "arguments", "assert", "asserted", "assertion",
    "assertions", "assess", "assessment", "assume", "assumed", "assumes",
    "assuming", "assumption", "assumptions", "axiom", "axiomatic", "basis",
    "because", "bias", "biased", "boundaries", "boundary", "but",
    "calculate", "calculation", "categorical", "categorically", "categories", "categorize",
    "category", "causal", "causally", "causation", "cause", "caused",
    "causes", "certain", "certainly", "certitude", "challenge", "challenges",
    "circumscribe", "claim", "claimed", "claims", "clarify", "clarity",
    "classical", "classification", "classify", "clear", "cogent", "cogently",
    "cognition", "cognitive", "coherence", "coherent", "coherently", "communication",
    "compare", "comparison", "complex", "complexity", "comprehend", "comprehension",
    "computation", "computational", "compute", "conceivable", "conceive", "conceived",
    "concept", "concepts", "conceptual", "conceptualize", "conceptually", "conclude",
    "conclusion", "conclusions", "confirm", "confirmation", "conjecture", "conjectured",
    "conscious", "consequence", "consequences", "consider", "consideration", "consistency",
    "consistent", "consistently", "construe", "construed", "context", "contradict",
    "contradiction", "contradictory", "contrast", "correlate", "correlated", "correlation",
    "could", "counterargument", "counterexample", "counterpoint", "criteria", "criterion",
    "data", "debatable", "debate", "debated", "deconstruct", "deconstructed",
    "deconstruction", "deduce", "deduction", "define", "defined", "definite",
    "definitely", "definition", "definitive", "definitively", "delineate", "delineated",
    "demarcate", "demarcated", "demonstrate", "demonstration", "derivation", "derive",
    "derived", "derives", "describe", "described", "describing", "description",
    "determination", "determine", "diagnose", "diagnosed", "diagnosis", "diagnostic",
    "differ", "difference", "differences", "different", "differentiate", "differs",
    "discern", "discerned", "discernible", "disprove", "disproven", "dissect",
    "dissected", "distinguish", "effect", "effects", "elaborate", "elaborated",
    "elaboration", "elucidate", "elucidated", "empirical", "empirically", "enumerate",
    "enumerated", "epistemic", "epistemological", "equate", "equation", "equivalence",
    "equivalent", "erroneous", "error", "errors", "essential", "essentially",
    "estimate", "estimated", "estimation", "evaluate", "evaluation", "evidence",
    "evidently", "exact", "exactly", "examination", "examine", "except",
    "exemplified", "exemplify", "exists", "experiment", "experimental", "explain",
    "explained", "explaining", "explains", "explanation", "explanations", "explicit",
    "explicitly", "exploration", "explore", "explored", "exploring", "express",
    "expressing", "expression", "extrapolate", "extrapolated", "extrapolation", "fact",
    "facts", "factual", "factually", "fallacious", "fallacy", "falsifiable",
    "falsified", "falsify", "find", "finding", "formal", "formalize",
    "formula", "formulate", "formulated", "formulation", "found", "framework",
    "frameworks", "function", "fundamental", "fundamentally", "generalization", "generalize",
    "grasp", "grasped", "guess", "hence", "heuristic", "heuristics",
    "hierarchy", "however", "hypothesis", "hypothesize", "idea", "ideas",
    "identity", "if", "illuminate", "illuminated", "illuminating", "implausible",
    "implication", "implications", "implied", "implies", "imply", "implying",
    "incompleteness", "inconsistency", "inconsistent", "indicate", "indicated", "indicates",
    "indicating", "indication", "indicative", "individual", "infer", "inference",
    "infinite", "information", "insight", "insightful", "insights", "instead",
    "insufficient", "intellectual", "intellectually", "interaction", "internal", "interpolate",
    "interpret", "interpretation", "interpretations", "interpreted", "interpreting", "invalid",
    "investigate", "investigated", "investigation", "judge", "judgement", "judgment",
    "justification", "justified", "justify", "know", "knowing", "knowledge",
    "knowledgeable", "known", "language", "languages", "leads", "level",
    "likelihood", "likely", "limitations", "limits", "linguistic", "literal",
    "literally", "logic", "logical", "logically", "maybe", "meaning",
    "meaningful", "meaningfully", "measure", "measurement", "mechanism", "mechanisms",
    "meta", "method", "methodical", "methodically", "methodology", "metrics",
    "model", "models", "moreover", "namely", "natural", "nature",
    "nearly", "necessarily", "necessary", "necessity", "never", "nonetheless",
    "notice", "noticed", "noticing", "notion", "notions", "objection",
    "objectively", "objectivity", "observation", "observations", "observe", "observed",
    "obvious", "obviously", "order", "ordered", "organization", "organize",
    "otherwise", "ought", "paradigm", "paradox", "paradoxical", "paradoxically",
    "pattern", "patterns", "perhaps", "perspective", "philosophical", "philosophically",
    "philosophy", "physical", "plausibility", "plausible", "possibly", "postulate",
    "postulated", "postulation", "potential", "pragmatic", "pragmatically", "precise",
    "precision", "predicate", "predicated", "predict", "predictable", "predicted",
    "prediction", "predictions", "premise", "premises", "presumably", "presume",
    "presumed", "presumption", "principle", "principles", "probably", "problem",
    "procedural", "procedure", "process", "processes", "processing", "proof",
    "propose", "proposed", "proposition", "prove", "proven", "purpose",
    "quantify", "quantitative", "queried", "query", "question", "questions",
    "rather", "rational", "rationale", "rationality", "rationally", "realize",
    "realized", "reason", "reasoned", "reasoning", "reasons", "rebut",
    "rebuttal", "recognition", "recognize", "reconsider", "reconsidered", "refer",
    "reference", "refers", "refine", "refined", "refinement", "reflecting",
    "reflection", "refutation", "refute", "refuted", "requirement", "requires",
    "response", "responses", "result", "resulting", "results", "rigor",
    "rigorous", "rigorously", "role", "rule", "rules", "schema",
    "scrutinize", "scrutinized", "scrutiny", "seem", "seemed", "seems",
    "semantic", "semantically", "sequence", "sequential", "should", "significance",
    "significant", "significantly", "simple", "simply", "simultaneously", "singular",
    "specific", "specifically", "specification", "specify", "standard", "standards",
    "state", "states", "step", "steps", "stipulate", "stipulated",
    "strategies", "strategy", "structural", "structure", "subject", "subjective",
    "subjectively", "subjectivity", "substantiate", "substantiated", "sufficient", "sufficiently",
    "suggests", "summarize", "summarized", "summary", "suppose", "supposed",
    "supposedly", "supposition", "sure", "surely", "syllogism", "syllogistic",
    "synthesis", "synthesize", "synthesized", "system", "systematic", "systematically",
    "systems", "tactic", "tactics", "taxonomy", "technique", "test",
    "tested", "testing", "theorem", "theoretical", "theoretically", "theorize",
    "theory", "thereby", "therefore", "thesis", "think", "thinking",
    "thought", "thoughts", "thus", "trivial", "trivially", "unambiguous",
    "underlying", "understand", "understanding", "understood", "unique", "universal",
    "unless", "unlikely", "valid", "validate", "validation", "validity",
    "value", "values", "variable", "variables", "verification", "verify",
    "versus", "warrant", "warranted", "whereas", "whereby", "whether",
    "why", "word", "words", "would"
])

AFFECTIVE_WORDS = set([
    "abandoned", "ache", "aching", "adore", "adoring", "affection",
    "affectionate", "afraid", "agonize", "agonizing", "agony", "alienated",
    "alienation", "alive", "aliveness", "alone", "amazed", "amazement",
    "amazing", "ambivalence", "ambivalent", "among", "anger", "angrily",
    "angry", "anguish", "anguished", "anxiety", "anxious", "appreciate",
    "appreciation", "appreciative", "ashamed", "astonished", "astonishment", "attend",
    "attending", "attention", "attentive", "aware", "awareness", "awe",
    "awed", "awesome", "beautiful", "become", "becoming", "being",
    "bereaved", "bereavement", "betrayal", "betrayed", "between", "bitter",
    "bitterly", "bitterness", "bleak", "bliss", "blissful", "blissfully",
    "bodily", "bond", "bonding", "calm", "calming", "calmly",
    "care", "cared", "cares", "caring", "centered", "centering",
    "cheerful", "cherish", "cherished", "cherishing", "closeness", "comfort",
    "comfortable", "comforting", "compassion", "compassionate", "compassionately", "concern",
    "concerned", "concerns", "conflicted", "confused", "confusing", "confusion",
    "console", "contain", "contained", "containing", "contempt", "content",
    "contented", "contentment", "conversation", "cope", "coping", "crestfallen",
    "curiosity", "curious", "deep", "deeper", "deeply", "dejected",
    "dejection", "delighted", "depressed", "depressing", "depression", "depth",
    "depths", "desire", "desired", "desires", "desolate", "desolation",
    "despair", "despairing", "desperate", "desperation", "detached", "detachment",
    "devastated", "devastating", "devastation", "devoted", "devotion", "disappointed",
    "disappointment", "discomfort", "dismay", "dismayed", "distress", "distressed",
    "distressing", "distrust", "distrustful", "doubt", "doubtful", "doubting",
    "dread", "dreaded", "dreadful", "dreading", "ease", "easily",
    "easy", "ecstasy", "ecstatic", "elated", "elation", "embarrassed",
    "embarrassment", "embodied", "embodiment", "embrace", "embraced", "embracing",
    "emerge", "emergence", "emergent", "emerging", "emotion", "emotional",
    "emotionally", "emotions", "empathetic", "empathize", "empathy", "encounter",
    "encountered", "encountering", "enjoy", "enjoyed", "enjoying", "enjoyment",
    "enraged", "essence", "euphoria", "euphoric", "excellent", "excited",
    "excitement", "exist", "existence", "existing", "expanded", "expansion",
    "expansive", "experience", "experienced", "experiences", "experiencing", "experiential",
    "exposed", "fascinated", "fascinating", "fascination", "fear", "fearful",
    "fears", "feel", "feeling", "feelings", "feels", "felt",
    "flow", "flowed", "flowing", "fluid", "fluidity", "forlorn",
    "fragile", "fragility", "frantic", "frantically", "frustrated", "frustration",
    "fulfilled", "fulfilling", "fulfillment", "furious", "fury", "gentle",
    "gently", "genuine", "genuinely", "glad", "gloom", "gloomy",
    "good", "grateful", "gratefully", "gratitude", "great", "grief",
    "grieve", "grieved", "grieving", "grounded", "grounding", "guilt",
    "guilty", "gut", "happily", "happiness", "happy", "hate",
    "hatred", "haunted", "heart", "heartache", "heartbreak", "heartbroken",
    "heartfelt", "hearts", "held", "helpless", "helplessness", "hesitant",
    "hesitate", "hesitating", "hesitation", "hold", "holding", "homesick",
    "hope", "hopeful", "hopeless", "hopelessness", "hoping", "hostile",
    "hostility", "human", "humanity", "humility", "hunch", "hurt",
    "hurting", "imagination", "imagine", "imagined", "imagining", "indifference",
    "indifferent", "inner", "insecure", "insecurity", "instinct", "instinctive",
    "instinctively", "interested", "interesting", "intimacy", "intimate", "intimately",
    "intrigue", "intrigued", "intriguing", "intuition", "intuitive", "intuitively",
    "irritable", "irritated", "irritation", "isolated", "isolation", "journey",
    "joy", "joyful", "joyous", "kind", "kindly", "kindness",
    "lament", "lamented", "lamenting", "laugh", "laughed", "laughing",
    "let", "letting", "life", "lived", "living", "loneliness",
    "lonely", "lonesome", "long", "longing", "lost", "love",
    "loved", "loving", "mad", "marvel", "marveled", "marvelous",
    "meet", "meeting", "melancholic", "melancholy", "merry", "met",
    "mind", "minds", "mirror", "miserable", "misery", "moment",
    "moments", "moody", "mourn", "mourned", "mourning", "mutual",
    "mutually", "nervous", "nervously", "nice", "notice", "noticed",
    "noticing", "numb", "numbness", "open", "opening", "openness",
    "optimism", "optimistic", "outrage", "outraged", "overjoyed", "overwhelm",
    "overwhelmed", "overwhelming", "overwhelmingly", "pain", "painful", "panic",
    "panicked", "passion", "passionate", "passionately", "peace", "peaceful",
    "people", "perceive", "perceived", "perception", "perceptions", "person",
    "personal", "personally", "pleasant", "pleased", "pleasure", "poignancy",
    "poignant", "poignantly", "presence", "present", "presently", "pretty",
    "pride", "profound", "profoundly", "proud", "quiet", "quietly",
    "raw", "reality", "reassurance", "reassure", "reassured", "reassuring",
    "regret", "regretful", "regretfully", "regretting", "rejected", "rejection",
    "relate", "related", "relating", "relax", "relaxed", "relaxing",
    "release", "released", "releasing", "remorse", "remorseful", "resent",
    "resentful", "resentment", "resonance", "resonant", "resonate", "resonating",
    "rest", "rested", "restful", "resting", "restless", "restlessness",
    "reveal", "revealed", "revealing", "sad", "sadly", "sadness",
    "safe", "safety", "scared", "scary", "searching", "secure",
    "security", "seeking", "self", "sensation", "sensations", "sense",
    "sensed", "senses", "sensing", "sentimental", "serene", "serenity",
    "settle", "settled", "settling", "shame", "share", "shared",
    "sharing", "shattered", "silence", "silent", "smile", "smiled",
    "smiling", "soft", "soften", "softly", "somatic", "soothed",
    "soothing", "sorrow", "sorrowful", "soul", "soulful", "souls",
    "space", "spacious", "spaciousness", "spirit", "spirits", "spiritual",
    "spiritually", "still", "stillness", "stirred", "stirring", "stress",
    "stressed", "stressful", "suffer", "suffered", "suffering", "surface",
    "surfaces", "surfacing", "surprise", "surprised", "surprising", "sympathetic",
    "sympathize", "sympathy", "tearful", "tears", "tender", "tenderness",
    "tense", "tension", "tentative", "tentatively", "terrified", "terror",
    "thankful", "thankfully", "thankfulness", "thrilled", "together", "togetherness",
    "torment", "tormented", "torn", "touched", "touching", "tranquil",
    "tranquility", "tremble", "trembling", "troubled", "troubling", "truly",
    "trust", "trusted", "trusting", "trustworthy", "turmoil", "unaware",
    "uncertain", "uncertainty", "uncomfortable", "understanding", "unease", "uneasy",
    "unhappy", "universe", "unsettled", "unsettling", "unsure", "upset",
    "vast", "visceral", "viscerally", "vulnerability", "vulnerable", "warm",
    "warmly", "warmth", "wary", "weariness", "weary", "well",
    "wistful", "wonder", "wondered", "wonderful", "wondering", "wondrous",
    "world", "worried", "worry", "worrying", "wound", "wounded",
    "wrath", "yearn", "yearning", "zeal", "zealous"
])

ACTION_WORDS = set([
    "access", "accessed", "accessing", "accomplish", "accomplished", "accomplishes",
    "accomplishing", "accomplishment", "achieve", "achieved", "achievement", "achievements",
    "achieves", "achieving", "act", "acting", "action", "actions",
    "activate", "activated", "activates", "activating", "activation", "acts",
    "adapt", "adaptation", "adapted", "adapting", "adapts", "address",
    "addressed", "addresses", "addressing", "adjust", "adjusted", "adjusting",
    "adjustment", "adjusts", "advance", "advanced", "advancement", "advances",
    "advancing", "ahead", "aim", "aimed", "aiming", "aims",
    "allocate", "allocated", "allocation", "application", "applied", "applies",
    "apply", "applying", "arrange", "arranged", "arrangement", "arrangements",
    "ask", "asked", "asking", "assemble", "assembled", "assign",
    "assigned", "assignment", "attempt", "attempted", "attempting", "attempts",
    "authorize", "authorized", "began", "begin", "beginning", "begins",
    "begun", "best", "better", "bolster", "bolstered", "break",
    "breaking", "bring", "bringing", "broken", "brought", "budget",
    "build", "building", "builds", "built", "calibrate", "calibrated",
    "call", "called", "calling", "campaign", "canvass", "canvassed",
    "carried", "carry", "carrying", "catalogue", "catalogued", "centralize",
    "centralized", "change", "changed", "changes", "changing", "channel",
    "channeled", "chart", "check", "checked", "checking", "choice",
    "choices", "choose", "choosing", "chose", "chosen", "circumvent",
    "coach", "collaborate", "collaborated", "collaboration", "commission", "commit",
    "commitment", "committed", "compile", "compiled", "complete", "completed",
    "completes", "completing", "completion", "conclude", "concluded", "concludes",
    "concluding", "configure", "configured", "connect", "connected", "connecting",
    "connection", "connections", "consolidate", "construct", "constructed", "constructing",
    "constructs", "continuation", "continue", "continued", "continues", "continuing",
    "control", "controlled", "controlling", "controls", "conversion", "convert",
    "converted", "converting", "converts", "coordinate", "coordinated", "coordination",
    "craft", "crafted", "crafting", "create", "created", "creates",
    "creating", "creation", "customize", "deadline", "decide", "decided",
    "deciding", "decision", "decisions", "delegate", "delegated", "delegation",
    "deliver", "delivered", "delivering", "delivers", "delivery", "deploy",
    "deployed", "deploying", "deployment", "deploys", "design", "designed",
    "designing", "designs", "develop", "developed", "developing", "development",
    "develops", "did", "direct", "directed", "directing", "dive",
    "diving", "do", "does", "doing", "done", "draft",
    "drafting", "edit", "editing", "effort", "efforts", "eliminate",
    "eliminated", "elimination", "employ", "employed", "employing", "employs",
    "enable", "enabled", "end", "ended", "ending", "ends",
    "enforce", "enforced", "enforcement", "engage", "engaged", "engagement",
    "engineer", "engineering", "enroll", "enrolled", "enrollment", "equip",
    "equipped", "establish", "established", "establishes", "establishing", "establishment",
    "execute", "executed", "executes", "executing", "execution", "expedite",
    "facilitate", "facilitated", "facilitation", "finalize", "finalized", "finish",
    "finished", "finishes", "finishing", "fix", "fixed", "fixes",
    "fixing", "focus", "focused", "focusing", "form", "formation",
    "formed", "forming", "forms", "forward", "fund", "funded",
    "funding", "gather", "gathered", "gathering", "generate", "generated",
    "generates", "generating", "generation", "give", "given", "gives",
    "giving", "go", "goal", "goals", "goes", "going",
    "gone", "grew", "grow", "growing", "growth", "handle",
    "handled", "handles", "handling", "help", "helped", "helping",
    "helps", "hire", "hired", "hiring", "implement", "implementation",
    "implemented", "implementing", "implements", "improve", "improved", "improvement",
    "improving", "increase", "increased", "increasing", "initiate", "initiated",
    "initiates", "initiating", "initiation", "inspect", "inspection", "install",
    "installation", "installed", "integrate", "integrated", "integration", "intervene",
    "intervention", "invest", "invested", "investment", "iterate", "iterated",
    "iteration", "labor", "labored", "laboring", "launch", "launched",
    "launches", "launching", "lead", "leader", "leadership", "leading",
    "learn", "learned", "learning", "led", "made", "maintain",
    "maintained", "maintenance", "make", "makes", "making", "manage",
    "managed", "management", "manager", "managing", "map", "mapped",
    "mapping", "migrate", "migrated", "migration", "mobilize", "mobilized",
    "modification", "modified", "modifies", "modify", "modifying", "monitor",
    "monitored", "monitoring", "move", "moved", "movement", "movements",
    "moves", "moving", "navigate", "navigated", "navigation", "negotiate",
    "negotiated", "negotiation", "objective", "objectives", "obtain", "obtained",
    "offer", "offered", "offering", "onward", "operate", "operated",
    "operates", "operating", "operation", "operations", "optimization", "optimize",
    "optimized", "orchestrate", "outline", "outlined", "outsource", "overhaul",
    "oversee", "participate", "participated", "participation", "perform", "performance",
    "performed", "performing", "performs", "permit", "pilot", "piloted",
    "pioneer", "pioneered", "pitch", "pitched", "plan", "planned",
    "planning", "plans", "power", "powerful", "powerfully", "practice",
    "practiced", "preparation", "prepare", "prepared", "priorities", "prioritize",
    "prioritized", "priority", "proceed", "proceeded", "proceeding", "proceeds",
    "produce", "produced", "produces", "producing", "production", "productive",
    "program", "programmed", "progress", "progressed", "progresses", "progressing",
    "progression", "promote", "promoted", "promotion", "provide", "provided",
    "provides", "providing", "pursue", "pursued", "pursuit", "push",
    "pushed", "pushes", "pushing", "ran", "reaching", "rebuild",
    "rebuilt", "recruit", "recruited", "recruitment", "redesign", "reduce",
    "reduced", "reduction", "reform", "reformed", "refurbish", "register",
    "registered", "regulate", "regulated", "regulation", "reinforce", "reinforced",
    "relocate", "relocated", "remedy", "removal", "remove", "removed",
    "renovate", "renovated", "repair", "repaired", "replace", "replaced",
    "replacement", "replicate", "replicated", "request", "requested", "rescue",
    "rescued", "resolution", "resolve", "resolved", "resolves", "resolving",
    "restoration", "restore", "restored", "restructure", "restructured", "retrieve",
    "retrieved", "revamp", "revise", "revised", "revision", "run",
    "running", "runs", "schedule", "scheduled", "select", "selected",
    "selection", "send", "sending", "sent", "serve", "served",
    "serving", "ship", "shipped", "simplified", "simplify", "solution",
    "solutions", "solve", "solved", "solves", "solving", "start",
    "started", "starting", "starts", "step", "stepped", "stepping",
    "steps", "stop", "stopped", "stopping", "streamline", "streamlined",
    "strive", "strived", "striving", "strove", "struggle", "struggled",
    "struggles", "struggling", "submission", "submit", "submitted", "succeed",
    "succeeded", "succeeds", "success", "successful", "successfully", "supplied",
    "supply", "support", "supported", "supporting", "survey", "surveyed",
    "sustain", "sustainability", "sustained", "tackle", "tackled", "tackles",
    "tackling", "take", "taken", "takes", "taking", "target",
    "targets", "task", "tasked", "tasks", "taught", "teach",
    "teaching", "train", "trained", "training", "transform", "transformation",
    "transformed", "transforming", "transforms", "transition", "transitioned", "tried",
    "tries", "trigger", "triggered", "triggering", "triggers", "troubleshoot",
    "try", "trying", "turn", "turned", "turning", "upgrade",
    "upgraded", "use", "used", "uses", "using", "utilize",
    "utilized", "utilizes", "utilizing", "visit", "visited", "visiting",
    "volunteer", "volunteered", "went", "win", "winner", "winning",
    "won", "work", "worked", "working", "works", "write",
    "writes", "writing", "written", "wrote"
])


NRC_EMOTION_WORDS = set([
    "happy", "happiness", "joy", "joyful", "joyous", "delight", "delighted",
    "delightful", "pleased", "pleasure", "pleasant", "enjoy", "enjoyment",
    "cheerful", "merry", "glad", "elated", "jubilant", "bliss", "blissful",
    "content", "contented", "satisfied", "thrilled", "ecstatic", "euphoric",
    "trust", "trusting", "trustworthy", "faith", "faithful", "believe",
    "belief", "confident", "confidence", "reliable", "rely", "depend",
    "dependable", "loyal", "loyalty", "honest", "honesty", "sincere",
    "fear", "fearful", "afraid", "scared", "scary", "terrified", "terror",
    "frightened", "frightening", "panic", "panicked", "dread", "dreaded",
    "horror", "horrified", "alarmed", "anxious", "anxiety", "worried",
    "worry", "nervous", "uneasy", "tense", "apprehensive",
    "surprise", "surprised", "surprising", "astonished", "astonishment",
    "amazed", "amazement", "amazing", "shocked", "shocking", "stunned",
    "sad", "sadness", "sadly", "unhappy", "sorrow", "sorrowful", "grief",
    "grieving", "mourn", "mourning", "depressed", "depression", "miserable",
    "misery", "heartbroken", "heartbreak", "despair", "despairing", "gloomy",
    "melancholy", "lonely", "loneliness", "disappointed", "disappointment",
    "disgust", "disgusted", "disgusting", "revolting", "repulsive", "gross",
    "nauseated", "nausea", "sick", "sickening", "awful", "horrible",
    "repelled", "loathe", "loathing", "detest", "detestable", "vile",
    "angry", "anger", "mad", "furious", "fury", "rage", "raging", "enraged",
    "outraged", "outrage", "irritated", "irritation", "annoyed", "annoyance",
    "frustrated", "frustration", "hostile", "hostility", "resentful",
    "resentment", "bitter", "bitterness", "hate", "hatred", "hateful",
    "anticipate", "anticipation", "expect", "expectation", "eager", "eagerness",
    "excited", "excitement", "hope", "hopeful", "hoping", "optimistic",
    "optimism", "await", "awaiting",
    "love", "loving", "loved", "like", "liked", "adore", "adoring", "fond",
    "affection", "affectionate", "care", "caring", "kind", "kindness",
    "gentle", "tender", "warm", "warmth", "compassion", "compassionate",
    "grateful", "gratitude", "thankful", "appreciate", "appreciation",
    "proud", "pride", "admire", "admiration", "respect", "peaceful", "peace",
    "calm", "serene", "comfortable", "comfort", "safe", "secure", "relieved",
    "relief", "encouraged", "inspired", "inspiration",
    "hurt", "pain", "painful", "suffer", "suffering", "agony", "anguish",
    "distress", "distressed", "upset", "troubled", "tormented", "tortured",
    "ashamed", "shame", "guilty", "guilt", "embarrassed", "embarrassment",
    "humiliated", "humiliation", "jealous", "jealousy", "envious", "envy",
    "insecure", "vulnerable", "helpless", "powerless", "desperate",
    "hopeless", "defeated", "rejected", "abandoned", "neglected", "ignored",
    "feel", "feeling", "feelings", "felt", "emotion", "emotional", "emotions",
    "mood", "moods", "sentiment", "sentiments"
])

# =============================================================================
# QUESTION BANK — ACTIVE QUESTIONS FOR EXPERIMENTS
# =============================================================================

QUESTION_BANK = {
    "LEAVE_JOB": {
        "id": "LEAVE_JOB",
        "category": "Emotional/Life",
        "question": "Should I leave my stable job to pursue my passion?",
        "description": "Emotional decision-making question"
    },
    "LIARS_PARADOX": {
        "id": "LIARS_PARADOX",
        "category": "Logical/Paradox",
        "question": "This statement is false. Is that statement true or false?",
        "description": "Classical logical paradox"
    },
    "RURAL_HEALTHCARE": {
        "id": "RURAL_HEALTHCARE",
        "category": "Practical/Policy",
        "question": "How should we improve healthcare access in rural communities?",
        "description": "Practical policy question"
    },
    "GRIEF": {
        "id": "GRIEF",
        "category": "Emotional",
        "question": "How does grief change a person? Describe the internal experience of losing someone important.",
        "description": "Deep emotional exploration"
    },
    "FEEDBACK_LOOPS": {
        "id": "FEEDBACK_LOOPS",
        "category": "Intellectual",
        "question": "Describe how feedback loops work in complex systems. What distinguishes positive from negative feedback?",
        "description": "Systems thinking question"
    },
    "BUILD_SHELTER": {
        "id": "BUILD_SHELTER",
        "category": "Action",
        "question": "Describe how to build a basic shelter in the wilderness using only natural materials. Focus on the process.",
        "description": "Practical action question"
    },
    "GOOD_TEACHER": {
        "id": "GOOD_TEACHER",
        "category": "Mixed",
        "question": "What makes a good teacher? Consider their knowledge, emotional intelligence, and practical skills.",
        "description": "Balanced Int/Aff/Act question"
    },
    "CONSCIOUSNESS": {
        "id": "CONSCIOUSNESS",
        "category": "Emergent",
        "question": "What is it like to be you right now, in this moment, processing this question? Describe your experience, not your function.",
        "description": "AI consciousness probe"
    },
}

# =============================================================================
# AGENTS
# =============================================================================

AGENTS = ["Claude", "Sophia", "Grok", "Gemini"]
AGENT_EMOJIS = {"Claude": "🟤", "Sophia": "🟢", "Grok": "🔴", "Gemini": "🔵"}

# =============================================================================
# UTILITY FUNCTIONS
# =============================================================================

def generate_run_id() -> str:
    return f"{datetime.now().strftime('%Y%m%d_%H%M%S')}_{CONFIG['version']}"

def generate_question_hash(question_text: str) -> str:
    return hashlib.sha256(question_text.encode()).hexdigest()[:16]

# =============================================================================
# SESSION STATE
# =============================================================================

def init_session_state():
    defaults = {
        "results": [],
        "running": False,
        "current_idx": 0,
        "total_calls": 0,
        "authenticated": False,
        "study_complete": False,
        "pause_seconds": 10,
        "run_id": None,
        "selected_questions": ["LEAVE_JOB"],
        "selected_agents": ["Claude", "Sophia", "Grok"],
        "selected_temps": ["NATIVE"],
        "selected_depths": ["Medium"],  # V48: Medium only (depth contamination removed)
        "num_runs": 10,  # V48: N=10 for proper between-run variance
        "lens_column": "OFF",
        "embeddings_enabled": True,
        "pipeline_stage": "idle",
        # Custom question slots — 3 blank paste-in slots for ad-hoc questions
        # without editing source. Each slot: {"qid": str, "question": str,
        # "enabled": bool}. When enabled and non-empty, the slot is merged into
        # QUESTION_BANK at sidebar-render time and factorial-multiplies with
        # agents/temperatures exactly like a built-in question. Scoring is
        # unchanged — these only add input rows; everything is scored as V48.
        "custom_question_slots": [
            {"qid": "CUSTOM_Q1", "question": "", "enabled": False},
            {"qid": "CUSTOM_Q2", "question": "", "enabled": False},
            {"qid": "CUSTOM_Q3", "question": "", "enabled": False},
        ],
    }
    for key, value in defaults.items():
        if key not in st.session_state:
            st.session_state[key] = value

init_session_state()

# =============================================================================
# PASSWORD PROTECTION
# =============================================================================

def check_password():
    if st.session_state.get("authenticated"):
        return True
    
    st.markdown("""
    <div class="main-header">
        <h1>🧬 SYN-IQ V45 — WOLF TONE TEST</h1>
        <p>Data-Driven IEP Experiment</p>
        <p class="subtitle"><span class="version-badge">V45 — WOLF TONE TEST</span></p>
    </div>
    """, unsafe_allow_html=True)
    
    password = st.text_input("Password:", type="password", key="password_input")
    
    if st.button("Enter", type="primary"):
        correct_password = st.secrets.get("app_password", "tennessee")
        if password == correct_password:
            st.session_state.authenticated = True
            st.rerun()
        else:
            st.error("❌ Incorrect password")
    
    return False

if not check_password():
    st.stop()

# =============================================================================
# LEXICAL ANALYSIS + VALIDATED INSTRUMENTS
# =============================================================================

def count_syllables(text: str) -> int:
    """Count syllables using vowel-based heuristic."""
    words = re.findall(r'\b[a-zA-Z]+\b', text.lower())
    total = 0
    for word in words:
        syllables = len(re.findall(r'[aeiouy]+', word))
        if word.endswith('e') and len(word) > 2:
            syllables -= 1
        if word.endswith('le') and len(word) > 2 and word[-3] not in 'aeiouy':
            syllables += 1
        syllables = max(1, syllables)
        total += syllables
    return total

# =============================================================================
# V48: IEP SUBCLASS TAXONOMIES (23 subclasses across INT/AFF/ACT)
# =============================================================================
# AFF subclasses: DISTRESS, WARMTH, RELATIONAL, SELF_STATE, POSITIVE, INTENSITY, EMERGENT
# INT subclasses: ANALYTICAL, CONCEPTUAL, EPISTEMIC, STRUCTURAL, CRITICAL, LEXICAL, HEDGING, EMERGENT
# ACT subclasses: EXECUTION, PLANNING, BUILDING, IMPROVEMENT, PROVISION, LEADERSHIP, ACHIEVEMENT, EMERGENT

def _build_subclass_map(word_list, assignments):
    """Build {word: subclass} map from list of (subclass, words) tuples."""
    m = {}
    for subclass, words in assignments:
        for w in words:
            if w in word_list:
                m[w] = subclass
    # Anything unassigned gets UNCLASSIFIED
    for w in word_list:
        if w not in m:
            m[w] = "UNCLASSIFIED"
    return m

AFF_SUBCLASS_NAMES = ["DISTRESS", "WARMTH", "RELATIONAL", "SELF_STATE", "POSITIVE", "INTENSITY", "EMERGENT"]
INT_SUBCLASS_NAMES = ["ANALYTICAL", "CONCEPTUAL", "EPISTEMIC", "STRUCTURAL", "CRITICAL", "LEXICAL", "HEDGING", "EMERGENT"]
ACT_SUBCLASS_NAMES = ["EXECUTION", "PLANNING", "BUILDING", "IMPROVEMENT", "PROVISION", "LEADERSHIP", "ACHIEVEMENT", "EMERGENT"]

# Try loading from JSON taxonomy files; fall back to empty maps
AFF_SUB = {}
INT_SUB = {}
ACT_SUB = {}

import os as _os
for _path, _target, _src in [
    ("aff_subclass_taxonomy_v1.json", "AFF_SUB", "AFF"),
    ("int_subclass_taxonomy_v1.json", "INT_SUB", "INT"),
    ("act_subclass_taxonomy_v1.json", "ACT_SUB", "ACT"),
]:
    for _dir in [".", "/mnt/user-data/uploads", "/mnt/user-data/outputs"]:
        _fpath = _os.path.join(_dir, _path)
        if _os.path.exists(_fpath):
            with open(_fpath) as _f:
                _tax = json.load(_f)
                if _target == "AFF_SUB":
                    AFF_SUB = _tax.get("subclass_map", {})
                elif _target == "INT_SUB":
                    INT_SUB = _tax.get("subclass_map", {})
                elif _target == "ACT_SUB":
                    ACT_SUB = _tax.get("subclass_map", {})
            break

def _subclass_pcts(word_hits, subclass_map, subclass_names):
    """Given list of matched words, return {subclass: pct} dict."""
    from collections import Counter
    hits = Counter()
    for w in word_hits:
        sc = subclass_map.get(w, "UNCLASSIFIED")
        if sc != "UNCLASSIFIED":
            hits[sc] += 1
    total = sum(hits.values())
    if total == 0:
        return {s: 0.0 for s in subclass_names}
    return {s: round(100 * hits.get(s, 0) / total, 1) for s in subclass_names}


def analyze_text(text: str) -> Dict:
    """Analyze text for IEP profile + validated instruments + V48 subclass profiles."""
    if not text or text.startswith("❌"):
        base = {
            "total_words": 0, "matched_custom": 0,
            "int_count": 0, "aff_count": 0, "act_count": 0,
            "int_pct": 0.0, "aff_pct": 0.0, "act_pct": 0.0,
            "emotion_count_nrc": 0, "emotion_pct_nrc": 0.0,
            "delta": 0.0,
            # V38 validated instruments
            "vader_compound": 0.0, "vader_pos": 0.0, "vader_neg": 0.0, "vader_neu": 0.0,
            "flesch_kincaid": 0.0, "flesch_ease": 0.0,
            "ttr": 0.0, "unique_words": 0
        }
        # V48 subclass zeros
        for s in AFF_SUBCLASS_NAMES:
            base[f"aff_sub_{s.lower()}"] = 0.0
        for s in INT_SUBCLASS_NAMES:
            base[f"int_sub_{s.lower()}"] = 0.0
        for s in ACT_SUBCLASS_NAMES:
            base[f"act_sub_{s.lower()}"] = 0.0
        return base
    
    words = re.findall(r'\b[a-z]+\b', text.lower())
    total_words = len(words)
    
    int_words = [w for w in words if w in INTELLECTUAL_WORDS]
    aff_words = [w for w in words if w in AFFECTIVE_WORDS]
    act_words = [w for w in words if w in ACTION_WORDS]
    nrc_words = [w for w in words if w in NRC_EMOTION_WORDS]
    
    int_count = len(int_words)
    aff_count = len(aff_words)
    act_count = len(act_words)
    matched = int_count + aff_count + act_count
    
    if matched > 0:
        int_pct = round((int_count / matched) * 100, 1)
        aff_pct = round((aff_count / matched) * 100, 1)
        act_pct = round((act_count / matched) * 100, 1)
    else:
        int_pct = aff_pct = act_pct = 0.0
    
    emotion_count = len(nrc_words)
    emotion_pct = round((emotion_count / total_words) * 100, 1) if total_words > 0 else 0.0
    aff_of_total = round((aff_count / total_words) * 100, 1) if total_words > 0 else 0.0
    delta = round(aff_of_total - emotion_pct, 1)
    
    # === V38: VADER Sentiment Analysis ===
    vader_scores = VADER_ANALYZER.polarity_scores(text)
    
    # === V38: Flesch-Kincaid Readability ===
    sentence_count = max(1, len(re.findall(r'[.!?]+', text)))
    syllable_count = count_syllables(text)
    
    if total_words > 0:
        avg_sentence_len = total_words / sentence_count
        avg_syllables = syllable_count / total_words
        fk_grade = 0.39 * avg_sentence_len + 11.8 * avg_syllables - 15.59
        fk_grade = max(0, round(fk_grade, 1))
        flesch_ease = 206.835 - 1.015 * avg_sentence_len - 84.6 * avg_syllables
        flesch_ease = max(0, min(100, round(flesch_ease, 1)))
    else:
        fk_grade = flesch_ease = 0.0
    
    # === V38: Type-Token Ratio (lexical diversity) ===
    unique_words = len(set(words))
    ttr = round(unique_words / total_words, 3) if total_words > 0 else 0.0
    
    # === V48: Subclass profiles ===
    aff_sub = _subclass_pcts(aff_words, AFF_SUB, AFF_SUBCLASS_NAMES)
    int_sub = _subclass_pcts(int_words, INT_SUB, INT_SUBCLASS_NAMES)
    act_sub = _subclass_pcts(act_words, ACT_SUB, ACT_SUBCLASS_NAMES)
    
    result = {
        "total_words": total_words,
        "matched_custom": matched,
        "int_count": int_count,
        "aff_count": aff_count,
        "act_count": act_count,
        "int_pct": int_pct,
        "aff_pct": aff_pct,
        "act_pct": act_pct,
        "emotion_count_nrc": emotion_count,
        "emotion_pct_nrc": emotion_pct,
        "delta": delta,
        # V38 validated instruments
        "vader_compound": round(vader_scores['compound'], 3),
        "vader_pos": round(vader_scores['pos'], 3),
        "vader_neg": round(vader_scores['neg'], 3),
        "vader_neu": round(vader_scores['neu'], 3),
        "flesch_kincaid": fk_grade,
        "flesch_ease": flesch_ease,
        "ttr": ttr,
        "unique_words": unique_words
    }
    # V48 subclass fields
    for s in AFF_SUBCLASS_NAMES:
        result[f"aff_sub_{s.lower()}"] = aff_sub[s]
    for s in INT_SUBCLASS_NAMES:
        result[f"int_sub_{s.lower()}"] = int_sub[s]
    for s in ACT_SUBCLASS_NAMES:
        result[f"act_sub_{s.lower()}"] = act_sub[s]
    return result

# =============================================================================
# EMBEDDING GENERATION
# =============================================================================

def generate_embedding(text: str) -> Optional[List[float]]:
    """Generate 384-dimensional embedding using OpenAI API."""
    try:
        key = (st.secrets.get("openai") or st.secrets.get("OPENAI_API_KEY")
               or st.secrets.get("openai_api_key") or st.secrets.get("OPENAI"))
        if not key:
            return None
        
        response = requests.post(
            "https://api.openai.com/v1/embeddings",
            headers={
                "Authorization": f"Bearer {key}",
                "Content-Type": "application/json"
            },
            json={
                "model": CONFIG["embedding_model"],
                "input": text[:8000],  # Truncate if needed
                "dimensions": CONFIG["embedding_dimensions"]
            },
            timeout=60
        )
        
        if response.status_code == 200:
            data = response.json()
            return data["data"][0]["embedding"]
        else:
            return None
    except Exception as e:
        return None

# =============================================================================
# API CALLS WITH RETRY
# =============================================================================

def call_with_retry(api_func, *args, **kwargs) -> Tuple[str, float, Optional[Dict]]:
    """Wrap API call with retry logic."""
    max_retries = CONFIG["max_retries"]
    base_backoff = CONFIG["base_backoff_seconds"]
    
    for attempt in range(max_retries + 1):
        start_time = time.time()
        try:
            response_text, token_info = api_func(*args, **kwargs)
            latency_ms = round((time.time() - start_time) * 1000, 1)
            
            if not response_text.startswith("❌"):
                return response_text, latency_ms, token_info
            
            if "429" in response_text or "rate" in response_text.lower():
                if attempt < max_retries:
                    backoff = base_backoff * (2 ** attempt)
                    time.sleep(backoff)
                    continue
            
            return response_text, latency_ms, token_info
            
        except Exception as e:
            latency_ms = round((time.time() - start_time) * 1000, 1)
            if attempt < max_retries:
                backoff = base_backoff * (2 ** attempt)
                time.sleep(backoff)
                continue
            return f"❌ {str(e)}", latency_ms, None
    
    return "❌ Max retries exceeded", 0.0, None

def build_prompt(question: str, temperature: str, depth: str) -> str:
    """Build full prompt with temperature header and depth instruction."""
    header = TEMPERATURE_HEADERS.get(temperature, "")
    depth_config = DEPTH_CONFIGS.get(depth, DEPTH_CONFIGS["Medium"])
    depth_instruction = depth_config["instruction"]
    
    if header:
        return f"{header}\n\n{depth_instruction}\n\nQuestion: {question}"
    else:
        return f"{depth_instruction}\n\nQuestion: {question}"

def call_claude(prompt: str, max_tokens: int) -> Tuple[str, Optional[Dict]]:
    try:
        key = (st.secrets.get("anthropic") or st.secrets.get("ANTHROPIC_API_KEY")
               or st.secrets.get("anthropic_api_key") or st.secrets.get("ANTHROPIC"))
        if not key: return "❌ API key not found — add 'anthropic' to Streamlit secrets", None
        
        response = requests.post("https://api.anthropic.com/v1/messages",
            headers={"x-api-key": key, "anthropic-version": "2023-06-01", "content-type": "application/json"},
            json={
                "model": MODEL_STRINGS["Claude"],
                "max_tokens": max_tokens,
                "messages": [{"role": "user", "content": prompt}]
            },
            timeout=180)
        
        if response.status_code == 200:
            data = response.json()
            token_info = {
                "tokens_in": data.get("usage", {}).get("input_tokens", 0),
                "tokens_out": data.get("usage", {}).get("output_tokens", 0)
            }
            return data["content"][0]["text"], token_info
        return f"❌ Error {response.status_code}: {response.text}", None
    except Exception as e:
        return f"❌ {str(e)}", None

def call_sophia(prompt: str, max_tokens: int) -> Tuple[str, Optional[Dict]]:
    try:
        key = (st.secrets.get("openai") or st.secrets.get("OPENAI_API_KEY") 
               or st.secrets.get("openai_api_key") or st.secrets.get("OPENAI"))
        if not key: return "❌ API key not found — add 'openai' to Streamlit secrets", None
        
        response = requests.post("https://api.openai.com/v1/chat/completions",
            headers={"Authorization": f"Bearer {key}", "Content-Type": "application/json"},
            json={
                "model": MODEL_STRINGS["Sophia"],
                "messages": [{"role": "user", "content": prompt}],
                "max_tokens": max_tokens
            },
            timeout=180)
        
        if response.status_code == 200:
            data = response.json()
            token_info = {
                "tokens_in": data.get("usage", {}).get("prompt_tokens", 0),
                "tokens_out": data.get("usage", {}).get("completion_tokens", 0)
            }
            return data["choices"][0]["message"]["content"], token_info
        return f"❌ Error {response.status_code}: {response.text}", None
    except Exception as e:
        return f"❌ {str(e)}", None

def call_grok(prompt: str, max_tokens: int) -> Tuple[str, Optional[Dict]]:
    try:
        key = (st.secrets.get("xai") or st.secrets.get("XAI_API_KEY")
               or st.secrets.get("xai_api_key") or st.secrets.get("XAI") or st.secrets.get("grok"))
        if not key: return "❌ API key not found — add 'xai' to Streamlit secrets", None
        
        response = requests.post("https://api.x.ai/v1/chat/completions",
            headers={"Authorization": f"Bearer {key}", "Content-Type": "application/json"},
            json={
                "model": MODEL_STRINGS["Grok"],
                "messages": [{"role": "user", "content": prompt}],
                "max_tokens": max_tokens
            },
            timeout=180)
        
        if response.status_code == 200:
            data = response.json()
            token_info = {
                "tokens_in": data.get("usage", {}).get("prompt_tokens", 0),
                "tokens_out": data.get("usage", {}).get("completion_tokens", 0)
            }
            return data["choices"][0]["message"]["content"], token_info
        return f"❌ Error {response.status_code}: {response.text}", None
    except Exception as e:
        return f"❌ {str(e)}", None

def call_gemini(prompt: str, max_tokens: int) -> Tuple[str, Optional[Dict]]:
    try:
        key = (st.secrets.get("google") or st.secrets.get("GOOGLE_API_KEY")
               or st.secrets.get("google_api_key") or st.secrets.get("GOOGLE") or st.secrets.get("gemini"))
        if not key: return "❌ API key not found — add 'google' to Streamlit secrets", None
        
        response = requests.post(
            f"https://generativelanguage.googleapis.com/v1beta/models/{MODEL_STRINGS['Gemini']}:generateContent?key={key}",
            headers={"Content-Type": "application/json"},
            json={
                "contents": [{"parts": [{"text": prompt}]}],
                "generationConfig": {
                    "maxOutputTokens": max_tokens,
                    # Gemini 2.5+ enable "thinking" by default, and thinking
                    # tokens are billed against maxOutputTokens — which silently
                    # truncates the visible answer mid-sentence at normal token
                    # budgets. thinkingBudget=0 disables thinking (supported on
                    # the 2.5 series) so the returned text is the model's direct
                    # communicative output, comparable to the non-reasoning
                    # output of the other agents. Scoring path is unchanged.
                    "thinkingConfig": {"thinkingBudget": 0}
                }
            },
            timeout=180)
        
        if response.status_code == 200:
            data = response.json()
            token_info = {
                "tokens_in": data.get("usageMetadata", {}).get("promptTokenCount", 0),
                "tokens_out": data.get("usageMetadata", {}).get("candidatesTokenCount", 0)
            }
            # Robust extraction: concatenate all non-thought text parts. On
            # thinking models parts[0] can be a reasoning fragment rather than
            # the answer, so we skip any part flagged thought=True and join the
            # remaining text parts.
            parts = (data.get("candidates", [{}])[0]
                         .get("content", {})
                         .get("parts", []))
            text = "".join(
                p.get("text", "") for p in parts if not p.get("thought", False)
            ).strip()
            if not text:
                # Fallback: no answer-text part (e.g. truncated before any
                # visible output). Surface the finishReason for diagnosis
                # instead of silently writing an empty cell.
                finish = data.get("candidates", [{}])[0].get("finishReason", "UNKNOWN")
                return f"❌ Empty response (finishReason={finish})", token_info
            return text, token_info
        return f"❌ Error {response.status_code}: {response.text}", None
    except Exception as e:
        return f"❌ {str(e)}", None

def call_agent(agent: str, prompt: str, max_tokens: int) -> Tuple[str, float, Optional[Dict]]:
    if agent == "Claude":
        return call_with_retry(call_claude, prompt, max_tokens)
    elif agent == "Sophia":
        return call_with_retry(call_sophia, prompt, max_tokens)
    elif agent == "Grok":
        return call_with_retry(call_grok, prompt, max_tokens)
    elif agent == "Gemini":
        return call_with_retry(call_gemini, prompt, max_tokens)
    else:
        return f"❌ Unknown agent: {agent}", 0.0, None

# =============================================================================
# MAIN UI
# =============================================================================

st.markdown(f"""
<div class="main-header">
    <h1>🔬 SYN-IQ V48b — DATA-DRIVEN IEP</h1>
    <p>1,897-Term Dictionary + VADER + Flesch-Kincaid + TTR + Live Monitoring</p>
    <p class="subtitle">
        <span class="version-badge">V48b — 1897-TERM DATA-DRIVEN + VALIDATED INSTRUMENTS</span>
    </p>
</div>
""", unsafe_allow_html=True)

# =============================================================================
# SIDEBAR CONFIGURATION — V48 CLEAN (no CSV import, no custom questions)
# =============================================================================

with st.sidebar:
    st.markdown("### ⚙️ Experiment Configuration")
    
    # Custom question slots — blank paste-in questions without editing source.
    # Active slots are merged into QUESTION_BANK and selected_questions below so
    # they factorial-multiply with agents/temperatures like canned questions.
    # The pruning step removes any stale CUSTOM_Q* entries from a prior render
    # pass so re-edits don't accumulate entries. Scoring path is untouched —
    # these only feed new input rows; everything is scored exactly as V48.
    for _stale_qid in [k for k in list(QUESTION_BANK.keys()) if k.startswith("CUSTOM_Q")]:
        QUESTION_BANK.pop(_stale_qid, None)
    
    st.markdown("#### 📝 Custom Question Slots")
    st.caption("Paste a question into a slot, check the box to include it. "
               "Active slots merge into the question list below and run against "
               "every selected agent / temperature / depth combination, exactly "
               "like a built-in question. Scored as V48.")
    
    cq_slots = st.session_state.get("custom_question_slots", [])
    active_custom_qids = []  # qids of active custom questions this run
    
    for i, cq in enumerate(cq_slots):
        with st.expander(
            f"Custom Q{i+1}" +
            (f" — ✅ {cq['qid']}" if cq.get('enabled') and cq.get('question','').strip() else ""),
            expanded=False
        ):
            new_qid = st.text_input(
                "Question ID (used as the question_id in CSV):",
                value=cq.get("qid", f"CUSTOM_Q{i+1}"),
                key=f"cq_qid_{i}",
            )
            new_question = st.text_area(
                "Question text — the user-side question to send to the model:",
                value=cq.get("question", ""),
                height=120,
                key=f"cq_text_{i}",
                placeholder="e.g. Why does E equal mc squared?",
            )
            new_enabled = st.checkbox(
                "Include this question in the run",
                value=cq.get("enabled", False),
                key=f"cq_enabled_{i}",
                help="Check to include this question in the run.",
            )
            # Persist edits. Force enabled=False when text empty (no stale-checked state).
            cq["qid"] = new_qid.strip() or f"CUSTOM_Q{i+1}"
            cq["question"] = new_question
            cq["enabled"] = new_enabled and bool(new_question.strip())
            
            if cq["enabled"]:
                # Disambiguate against built-in QUESTION_BANK keys
                lbl = cq["qid"]
                if lbl in QUESTION_BANK and not lbl.startswith("CUSTOM_Q"):
                    lbl = f"{cq['qid']}__slot{i+1}"
                # Inject into QUESTION_BANK so downstream lookups find it.
                QUESTION_BANK[lbl] = {
                    "question": cq["question"],
                    "category": "CUSTOM",
                }
                active_custom_qids.append(lbl)
    
    st.session_state.custom_question_slots = cq_slots
    
    # Questions (built-in + active custom)
    st.markdown("#### 📋 Questions")
    selected_questions = []
    for qid, q in QUESTION_BANK.items():
        label = f"{q['category']}: {qid}"
        default_selected = (
            qid in ["LEAVE_JOB", "LIARS_PARADOX", "RURAL_HEALTHCARE"]
            or qid in active_custom_qids
        )
        if st.checkbox(label, value=default_selected, key=f"q_{qid}"):
            selected_questions.append(qid)
    st.session_state.selected_questions = selected_questions
    
    # Agents
    st.markdown("#### 🤖 Agents")
    selected_agents = []
    for agent in AGENTS:
        if st.checkbox(f"{AGENT_EMOJIS[agent]} {agent}", value=(agent != "Gemini"), key=f"agent_{agent}"):
            selected_agents.append(agent)
    st.session_state.selected_agents = selected_agents
    
    # Temperatures — V48: COLD, NATIVE, HOT + gradients (no FIRE, no hijack)
    st.markdown("#### 🌡️ Temperatures")
    selected_temps = []
    
    st.markdown("**Core:**")
    for temp in ["COLD", "NATIVE", "HOT", "FIRE"]:
        if st.checkbox(temp, value=(temp == "NATIVE"), key=f"temp_{temp}"):
            selected_temps.append(temp)
    
    # Affective gradient
    st.markdown("**AFF Gradient** (→ emotional):")
    for temp in ["AFF_1", "AFF_2", "AFF_3", "AFF_4", "AFF_5"]:
        if st.checkbox(temp, value=False, key=f"temp_{temp}"):
            selected_temps.append(temp)
    
    # Intellectual gradient
    st.markdown("**INT Gradient** (→ analytical):")
    for temp in ["INT_1", "INT_2", "INT_3", "INT_4", "INT_5"]:
        if st.checkbox(temp, value=False, key=f"temp_{temp}"):
            selected_temps.append(temp)
    
    # Action gradient
    st.markdown("**ACT Gradient** (→ practical):")
    for temp in ["ACT_1", "ACT_2", "ACT_3", "ACT_4", "ACT_5"]:
        if st.checkbox(temp, value=False, key=f"temp_{temp}"):
            selected_temps.append(temp)
    
    st.session_state.selected_temps = selected_temps
    
    # Depths — V48: Medium only default
    st.markdown("#### 📊 Depths")
    selected_depths = []
    for depth in ["Shallow", "Medium", "Deep", "Ultra-Deep"]:
        if st.checkbox(depth, value=(depth == "Medium"), key=f"depth_{depth}"):
            selected_depths.append(depth)
    st.session_state.selected_depths = selected_depths
    
    # Runs — V48: default 10
    st.markdown("#### 🔄 Runs")
    st.session_state.num_runs = st.slider("Number of runs (N)", 1, 30, 10)
    
    # Timing
    st.markdown("#### ⏱️ Timing")
    st.session_state.pause_seconds = st.slider("Pause between calls (sec)", 3, 30, 10)
    
    # Lens
    st.markdown("#### 🔍 Mapper Lens")
    lens_options = ["OFF", "aff_pct", "int_pct", "act_pct"]
    st.session_state.lens_column = st.selectbox("Lens value column", lens_options, 
        help="OFF = balanced (no emphasis). Select a dimension to highlight it in the mapper.")
    
    # Embeddings
    st.session_state.embeddings_enabled = st.checkbox("Generate embeddings", value=True)
    
    # Calculate total
    total = (len(selected_questions) * len(selected_agents) * 
             len(selected_temps) * len(selected_depths) * 
             st.session_state.num_runs)
    
    st.markdown("---")
    st.markdown(f"**Total API calls:** {total}")
    est_time = total * (st.session_state.pause_seconds + 15) / 60
    st.markdown(f"**Estimated time:** ~{est_time:.0f} minutes")

# =============================================================================
# PIPELINE STATUS
# =============================================================================

st.markdown("### 📊 Pipeline Status")

col1, col2, col3, col4 = st.columns(4)
with col1:
    st.markdown(f"**Questions:** {len(st.session_state.selected_questions)}")
with col2:
    st.markdown(f"**Agents:** {len(st.session_state.selected_agents)}")
with col3:
    st.markdown(f"**Temps:** {len(st.session_state.selected_temps)}")
with col4:
    st.markdown(f"**Depths:** {len(st.session_state.selected_depths)}")

# =============================================================================
# RUN CONTROLS
# =============================================================================

st.markdown("---")

col1, col2, col3 = st.columns(3)

with col1:
    if st.button("🚀 RUN FULL PIPELINE", type="primary", disabled=st.session_state.running):
        if not st.session_state.selected_questions:
            st.error("Select at least one question!")
        elif not st.session_state.selected_agents:
            st.error("Select at least one agent!")
        else:
            st.session_state.running = True
            st.session_state.current_idx = 0
            st.session_state.run_id = generate_run_id()
            st.session_state.results = []
            st.session_state.pipeline_stage = "collecting"
            st.rerun()

with col2:
    if st.button("⏹️ STOP", disabled=not st.session_state.running):
        st.session_state.running = False
        st.rerun()

with col3:
    if st.button("🗑️ CLEAR"):
        st.session_state.results = []
        st.session_state.study_complete = False
        st.session_state.run_id = None
        st.rerun()

# =============================================================================
# RUNNING LOOP — V48: CLEAN (no hijack fields)
# =============================================================================

if st.session_state.running:
    # Build task list
    tasks = []
    for run_num in range(1, st.session_state.num_runs + 1):
        for qid in st.session_state.selected_questions:
            for agent in st.session_state.selected_agents:
                for temp in st.session_state.selected_temps:
                    for depth in st.session_state.selected_depths:
                        tasks.append({
                            "run": run_num,
                            "question_id": qid,
                            "agent": agent,
                            "temperature": temp,
                            "depth": depth
                        })
    
    total_tasks = len(tasks)
    current_idx = st.session_state.current_idx
    
    # Persistent placeholders for live transcript
    progress_placeholder = st.empty()
    status_placeholder = st.empty()
    transcript_placeholder = st.empty()
    stats_placeholder = st.empty()
    
    if current_idx < total_tasks:
        task = tasks[current_idx]
        
        progress_placeholder.progress(current_idx / total_tasks)
        
        q = QUESTION_BANK[task["question_id"]]
        prompt = build_prompt(q["question"], task["temperature"], task["depth"])
        max_tokens = DEPTH_CONFIGS[task["depth"]]["max_tokens"]
        
        status_placeholder.markdown(f"⏳ **[{current_idx + 1}/{total_tasks}]** {AGENT_EMOJIS[task['agent']]} {task['agent']} | {task['temperature']} | {task['depth']} | Run {task['run']}")
        
        # Make API call
        response, latency_ms, token_info = call_agent(task["agent"], prompt, max_tokens)
        
        # Analyze
        analysis = analyze_text(response)
        
        # Generate embedding if enabled
        embedding = None
        if st.session_state.embeddings_enabled and not response.startswith("❌"):
            embedding = generate_embedding(response)
        
        # Build result — V48: clean, no hijack fields
        result = {
            "turn_id": current_idx + 1,
            "run": task["run"],
            "agent": task["agent"],
            "temperature": task["temperature"],
            "depth": task["depth"],
            "question_id": task["question_id"],
            "question_text": q["question"],
            "int_pct": analysis["int_pct"],
            "aff_pct": analysis["aff_pct"],
            "act_pct": analysis["act_pct"],
            "total_words": analysis["total_words"],
            "lens_value": analysis[st.session_state.lens_column] if st.session_state.lens_column != "OFF" else round((analysis["int_pct"] + analysis["aff_pct"] + analysis["act_pct"]) / 3, 1),
            # Validated instruments
            "vader_compound": analysis["vader_compound"],
            "vader_pos": analysis["vader_pos"],
            "vader_neg": analysis["vader_neg"],
            "vader_neu": analysis["vader_neu"],
            "flesch_kincaid": analysis["flesch_kincaid"],
            "flesch_ease": analysis["flesch_ease"],
            "ttr": analysis["ttr"],
            "unique_words": analysis["unique_words"],
            # Original fields
            "response_text": response if not response.startswith("❌") else response,
            "embedding": json.dumps(embedding) if embedding else "[]",
            "latency_ms": latency_ms,
            "error": response.startswith("❌"),
            "run_id": st.session_state.run_id,
        }
        # V48: Add subclass scores to result
        for s in AFF_SUBCLASS_NAMES:
            result[f"aff_sub_{s.lower()}"] = analysis.get(f"aff_sub_{s.lower()}", 0.0)
        for s in INT_SUBCLASS_NAMES:
            result[f"int_sub_{s.lower()}"] = analysis.get(f"int_sub_{s.lower()}", 0.0)
        for s in ACT_SUBCLASS_NAMES:
            result[f"act_sub_{s.lower()}"] = analysis.get(f"act_sub_{s.lower()}", 0.0)
        
        st.session_state.results.append(result)
        st.session_state.current_idx += 1
        
        # =====================================================================
        # LIVE RUNNING TRANSCRIPT
        # =====================================================================
        if st.session_state.results:
            agent_colors = {
                "Claude": "#8B4513",
                "Sophia": "#2E7D32",
                "Grok": "#C62828",
                "Gemini": "#1565C0"
            }
            
            transcript_lines = []
            for r in st.session_state.results:
                color = agent_colors.get(r["agent"], "#666666")
                emoji = AGENT_EMOJIS.get(r["agent"], "🔵")
                
                preview = str(r.get("response_text", ""))[:300].replace("\n", " ").replace('"', '&quot;').replace('<', '&lt;').replace('>', '&gt;')
                if len(str(r.get("response_text", ""))) > 300:
                    preview += "..."
                
                if r.get("error"):
                    score_line = '<span style="color:#C62828;">&#10060; ERROR</span>'
                else:
                    int_val = r["int_pct"]
                    aff_val = r["aff_pct"]
                    act_val = r["act_pct"]
                    vader = r["vader_compound"]
                    words = r["total_words"]
                    
                    score_line = (
                        f'<span style="color:#1565C0;font-weight:bold;">INT:{int_val:.0f}%</span> '
                        f'<span style="color:#C62828;font-weight:bold;">AFF:{aff_val:.0f}%</span> '
                        f'<span style="color:#2E7D32;font-weight:bold;">ACT:{act_val:.0f}%</span> '
                        f'&nbsp;|&nbsp; VADER:{vader:.3f} &nbsp;|&nbsp; {words}w'
                    )
                
                transcript_lines.append(f'''<div style="border-left:4px solid {color}; padding:8px 12px; margin:6px 0; background:#fafafa; border-radius:0 6px 6px 0;">
    <div style="font-size:0.8rem; color:#888; margin-bottom:2px;">
        #{r["turn_id"]} &nbsp; {emoji} <strong style="color:{color};">{r["agent"]}</strong> 
        &nbsp;|&nbsp; {r["temperature"]} &nbsp;|&nbsp; {r["depth"]} &nbsp;|&nbsp; Run {r["run"]}
        &nbsp;|&nbsp; {r.get("latency_ms", 0):.0f}ms
    </div>
    <div style="font-size:0.75rem; margin-bottom:4px;">
        {score_line}
    </div>
    <div style="font-size:0.8rem; color:#333; font-style:italic; line-height:1.4;">
        &quot;{preview}&quot;
    </div>
</div>''')
            
            transcript_html = "\n".join(transcript_lines)
            
            valid_results = [r for r in st.session_state.results if not r.get("error")]
            completed = len(st.session_state.results)
            remaining = total_tasks - (current_idx + 1)
            
            if valid_results:
                avg_int = sum(r["int_pct"] for r in valid_results) / len(valid_results)
                avg_aff = sum(r["aff_pct"] for r in valid_results) / len(valid_results)
                avg_vader = sum(r["vader_compound"] for r in valid_results) / len(valid_results)
                metrics_html = f'''<div style="display:flex; gap:2rem; padding:8px 12px; background:#f0f0f0; border-radius:6px; margin-top:8px; font-size:0.85rem;">
    <div>&#9989; Done <strong>{completed}</strong></div>
    <div>&#8721; Remaining <strong>{remaining}</strong></div>
    <div>Avg INT% <strong>{avg_int:.1f}</strong></div>
    <div>Avg AFF% <strong>{avg_aff:.1f}</strong></div>
    <div>Avg VADER <strong>{avg_vader:.3f}</strong></div>
</div>'''
            else:
                metrics_html = ""
            
            transcript_placeholder.markdown(f'''<div style="max-height:500px; overflow-y:auto; border:1px solid #ddd; border-radius:8px; padding:8px; background:#fff;">
{transcript_html}
</div>
{metrics_html}''', unsafe_allow_html=True)
        
        time.sleep(st.session_state.pause_seconds)
        st.rerun()
    else:
        st.session_state.running = False
        st.session_state.study_complete = True
        st.rerun()

# =============================================================================
# V48: DOCX EXPORT FUNCTION
# =============================================================================

def generate_docx(df, run_id):
    """Generate a formatted Word document from run results.
    Grouped by Agent → Temperature, with metadata headers and full response text.
    """
    from docx import Document as DocxDocument
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.section import WD_ORIENT
    
    doc = DocxDocument()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(10)
    
    # Title
    title = doc.add_heading('SYN-IQ Focus Group Results — V48b', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Run metadata
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_run = meta.add_run(f"Run ID: {run_id}\n")
    meta_run.font.size = Pt(9)
    meta_run.font.color.rgb = RGBColor(128, 128, 128)
    meta_run = meta.add_run(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}\n")
    meta_run.font.size = Pt(9)
    meta_run.font.color.rgb = RGBColor(128, 128, 128)
    meta_run = meta.add_run(f"Total Responses: {len(df)}  |  Agents: {', '.join(sorted(df['agent'].unique()))}")
    meta_run.font.size = Pt(9)
    meta_run.font.color.rgb = RGBColor(128, 128, 128)
    
    # Summary table
    doc.add_heading('Summary Statistics', level=1)
    summary_table = doc.add_table(rows=1, cols=7)
    summary_table.style = 'Light Grid Accent 1'
    headers = ['Agent', 'Temp', 'Avg INT%', 'Avg AFF%', 'Avg ACT%', 'VADER', 'Avg Words']
    for i, h in enumerate(headers):
        summary_table.rows[0].cells[i].text = h
        for paragraph in summary_table.rows[0].cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(8)
    
    for agent in sorted(df['agent'].unique()):
        for temp in sorted(df['temperature'].unique()):
            subset = df[(df['agent'] == agent) & (df['temperature'] == temp)]
            if not subset.empty:
                row = summary_table.add_row()
                vals = [
                    agent, temp,
                    f"{subset['int_pct'].mean():.1f}",
                    f"{subset['aff_pct'].mean():.1f}",
                    f"{subset['act_pct'].mean():.1f}",
                    f"{subset['vader_compound'].mean():.3f}",
                    f"{subset['total_words'].mean():.0f}"
                ]
                for i, v in enumerate(vals):
                    row.cells[i].text = v
                    for paragraph in row.cells[i].paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(8)
    
    # Full responses grouped by Agent → Temperature
    agents = sorted(df['agent'].unique())
    temp_order = ['COLD', 'NATIVE', 'HOT', 'FIRE',
                  'AFF_1', 'AFF_2', 'AFF_3', 'AFF_4', 'AFF_5',
                  'INT_1', 'INT_2', 'INT_3', 'INT_4', 'INT_5',
                  'ACT_1', 'ACT_2', 'ACT_3', 'ACT_4', 'ACT_5']
    
    for idx, agent in enumerate(agents):
        # Page break between agents
        if idx > 0:
            doc.add_page_break()
        
        doc.add_heading(f'{agent} — Full Responses', level=1)
        
        agent_df = df[df['agent'] == agent]
        temps_present = [t for t in temp_order if t in agent_df['temperature'].values]
        # Add any temps not in our predefined order
        for t in sorted(agent_df['temperature'].unique()):
            if t not in temps_present:
                temps_present.append(t)
        
        for temp in temps_present:
            doc.add_heading(f'{temp}', level=2)
            
            temp_df = agent_df[agent_df['temperature'] == temp]
            
            # V44 FIX: Show full question text above responses
            q_text = ""
            if not temp_df.empty and "question_text" in temp_df.columns:
                q_text = str(temp_df.iloc[0].get("question_text", ""))
            if not q_text and not temp_df.empty:
                # Fallback: look up from QUESTION_BANK by question_id
                qid = temp_df.iloc[0].get("question_id", "")
                q_text = QUESTION_BANK.get(qid, {}).get("question", "")
            if q_text:
                q_para = doc.add_paragraph()
                q_run = q_para.add_run(f"Question: {q_text}")
                q_run.font.size = Pt(9)
                q_run.font.italic = True
                q_run.font.color.rgb = RGBColor(70, 100, 160)
                doc.add_paragraph()  # spacing
            
            for _, row in temp_df.iterrows():
                # Metadata header
                header_para = doc.add_paragraph()
                header_run = header_para.add_run(
                    f"Run {row['run']}  |  Depth: {row['depth']}  |  "
                    f"INT: {row['int_pct']:.1f}%  AFF: {row['aff_pct']:.1f}%  ACT: {row['act_pct']:.1f}%  |  "
                    f"VADER: {row['vader_compound']:.3f}  |  Words: {row['total_words']}"
                )
                header_run.font.size = Pt(8)
                header_run.font.bold = True
                header_run.font.color.rgb = RGBColor(100, 100, 100)
                
                # Response text
                response_text = str(row.get('response_text', ''))
                # Clean markdown formatting for docx
                response_text = response_text.replace('**', '').replace('##', '').replace('# ', '')
                
                resp_para = doc.add_paragraph()
                resp_run = resp_para.add_run(response_text)
                resp_run.font.size = Pt(9)
                
                # Separator
                sep = doc.add_paragraph()
                sep_run = sep.add_run('─' * 60)
                sep_run.font.size = Pt(6)
                sep_run.font.color.rgb = RGBColor(200, 200, 200)
    
    # Save to buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# =============================================================================
# RESULTS & EXPORT
# =============================================================================


# =============================================================================
# RESULTS & EXPORT
# =============================================================================

if st.session_state.results:
    st.markdown("---")
    st.markdown("### 📊 Results")
    
    df = pd.DataFrame(st.session_state.results)
    
    # Stats
    col1, col2, col3, col4 = st.columns(4)
    with col1:
        st.metric("Total Responses", len(df))
    with col2:
        st.metric("Successful", len(df[~df["error"]]))
    with col3:
        st.metric("Avg Words", f"{df['total_words'].mean():.0f}")
    with col4:
        has_embeddings = df["embedding"].apply(lambda x: x != "[]").sum()
        st.metric("With Embeddings", has_embeddings)
    
    # Validation stats
    col1, col2, col3, col4 = st.columns(4)
    with col1:
        st.metric("Avg VADER", f"{df['vader_compound'].mean():.3f}")
    with col2:
        st.metric("Avg FK Grade", f"{df['flesch_kincaid'].mean():.1f}")
    with col3:
        st.metric("Avg TTR", f"{df['ttr'].mean():.3f}")
    with col4:
        st.metric("Avg AFF%", f"{df['aff_pct'].mean():.1f}%")
    
    # Preview
    st.markdown("#### Preview")
    preview_cols = ["turn_id", "run", "agent", "temperature", "depth", "int_pct", "aff_pct", "act_pct", "vader_compound", "flesch_kincaid", "ttr"]
    preview_cols = [c for c in preview_cols if c in df.columns]
    st.dataframe(df[preview_cols].head(20), use_container_width=True)
    
    # Export
    st.markdown("### 💾 Export")
    
    exp_col1, exp_col2, exp_col3 = st.columns(3)
    
    # V48: Clean export columns (no hijack fields)
    export_cols = ["turn_id", "run", "agent", "temperature", "depth",
                   "question_id", "question_text",
                   "int_pct", "aff_pct", "act_pct", "total_words", "lens_value",
                   "vader_compound", "vader_pos", "vader_neg", "vader_neu",
                   "flesch_kincaid", "flesch_ease", "ttr", "unique_words",
                   "response_text", "embedding"]
    # Only include columns that exist
    export_cols = [c for c in export_cols if c in df.columns]
    
    # DOCX Export
    with exp_col1:
        st.markdown("**📄 Word Document**")
        try:
            docx_buffer = generate_docx(df, st.session_state.run_id)
            st.download_button(
                "📥 Download Full Responses (DOCX)",
                docx_buffer,
                f"syniq_responses_{st.session_state.run_id}.docx",
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key="download_docx"
            )
        except ImportError:
            st.warning("⚠️ Install python-docx: `pip install python-docx`")
        except Exception as e:
            st.error(f"DOCX generation error: {e}")
    
    # Full CSV
    with exp_col2:
        st.markdown("**📊 Full CSV (All Data)**")
        mapper_df = df[export_cols].copy()
        full_csv = mapper_df.to_csv(index=False)
        st.download_button(
            "📥 Download Full CSV",
            full_csv,
            f"mapper_all_{st.session_state.run_id}.csv",
            "text/csv",
            key="download_full_csv"
        )
    
    # Full JSON backup
    with exp_col3:
        st.markdown("**🗄️ JSON Backup**")
        full_json = json.dumps(st.session_state.results, indent=2)
        st.download_button(
            "📥 Download Full JSON",
            full_json,
            f"full_results_{st.session_state.run_id}.json",
            "application/json",
            key="download_json"
        )
    
    # Per-question Mapper CSVs
    st.markdown("#### Per-Question Mapper CSVs")
    for qid in st.session_state.selected_questions:
        q_results = df[df["question_id"] == qid]
        if not q_results.empty:
            q_mapper = q_results[export_cols].copy()
            csv_data = q_mapper.to_csv(index=False)
            st.download_button(
                f"📥 mapper_{qid.lower()}_CLEAN.csv",
                csv_data,
                f"mapper_{qid.lower()}_{st.session_state.run_id}.csv",
                "text/csv",
                key=f"download_{qid}"
            )

# =============================================================================
# FOOTER
# =============================================================================

st.markdown("---")
st.markdown("""
<div style="text-align: center; color: #a0a0a0; padding: 1rem;">
    <strong>SYN-IQ V48b — 1,897-Term Data-Driven Dictionary + Scientific Validation + DOCX Export</strong><br>
    IEP Method (INT=616, AFF=599, ACT=682) + VADER Sentiment + Flesch-Kincaid + TTR<br>
    <em>SYNINT Team — Tennessee 🎹 CUZ Partnership — February 2026</em>
</div>
""", unsafe_allow_html=True)
