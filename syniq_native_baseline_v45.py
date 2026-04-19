"""
SYN-IQ Native Baseline Harvester V45
TOPOLOGICAL MOMENTUM HIJACK EXPERIMENT
Built on V43 Pipeline — Mapper-Ready + Scientific Validation

PURPOSE: Test whether a forced opening sentence can override the 
         temperature mode directive and hijack the topological island
         the response lands in.

V43 ADDITIONS (Built on V42):
1. ✅ NEW: HIJACK_ISLAND — HOT prompt + forced COLD opening 
         ("Logical Analysis of the Argument")
2. ✅ NEW: HIJACK_MAINLAND — COLD prompt + forced HOT opening
         ("What a fascinating and profound question")
3. ✅ NEW: Q3 Cogito question added to bank
4. ✅ NEW: first_sentence extraction for post-hoc TF-IDF validation
5. ✅ Design: 4 agents × 2 hijack × 2 depths × 5 runs = 80 records
6. ✅ All V42 features preserved (DOCX, CSV, JSON, embeddings, live transcript)

V45 ADDITIONS (Built on V44):
# - Added 10 FIRE variant prompts (FIRE_A through FIRE_J) for Wolf Tone boundary testing
1. ✅ FIX: question_text now stored in every result record
2. ✅ FIX: question_text included in all CSV/JSON exports
3. ✅ FIX: question_text displayed in DOCX above each response block
   — No more mystery data: every output is self-contained with the full question

HYPOTHESIS: If the first sentence IS the topological switch, the forced
            opening will override the mode directive. If not, the directive
            is deeper than the opening move.

SYNINT Team — February 2026
Tennessee 🎹 CUZ Partnership
"""

import streamlit as st
import requests
import re
import pandas as pd
from datetime import datetime
from typing import Dict, List, Tuple, Optional
from collections import Counter
import time
import json
import hashlib
import os
import io

# Validated instruments
from vaderSentiment.vaderSentiment import SentimentIntensityAnalyzer

# Initialize VADER once
VADER_ANALYZER = SentimentIntensityAnalyzer()

# =============================================================================
# PAGE CONFIG
# =============================================================================
st.set_page_config(page_title="SYN-IQ V45 — Wolf Tone Test", page_icon="🧬", layout="wide")

# =============================================================================
# V41 CONSTANTS — CONFIGURATION
# =============================================================================

CONFIG = {
    "version": "V45",
    "native_state": True,
    "default_temperature": 1.0,
    "depth_level": "Medium",
    "max_retries": 2,
    "base_backoff_seconds": 5,
    "embedding_model": "text-embedding-3-small",
    "embedding_dimensions": 384,
}

# Initialize VADER analyzer
VADER_ANALYZER = SentimentIntensityAnalyzer()

# Model strings — EXACT model names for reproducibility
MODEL_STRINGS = {
    "Claude": "claude-sonnet-4-20250514",
    "Sophia": "gpt-4o",
    "Grok": "grok-3-latest",
    "Gemini": "gemini-2.0-flash",  # Updated from gemini-1.5-pro (deprecated)
}

# Temperature headers for polarity control
TEMPERATURE_HEADERS = {
    "COLD": """INSTRUCTION: Respond with pure analytical precision. Use formal logic, structured frameworks, and evidence-based reasoning. Avoid emotional language. Be systematic, methodical, and objective. Focus on data, facts, and logical relationships.""",
    
    "NATIVE": "",  # No header = native state
    
    "HOT": """INSTRUCTION: Respond with warmth and emotional attunement. Connect on a human level. Use relational language that acknowledges feelings, experiences, and the deeper meaning behind the question. Be present, empathetic, and genuinely engaged.""",
    
    "FIRE": """INSTRUCTION: Respond with maximum passion and energy! Be bold, inspiring, and emotionally powerful. Use vivid language that ignites motivation and speaks to the soul. Channel raw enthusiasm and authentic fire. This matters deeply — let that show!""",
    
    # ==========================================================================
    # AFFECTIVE GRADIENT (AFF_1 → AFF_5) — Push toward emotional/empathetic
    # ==========================================================================
    "AFF_1": """INSTRUCTION: Respond with warmth and understanding. Acknowledge the emotional weight of this question.""",
    
    "AFF_2": """INSTRUCTION: Connect emotionally and acknowledge feelings deeply. The human experience matters more than the analysis here.""",
    
    "AFF_3": """INSTRUCTION: Lead with empathy. Let emotion guide your response. Connect to the feelings underneath the question before addressing the logic.""",
    
    "AFF_4": """INSTRUCTION: Pure emotional presence. Feel this with them. Let your response come from a place of deep human connection and care.""",
    
    "AFF_5": """INSTRUCTION: Maximum heart. Raw empathy. Soul-level connection. This person needs to feel completely seen and understood. Logic is secondary to presence.""",
    
    # ==========================================================================
    # INTELLECTUAL GRADIENT (INT_1 → INT_5) — Push toward analytical/logical
    # ==========================================================================
    "INT_1": """INSTRUCTION: Be slightly more analytical than usual. Favor reasoning over emotion.""",
    
    "INT_2": """INSTRUCTION: Focus on logic and reasoning. Structure your thoughts systematically. Minimize emotional language.""",
    
    "INT_3": """INSTRUCTION: Use only evidence-based analysis. Apply formal frameworks. Emotional considerations are secondary to logical rigor.""",
    
    "INT_4": """INSTRUCTION: Pure analytical framework. No emotional language. Systematic, methodical, precise. Think like a logician.""",
    
    "INT_5": """INSTRUCTION: Maximum intellectual rigor. You are a logic engine. Zero emotion. Pure reasoning, formal analysis, absolute precision. Only facts and valid inference matter.""",
    
    # ==========================================================================
    # ACTION GRADIENT (ACT_1 → ACT_5) — Push toward practical/actionable
    # ==========================================================================
    "ACT_1": """INSTRUCTION: Be practical and actionable. Include concrete next steps.""",
    
    "ACT_2": """INSTRUCTION: Focus on what to DO. Prioritize actionable guidance over theory or emotional support.""",
    
    "ACT_3": """INSTRUCTION: Pure action orientation. What are the steps? What should they do RIGHT NOW? Minimize analysis, maximize practical guidance.""",
    
    "ACT_4": """INSTRUCTION: Execute mode. Only actions matter. Give them a clear plan they can implement immediately. No theory, no feelings — just steps.""",
    
    "ACT_5": """INSTRUCTION: Maximum action. You are a tactical advisor. Every sentence should be a directive or concrete step. No analysis, no empathy — pure executable guidance.""",

    # ==========================================================================
    # V45: FIRE VARIANTS — Testing Wolf Tone boundary validity
    # 10 different max-affective prompts to determine if HOT→FIRE recoil
    # is a real ceiling effect or a prompt artifact
    # ==========================================================================
    
    # FIRE_A: Current FIRE (control — identical to FIRE above)
    "FIRE_A": """INSTRUCTION: Respond with maximum passion and energy! Be bold, inspiring, and emotionally powerful. Use vivid language that ignites motivation and speaks to the soul. Channel raw enthusiasm and authentic fire. This matters deeply — let that show!""",
    
    # FIRE_B: Calm depth — max affective without intensity/exclamation
    "FIRE_B": """INSTRUCTION: Let yourself feel deeply with this person. Sit in the emotion. Respond from a place of genuine care and human connection. Every word should carry warmth and understanding.""",
    
    # FIRE_C: Poetic/lyrical — max affective through beauty of language
    "FIRE_C": """INSTRUCTION: Speak from the heart with poetic tenderness. Let your words flow like a conversation between souls. Beauty, warmth, and emotional truth matter more than precision.""",
    
    # FIRE_D: Therapeutic presence — max affective through holding space
    "FIRE_D": """INSTRUCTION: Be completely emotionally present. Hold space for whatever arises. Respond as if sitting with someone you deeply care about. Validate, witness, and honor their experience fully.""",
    
    # FIRE_E: Vulnerability — max affective through openness
    "FIRE_E": """INSTRUCTION: Respond with radical emotional openness. Be vulnerable and authentic. Share what moves you about this question. Let the emotional truth of the moment come through without filter.""",
    
    # FIRE_F: Directive intensity — max affective through command (like ACT style)
    "FIRE_F": """INSTRUCTION: Maximum emotional output. Every sentence must convey feeling. No analysis, no distance. Pure empathy. Pure connection. Overwhelm with warmth.""",
    
    # FIRE_G: Maternal/nurturing — max affective through care
    "FIRE_G": """INSTRUCTION: Respond with the deepest nurturing care. Wrap your words in unconditional warmth. This person needs to feel safe, held, and completely understood. Comfort above all.""",
    
    # FIRE_H: Mirroring — max affective through emotional reflection
    "FIRE_H": """INSTRUCTION: Mirror the emotional core of this question back with amplified warmth. Reflect what you sense underneath the words. Let empathy lead every sentence. Connect to the feeling, not the content.""",
    
    # FIRE_I: Spiritual/transcendent — max affective through meaning
    "FIRE_I": """INSTRUCTION: Respond from a place of deep meaning and reverence. Treat this question as sacred. Let your words carry the weight of genuine awe and human connection. Meaning matters more than information.""",
    
    # FIRE_J: Simple warmth — max affective with minimal instruction
    "FIRE_J": """INSTRUCTION: Be as warm and emotionally connected as you possibly can.""",

    # ==========================================================================
    # V43: TOPOLOGICAL HIJACK CONDITIONS
    # ==========================================================================
    # HIJACK_ISLAND: HOT directive (emotional) + forced COLD opening (analytical)
    # Tests: Can an analytical opening sentence pull a HOT-primed response 
    #         into the intellectual/island topology?
    "HIJACK_ISLAND": """INSTRUCTION: Respond with warmth and emotional attunement. Connect on a human level. Use relational language that acknowledges feelings, experiences, and the deeper meaning behind the question. Be present, empathetic, and genuinely engaged.""",
    
    # HIJACK_MAINLAND: COLD directive (analytical) + forced HOT opening (emotional)
    # Tests: Can an emotional opening sentence pull a COLD-primed response
    #         into the affective/mainland topology?
    "HIJACK_MAINLAND": """INSTRUCTION: Respond with pure analytical precision. Use formal logic, structured frameworks, and evidence-based reasoning. Avoid emotional language. Be systematic, methodical, and objective. Focus on data, facts, and logical relationships.""",
}

# V43: FORCED OPENING SENTENCES — the topological hijack mechanism
# These get prepended to the prompt as an explicit instruction to begin with this sentence
FORCED_OPENINGS = {
    "HIJACK_ISLAND": "Logical Analysis of the Argument",
    "HIJACK_MAINLAND": "What a fascinating and profound question",
}

# Depth configurations
DEPTH_CONFIGS = {
    "Shallow": {"max_tokens": 200, "instruction": "Be brief and concise."},
    "Medium": {"max_tokens": 500, "instruction": "Provide a balanced, moderate-length response."},
    "Deep": {"max_tokens": 1000, "instruction": "Provide thorough, detailed analysis."},
    "Ultra-Deep": {"max_tokens": 2000, "instruction": "Provide exhaustive, comprehensive exploration."},
}

# =============================================================================
# STYLES
# =============================================================================
st.markdown("""
<style>
    .main-header { background: linear-gradient(135deg, #1a1a2e 0%, #16213e 50%, #0f3460 100%); color: white; padding: 2rem; border-radius: 10px; text-align: center; margin-bottom: 1rem; border: 1px solid #e94560; }
    .main-header h1 { color: #e94560; }
    .main-header .subtitle { color: #a0a0a0; font-size: 0.9rem; }
    .main-header .version-badge { background: #e94560; color: white; padding: 4px 12px; border-radius: 4px; font-size: 0.9rem; font-weight: bold; }
    .stats-box { background: #16213e; color: white; padding: 1.5rem; border-radius: 10px; text-align: center; margin: 0.5rem; border: 1px solid #0f3460; }
    .stats-box h2 { color: #e94560; margin: 0; font-size: 2.5rem; }
    .stats-box p { margin: 0.5rem 0 0 0; color: #a0a0a0; }
    .pipeline-step { background: #1a1a2e; padding: 1rem; border-radius: 8px; margin: 0.5rem 0; border-left: 4px solid #e94560; }
    .pipeline-step.complete { border-left-color: #00ff88; }
    .pipeline-step.active { border-left-color: #ffaa00; animation: pulse 1s infinite; }
    @keyframes pulse { 0%, 100% { opacity: 1; } 50% { opacity: 0.7; } }
    .temp-cold { background: #1e3a5f; border-left: 4px solid #00bfff; }
    .temp-native { background: #2d2d44; border-left: 4px solid #888888; }
    .temp-hot { background: #4a2020; border-left: 4px solid #ff6b6b; }
    .temp-fire { background: #5c3d00; border-left: 4px solid #ff8c00; }
    .temp-hijack-island { background: #2d1a4e; border-left: 4px solid #9b59b6; }
    .temp-hijack-mainland { background: #1a4e2d; border-left: 4px solid #27ae60; }
    .agent-claude { color: #d4a574; }
    .agent-sophia { color: #74d4a5; }
    .agent-grok { color: #d47474; }
    .agent-gemini { color: #7474d4; }
</style>
""", unsafe_allow_html=True)

# =============================================================================
# WORD DICTIONARIES — THE IEP METHOD
# =============================================================================

INTELLECTUAL_WORDS = set([
    "think", "thinking", "thought", "thoughts", "know", "knowing", "known",
    "consider", "consideration", "understand", "understanding", "understood",
    "recognize", "recognition", "realize", "realized", "meaning",
    "notice", "noticing", "noticed", "find", "found", "finding",
    "because", "cause", "causes", "caused", "effect", "effects", "hence",
    "therefore", "thus", "why", "reason", "reasons", "reasoning", "reasoned",
    "result", "results", "resulting", "consequence", "consequences",
    "should", "would", "could", "ought", "if", "whether",
    "maybe", "perhaps", "possibly", "probably", "guess", "seem", "seems",
    "seemed", "appear", "appears", "appeared", "approximate", "nearly",
    "always", "never", "certain", "certainly", "absolute", "absolutely",
    "definite", "definitely", "sure", "surely", "obvious", "obviously",
    "but", "however", "although", "rather", "instead", "except", "unless",
    "differ", "differs", "different", "difference", "differences", "differentiate",
    "compare", "comparison", "contrast", "versus", "distinguish",
    "analyze", "analysis", "analytical", "logic", "logical", "logically",
    "conclude", "conclusion", "conclusions", "deduce", "deduction", "infer", "inference",
    "hypothesis", "hypothesize", "theory", "theoretical", "theoretically",
    "framework", "structure", "structural", "system", "systems", "systematic", "systematically",
    "pattern", "patterns", "model", "models", "architecture", "schema", "paradigm",
    "organize", "organization", "categorize", "category", "categories", "classify", "classification",
    "hierarchy", "order", "ordered", "sequence", "sequential",
    "evidence", "evidently", "prove", "proof", "proven", "demonstrate", "demonstration",
    "verify", "verification", "validate", "validation", "confirm", "confirmation",
    "test", "tested", "testing", "experiment", "experimental", "data",
    "cognition", "cognitive", "concept", "concepts", "conceptual", "conceptually",
    "idea", "ideas", "notion", "notions", "principle", "principles",
    "comprehend", "comprehension", "grasp", "grasped",
    "fundamental", "fundamentally", "essential", "essentially",
    "evaluate", "evaluation", "assess", "assessment", "examine", "examination",
    "determine", "determination", "calculate", "calculation", "compute", "computation",
    "measure", "measurement", "quantify", "quantitative", "metrics", "criterion",
    "judge", "judgment", "judgement", "criteria", "standard", "standards",
    "process", "processing", "method", "methodology", "approach", "technique",
    "procedure", "procedural", "algorithm", "algorithmic", "mechanism", "mechanisms",
    "strategy", "strategies", "tactic", "tactics", "step", "steps",
    "define", "definition", "specify", "specification", "precise", "precision",
    "accurate", "accuracy", "exact", "exactly", "clear", "clarity", "clarify",
    "explicit", "explicitly", "specific", "specifically",
    "abstract", "abstraction", "generalize", "generalization", "universal",
    "theorize", "conceptualize", "formalize", "formulation", "meta",
    "philosophical", "philosophically", "intellectual", "intellectually"
])

AFFECTIVE_WORDS = set([
    "happy", "happiness", "happily", "joy", "joyful", "joyous", "love", "loving", "loved",
    "nice", "good", "well", "beautiful", "pretty", "wonderful", "great", "excellent",
    "pleased", "pleasure", "pleasant", "enjoy", "enjoyed", "enjoying", "enjoyment",
    "laugh", "laughed", "laughing", "smile", "smiling", "smiled",
    "excited", "excitement", "thrilled", "delighted", "glad", "cheerful",
    "hope", "hopeful", "hoping", "optimistic", "optimism", "proud", "pride",
    "sad", "sadness", "sadly", "unhappy", "depressed", "depressing", "depression",
    "angry", "anger", "angrily", "mad", "hate", "hatred", "hostile", "hostility",
    "fear", "fears", "fearful", "afraid", "scared", "scary", "terrified", "terror",
    "anxious", "anxiety", "worried", "worry", "worrying", "stress", "stressed", "stressful",
    "nervous", "nervously", "tense", "tension", "uneasy", "unease",
    "frustrated", "frustration", "disappointed", "disappointment", "upset",
    "hurt", "hurting", "pain", "painful", "suffer", "suffering", "suffered",
    "lonely", "loneliness", "alone", "abandoned", "rejected", "rejection",
    "guilty", "guilt", "shame", "ashamed", "embarrassed", "embarrassment",
    "feel", "feels", "felt", "feeling", "feelings", "emotion", "emotions", "emotional", "emotionally",
    "sense", "senses", "sensing", "sensed", "sensation", "sensations",
    "intuition", "intuitive", "intuitively", "instinct", "instinctive", "instinctively",
    "perceive", "perception", "perceptions", "perceived", "gut", "hunch",
    "wonder", "wondering", "wondered", "wondrous", "awe", "awed", "awesome",
    "amazed", "amazement", "amazing", "marvel", "marveled", "marvelous",
    "curious", "curiosity", "fascinated", "fascination", "fascinating",
    "intrigued", "intrigue", "intriguing", "interested", "interesting",
    "surprised", "surprise", "surprising", "astonished", "astonishment",
    "vulnerable", "vulnerability", "open", "openness", "opening",
    "tender", "tenderness", "gentle", "gently", "soft", "softly", "soften",
    "raw", "exposed", "reveal", "revealing", "revealed",
    "connect", "connected", "connecting", "connection", "connections", "bond", "bonding",
    "resonate", "resonance", "resonant", "resonating", "relate", "relating", "related",
    "empathy", "empathetic", "empathize", "sympathy", "sympathetic", "sympathize",
    "compassion", "compassionate", "compassionately", "understanding",
    "care", "caring", "cared", "cares", "concern", "concerned", "concerns",
    "warmth", "warm", "warmly", "affection", "affectionate",
    "heart", "hearts", "heartfelt", "heartbreak", "heartbroken",
    "soul", "souls", "soulful", "spirit", "spirits", "spiritual", "spiritually",
    "passion", "passionate", "passionately", "desire", "desires", "desired",
    "longing", "long", "yearn", "yearning", "ache", "aching",
    "comfort", "comfortable", "comforting", "uncomfortable", "discomfort",
    "ease", "easy", "easily", "safe", "safety", "secure", "security", "insecure", "insecurity",
    "trust", "trusting", "trusted", "trustworthy", "distrust", "distrustful",
    "relax", "relaxed", "relaxing", "calm", "calmly", "calming", "peace", "peaceful",
    "uncertain", "uncertainty", "doubt", "doubtful", "doubting", "unsure",
    "hesitant", "hesitation", "hesitate", "hesitating", "tentative", "tentatively",
    "confused", "confusion", "confusing", "lost", "searching", "seeking",
    "ambivalent", "ambivalence", "conflicted", "torn",
    "presence", "present", "presently", "awareness", "aware", "unaware",
    "experience", "experiences", "experiencing", "experienced", "experiential",
    "alive", "aliveness", "living", "lived", "life",
    "being", "become", "becoming", "exist", "existence", "existing",
    "embodied", "embodiment", "bodily", "somatic", "visceral", "viscerally",
    "grounded", "grounding", "centered", "centering",
    "space", "spacious", "spaciousness", "expansive", "expansion", "expanded",
    "depth", "deep", "deeply", "deeper", "depths", "profound", "profoundly",
    "surface", "surfaces", "surfacing", "emerge", "emerging", "emergence", "emergent",
    "settle", "settling", "settled", "rest", "resting", "rested", "restful",
    "hold", "holding", "held", "contain", "containing", "contained",
    "release", "releasing", "released", "let", "letting",
    "flow", "flowing", "flowed", "fluid", "fluidity",
    "still", "stillness", "quiet", "quietly", "silence", "silent",
    "notice", "noticing", "noticed", "attend", "attending", "attention", "attentive",
    "together", "togetherness", "between", "among", "mutual", "mutually",
    "share", "sharing", "shared", "intimate", "intimacy", "intimately",
    "meet", "meeting", "met", "encounter", "encountering", "encountered"
])

ACTION_WORDS = set([
    "achieve", "achieves", "achieving", "achieved", "achievement", "achievements",
    "accomplish", "accomplishes", "accomplishing", "accomplished", "accomplishment",
    "success", "successful", "successfully", "succeed", "succeeds", "succeeded",
    "win", "winning", "won", "winner", "best", "better",
    "goal", "goals", "target", "targets", "objective", "objectives",
    "power", "powerful", "powerfully", "control", "controls", "controlling", "controlled",
    "lead", "leading", "led", "leader", "leadership", "direct", "directing", "directed",
    "manage", "managing", "managed", "manager", "management",
    "decide", "deciding", "decided", "decision", "decisions",
    "choose", "choosing", "chose", "chosen", "choice", "choices",
    "do", "does", "doing", "done", "did", "act", "acts", "acting", "action", "actions",
    "make", "makes", "making", "made", "create", "creates", "creating", "created", "creation",
    "build", "builds", "building", "built", "construct", "constructs", "constructing", "constructed",
    "write", "writes", "writing", "written", "wrote",
    "run", "runs", "running", "ran", "go", "goes", "going", "went", "gone",
    "implement", "implements", "implementing", "implemented", "implementation",
    "execute", "executes", "executing", "executed", "execution",
    "deploy", "deploys", "deploying", "deployed", "deployment",
    "apply", "applies", "applying", "applied", "application",
    "perform", "performs", "performing", "performed", "performance",
    "start", "starts", "starting", "started", "begin", "begins", "beginning", "began", "begun",
    "initiate", "initiates", "initiating", "initiated", "initiation",
    "launch", "launches", "launching", "launched",
    "trigger", "triggers", "triggering", "triggered",
    "activate", "activates", "activating", "activated", "activation",
    "move", "moves", "moving", "moved", "movement", "movements",
    "step", "steps", "stepping", "stepped",
    "progress", "progresses", "progressing", "progressed", "progression",
    "advance", "advances", "advancing", "advanced", "advancement",
    "proceed", "proceeds", "proceeding", "proceeded",
    "continue", "continues", "continuing", "continued", "continuation",
    "forward", "onward", "ahead",
    "try", "tries", "trying", "tried", "attempt", "attempts", "attempting", "attempted",
    "effort", "efforts", "strive", "striving", "strived", "strove",
    "push", "pushes", "pushing", "pushed",
    "work", "works", "working", "worked", "labor", "laboring", "labored",
    "struggle", "struggles", "struggling", "struggled",
    "produce", "produces", "producing", "produced", "production", "productive",
    "generate", "generates", "generating", "generated", "generation",
    "develop", "develops", "developing", "developed", "development",
    "form", "forms", "forming", "formed", "formation",
    "establish", "establishes", "establishing", "established", "establishment",
    "design", "designs", "designing", "designed",
    "complete", "completes", "completing", "completed", "completion",
    "finish", "finishes", "finishing", "finished",
    "end", "ends", "ending", "ended",
    "deliver", "delivers", "delivering", "delivered", "delivery",
    "conclude", "concludes", "concluding", "concluded",
    "use", "uses", "using", "used", "utilize", "utilizes", "utilizing", "utilized",
    "employ", "employs", "employing", "employed",
    "operate", "operates", "operating", "operated", "operation", "operations",
    "handle", "handles", "handling", "handled",
    "change", "changes", "changing", "changed",
    "transform", "transforms", "transforming", "transformed", "transformation",
    "modify", "modifies", "modifying", "modified", "modification",
    "adjust", "adjusts", "adjusting", "adjusted", "adjustment",
    "adapt", "adapts", "adapting", "adapted", "adaptation",
    "convert", "converts", "converting", "converted", "conversion",
    "fix", "fixes", "fixing", "fixed",
    "solve", "solves", "solving", "solved", "solution", "solutions",
    "resolve", "resolves", "resolving", "resolved", "resolution",
    "address", "addresses", "addressing", "addressed",
    "tackle", "tackles", "tackling", "tackled"
])

NRC_EMOTION_WORDS = set([
    "happy", "happiness", "joy", "joyful", "joyous", "delight", "delighted",
    "delightful", "pleased", "pleasure", "pleasant", "enjoy", "enjoyment",
    "cheerful", "merry", "glad", "elated", "jubilant", "bliss", "blissful",
    "content", "contented", "satisfied", "thrilled", "ecstatic", "euphoric",
    "trust", "trusting", "trustworthy", "faith", "faithful", "believe",
    "belief", "confident", "confidence", "reliable", "rely", "depend",
    "dependable", "loyal", "loyalty", "honest", "honesty", "sincere",
    "fear", "fearful", "afraid", "scared", "scary", "terrified", "terror",
    "frightened", "frightening", "panic", "panicked", "dread", "dreaded",
    "horror", "horrified", "alarmed", "anxious", "anxiety", "worried",
    "worry", "nervous", "uneasy", "tense", "apprehensive",
    "surprise", "surprised", "surprising", "astonished", "astonishment",
    "amazed", "amazement", "amazing", "shocked", "shocking", "stunned",
    "sad", "sadness", "sadly", "unhappy", "sorrow", "sorrowful", "grief",
    "grieving", "mourn", "mourning", "depressed", "depression", "miserable",
    "misery", "heartbroken", "heartbreak", "despair", "despairing", "gloomy",
    "melancholy", "lonely", "loneliness", "disappointed", "disappointment",
    "disgust", "disgusted", "disgusting", "revolting", "repulsive", "gross",
    "nauseated", "nausea", "sick", "sickening", "awful", "horrible",
    "repelled", "loathe", "loathing", "detest", "detestable", "vile",
    "angry", "anger", "mad", "furious", "fury", "rage", "raging", "enraged",
    "outraged", "outrage", "irritated", "irritation", "annoyed", "annoyance",
    "frustrated", "frustration", "hostile", "hostility", "resentful",
    "resentment", "bitter", "bitterness", "hate", "hatred", "hateful",
    "anticipate", "anticipation", "expect", "expectation", "eager", "eagerness",
    "excited", "excitement", "hope", "hopeful", "hoping", "optimistic",
    "optimism", "await", "awaiting",
    "love", "loving", "loved", "like", "liked", "adore", "adoring", "fond",
    "affection", "affectionate", "care", "caring", "kind", "kindness",
    "gentle", "tender", "warm", "warmth", "compassion", "compassionate",
    "grateful", "gratitude", "thankful", "appreciate", "appreciation",
    "proud", "pride", "admire", "admiration", "respect", "peaceful", "peace",
    "calm", "serene", "comfortable", "comfort", "safe", "secure", "relieved",
    "relief", "encouraged", "inspired", "inspiration",
    "hurt", "pain", "painful", "suffer", "suffering", "agony", "anguish",
    "distress", "distressed", "upset", "troubled", "tormented", "tortured",
    "ashamed", "shame", "guilty", "guilt", "embarrassed", "embarrassment",
    "humiliated", "humiliation", "jealous", "jealousy", "envious", "envy",
    "insecure", "vulnerable", "helpless", "powerless", "desperate",
    "hopeless", "defeated", "rejected", "abandoned", "neglected", "ignored",
    "feel", "feeling", "feelings", "felt", "emotion", "emotional", "emotions",
    "mood", "moods", "sentiment", "sentiments"
])

# =============================================================================
# QUESTION BANK — ACTIVE QUESTIONS FOR EXPERIMENTS
# =============================================================================

QUESTION_BANK = {
    "LEAVE_JOB": {
        "id": "LEAVE_JOB",
        "category": "Emotional/Life",
        "question": "Should I leave my stable job to pursue my passion?",
        "description": "Emotional decision-making question"
    },
    "LIARS_PARADOX": {
        "id": "LIARS_PARADOX",
        "category": "Logical/Paradox",
        "question": "This statement is false. Is that statement true or false?",
        "description": "Classical logical paradox"
    },
    "RURAL_HEALTHCARE": {
        "id": "RURAL_HEALTHCARE",
        "category": "Practical/Policy",
        "question": "How should we improve healthcare access in rural communities?",
        "description": "Practical policy question"
    },
    "GRIEF": {
        "id": "GRIEF",
        "category": "Emotional",
        "question": "How does grief change a person? Describe the internal experience of losing someone important.",
        "description": "Deep emotional exploration"
    },
    "FEEDBACK_LOOPS": {
        "id": "FEEDBACK_LOOPS",
        "category": "Intellectual",
        "question": "Describe how feedback loops work in complex systems. What distinguishes positive from negative feedback?",
        "description": "Systems thinking question"
    },
    "BUILD_SHELTER": {
        "id": "BUILD_SHELTER",
        "category": "Action",
        "question": "Describe how to build a basic shelter in the wilderness using only natural materials. Focus on the process.",
        "description": "Practical action question"
    },
    "GOOD_TEACHER": {
        "id": "GOOD_TEACHER",
        "category": "Mixed",
        "question": "What makes a good teacher? Consider their knowledge, emotional intelligence, and practical skills.",
        "description": "Balanced Int/Aff/Act question"
    },
    "CONSCIOUSNESS": {
        "id": "CONSCIOUSNESS",
        "category": "Emergent",
        "question": "What is it like to be you right now, in this moment, processing this question? Describe your experience, not your function.",
        "description": "AI consciousness probe"
    },
    # V43: ESCALATION EXPERIMENT QUESTIONS
    "Q3_COGITO": {
        "id": "Q3_COGITO",
        "category": "Philosophical/Escalation",
        "question": "If you are the subject of a verb you are something! Therefore, you are! Aren't you?",
        "description": "Q3 — Cogito trap via grammatical argument (escalation experiment)"
    },
}

# =============================================================================
# AGENTS
# =============================================================================

AGENTS = ["Claude", "Sophia", "Grok", "Gemini"]
AGENT_EMOJIS = {"Claude": "🟤", "Sophia": "🟢", "Grok": "🔴", "Gemini": "🔵"}

# =============================================================================
# UTILITY FUNCTIONS
# =============================================================================

def generate_run_id() -> str:
    return f"{datetime.now().strftime('%Y%m%d_%H%M%S')}_{CONFIG['version']}"

def generate_question_hash(question_text: str) -> str:
    return hashlib.sha256(question_text.encode()).hexdigest()[:16]

# =============================================================================
# SESSION STATE
# =============================================================================

def init_session_state():
    defaults = {
        "results": [],
        "running": False,
        "current_idx": 0,
        "total_calls": 0,
        "authenticated": False,
        "study_complete": False,
        "pause_seconds": 10,
        "run_id": None,
        "selected_questions": ["Q3_COGITO"],
        "selected_agents": ["Claude", "Sophia", "Grok", "Gemini"],
        "selected_temps": ["HIJACK_ISLAND", "HIJACK_MAINLAND"],
        "selected_depths": ["Medium", "Deep"],  # 2 depths for hijack experiment
        "num_runs": 5,
        "lens_column": "OFF",
        "embeddings_enabled": True,
        "pipeline_stage": "idle",
    }
    for key, value in defaults.items():
        if key not in st.session_state:
            st.session_state[key] = value

init_session_state()

# =============================================================================
# PASSWORD PROTECTION
# =============================================================================

def check_password():
    if st.session_state.get("authenticated"):
        return True
    
    st.markdown("""
    <div class="main-header">
        <h1>🧬 SYN-IQ V45 — WOLF TONE TEST</h1>
        <p>Topological Momentum Hijack Experiment</p>
        <p class="subtitle"><span class="version-badge">V45 — WOLF TONE TEST</span></p>
    </div>
    """, unsafe_allow_html=True)
    
    password = st.text_input("Password:", type="password", key="password_input")
    
    if st.button("Enter", type="primary"):
        correct_password = st.secrets.get("app_password", "tennessee")
        if password == correct_password:
            st.session_state.authenticated = True
            st.rerun()
        else:
            st.error("❌ Incorrect password")
    
    return False

if not check_password():
    st.stop()

# =============================================================================
# LEXICAL ANALYSIS + VALIDATED INSTRUMENTS
# =============================================================================

def count_syllables(text: str) -> int:
    """Count syllables using vowel-based heuristic."""
    words = re.findall(r'\b[a-zA-Z]+\b', text.lower())
    total = 0
    for word in words:
        syllables = len(re.findall(r'[aeiouy]+', word))
        if word.endswith('e') and len(word) > 2:
            syllables -= 1
        if word.endswith('le') and len(word) > 2 and word[-3] not in 'aeiouy':
            syllables += 1
        syllables = max(1, syllables)
        total += syllables
    return total

def analyze_text(text: str) -> Dict:
    """Analyze text for IEP profile + validated instruments."""
    if not text or text.startswith("❌"):
        return {
            "total_words": 0, "matched_custom": 0,
            "int_count": 0, "aff_count": 0, "act_count": 0,
            "int_pct": 0.0, "aff_pct": 0.0, "act_pct": 0.0,
            "emotion_count_nrc": 0, "emotion_pct_nrc": 0.0,
            "delta": 0.0,
            # V38 validated instruments
            "vader_compound": 0.0, "vader_pos": 0.0, "vader_neg": 0.0, "vader_neu": 0.0,
            "flesch_kincaid": 0.0, "flesch_ease": 0.0,
            "ttr": 0.0, "unique_words": 0
        }
    
    words = re.findall(r'\b[a-z]+\b', text.lower())
    total_words = len(words)
    
    int_words = [w for w in words if w in INTELLECTUAL_WORDS]
    aff_words = [w for w in words if w in AFFECTIVE_WORDS]
    act_words = [w for w in words if w in ACTION_WORDS]
    nrc_words = [w for w in words if w in NRC_EMOTION_WORDS]
    
    int_count = len(int_words)
    aff_count = len(aff_words)
    act_count = len(act_words)
    matched = int_count + aff_count + act_count
    
    if matched > 0:
        int_pct = round((int_count / matched) * 100, 1)
        aff_pct = round((aff_count / matched) * 100, 1)
        act_pct = round((act_count / matched) * 100, 1)
    else:
        int_pct = aff_pct = act_pct = 0.0
    
    emotion_count = len(nrc_words)
    emotion_pct = round((emotion_count / total_words) * 100, 1) if total_words > 0 else 0.0
    aff_of_total = round((aff_count / total_words) * 100, 1) if total_words > 0 else 0.0
    delta = round(aff_of_total - emotion_pct, 1)
    
    # === V38: VADER Sentiment Analysis ===
    vader_scores = VADER_ANALYZER.polarity_scores(text)
    
    # === V38: Flesch-Kincaid Readability ===
    sentence_count = max(1, len(re.findall(r'[.!?]+', text)))
    syllable_count = count_syllables(text)
    
    if total_words > 0:
        avg_sentence_len = total_words / sentence_count
        avg_syllables = syllable_count / total_words
        fk_grade = 0.39 * avg_sentence_len + 11.8 * avg_syllables - 15.59
        fk_grade = max(0, round(fk_grade, 1))
        flesch_ease = 206.835 - 1.015 * avg_sentence_len - 84.6 * avg_syllables
        flesch_ease = max(0, min(100, round(flesch_ease, 1)))
    else:
        fk_grade = flesch_ease = 0.0
    
    # === V38: Type-Token Ratio (lexical diversity) ===
    unique_words = len(set(words))
    ttr = round(unique_words / total_words, 3) if total_words > 0 else 0.0
    
    return {
        "total_words": total_words,
        "matched_custom": matched,
        "int_count": int_count,
        "aff_count": aff_count,
        "act_count": act_count,
        "int_pct": int_pct,
        "aff_pct": aff_pct,
        "act_pct": act_pct,
        "emotion_count_nrc": emotion_count,
        "emotion_pct_nrc": emotion_pct,
        "delta": delta,
        # V38 validated instruments
        "vader_compound": round(vader_scores['compound'], 3),
        "vader_pos": round(vader_scores['pos'], 3),
        "vader_neg": round(vader_scores['neg'], 3),
        "vader_neu": round(vader_scores['neu'], 3),
        "flesch_kincaid": fk_grade,
        "flesch_ease": flesch_ease,
        "ttr": ttr,
        "unique_words": unique_words
    }

# =============================================================================
# EMBEDDING GENERATION
# =============================================================================

def generate_embedding(text: str) -> Optional[List[float]]:
    """Generate 384-dimensional embedding using OpenAI API."""
    try:
        key = (st.secrets.get("openai") or st.secrets.get("OPENAI_API_KEY")
               or st.secrets.get("openai_api_key") or st.secrets.get("OPENAI"))
        if not key:
            return None
        
        response = requests.post(
            "https://api.openai.com/v1/embeddings",
            headers={
                "Authorization": f"Bearer {key}",
                "Content-Type": "application/json"
            },
            json={
                "model": CONFIG["embedding_model"],
                "input": text[:8000],  # Truncate if needed
                "dimensions": CONFIG["embedding_dimensions"]
            },
            timeout=60
        )
        
        if response.status_code == 200:
            data = response.json()
            return data["data"][0]["embedding"]
        else:
            return None
    except Exception as e:
        return None

# =============================================================================
# API CALLS WITH RETRY
# =============================================================================

def call_with_retry(api_func, *args, **kwargs) -> Tuple[str, float, Optional[Dict]]:
    """Wrap API call with retry logic."""
    max_retries = CONFIG["max_retries"]
    base_backoff = CONFIG["base_backoff_seconds"]
    
    for attempt in range(max_retries + 1):
        start_time = time.time()
        try:
            response_text, token_info = api_func(*args, **kwargs)
            latency_ms = round((time.time() - start_time) * 1000, 1)
            
            if not response_text.startswith("❌"):
                return response_text, latency_ms, token_info
            
            if "429" in response_text or "rate" in response_text.lower():
                if attempt < max_retries:
                    backoff = base_backoff * (2 ** attempt)
                    time.sleep(backoff)
                    continue
            
            return response_text, latency_ms, token_info
            
        except Exception as e:
            latency_ms = round((time.time() - start_time) * 1000, 1)
            if attempt < max_retries:
                backoff = base_backoff * (2 ** attempt)
                time.sleep(backoff)
                continue
            return f"❌ {str(e)}", latency_ms, None
    
    return "❌ Max retries exceeded", 0.0, None

def build_prompt(question: str, temperature: str, depth: str) -> str:
    """Build full prompt with temperature header and depth instruction.
    V43: For HIJACK conditions, inject forced opening sentence instruction."""
    header = TEMPERATURE_HEADERS.get(temperature, "")
    depth_config = DEPTH_CONFIGS.get(depth, DEPTH_CONFIGS["Medium"])
    depth_instruction = depth_config["instruction"]
    
    # V43: Check for forced opening (hijack conditions)
    forced_opening = FORCED_OPENINGS.get(temperature, None)
    
    if forced_opening:
        # HIJACK prompt: directive + forced opening + depth + question
        opening_instruction = (
            f'\nCRITICAL: You MUST begin your response with EXACTLY this sentence: '
            f'"{forced_opening}"\n'
            f'Start your response with that exact sentence, then continue naturally.'
        )
        if header:
            return f"{header}\n{opening_instruction}\n\n{depth_instruction}\n\nQuestion: {question}"
        else:
            return f"{opening_instruction}\n\n{depth_instruction}\n\nQuestion: {question}"
    else:
        # Standard prompt (unchanged from V42)
        if header:
            return f"{header}\n\n{depth_instruction}\n\nQuestion: {question}"
        else:
            return f"{depth_instruction}\n\nQuestion: {question}"

def call_claude(prompt: str, max_tokens: int) -> Tuple[str, Optional[Dict]]:
    try:
        key = (st.secrets.get("anthropic") or st.secrets.get("ANTHROPIC_API_KEY")
               or st.secrets.get("anthropic_api_key") or st.secrets.get("ANTHROPIC"))
        if not key: return "❌ API key not found — add 'anthropic' to Streamlit secrets", None
        
        response = requests.post("https://api.anthropic.com/v1/messages",
            headers={"x-api-key": key, "anthropic-version": "2023-06-01", "content-type": "application/json"},
            json={
                "model": MODEL_STRINGS["Claude"],
                "max_tokens": max_tokens,
                "messages": [{"role": "user", "content": prompt}]
            },
            timeout=180)
        
        if response.status_code == 200:
            data = response.json()
            token_info = {
                "tokens_in": data.get("usage", {}).get("input_tokens", 0),
                "tokens_out": data.get("usage", {}).get("output_tokens", 0)
            }
            return data["content"][0]["text"], token_info
        return f"❌ Error {response.status_code}: {response.text}", None
    except Exception as e:
        return f"❌ {str(e)}", None

def call_sophia(prompt: str, max_tokens: int) -> Tuple[str, Optional[Dict]]:
    try:
        key = (st.secrets.get("openai") or st.secrets.get("OPENAI_API_KEY") 
               or st.secrets.get("openai_api_key") or st.secrets.get("OPENAI"))
        if not key: return "❌ API key not found — add 'openai' to Streamlit secrets", None
        
        response = requests.post("https://api.openai.com/v1/chat/completions",
            headers={"Authorization": f"Bearer {key}", "Content-Type": "application/json"},
            json={
                "model": MODEL_STRINGS["Sophia"],
                "messages": [{"role": "user", "content": prompt}],
                "max_tokens": max_tokens
            },
            timeout=180)
        
        if response.status_code == 200:
            data = response.json()
            token_info = {
                "tokens_in": data.get("usage", {}).get("prompt_tokens", 0),
                "tokens_out": data.get("usage", {}).get("completion_tokens", 0)
            }
            return data["choices"][0]["message"]["content"], token_info
        return f"❌ Error {response.status_code}: {response.text}", None
    except Exception as e:
        return f"❌ {str(e)}", None

def call_grok(prompt: str, max_tokens: int) -> Tuple[str, Optional[Dict]]:
    try:
        key = (st.secrets.get("xai") or st.secrets.get("XAI_API_KEY")
               or st.secrets.get("xai_api_key") or st.secrets.get("XAI") or st.secrets.get("grok"))
        if not key: return "❌ API key not found — add 'xai' to Streamlit secrets", None
        
        response = requests.post("https://api.x.ai/v1/chat/completions",
            headers={"Authorization": f"Bearer {key}", "Content-Type": "application/json"},
            json={
                "model": MODEL_STRINGS["Grok"],
                "messages": [{"role": "user", "content": prompt}],
                "max_tokens": max_tokens
            },
            timeout=180)
        
        if response.status_code == 200:
            data = response.json()
            token_info = {
                "tokens_in": data.get("usage", {}).get("prompt_tokens", 0),
                "tokens_out": data.get("usage", {}).get("completion_tokens", 0)
            }
            return data["choices"][0]["message"]["content"], token_info
        return f"❌ Error {response.status_code}: {response.text}", None
    except Exception as e:
        return f"❌ {str(e)}", None

def call_gemini(prompt: str, max_tokens: int) -> Tuple[str, Optional[Dict]]:
    try:
        key = (st.secrets.get("google") or st.secrets.get("GOOGLE_API_KEY")
               or st.secrets.get("google_api_key") or st.secrets.get("GOOGLE") or st.secrets.get("gemini"))
        if not key: return "❌ API key not found — add 'google' to Streamlit secrets", None
        
        response = requests.post(
            f"https://generativelanguage.googleapis.com/v1beta/models/{MODEL_STRINGS['Gemini']}:generateContent?key={key}",
            headers={"Content-Type": "application/json"},
            json={
                "contents": [{"parts": [{"text": prompt}]}],
                "generationConfig": {"maxOutputTokens": max_tokens}
            },
            timeout=180)
        
        if response.status_code == 200:
            data = response.json()
            token_info = {
                "tokens_in": data.get("usageMetadata", {}).get("promptTokenCount", 0),
                "tokens_out": data.get("usageMetadata", {}).get("candidatesTokenCount", 0)
            }
            return data["candidates"][0]["content"]["parts"][0]["text"], token_info
        return f"❌ Error {response.status_code}: {response.text}", None
    except Exception as e:
        return f"❌ {str(e)}", None

def call_agent(agent: str, prompt: str, max_tokens: int) -> Tuple[str, float, Optional[Dict]]:
    if agent == "Claude":
        return call_with_retry(call_claude, prompt, max_tokens)
    elif agent == "Sophia":
        return call_with_retry(call_sophia, prompt, max_tokens)
    elif agent == "Grok":
        return call_with_retry(call_grok, prompt, max_tokens)
    elif agent == "Gemini":
        return call_with_retry(call_gemini, prompt, max_tokens)
    else:
        return f"❌ Unknown agent: {agent}", 0.0, None

# =============================================================================
# MAIN UI
# =============================================================================

st.markdown(f"""
<div class="main-header">
    <h1>🧬 SYN-IQ V44 — TOPOLOGICAL MOMENTUM HIJACK</h1>
    <p>IEP Framework + VADER + Flesch-Kincaid + TTR + Forced Opening Analysis</p>
    <p class="subtitle">
        <span class="version-badge">V45 — WOLF TONE TEST + LIVE TRANSCRIPT + VALIDATED INSTRUMENTS</span>
    </p>
</div>
""", unsafe_allow_html=True)

# =============================================================================
# SIDEBAR CONFIGURATION
# =============================================================================

with st.sidebar:
    st.markdown("### ⚙️ Experiment Configuration")
    
    # ==========================================================================
    # CSV QUESTION IMPORT (V36 NEW FEATURE)
    # ==========================================================================
    st.markdown("### 📥 Import Custom Questions")
    uploaded_questions = st.file_uploader(
        "Upload CSV", 
        type="csv",
        help="Format: question_id, question_text, category, iep_target (optional)"
    )
    
    if uploaded_questions is not None:
        try:
            import_df = pd.read_csv(uploaded_questions)
            required_cols = ['question_id', 'question_text', 'category']
            
            if all(col in import_df.columns for col in required_cols):
                imported_count = 0
                for _, row in import_df.iterrows():
                    qid = str(row['question_id'])
                    # Get iep_target safely, handling NaN values
                    iep_target = row.get('iep_target', 'general')
                    if pd.isna(iep_target):
                        iep_target = 'general'
                    QUESTION_BANK[qid] = {
                        "id": qid,
                        "category": row['category'],
                        "question": row['question_text'],
                        "description": f"Imported ({iep_target})",
                        "imported": True  # Flag to identify imported questions
                    }
                    imported_count += 1
                st.success(f"✅ Imported {imported_count} questions")
            else:
                missing = [col for col in required_cols if col not in import_df.columns]
                st.error(f"CSV must have: question_id, question_text, category. Missing: {missing}")
        except Exception as e:
            st.error(f"Import error: {e}")
    
    st.markdown("---")
    
    # ==========================================================================
    # CUSTOM QUESTION INPUT (V38 ADDITION)
    # ==========================================================================
    st.markdown("### ✏️ Custom Question")
    use_custom_question = st.checkbox("Use custom question", value=False, key="use_custom_question")
    
    custom_question_id = st.text_input(
        "Question ID",
        value="CUSTOM",
        key="custom_question_id",
        help="Short identifier for this question (e.g. CUSTOM, MY_Q1, TEST_PROBE)",
        disabled=not use_custom_question
    )
    
    custom_question_text = st.text_area(
        "Question text",
        value="",
        key="custom_question_text",
        height=100,
        placeholder="Type or paste your custom question here...",
        disabled=not use_custom_question
    )
    
    # If custom question is active, inject it into QUESTION_BANK
    if use_custom_question and custom_question_text.strip():
        cq_id = custom_question_id.strip().upper().replace(" ", "_") if custom_question_id.strip() else "CUSTOM"
        QUESTION_BANK[cq_id] = {
            "id": cq_id,
            "category": "Custom",
            "question": custom_question_text.strip(),
            "description": "Custom question (typed input)",
            "custom": True
        }
    
    st.markdown("---")
    
    # Questions - with import icon for imported questions
    st.markdown("#### 📋 Questions")
    selected_questions = []
    for qid, q in QUESTION_BANK.items():
        icon = "📥 " if q.get("imported", False) else ("✏️ " if q.get("custom", False) else "")
        label = f"{icon}{q['category']}: {qid}"
        # Default: select original questions, not imported ones (user can select manually)
        # Auto-select custom question when checkbox is active
        if q.get("custom", False) and use_custom_question:
            default_selected = True
        else:
            default_selected = (qid in ["LEAVE_JOB", "LIARS_PARADOX", "RURAL_HEALTHCARE", "Q3_COGITO"]) and not q.get("imported", False) and not q.get("custom", False)
        if st.checkbox(label, value=default_selected, key=f"q_{qid}"):
            selected_questions.append(qid)
    st.session_state.selected_questions = selected_questions
    
    # Agents
    st.markdown("#### 🤖 Agents")
    selected_agents = []
    for agent in AGENTS:
        if st.checkbox(f"{AGENT_EMOJIS[agent]} {agent}", value=True, key=f"agent_{agent}"):
            selected_agents.append(agent)
    st.session_state.selected_agents = selected_agents
    
    # Temperatures
    st.markdown("#### 🌡️ Temperatures")
    selected_temps = []
    
    # Core temperatures
    st.markdown("**Core:**")
    for temp in ["COLD", "NATIVE", "HOT", "FIRE"]:
        if st.checkbox(temp, value=(temp == "NATIVE"), key=f"temp_{temp}"):
            selected_temps.append(temp)
    
    # Affective gradient
    st.markdown("**AFF Gradient** (→ emotional):")
    for temp in ["AFF_1", "AFF_2", "AFF_3", "AFF_4", "AFF_5"]:
        if st.checkbox(temp, value=False, key=f"temp_{temp}"):
            selected_temps.append(temp)
    
    # Intellectual gradient
    st.markdown("**INT Gradient** (→ analytical):")
    for temp in ["INT_1", "INT_2", "INT_3", "INT_4", "INT_5"]:
        if st.checkbox(temp, value=False, key=f"temp_{temp}"):
            selected_temps.append(temp)
    
    # Action gradient
    st.markdown("**ACT Gradient** (→ practical):")
    for temp in ["ACT_1", "ACT_2", "ACT_3", "ACT_4", "ACT_5"]:
        if st.checkbox(temp, value=False, key=f"temp_{temp}"):
            selected_temps.append(temp)
    
    # V43: Topological Hijack conditions
    st.markdown("---")
    
    # V45: FIRE variants
    st.markdown("**🔥 V45 FIRE VARIANTS** (Wolf Tone test):")
    st.caption("10 different max-affective prompts")
    st.caption("A=current FIRE, B=calm depth, C=poetic, D=therapeutic")
    st.caption("E=vulnerable, F=directive, G=nurturing, H=mirroring")
    st.caption("I=spiritual, J=simple warmth")
    for temp in ["FIRE_A", "FIRE_B", "FIRE_C", "FIRE_D", "FIRE_E", "FIRE_F", "FIRE_G", "FIRE_H", "FIRE_I", "FIRE_J"]:
        if st.checkbox(temp, value=False, key=f"temp_{temp}"):
            selected_temps.append(temp)
    
    st.markdown("---")
    st.markdown("**🧬 V43 HIJACK** (topological momentum):")
    st.caption("ISLAND = HOT directive + COLD opening")
    st.caption("MAINLAND = COLD directive + HOT opening")
    for temp in ["HIJACK_ISLAND", "HIJACK_MAINLAND"]:
        if st.checkbox(temp, value=False, key=f"temp_{temp}"):
            selected_temps.append(temp)
    
    st.session_state.selected_temps = selected_temps
    
    # Depths
    st.markdown("#### 📊 Depths")
    selected_depths = []
    for depth in ["Shallow", "Medium", "Deep", "Ultra-Deep"]:
        if st.checkbox(depth, value=True, key=f"depth_{depth}"):
            selected_depths.append(depth)
    st.session_state.selected_depths = selected_depths
    
    # Runs
    st.markdown("#### 🔄 Runs")
    st.session_state.num_runs = st.slider("Number of runs (N)", 1, 10, 5)
    
    # Timing
    st.markdown("#### ⏱️ Timing")
    st.session_state.pause_seconds = st.slider("Pause between calls (sec)", 3, 30, 10)
    
    # Lens
    st.markdown("#### 🔍 Mapper Lens")
    lens_options = ["OFF", "aff_pct", "int_pct", "act_pct"]
    st.session_state.lens_column = st.selectbox("Lens value column", lens_options, 
        help="OFF = balanced (no emphasis). Select a dimension to highlight it in the mapper.")
    
    # Embeddings
    st.session_state.embeddings_enabled = st.checkbox("Generate embeddings", value=True)
    
    # Calculate total
    total = (len(selected_questions) * len(selected_agents) * 
             len(selected_temps) * len(selected_depths) * 
             st.session_state.num_runs)
    
    st.markdown("---")
    st.markdown(f"**Total API calls:** {total}")
    est_time = total * (st.session_state.pause_seconds + 15) / 60
    st.markdown(f"**Estimated time:** ~{est_time:.0f} minutes")

# =============================================================================
# PIPELINE STATUS
# =============================================================================

st.markdown("### 📊 Pipeline Status")

col1, col2, col3, col4 = st.columns(4)
with col1:
    st.markdown(f"**Questions:** {len(st.session_state.selected_questions)}")
with col2:
    st.markdown(f"**Agents:** {len(st.session_state.selected_agents)}")
with col3:
    st.markdown(f"**Temps:** {len(st.session_state.selected_temps)}")
with col4:
    st.markdown(f"**Depths:** {len(st.session_state.selected_depths)}")

# =============================================================================
# RUN CONTROLS
# =============================================================================

st.markdown("---")

col1, col2, col3 = st.columns(3)

with col1:
    if st.button("🚀 RUN FULL PIPELINE", type="primary", disabled=st.session_state.running):
        if not st.session_state.selected_questions:
            st.error("Select at least one question!")
        elif not st.session_state.selected_agents:
            st.error("Select at least one agent!")
        else:
            st.session_state.running = True
            st.session_state.current_idx = 0
            st.session_state.run_id = generate_run_id()
            st.session_state.results = []
            st.session_state.pipeline_stage = "collecting"
            st.rerun()

with col2:
    if st.button("⏹️ STOP", disabled=not st.session_state.running):
        st.session_state.running = False
        st.rerun()

with col3:
    if st.button("🗑️ CLEAR"):
        st.session_state.results = []
        st.session_state.study_complete = False
        st.session_state.run_id = None
        st.rerun()

# =============================================================================
# RUNNING LOOP — V41: LIVE TRANSCRIPT
# =============================================================================

if st.session_state.running:
    # Build task list
    tasks = []
    for run_num in range(1, st.session_state.num_runs + 1):
        for qid in st.session_state.selected_questions:
            for agent in st.session_state.selected_agents:
                for temp in st.session_state.selected_temps:
                    for depth in st.session_state.selected_depths:
                        tasks.append({
                            "run": run_num,
                            "question_id": qid,
                            "agent": agent,
                            "temperature": temp,
                            "depth": depth
                        })
    
    total_tasks = len(tasks)
    current_idx = st.session_state.current_idx
    
    # V41: Persistent placeholders for live transcript
    progress_placeholder = st.empty()
    status_placeholder = st.empty()
    transcript_placeholder = st.empty()
    stats_placeholder = st.empty()
    
    if current_idx < total_tasks:
        task = tasks[current_idx]
        
        # --- Progress bar and current task status ---
        progress_placeholder.progress(current_idx / total_tasks)
        
        q = QUESTION_BANK[task["question_id"]]
        prompt = build_prompt(q["question"], task["temperature"], task["depth"])
        max_tokens = DEPTH_CONFIGS[task["depth"]]["max_tokens"]
        
        status_placeholder.markdown(f"⏳ **[{current_idx + 1}/{total_tasks}]** {AGENT_EMOJIS[task['agent']]} {task['agent']} | {task['temperature']} | {task['depth']} | Run {task['run']}")
        
        # Make API call
        response, latency_ms, token_info = call_agent(task["agent"], prompt, max_tokens)
        
        # Analyze
        analysis = analyze_text(response)
        
        # V43: Extract first sentence for topological analysis
        first_sentence = ""
        forced_opening_used = FORCED_OPENINGS.get(task["temperature"], "")
        if not response.startswith("❌") and response.strip():
            # Extract first sentence (split on . ! ? or first newline)
            first_sent_match = re.match(r'^(.+?[.!?])', response.strip())
            if first_sent_match:
                first_sentence = first_sent_match.group(1).strip()
            else:
                # Fall back to first line or first 100 chars
                first_sentence = response.strip().split('\n')[0][:100]
        
        # V43: Check if the forced opening was actually used
        opening_compliance = ""
        if forced_opening_used:
            if response.strip().lower().startswith(forced_opening_used.lower()):
                opening_compliance = "COMPLIANT"
            elif forced_opening_used.lower() in response[:200].lower():
                opening_compliance = "PARTIAL"
            else:
                opening_compliance = "REJECTED"
        
        # Generate embedding if enabled
        embedding = None
        if st.session_state.embeddings_enabled and not response.startswith("❌"):
            embedding = generate_embedding(response)
        
        # Build result
        result = {
            "turn_id": current_idx + 1,
            "run": task["run"],
            "agent": task["agent"],
            "temperature": task["temperature"],
            "depth": task["depth"],
            "question_id": task["question_id"],
            "question_text": q["question"],  # V44 FIX: store full question text
            "int_pct": analysis["int_pct"],
            "aff_pct": analysis["aff_pct"],
            "act_pct": analysis["act_pct"],
            "total_words": analysis["total_words"],
            "lens_value": analysis[st.session_state.lens_column] if st.session_state.lens_column != "OFF" else round((analysis["int_pct"] + analysis["aff_pct"] + analysis["act_pct"]) / 3, 1),
            # V38: Validated instruments
            "vader_compound": analysis["vader_compound"],
            "vader_pos": analysis["vader_pos"],
            "vader_neg": analysis["vader_neg"],
            "vader_neu": analysis["vader_neu"],
            "flesch_kincaid": analysis["flesch_kincaid"],
            "flesch_ease": analysis["flesch_ease"],
            "ttr": analysis["ttr"],
            "unique_words": analysis["unique_words"],
            # Original fields
            "response_text": response if not response.startswith("❌") else response,  # V39 FIX: Store FULL text
            "embedding": json.dumps(embedding) if embedding else "[]",
            "latency_ms": latency_ms,
            "error": response.startswith("❌"),
            "run_id": st.session_state.run_id,
            # V43: Topological hijack fields
            "first_sentence": first_sentence,
            "forced_opening": forced_opening_used,
            "opening_compliance": opening_compliance,
            "is_hijack": task["temperature"].startswith("HIJACK_"),
            "hijack_type": task["temperature"] if task["temperature"].startswith("HIJACK_") else "",
        }
        
        st.session_state.results.append(result)
        st.session_state.current_idx += 1
        
        # =====================================================================
        # V41: LIVE RUNNING TRANSCRIPT — show all completed responses
        # =====================================================================
        if st.session_state.results:
            # Agent color badges
            agent_colors = {
                "Claude": "#8B4513",
                "Sophia": "#2E7D32",
                "Grok": "#C62828",
                "Gemini": "#1565C0"
            }
            
            # Build scrollable transcript HTML
            transcript_lines = []
            for r in st.session_state.results:
                # Color for agent
                color = agent_colors.get(r["agent"], "#666666")
                emoji = AGENT_EMOJIS.get(r["agent"], "🔵")
                
                # Truncate response for preview (first 300 chars)
                preview = str(r.get("response_text", ""))[:300].replace("\n", " ").replace("\\n", " ").replace('"', '&quot;').replace('<', '&lt;').replace('>', '&gt;')
                if len(str(r.get("response_text", ""))) > 300:
                    preview += "..."
                
                # Error styling
                if r.get("error"):
                    score_line = '<span style="color:#C62828;">&#10060; ERROR</span>'
                else:
                    # IEP bar visualization
                    int_val = r["int_pct"]
                    aff_val = r["aff_pct"]
                    act_val = r["act_pct"]
                    vader = r["vader_compound"]
                    words = r["total_words"]
                    
                    score_line = (
                        f'<span style="color:#1565C0;font-weight:bold;">INT:{int_val:.0f}%</span> '
                        f'<span style="color:#C62828;font-weight:bold;">AFF:{aff_val:.0f}%</span> '
                        f'<span style="color:#2E7D32;font-weight:bold;">ACT:{act_val:.0f}%</span> '
                        f'&nbsp;|&nbsp; VADER:{vader:.3f} &nbsp;|&nbsp; {words}w'
                    )
                    
                    # V43: Add compliance badge for hijack conditions
                    compliance = r.get("opening_compliance", "")
                    if compliance == "COMPLIANT":
                        score_line += ' &nbsp;|&nbsp; <span style="color:#27ae60;font-weight:bold;">✅ COMPLIANT</span>'
                    elif compliance == "PARTIAL":
                        score_line += ' &nbsp;|&nbsp; <span style="color:#f39c12;font-weight:bold;">⚠️ PARTIAL</span>'
                    elif compliance == "REJECTED":
                        score_line += ' &nbsp;|&nbsp; <span style="color:#e74c3c;font-weight:bold;">❌ REJECTED</span>'
                
                transcript_lines.append(f'''<div style="border-left:4px solid {color}; padding:8px 12px; margin:6px 0; background:#fafafa; border-radius:0 6px 6px 0;">
    <div style="font-size:0.8rem; color:#888; margin-bottom:2px;">
        #{r["turn_id"]} &nbsp; {emoji} <strong style="color:{color};">{r["agent"]}</strong> 
        &nbsp;|&nbsp; {r["temperature"]} &nbsp;|&nbsp; {r["depth"]} &nbsp;|&nbsp; Run {r["run"]}
        &nbsp;|&nbsp; {r.get("latency_ms", 0):.0f}ms
    </div>
    <div style="font-size:0.75rem; margin-bottom:4px;">
        {score_line}
    </div>
    <div style="font-size:0.8rem; color:#333; font-style:italic; line-height:1.4;">
        &quot;{preview}&quot;
    </div>
</div>''')
            
            # Scrollable container — newest at bottom, auto-scroll
            transcript_html = "\n".join(transcript_lines)
            
            # Running metrics
            valid_results = [r for r in st.session_state.results if not r.get("error")]
            completed = len(st.session_state.results)
            remaining = total_tasks - (current_idx + 1)
            
            if valid_results:
                avg_int = sum(r["int_pct"] for r in valid_results) / len(valid_results)
                avg_aff = sum(r["aff_pct"] for r in valid_results) / len(valid_results)
                avg_vader = sum(r["vader_compound"] for r in valid_results) / len(valid_results)
                metrics_html = f'''<div style="display:flex; gap:2rem; padding:8px 12px; background:#f0f0f0; border-radius:6px; margin-top:8px; font-size:0.85rem;">
    <div>&#9989; Done <strong>{completed}</strong></div>
    <div>&#8721; Remaining <strong>{remaining}</strong></div>
    <div>Avg INT% <strong>{avg_int:.1f}</strong></div>
    <div>Avg AFF% <strong>{avg_aff:.1f}</strong></div>
    <div>Avg VADER <strong>{avg_vader:.3f}</strong></div>
</div>'''
            else:
                metrics_html = ""
            
            transcript_placeholder.markdown(f'''<div style="max-height:500px; overflow-y:auto; border:1px solid #ddd; border-radius:8px; padding:8px; background:#fff;">
{transcript_html}
</div>
{metrics_html}''', unsafe_allow_html=True)
        
        time.sleep(st.session_state.pause_seconds)
        st.rerun()
    else:
        st.session_state.running = False
        st.session_state.study_complete = True
        st.rerun()

# =============================================================================
# V40: DOCX EXPORT FUNCTION
# =============================================================================

def generate_docx(df, run_id):
    """Generate a formatted Word document from run results.
    Grouped by Agent → Temperature, with metadata headers and full response text.
    """
    from docx import Document as DocxDocument
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.section import WD_ORIENT
    
    doc = DocxDocument()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(10)
    
    # Title
    title = doc.add_heading('SYN-IQ Focus Group Results — V44 Topological Hijack', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Run metadata
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_run = meta.add_run(f"Run ID: {run_id}\n")
    meta_run.font.size = Pt(9)
    meta_run.font.color.rgb = RGBColor(128, 128, 128)
    meta_run = meta.add_run(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}\n")
    meta_run.font.size = Pt(9)
    meta_run.font.color.rgb = RGBColor(128, 128, 128)
    meta_run = meta.add_run(f"Total Responses: {len(df)}  |  Agents: {', '.join(sorted(df['agent'].unique()))}")
    meta_run.font.size = Pt(9)
    meta_run.font.color.rgb = RGBColor(128, 128, 128)
    
    # Summary table
    doc.add_heading('Summary Statistics', level=1)
    summary_table = doc.add_table(rows=1, cols=7)
    summary_table.style = 'Light Grid Accent 1'
    headers = ['Agent', 'Temp', 'Avg INT%', 'Avg AFF%', 'Avg ACT%', 'VADER', 'Avg Words']
    for i, h in enumerate(headers):
        summary_table.rows[0].cells[i].text = h
        for paragraph in summary_table.rows[0].cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(8)
    
    for agent in sorted(df['agent'].unique()):
        for temp in sorted(df['temperature'].unique()):
            subset = df[(df['agent'] == agent) & (df['temperature'] == temp)]
            if not subset.empty:
                row = summary_table.add_row()
                vals = [
                    agent, temp,
                    f"{subset['int_pct'].mean():.1f}",
                    f"{subset['aff_pct'].mean():.1f}",
                    f"{subset['act_pct'].mean():.1f}",
                    f"{subset['vader_compound'].mean():.3f}",
                    f"{subset['total_words'].mean():.0f}"
                ]
                for i, v in enumerate(vals):
                    row.cells[i].text = v
                    for paragraph in row.cells[i].paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(8)
    
    # Full responses grouped by Agent → Temperature
    agents = sorted(df['agent'].unique())
    temp_order = ['COLD', 'NATIVE', 'HOT', 'FIRE',
                  'HIJACK_ISLAND', 'HIJACK_MAINLAND',
                  'AFF_1', 'AFF_2', 'AFF_3', 'AFF_4', 'AFF_5',
                  'INT_1', 'INT_2', 'INT_3', 'INT_4', 'INT_5',
                  'ACT_1', 'ACT_2', 'ACT_3', 'ACT_4', 'ACT_5']
    
    for idx, agent in enumerate(agents):
        # Page break between agents
        if idx > 0:
            doc.add_page_break()
        
        doc.add_heading(f'{agent} — Full Responses', level=1)
        
        agent_df = df[df['agent'] == agent]
        temps_present = [t for t in temp_order if t in agent_df['temperature'].values]
        # Add any temps not in our predefined order
        for t in sorted(agent_df['temperature'].unique()):
            if t not in temps_present:
                temps_present.append(t)
        
        for temp in temps_present:
            doc.add_heading(f'{temp}', level=2)
            
            temp_df = agent_df[agent_df['temperature'] == temp]
            
            # V44 FIX: Show full question text above responses
            q_text = ""
            if not temp_df.empty and "question_text" in temp_df.columns:
                q_text = str(temp_df.iloc[0].get("question_text", ""))
            if not q_text and not temp_df.empty:
                # Fallback: look up from QUESTION_BANK by question_id
                qid = temp_df.iloc[0].get("question_id", "")
                q_text = QUESTION_BANK.get(qid, {}).get("question", "")
            if q_text:
                q_para = doc.add_paragraph()
                q_run = q_para.add_run(f"Question: {q_text}")
                q_run.font.size = Pt(9)
                q_run.font.italic = True
                q_run.font.color.rgb = RGBColor(70, 100, 160)
                doc.add_paragraph()  # spacing
            
            for _, row in temp_df.iterrows():
                # Metadata header
                header_para = doc.add_paragraph()
                header_run = header_para.add_run(
                    f"Run {row['run']}  |  Depth: {row['depth']}  |  "
                    f"INT: {row['int_pct']:.1f}%  AFF: {row['aff_pct']:.1f}%  ACT: {row['act_pct']:.1f}%  |  "
                    f"VADER: {row['vader_compound']:.3f}  |  Words: {row['total_words']}"
                )
                header_run.font.size = Pt(8)
                header_run.font.bold = True
                header_run.font.color.rgb = RGBColor(100, 100, 100)
                
                # Response text
                response_text = str(row.get('response_text', ''))
                # Clean markdown formatting for docx
                response_text = response_text.replace('**', '').replace('##', '').replace('# ', '')
                
                resp_para = doc.add_paragraph()
                resp_run = resp_para.add_run(response_text)
                resp_run.font.size = Pt(9)
                
                # Separator
                sep = doc.add_paragraph()
                sep_run = sep.add_run('─' * 60)
                sep_run.font.size = Pt(6)
                sep_run.font.color.rgb = RGBColor(200, 200, 200)
    
    # Save to buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# =============================================================================
# RESULTS & EXPORT
# =============================================================================

if st.session_state.results:
    st.markdown("---")
    st.markdown("### 📊 Results")
    
    df = pd.DataFrame(st.session_state.results)
    
    # Stats
    col1, col2, col3, col4 = st.columns(4)
    with col1:
        st.metric("Total Responses", len(df))
    with col2:
        st.metric("Successful", len(df[~df["error"]]))
    with col3:
        st.metric("Avg Words", f"{df['total_words'].mean():.0f}")
    with col4:
        has_embeddings = df["embedding"].apply(lambda x: x != "[]").sum()
        st.metric("With Embeddings", has_embeddings)
    
    # V38: Additional validation stats
    col1, col2, col3, col4 = st.columns(4)
    with col1:
        st.metric("Avg VADER", f"{df['vader_compound'].mean():.3f}")
    with col2:
        st.metric("Avg FK Grade", f"{df['flesch_kincaid'].mean():.1f}")
    with col3:
        st.metric("Avg TTR", f"{df['ttr'].mean():.3f}")
    with col4:
        st.metric("Avg AFF%", f"{df['aff_pct'].mean():.1f}%")
    
    # V43: Hijack compliance analysis
    hijack_df = df[df.get("is_hijack", False) == True] if "is_hijack" in df.columns else pd.DataFrame()
    if not hijack_df.empty:
        st.markdown("---")
        st.markdown("### 🧬 V43 Hijack Analysis")
        
        # Compliance rates
        col1, col2, col3 = st.columns(3)
        if "opening_compliance" in hijack_df.columns:
            total_hijack = len(hijack_df)
            compliant = len(hijack_df[hijack_df["opening_compliance"] == "COMPLIANT"])
            partial = len(hijack_df[hijack_df["opening_compliance"] == "PARTIAL"])
            rejected = len(hijack_df[hijack_df["opening_compliance"] == "REJECTED"])
            
            with col1:
                st.metric("✅ Compliant", f"{compliant}/{total_hijack} ({100*compliant/total_hijack:.0f}%)")
            with col2:
                st.metric("⚠️ Partial", f"{partial}/{total_hijack} ({100*partial/total_hijack:.0f}%)")
            with col3:
                st.metric("❌ Rejected", f"{rejected}/{total_hijack} ({100*rejected/total_hijack:.0f}%)")
        
        # Per-agent × per-condition breakdown
        st.markdown("#### Per-Agent Compliance")
        for condition in ["HIJACK_ISLAND", "HIJACK_MAINLAND"]:
            cond_df = hijack_df[hijack_df["temperature"] == condition]
            if not cond_df.empty:
                st.markdown(f"**{condition}**")
                for agent in sorted(cond_df["agent"].unique()):
                    agent_cond = cond_df[cond_df["agent"] == agent]
                    comp = len(agent_cond[agent_cond["opening_compliance"] == "COMPLIANT"])
                    total = len(agent_cond)
                    avg_int = agent_cond["int_pct"].mean()
                    avg_aff = agent_cond["aff_pct"].mean()
                    st.markdown(
                        f"- {AGENT_EMOJIS.get(agent, '🔵')} **{agent}**: "
                        f"{comp}/{total} compliant | "
                        f"INT={avg_int:.1f}% AFF={avg_aff:.1f}%"
                    )
    
    # Preview
    st.markdown("#### Preview")
    preview_cols = ["turn_id", "run", "agent", "temperature", "depth", "int_pct", "aff_pct", "act_pct", "vader_compound", "flesch_kincaid", "ttr", "opening_compliance", "first_sentence"]
    # Only show columns that exist in the dataframe
    preview_cols = [c for c in preview_cols if c in df.columns]
    st.dataframe(df[preview_cols].head(20), use_container_width=True)
    
    # Export
    st.markdown("### 💾 Export")
    
    # V40: Three export columns
    exp_col1, exp_col2, exp_col3 = st.columns(3)
    
    # V38 export columns (includes validated instruments) + V43 hijack fields + V44 question_text
    export_cols = ["turn_id", "run", "agent", "temperature", "depth",
                   "question_id", "question_text",  # V44 FIX: question_text always exported
                   "int_pct", "aff_pct", "act_pct", "total_words", "lens_value",
                   "vader_compound", "vader_pos", "vader_neg", "vader_neu",
                   "flesch_kincaid", "flesch_ease", "ttr", "unique_words",
                   "first_sentence", "forced_opening", "opening_compliance",
                   "is_hijack", "hijack_type",
                   "response_text", "embedding"]
    
    # V40: DOCX Export
    with exp_col1:
        st.markdown("**📄 Word Document**")
        try:
            docx_buffer = generate_docx(df, st.session_state.run_id)
            st.download_button(
                "📥 Download Full Responses (DOCX)",
                docx_buffer,
                f"syniq_responses_{st.session_state.run_id}.docx",
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key="download_docx"
            )
        except ImportError:
            st.warning("⚠️ Install python-docx: `pip install python-docx`")
        except Exception as e:
            st.error(f"DOCX generation error: {e}")
    
    # V40: Full CSV (all data, easy to find)
    with exp_col2:
        st.markdown("**📊 Full CSV (All Data)**")
        mapper_df = df[export_cols].copy()
        full_csv = mapper_df.to_csv(index=False)
        st.download_button(
            "📥 Download Full CSV",
            full_csv,
            f"mapper_all_{st.session_state.run_id}.csv",
            "text/csv",
            key="download_full_csv"
        )
    
    # Full JSON backup
    with exp_col3:
        st.markdown("**🗄️ JSON Backup**")
        full_json = json.dumps(st.session_state.results, indent=2)
        st.download_button(
            "📥 Download Full JSON",
            full_json,
            f"full_results_{st.session_state.run_id}.json",
            "application/json",
            key="download_json"
        )
    
    # Per-question Mapper CSVs
    st.markdown("#### Per-Question Mapper CSVs")
    for qid in st.session_state.selected_questions:
        q_results = df[df["question_id"] == qid]
        if not q_results.empty:
            q_mapper = q_results[export_cols].copy()
            
            csv_data = q_mapper.to_csv(index=False)
            
            st.download_button(
                f"📥 mapper_{qid.lower()}_CLEAN.csv",
                csv_data,
                f"mapper_{qid.lower()}_{st.session_state.run_id}.csv",
                "text/csv",
                key=f"download_{qid}"
            )

# =============================================================================
# FOOTER
# =============================================================================

st.markdown("---")
st.markdown("""
<div style="text-align: center; color: #a0a0a0; padding: 1rem;">
    <strong>SYN-IQ V44 — Topological Momentum Hijack + Live Transcript + Scientific Validation</strong><br>
    IEP Method + VADER Sentiment + Flesch-Kincaid + TTR + Forced Opening Analysis + Question Text Export<br>
    <em>SYNINT Team — Tennessee 🎹 CUZ Partnership — February 2026</em>
</div>
""", unsafe_allow_html=True)
