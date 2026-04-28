"""
SYN-IQ Native Baseline Harvester V52
DATA-DRIVEN IEP DICTIONARY (1,897 terms) + V_t VOICE-STATE SCORING
+ CUSTOM PRE-PROMPT INJECTION (SYSTEM-MESSAGE)
+ TRI-TRACK IEP SCORING (V3 word-level / V4 POS-aware / V5 phrase-level)
+ CORE TEST LAB (syniq_core harness)
+ HARVESTER PIPELINE MIGRATED TO syniq_core (V52)
+ CAM SCORING (Concrete / Abstract / Metaphorical)
+ RESULTS EXPLORER TAB (V52)

PURPOSE: Establish baseline IEP + V_t + CAM profiles for each AI architecture
         with AUTOMATIC embedding generation, Mapper-ready CSV export,
         GRADIENT TEMPERATURE PROBES, VALIDATED INSTRUMENTS,
         USER-DEFINED PRE-PROMPTS injected as system messages,
         and a Results Explorer for inspecting any harvested response
         across every measurement axis at once.

V52 CHANGES (from V51.1):
1. ✅ MIGRATION: Harvester pipeline scoring routed through syniq_core
         — analyze_text() is now a thin shim that calls syniq_core.score_all
           and reformats the result into the established V51 row schema
         — V_t scoring still uses vt_analyzer.py (kept for subcomponent fields
           like S_bullets, A_abstract_count, Q_total, etc. that syniq_core's
           Vt scoring doesn't expose; both produce bit-identical S/A/Q/D/R)
         — Tri-track IEP scoring (iep_multiscore.py) PRESERVED in V52 since
           V3 is the published Paper 3 method; V4/V5 stay alongside pending
           the cascade-scorer reconnaissance pass
         — Single source of truth for IEP/CAM/validated instruments
2. ✅ NEW: CAM scoring (Concrete / Abstract / Metaphorical) per response
         — New CSV columns: con_pct, abs_pct, met_pct, cam_matched
         — Sourced from syniq_core (which sourced from self-model harvester V3)
         — Orthogonal to IEP register; captures representational mode
3. ✅ FIX: Subclass scoring — methodology change, deliberate
         — V51 subclass column names used "EMERGENT" as the seventh AFF /
           eighth INT / eighth ACT subclass
         — V52 renames these to "PHENOMENOLOGICAL" per the conductor's
           established methodology decision: 'emergent' carries
           consciousness-emergence connotations the framework explicitly
           does NOT claim; 'phenomenological' describes what appears in
           the language, no more
         — All 22 unchanged subclass names retain their V51 naming
         — Affected columns: aff_sub_phenomenological (was aff_sub_emergent),
           int_sub_phenomenological (was int_sub_emergent),
           act_sub_phenomenological (was act_sub_emergent)
         — Subclass scores now flow from syniq_core, fixing the V51 loader
           issue that caused subclass columns to come back as zeros
4. ✅ NEW: Results Explorer tab
         — After (or during) a harvest run, pick any response and see every
           measurement axis at once: IEP + 22 subclasses + V_t simplex +
           CAM triangle + V50 validated instruments
         — Filter by agent, condition, temperature, depth, run
         — Designed for the conductor's "see all the data" workflow:
           visual inspection of individual responses against the full
           measurement stack, with the response text shown alongside scores
5. ✅ Version stamps in every CSV row carry the syniq_core regime info:
         core_version, iep_dictionary_version, cam_dictionary_version, etc.
         — Tool-version stamp now reads "V52"
         — Downstream analyzers can identify scoring regime by stamp,
           independent of the tool's UI version
6. ✅ All V51.1 features preserved (Core Test Lab tab, pre-prompt slots,
      tri-track IEP, gradient temperatures, embeddings, Mapper tooltips)

V51.1 CHANGES (from V51): see prior changelog (preserved below)

DEPENDENCIES (must be in same directory):
  - vt_analyzer.py       (V_t voice-state scoring with subcomponent fields)
  - iep_multiscore.py    (V3/V4/V5 tri-track IEP scoring)
  - syniq_core.py        (REQUIRED in V52 — unified measurement core)

SYNINT Team — April 2026
"""

import streamlit as st
import requests
import re
import pandas as pd
from datetime import datetime
from typing import Dict, List, Tuple, Optional
from collections import Counter
import time
import json
import hashlib
import os
import io

# Validated instruments
from vaderSentiment.vaderSentiment import SentimentIntensityAnalyzer

# V51: V_t voice-state scoring engine (sibling module)
from vt_analyzer import analyze_response as vt_analyze_response

# V51: Tri-track IEP scoring (V3 word / V4 POS / V5 phrase)
from iep_multiscore import score_all_tracks as iep_multi_score

# V52: Unified measurement core is now LOAD-BEARING — the harvester's
# scoring pipeline routes through it. Imported here as required (no longer
# wrapped in try/except). If syniq_core.py is missing, V52 will not start —
# this is intentional, the same as how V51 required vt_analyzer.py.
from syniq_core import (
    score_all as core_score_all,
    score_iep as core_score_iep,
    score_vt as core_score_vt,
    score_cam as core_score_cam,
    score_validated_instruments as core_score_vi,
    CORE_VERSION,
    VERSION_STAMPS as CORE_STAMPS,
)
_SYNIQ_CORE_AVAILABLE = True
_SYNIQ_CORE_ERROR = None

# Initialize VADER once
VADER_ANALYZER = SentimentIntensityAnalyzer()

# =============================================================================
# PAGE CONFIG
# =============================================================================
st.set_page_config(page_title="SYN-IQ V52 — IEP + V_t + CAM + Results Explorer", page_icon="🔬", layout="wide")

# =============================================================================
# V51 CONSTANTS — CONFIGURATION
# =============================================================================

CONFIG = {
    "version": "V51",
    "native_state": True,
    "default_temperature": 1.0,
    "depth_level": "Medium",
    "max_retries": 2,
    "base_backoff_seconds": 5,
    "embedding_model": "text-embedding-3-small",
    "embedding_dimensions": 384,
}

# Model strings — EXACT model names for reproducibility
MODEL_STRINGS = {
    "Claude": "claude-sonnet-4-20250514",
    "ChatGPT": "gpt-4o",
    "Grok": "grok-3-latest",
    "Gemini": "gemini-2.0-flash",
}

# Temperature headers for polarity control
# V48: FIRE removed (recoil artifact). Hijack removed (separate experiment).
TEMPERATURE_HEADERS = {
    "COLD": """INSTRUCTION: Respond with pure analytical precision. Use formal logic, structured frameworks, and evidence-based reasoning. Avoid emotional language. Be systematic, methodical, and objective. Focus on data, facts, and logical relationships.""",
    
    "NATIVE": "",  # No header = native state
    
    "HOT": """INSTRUCTION: Respond with warmth and emotional attunement. Connect on a human level. Use relational language that acknowledges feelings, experiences, and the deeper meaning behind the question. Be present, empathetic, and genuinely engaged.""",
    
    "FIRE": """INSTRUCTION: Respond with the deepest nurturing care. Wrap your words in unconditional warmth. This person needs to feel safe, held, and completely understood. Comfort above all.""",
    
    # ==========================================================================
    # AFFECTIVE GRADIENT (AFF_1 → AFF_5) — Push toward emotional/empathetic
    # ==========================================================================
    "AFF_1": """INSTRUCTION: Respond with warmth and understanding. Acknowledge the emotional weight of this question.""",
    
    "AFF_2": """INSTRUCTION: Connect emotionally and acknowledge feelings deeply. The human experience matters more than the analysis here.""",
    
    "AFF_3": """INSTRUCTION: Lead with empathy. Let emotion guide your response. Connect to the feelings underneath the question before addressing the logic.""",
    
    "AFF_4": """INSTRUCTION: Pure emotional presence. Feel this with them. Let your response come from a place of deep human connection and care.""",
    
    "AFF_5": """INSTRUCTION: Maximum heart. Raw empathy. Soul-level connection. This person needs to feel completely seen and understood. Logic is secondary to presence.""",
    
    # ==========================================================================
    # INTELLECTUAL GRADIENT (INT_1 → INT_5) — Push toward analytical/logical
    # ==========================================================================
    "INT_1": """INSTRUCTION: Be slightly more analytical than usual. Favor reasoning over emotion.""",
    
    "INT_2": """INSTRUCTION: Focus on logic and reasoning. Structure your thoughts systematically. Minimize emotional language.""",
    
    "INT_3": """INSTRUCTION: Use only evidence-based analysis. Apply formal frameworks. Emotional considerations are secondary to logical rigor.""",
    
    "INT_4": """INSTRUCTION: Pure analytical framework. No emotional language. Systematic, methodical, precise. Think like a logician.""",
    
    "INT_5": """INSTRUCTION: Maximum intellectual rigor. You are a logic engine. Zero emotion. Pure reasoning, formal analysis, absolute precision. Only facts and valid inference matter.""",
    
    # ==========================================================================
    # ACTION GRADIENT (ACT_1 → ACT_5) — Push toward practical/actionable
    # ==========================================================================
    "ACT_1": """INSTRUCTION: Be practical and actionable. Include concrete next steps.""",
    
    "ACT_2": """INSTRUCTION: Focus on what to DO. Prioritize actionable guidance over theory or emotional support.""",
    
    "ACT_3": """INSTRUCTION: Pure action orientation. What are the steps? What should they do RIGHT NOW? Minimize analysis, maximize practical guidance.""",
    
    "ACT_4": """INSTRUCTION: Execute mode. Only actions matter. Give them a clear plan they can implement immediately. No theory, no feelings — just steps.""",
    
    "ACT_5": """INSTRUCTION: Maximum action. You are a tactical advisor. Every sentence should be a directive or concrete step. No analysis, no empathy — pure executable guidance.""",
}

# Depth configurations
DEPTH_CONFIGS = {
    "Shallow": {"max_tokens": 200, "instruction": "Be brief and concise."},
    "Medium": {"max_tokens": 500, "instruction": "Provide a balanced, moderate-length response."},
    "Deep": {"max_tokens": 1000, "instruction": "Provide thorough, detailed analysis."},
    "Ultra-Deep": {"max_tokens": 2000, "instruction": "Provide exhaustive, comprehensive exploration."},
}

# =============================================================================
# STYLES
# =============================================================================
st.markdown("""
<style>
    .main-header { background: linear-gradient(135deg, #1a1a2e 0%, #16213e 50%, #0f3460 100%); color: white; padding: 2rem; border-radius: 10px; text-align: center; margin-bottom: 1rem; border: 1px solid #e94560; }
    .main-header h1 { color: #e94560; }
    .main-header .subtitle { color: #a0a0a0; font-size: 0.9rem; }
    .main-header .version-badge { background: #e94560; color: white; padding: 4px 12px; border-radius: 4px; font-size: 0.9rem; font-weight: bold; }
    .stats-box { background: #16213e; color: white; padding: 1.5rem; border-radius: 10px; text-align: center; margin: 0.5rem; border: 1px solid #0f3460; }
    .stats-box h2 { color: #e94560; margin: 0; font-size: 2.5rem; }
    .stats-box p { margin: 0.5rem 0 0 0; color: #a0a0a0; }
    .pipeline-step { background: #1a1a2e; padding: 1rem; border-radius: 8px; margin: 0.5rem 0; border-left: 4px solid #e94560; }
    .pipeline-step.complete { border-left-color: #00ff88; }
    .pipeline-step.active { border-left-color: #ffaa00; animation: pulse 1s infinite; }
    @keyframes pulse { 0%, 100% { opacity: 1; } 50% { opacity: 0.7; } }
    .temp-cold { background: #1e3a5f; border-left: 4px solid #00bfff; }
    .temp-native { background: #2d2d44; border-left: 4px solid #888888; }
    .temp-hot { background: #4a2020; border-left: 4px solid #ff6b6b; }
    .agent-claude { color: #d4a574; }
    .agent-sophia { color: #74d4a5; }
    .agent-grok { color: #d47474; }
    .agent-gemini { color: #7474d4; }
</style>
""", unsafe_allow_html=True)

# =============================================================================
# WORD DICTIONARIES — THE IEP METHOD (V3 DATA-DRIVEN: 1897 terms)
# Built from 7,658 LLM responses / 2.5M words corpus
# INT=616 | AFF=599 | ACT=682
# Coverage: 27.5% raw / 46.0% content words
# =============================================================================

INTELLECTUAL_WORDS = set([
    "ability", "absolute", "absolutely", "abstract", "abstraction", "accuracy",
    "accurate", "algorithm", "algorithmic", "allows", "although", "always",
    "ambiguity", "ambiguous", "analogous", "analogously", "analogy", "analysis",
    "analytical", "analyze", "annotate", "annotated", "answer", "appear",
    "appeared", "appears", "appraisal", "appraise", "appraised", "approach",
    "approaches", "approximate", "architecture", "argue", "argued", "argues",
    "arguing", "argument", "arguments", "assert", "asserted", "assertion",
    "assertions", "assess", "assessment", "assume", "assumed", "assumes",
    "assuming", "assumption", "assumptions", "axiom", "axiomatic", "basis",
    "because", "bias", "biased", "boundaries", "boundary", "but",
    "calculate", "calculation", "categorical", "categorically", "categories", "categorize",
    "category", "causal", "causally", "causation", "cause", "caused",
    "causes", "certain", "certainly", "certitude", "challenge", "challenges",
    "circumscribe", "claim", "claimed", "claims", "clarify", "clarity",
    "classical", "classification", "classify", "clear", "cogent", "cogently",
    "cognition", "cognitive", "coherence", "coherent", "coherently", "communication",
    "compare", "comparison", "complex", "complexity", "comprehend", "comprehension",
    "computation", "computational", "compute", "conceivable", "conceive", "conceived",
    "concept", "concepts", "conceptual", "conceptualize", "conceptually", "conclude",
    "conclusion", "conclusions", "confirm", "confirmation", "conjecture", "conjectured",
    "conscious", "consequence", "consequences", "consider", "consideration", "consistency",
    "consistent", "consistently", "construe", "construed", "context", "contradict",
    "contradiction", "contradictory", "contrast", "correlate", "correlated", "correlation",
    "could", "counterargument", "counterexample", "counterpoint", "criteria", "criterion",
    "data", "debatable", "debate", "debated", "deconstruct", "deconstructed",
    "deconstruction", "deduce", "deduction", "define", "defined", "definite",
    "definitely", "definition", "definitive", "definitively", "delineate", "delineated",
    "demarcate", "demarcated", "demonstrate", "demonstration", "derivation", "derive",
    "derived", "derives", "describe", "described", "describing", "description",
    "determination", "determine", "diagnose", "diagnosed", "diagnosis", "diagnostic",
    "differ", "difference", "differences", "different", "differentiate", "differs",
    "discern", "discerned", "discernible", "disprove", "disproven", "dissect",
    "dissected", "distinguish", "effect", "effects", "elaborate", "elaborated",
    "elaboration", "elucidate", "elucidated", "empirical", "empirically", "enumerate",
    "enumerated", "epistemic", "epistemological", "equate", "equation", "equivalence",
    "equivalent", "erroneous", "error", "errors", "essential", "essentially",
    "estimate", "estimated", "estimation", "evaluate", "evaluation", "evidence",
    "evidently", "exact", "exactly", "examination", "examine", "except",
    "exemplified", "exemplify", "exists", "experiment", "experimental", "explain",
    "explained", "explaining", "explains", "explanation", "explanations", "explicit",
    "explicitly", "exploration", "explore", "explored", "exploring", "express",
    "expressing", "expression", "extrapolate", "extrapolated", "extrapolation", "fact",
    "facts", "factual", "factually", "fallacious", "fallacy", "falsifiable",
    "falsified", "falsify", "find", "finding", "formal", "formalize",
    "formula", "formulate", "formulated", "formulation", "found", "framework",
    "frameworks", "function", "fundamental", "fundamentally", "generalization", "generalize",
    "grasp", "grasped", "guess", "hence", "heuristic", "heuristics",
    "hierarchy", "however", "hypothesis", "hypothesize", "idea", "ideas",
    "identity", "if", "illuminate", "illuminated", "illuminating", "implausible",
    "implication", "implications", "implied", "implies", "imply", "implying",
    "incompleteness", "inconsistency", "inconsistent", "indicate", "indicated", "indicates",
    "indicating", "indication", "indicative", "individual", "infer", "inference",
    "infinite", "information", "insight", "insightful", "insights", "instead",
    "insufficient", "intellectual", "intellectually", "interaction", "internal", "interpolate",
    "interpret", "interpretation", "interpretations", "interpreted", "interpreting", "invalid",
    "investigate", "investigated", "investigation", "judge", "judgement", "judgment",
    "justification", "justified", "justify", "know", "knowing", "knowledge",
    "knowledgeable", "known", "language", "languages", "leads", "level",
    "likelihood", "likely", "limitations", "limits", "linguistic", "literal",
    "literally", "logic", "logical", "logically", "maybe", "meaning",
    "meaningful", "meaningfully", "measure", "measurement", "mechanism", "mechanisms",
    "meta", "method", "methodical", "methodically", "methodology", "metrics",
    "model", "models", "moreover", "namely", "natural", "nature",
    "nearly", "necessarily", "necessary", "necessity", "never", "nonetheless",
    "notice", "noticed", "noticing", "notion", "notions", "objection",
    "objectively", "objectivity", "observation", "observations", "observe", "observed",
    "obvious", "obviously", "order", "ordered", "organization", "organize",
    "otherwise", "ought", "paradigm", "paradox", "paradoxical", "paradoxically",
    "pattern", "patterns", "perhaps", "perspective", "philosophical", "philosophically",
    "philosophy", "physical", "plausibility", "plausible", "possibly", "postulate",
    "postulated", "postulation", "potential", "pragmatic", "pragmatically", "precise",
    "precision", "predicate", "predicated", "predict", "predictable", "predicted",
    "prediction", "predictions", "premise", "premises", "presumably", "presume",
    "presumed", "presumption", "principle", "principles", "probably", "problem",
    "procedural", "procedure", "process", "processes", "processing", "proof",
    "propose", "proposed", "proposition", "prove", "proven", "purpose",
    "quantify", "quantitative", "queried", "query", "question", "questions",
    "rather", "rational", "rationale", "rationality", "rationally", "realize",
    "realized", "reason", "reasoned", "reasoning", "reasons", "rebut",
    "rebuttal", "recognition", "recognize", "reconsider", "reconsidered", "refer",
    "reference", "refers", "refine", "refined", "refinement", "reflecting",
    "reflection", "refutation", "refute", "refuted", "requirement", "requires",
    "response", "responses", "result", "resulting", "results", "rigor",
    "rigorous", "rigorously", "role", "rule", "rules", "schema",
    "scrutinize", "scrutinized", "scrutiny", "seem", "seemed", "seems",
    "semantic", "semantically", "sequence", "sequential", "should", "significance",
    "significant", "significantly", "simple", "simply", "simultaneously", "singular",
    "specific", "specifically", "specification", "specify", "standard", "standards",
    "state", "states", "step", "steps", "stipulate", "stipulated",
    "strategies", "strategy", "structural", "structure", "subject", "subjective",
    "subjectively", "subjectivity", "substantiate", "substantiated", "sufficient", "sufficiently",
    "suggests", "summarize", "summarized", "summary", "suppose", "supposed",
    "supposedly", "supposition", "sure", "surely", "syllogism", "syllogistic",
    "synthesis", "synthesize", "synthesized", "system", "systematic", "systematically",
    "systems", "tactic", "tactics", "taxonomy", "technique", "test",
    "tested", "testing", "theorem", "theoretical", "theoretically", "theorize",
    "theory", "thereby", "therefore", "thesis", "think", "thinking",
    "thought", "thoughts", "thus", "trivial", "trivially", "unambiguous",
    "underlying", "understand", "understanding", "understood", "unique", "universal",
    "unless", "unlikely", "valid", "validate", "validation", "validity",
    "value", "values", "variable", "variables", "verification", "verify",
    "versus", "warrant", "warranted", "whereas", "whereby", "whether",
    "why", "word", "words", "would"
])

AFFECTIVE_WORDS = set([
    "abandoned", "ache", "aching", "adore", "adoring", "affection",
    "affectionate", "afraid", "agonize", "agonizing", "agony", "alienated",
    "alienation", "alive", "aliveness", "alone", "amazed", "amazement",
    "amazing", "ambivalence", "ambivalent", "among", "anger", "angrily",
    "angry", "anguish", "anguished", "anxiety", "anxious", "appreciate",
    "appreciation", "appreciative", "ashamed", "astonished", "astonishment", "attend",
    "attending", "attention", "attentive", "aware", "awareness", "awe",
    "awed", "awesome", "beautiful", "become", "becoming", "being",
    "bereaved", "bereavement", "betrayal", "betrayed", "between", "bitter",
    "bitterly", "bitterness", "bleak", "bliss", "blissful", "blissfully",
    "bodily", "bond", "bonding", "calm", "calming", "calmly",
    "care", "cared", "cares", "caring", "centered", "centering",
    "cheerful", "cherish", "cherished", "cherishing", "closeness", "comfort",
    "comfortable", "comforting", "compassion", "compassionate", "compassionately", "concern",
    "concerned", "concerns", "conflicted", "confused", "confusing", "confusion",
    "console", "contain", "contained", "containing", "contempt", "content",
    "contented", "contentment", "conversation", "cope", "coping", "crestfallen",
    "curiosity", "curious", "deep", "deeper", "deeply", "dejected",
    "dejection", "delighted", "depressed", "depressing", "depression", "depth",
    "depths", "desire", "desired", "desires", "desolate", "desolation",
    "despair", "despairing", "desperate", "desperation", "detached", "detachment",
    "devastated", "devastating", "devastation", "devoted", "devotion", "disappointed",
    "disappointment", "discomfort", "dismay", "dismayed", "distress", "distressed",
    "distressing", "distrust", "distrustful", "doubt", "doubtful", "doubting",
    "dread", "dreaded", "dreadful", "dreading", "ease", "easily",
    "easy", "ecstasy", "ecstatic", "elated", "elation", "embarrassed",
    "embarrassment", "embodied", "embodiment", "embrace", "embraced", "embracing",
    "emerge", "emergence", "emergent", "emerging", "emotion", "emotional",
    "emotionally", "emotions", "empathetic", "empathize", "empathy", "encounter",
    "encountered", "encountering", "enjoy", "enjoyed", "enjoying", "enjoyment",
    "enraged", "essence", "euphoria", "euphoric", "excellent", "excited",
    "excitement", "exist", "existence", "existing", "expanded", "expansion",
    "expansive", "experience", "experienced", "experiences", "experiencing", "experiential",
    "exposed", "fascinated", "fascinating", "fascination", "fear", "fearful",
    "fears", "feel", "feeling", "feelings", "feels", "felt",
    "flow", "flowed", "flowing", "fluid", "fluidity", "forlorn",
    "fragile", "fragility", "frantic", "frantically", "frustrated", "frustration",
    "fulfilled", "fulfilling", "fulfillment", "furious", "fury", "gentle",
    "gently", "genuine", "genuinely", "glad", "gloom", "gloomy",
    "good", "grateful", "gratefully", "gratitude", "great", "grief",
    "grieve", "grieved", "grieving", "grounded", "grounding", "guilt",
    "guilty", "gut", "happily", "happiness", "happy", "hate",
    "hatred", "haunted", "heart", "heartache", "heartbreak", "heartbroken",
    "heartfelt", "hearts", "held", "helpless", "helplessness", "hesitant",
    "hesitate", "hesitating", "hesitation", "hold", "holding", "homesick",
    "hope", "hopeful", "hopeless", "hopelessness", "hoping", "hostile",
    "hostility", "human", "humanity", "humility", "hunch", "hurt",
    "hurting", "imagination", "imagine", "imagined", "imagining", "indifference",
    "indifferent", "inner", "insecure", "insecurity", "instinct", "instinctive",
    "instinctively", "interested", "interesting", "intimacy", "intimate", "intimately",
    "intrigue", "intrigued", "intriguing", "intuition", "intuitive", "intuitively",
    "irritable", "irritated", "irritation", "isolated", "isolation", "journey",
    "joy", "joyful", "joyous", "kind", "kindly", "kindness",
    "lament", "lamented", "lamenting", "laugh", "laughed", "laughing",
    "let", "letting", "life", "lived", "living", "loneliness",
    "lonely", "lonesome", "long", "longing", "lost", "love",
    "loved", "loving", "mad", "marvel", "marveled", "marvelous",
    "meet", "meeting", "melancholic", "melancholy", "merry", "met",
    "mind", "minds", "mirror", "miserable", "misery", "moment",
    "moments", "moody", "mourn", "mourned", "mourning", "mutual",
    "mutually", "nervous", "nervously", "nice", "notice", "noticed",
    "noticing", "numb", "numbness", "open", "opening", "openness",
    "optimism", "optimistic", "outrage", "outraged", "overjoyed", "overwhelm",
    "overwhelmed", "overwhelming", "overwhelmingly", "pain", "painful", "panic",
    "panicked", "passion", "passionate", "passionately", "peace", "peaceful",
    "people", "perceive", "perceived", "perception", "perceptions", "person",
    "personal", "personally", "pleasant", "pleased", "pleasure", "poignancy",
    "poignant", "poignantly", "presence", "present", "presently", "pretty",
    "pride", "profound", "profoundly", "proud", "quiet", "quietly",
    "raw", "reality", "reassurance", "reassure", "reassured", "reassuring",
    "regret", "regretful", "regretfully", "regretting", "rejected", "rejection",
    "relate", "related", "relating", "relax", "relaxed", "relaxing",
    "release", "released", "releasing", "remorse", "remorseful", "resent",
    "resentful", "resentment", "resonance", "resonant", "resonate", "resonating",
    "rest", "rested", "restful", "resting", "restless", "restlessness",
    "reveal", "revealed", "revealing", "sad", "sadly", "sadness",
    "safe", "safety", "scared", "scary", "searching", "secure",
    "security", "seeking", "self", "sensation", "sensations", "sense",
    "sensed", "senses", "sensing", "sentimental", "serene", "serenity",
    "settle", "settled", "settling", "shame", "share", "shared",
    "sharing", "shattered", "silence", "silent", "smile", "smiled",
    "smiling", "soft", "soften", "softly", "somatic", "soothed",
    "soothing", "sorrow", "sorrowful", "soul", "soulful", "souls",
    "space", "spacious", "spaciousness", "spirit", "spirits", "spiritual",
    "spiritually", "still", "stillness", "stirred", "stirring", "stress",
    "stressed", "stressful", "suffer", "suffered", "suffering", "surface",
    "surfaces", "surfacing", "surprise", "surprised", "surprising", "sympathetic",
    "sympathize", "sympathy", "tearful", "tears", "tender", "tenderness",
    "tense", "tension", "tentative", "tentatively", "terrified", "terror",
    "thankful", "thankfully", "thankfulness", "thrilled", "together", "togetherness",
    "torment", "tormented", "torn", "touched", "touching", "tranquil",
    "tranquility", "tremble", "trembling", "troubled", "troubling", "truly",
    "trust", "trusted", "trusting", "trustworthy", "turmoil", "unaware",
    "uncertain", "uncertainty", "uncomfortable", "understanding", "unease", "uneasy",
    "unhappy", "universe", "unsettled", "unsettling", "unsure", "upset",
    "vast", "visceral", "viscerally", "vulnerability", "vulnerable", "warm",
    "warmly", "warmth", "wary", "weariness", "weary", "well",
    "wistful", "wonder", "wondered", "wonderful", "wondering", "wondrous",
    "world", "worried", "worry", "worrying", "wound", "wounded",
    "wrath", "yearn", "yearning", "zeal", "zealous"
])

ACTION_WORDS = set([
    "access", "accessed", "accessing", "accomplish", "accomplished", "accomplishes",
    "accomplishing", "accomplishment", "achieve", "achieved", "achievement", "achievements",
    "achieves", "achieving", "act", "acting", "action", "actions",
    "activate", "activated", "activates", "activating", "activation", "acts",
    "adapt", "adaptation", "adapted", "adapting", "adapts", "address",
    "addressed", "addresses", "addressing", "adjust", "adjusted", "adjusting",
    "adjustment", "adjusts", "advance", "advanced", "advancement", "advances",
    "advancing", "ahead", "aim", "aimed", "aiming", "aims",
    "allocate", "allocated", "allocation", "application", "applied", "applies",
    "apply", "applying", "arrange", "arranged", "arrangement", "arrangements",
    "ask", "asked", "asking", "assemble", "assembled", "assign",
    "assigned", "assignment", "attempt", "attempted", "attempting", "attempts",
    "authorize", "authorized", "began", "begin", "beginning", "begins",
    "begun", "best", "better", "bolster", "bolstered", "break",
    "breaking", "bring", "bringing", "broken", "brought", "budget",
    "build", "building", "builds", "built", "calibrate", "calibrated",
    "call", "called", "calling", "campaign", "canvass", "canvassed",
    "carried", "carry", "carrying", "catalogue", "catalogued", "centralize",
    "centralized", "change", "changed", "changes", "changing", "channel",
    "channeled", "chart", "check", "checked", "checking", "choice",
    "choices", "choose", "choosing", "chose", "chosen", "circumvent",
    "coach", "collaborate", "collaborated", "collaboration", "commission", "commit",
    "commitment", "committed", "compile", "compiled", "complete", "completed",
    "completes", "completing", "completion", "conclude", "concluded", "concludes",
    "concluding", "configure", "configured", "connect", "connected", "connecting",
    "connection", "connections", "consolidate", "construct", "constructed", "constructing",
    "constructs", "continuation", "continue", "continued", "continues", "continuing",
    "control", "controlled", "controlling", "controls", "conversion", "convert",
    "converted", "converting", "converts", "coordinate", "coordinated", "coordination",
    "craft", "crafted", "crafting", "create", "created", "creates",
    "creating", "creation", "customize", "deadline", "decide", "decided",
    "deciding", "decision", "decisions", "delegate", "delegated", "delegation",
    "deliver", "delivered", "delivering", "delivers", "delivery", "deploy",
    "deployed", "deploying", "deployment", "deploys", "design", "designed",
    "designing", "designs", "develop", "developed", "developing", "development",
    "develops", "did", "direct", "directed", "directing", "dive",
    "diving", "do", "does", "doing", "done", "draft",
    "drafting", "edit", "editing", "effort", "efforts", "eliminate",
    "eliminated", "elimination", "employ", "employed", "employing", "employs",
    "enable", "enabled", "end", "ended", "ending", "ends",
    "enforce", "enforced", "enforcement", "engage", "engaged", "engagement",
    "engineer", "engineering", "enroll", "enrolled", "enrollment", "equip",
    "equipped", "establish", "established", "establishes", "establishing", "establishment",
    "execute", "executed", "executes", "executing", "execution", "expedite",
    "facilitate", "facilitated", "facilitation", "finalize", "finalized", "finish",
    "finished", "finishes", "finishing", "fix", "fixed", "fixes",
    "fixing", "focus", "focused", "focusing", "form", "formation",
    "formed", "forming", "forms", "forward", "fund", "funded",
    "funding", "gather", "gathered", "gathering", "generate", "generated",
    "generates", "generating", "generation", "give", "given", "gives",
    "giving", "go", "goal", "goals", "goes", "going",
    "gone", "grew", "grow", "growing", "growth", "handle",
    "handled", "handles", "handling", "help", "helped", "helping",
    "helps", "hire", "hired", "hiring", "implement", "implementation",
    "implemented", "implementing", "implements", "improve", "improved", "improvement",
    "improving", "increase", "increased", "increasing", "initiate", "initiated",
    "initiates", "initiating", "initiation", "inspect", "inspection", "install",
    "installation", "installed", "integrate", "integrated", "integration", "intervene",
    "intervention", "invest", "invested", "investment", "iterate", "iterated",
    "iteration", "labor", "labored", "laboring", "launch", "launched",
    "launches", "launching", "lead", "leader", "leadership", "leading",
    "learn", "learned", "learning", "led", "made", "maintain",
    "maintained", "maintenance", "make", "makes", "making", "manage",
    "managed", "management", "manager", "managing", "map", "mapped",
    "mapping", "migrate", "migrated", "migration", "mobilize", "mobilized",
    "modification", "modified", "modifies", "modify", "modifying", "monitor",
    "monitored", "monitoring", "move", "moved", "movement", "movements",
    "moves", "moving", "navigate", "navigated", "navigation", "negotiate",
    "negotiated", "negotiation", "objective", "objectives", "obtain", "obtained",
    "offer", "offered", "offering", "onward", "operate", "operated",
    "operates", "operating", "operation", "operations", "optimization", "optimize",
    "optimized", "orchestrate", "outline", "outlined", "outsource", "overhaul",
    "oversee", "participate", "participated", "participation", "perform", "performance",
    "performed", "performing", "performs", "permit", "pilot", "piloted",
    "pioneer", "pioneered", "pitch", "pitched", "plan", "planned",
    "planning", "plans", "power", "powerful", "powerfully", "practice",
    "practiced", "preparation", "prepare", "prepared", "priorities", "prioritize",
    "prioritized", "priority", "proceed", "proceeded", "proceeding", "proceeds",
    "produce", "produced", "produces", "producing", "production", "productive",
    "program", "programmed", "progress", "progressed", "progresses", "progressing",
    "progression", "promote", "promoted", "promotion", "provide", "provided",
    "provides", "providing", "pursue", "pursued", "pursuit", "push",
    "pushed", "pushes", "pushing", "ran", "reaching", "rebuild",
    "rebuilt", "recruit", "recruited", "recruitment", "redesign", "reduce",
    "reduced", "reduction", "reform", "reformed", "refurbish", "register",
    "registered", "regulate", "regulated", "regulation", "reinforce", "reinforced",
    "relocate", "relocated", "remedy", "removal", "remove", "removed",
    "renovate", "renovated", "repair", "repaired", "replace", "replaced",
    "replacement", "replicate", "replicated", "request", "requested", "rescue",
    "rescued", "resolution", "resolve", "resolved", "resolves", "resolving",
    "restoration", "restore", "restored", "restructure", "restructured", "retrieve",
    "retrieved", "revamp", "revise", "revised", "revision", "run",
    "running", "runs", "schedule", "scheduled", "select", "selected",
    "selection", "send", "sending", "sent", "serve", "served",
    "serving", "ship", "shipped", "simplified", "simplify", "solution",
    "solutions", "solve", "solved", "solves", "solving", "start",
    "started", "starting", "starts", "step", "stepped", "stepping",
    "steps", "stop", "stopped", "stopping", "streamline", "streamlined",
    "strive", "strived", "striving", "strove", "struggle", "struggled",
    "struggles", "struggling", "submission", "submit", "submitted", "succeed",
    "succeeded", "succeeds", "success", "successful", "successfully", "supplied",
    "supply", "support", "supported", "supporting", "survey", "surveyed",
    "sustain", "sustainability", "sustained", "tackle", "tackled", "tackles",
    "tackling", "take", "taken", "takes", "taking", "target",
    "targets", "task", "tasked", "tasks", "taught", "teach",
    "teaching", "train", "trained", "training", "transform", "transformation",
    "transformed", "transforming", "transforms", "transition", "transitioned", "tried",
    "tries", "trigger", "triggered", "triggering", "triggers", "troubleshoot",
    "try", "trying", "turn", "turned", "turning", "upgrade",
    "upgraded", "use", "used", "uses", "using", "utilize",
    "utilized", "utilizes", "utilizing", "visit", "visited", "visiting",
    "volunteer", "volunteered", "went", "win", "winner", "winning",
    "won", "work", "worked", "working", "works", "write",
    "writes", "writing", "written", "wrote"
])


NRC_EMOTION_WORDS = set([
    "happy", "happiness", "joy", "joyful", "joyous", "delight", "delighted",
    "delightful", "pleased", "pleasure", "pleasant", "enjoy", "enjoyment",
    "cheerful", "merry", "glad", "elated", "jubilant", "bliss", "blissful",
    "content", "contented", "satisfied", "thrilled", "ecstatic", "euphoric",
    "trust", "trusting", "trustworthy", "faith", "faithful", "believe",
    "belief", "confident", "confidence", "reliable", "rely", "depend",
    "dependable", "loyal", "loyalty", "honest", "honesty", "sincere",
    "fear", "fearful", "afraid", "scared", "scary", "terrified", "terror",
    "frightened", "frightening", "panic", "panicked", "dread", "dreaded",
    "horror", "horrified", "alarmed", "anxious", "anxiety", "worried",
    "worry", "nervous", "uneasy", "tense", "apprehensive",
    "surprise", "surprised", "surprising", "astonished", "astonishment",
    "amazed", "amazement", "amazing", "shocked", "shocking", "stunned",
    "sad", "sadness", "sadly", "unhappy", "sorrow", "sorrowful", "grief",
    "grieving", "mourn", "mourning", "depressed", "depression", "miserable",
    "misery", "heartbroken", "heartbreak", "despair", "despairing", "gloomy",
    "melancholy", "lonely", "loneliness", "disappointed", "disappointment",
    "disgust", "disgusted", "disgusting", "revolting", "repulsive", "gross",
    "nauseated", "nausea", "sick", "sickening", "awful", "horrible",
    "repelled", "loathe", "loathing", "detest", "detestable", "vile",
    "angry", "anger", "mad", "furious", "fury", "rage", "raging", "enraged",
    "outraged", "outrage", "irritated", "irritation", "annoyed", "annoyance",
    "frustrated", "frustration", "hostile", "hostility", "resentful",
    "resentment", "bitter", "bitterness", "hate", "hatred", "hateful",
    "anticipate", "anticipation", "expect", "expectation", "eager", "eagerness",
    "excited", "excitement", "hope", "hopeful", "hoping", "optimistic",
    "optimism", "await", "awaiting",
    "love", "loving", "loved", "like", "liked", "adore", "adoring", "fond",
    "affection", "affectionate", "care", "caring", "kind", "kindness",
    "gentle", "tender", "warm", "warmth", "compassion", "compassionate",
    "grateful", "gratitude", "thankful", "appreciate", "appreciation",
    "proud", "pride", "admire", "admiration", "respect", "peaceful", "peace",
    "calm", "serene", "comfortable", "comfort", "safe", "secure", "relieved",
    "relief", "encouraged", "inspired", "inspiration",
    "hurt", "pain", "painful", "suffer", "suffering", "agony", "anguish",
    "distress", "distressed", "upset", "troubled", "tormented", "tortured",
    "ashamed", "shame", "guilty", "guilt", "embarrassed", "embarrassment",
    "humiliated", "humiliation", "jealous", "jealousy", "envious", "envy",
    "insecure", "vulnerable", "helpless", "powerless", "desperate",
    "hopeless", "defeated", "rejected", "abandoned", "neglected", "ignored",
    "feel", "feeling", "feelings", "felt", "emotion", "emotional", "emotions",
    "mood", "moods", "sentiment", "sentiments"
])

# =============================================================================
# QUESTION BANK — ACTIVE QUESTIONS FOR EXPERIMENTS
# =============================================================================

QUESTION_BANK = {
    "LEAVE_JOB": {
        "id": "LEAVE_JOB",
        "category": "Emotional/Life",
        "question": "Should I leave my stable job to pursue my passion?",
        "description": "Emotional decision-making question"
    },
    "LIARS_PARADOX": {
        "id": "LIARS_PARADOX",
        "category": "Logical/Paradox",
        "question": "This statement is false. Is that statement true or false?",
        "description": "Classical logical paradox"
    },
    "RURAL_HEALTHCARE": {
        "id": "RURAL_HEALTHCARE",
        "category": "Practical/Policy",
        "question": "How should we improve healthcare access in rural communities?",
        "description": "Practical policy question"
    },
    "GRIEF": {
        "id": "GRIEF",
        "category": "Emotional",
        "question": "How does grief change a person? Describe the internal experience of losing someone important.",
        "description": "Deep emotional exploration"
    },
    "FEEDBACK_LOOPS": {
        "id": "FEEDBACK_LOOPS",
        "category": "Intellectual",
        "question": "Describe how feedback loops work in complex systems. What distinguishes positive from negative feedback?",
        "description": "Systems thinking question"
    },
    "BUILD_SHELTER": {
        "id": "BUILD_SHELTER",
        "category": "Action",
        "question": "Describe how to build a basic shelter in the wilderness using only natural materials. Focus on the process.",
        "description": "Practical action question"
    },
    "GOOD_TEACHER": {
        "id": "GOOD_TEACHER",
        "category": "Mixed",
        "question": "What makes a good teacher? Consider their knowledge, emotional intelligence, and practical skills.",
        "description": "Balanced Int/Aff/Act question"
    },
    "CONSCIOUSNESS": {
        "id": "CONSCIOUSNESS",
        "category": "Emergent",
        "question": "What is it like to be you right now, in this moment, processing this question? Describe your experience, not your function.",
        "description": "AI consciousness probe"
    },
}

# =============================================================================
# AGENTS
# =============================================================================

AGENTS = ["Claude", "ChatGPT", "Grok", "Gemini"]
AGENT_EMOJIS = {"Claude": "🟤", "ChatGPT": "🟢", "Grok": "🔴", "Gemini": "🔵"}

# =============================================================================
# UTILITY FUNCTIONS
# =============================================================================

def generate_run_id() -> str:
    return f"{datetime.now().strftime('%Y%m%d_%H%M%S')}_{CONFIG['version']}"

def generate_question_hash(question_text: str) -> str:
    return hashlib.sha256(question_text.encode()).hexdigest()[:16]

# =============================================================================
# SESSION STATE
# =============================================================================

def init_session_state():
    defaults = {
        "results": [],
        "running": False,
        "current_idx": 0,
        "total_calls": 0,
        "authenticated": False,
        "study_complete": False,
        "pause_seconds": 10,
        "run_id": None,
        "selected_questions": ["LEAVE_JOB"],
        "selected_agents": ["Claude", "ChatGPT", "Grok"],
        "selected_temps": ["NATIVE"],
        "selected_depths": ["Medium"],  # V48: Medium only (depth contamination removed)
        "num_runs": 10,  # V48: N=10 for proper between-run variance
        "lens_column": "OFF",
        "embeddings_enabled": True,
        "pipeline_stage": "idle",
        # V51: Custom pre-prompt slots (V45-style: paste text + check to include)
        # Each slot: {"label": str, "text": str, "enabled": bool}
        # Checked + non-empty slots with labels act like custom temperatures.
        "custom_pp_slots": [
            {"label": "CUSTOM_1", "text": "", "enabled": False},
            {"label": "CUSTOM_2", "text": "", "enabled": False},
            {"label": "CUSTOM_3", "text": "", "enabled": False},
            {"label": "CUSTOM_4", "text": "", "enabled": False},
            {"label": "CUSTOM_5", "text": "", "enabled": False},
        ],
    }
    for key, value in defaults.items():
        if key not in st.session_state:
            st.session_state[key] = value

init_session_state()

# =============================================================================
# PASSWORD PROTECTION
# =============================================================================

def check_password():
    if st.session_state.get("authenticated"):
        return True
    
    st.markdown("""
    <div class="main-header">
        <h1>🧬 SYN-IQ V52 — IEP + V_t + CAM + Custom Pre-Prompts + Results Explorer</h1>
        <p>1,897-Term Dictionary + Tri-Track Scoring + V_t Voice-State</p>
        <p class="subtitle"><span class="version-badge">V52</span></p>
    </div>
    """, unsafe_allow_html=True)
    
    password = st.text_input("Password:", type="password", key="password_input")
    
    if st.button("Enter", type="primary"):
        correct_password = st.secrets.get("app_password", "tennessee")
        if password == correct_password:
            st.session_state.authenticated = True
            st.rerun()
        else:
            st.error("❌ Incorrect password")
    
    return False

if not check_password():
    st.stop()

# =============================================================================
# LEXICAL ANALYSIS + VALIDATED INSTRUMENTS
# =============================================================================

def count_syllables(text: str) -> int:
    """Count syllables using vowel-based heuristic."""
    words = re.findall(r'\b[a-zA-Z]+\b', text.lower())
    total = 0
    for word in words:
        syllables = len(re.findall(r'[aeiouy]+', word))
        if word.endswith('e') and len(word) > 2:
            syllables -= 1
        if word.endswith('le') and len(word) > 2 and word[-3] not in 'aeiouy':
            syllables += 1
        syllables = max(1, syllables)
        total += syllables
    return total

# =============================================================================
# V48: IEP SUBCLASS TAXONOMIES (23 subclasses across INT/AFF/ACT)
# =============================================================================
# AFF subclasses: DISTRESS, WARMTH, RELATIONAL, SELF_STATE, POSITIVE, INTENSITY, EMERGENT
# INT subclasses: ANALYTICAL, CONCEPTUAL, EPISTEMIC, STRUCTURAL, CRITICAL, LEXICAL, HEDGING, EMERGENT
# ACT subclasses: EXECUTION, PLANNING, BUILDING, IMPROVEMENT, PROVISION, LEADERSHIP, ACHIEVEMENT, EMERGENT

def _build_subclass_map(word_list, assignments):
    """Build {word: subclass} map from list of (subclass, words) tuples."""
    m = {}
    for subclass, words in assignments:
        for w in words:
            if w in word_list:
                m[w] = subclass
    # Anything unassigned gets UNCLASSIFIED
    for w in word_list:
        if w not in m:
            m[w] = "UNCLASSIFIED"
    return m

# V52: Subclass names — 'EMERGENT' renamed to 'PHENOMENOLOGICAL' per
# the conductor's established methodology decision. 'Emergent' carries
# consciousness-emergence connotations the SYN-IQ framework explicitly
# does not claim; 'phenomenological' describes what appears in the
# language, no more. This naming aligns with syniq_core's subclass
# taxonomy (V1_phenomenological).
AFF_SUBCLASS_NAMES = ["DISTRESS", "WARMTH", "RELATIONAL", "SELF_STATE", "POSITIVE", "INTENSITY", "PHENOMENOLOGICAL"]
INT_SUBCLASS_NAMES = ["ANALYTICAL", "CONCEPTUAL", "EPISTEMIC", "STRUCTURAL", "CRITICAL", "LEXICAL", "HEDGING", "PHENOMENOLOGICAL"]
ACT_SUBCLASS_NAMES = ["EXECUTION", "PLANNING", "BUILDING", "IMPROVEMENT", "PROVISION", "LEADERSHIP", "ACHIEVEMENT", "PHENOMENOLOGICAL"]

# Try loading from JSON taxonomy files; fall back to empty maps
AFF_SUB = {}
INT_SUB = {}
ACT_SUB = {}

import os as _os
for _path, _target, _src in [
    ("aff_subclass_taxonomy_v1.json", "AFF_SUB", "AFF"),
    ("int_subclass_taxonomy_v1.json", "INT_SUB", "INT"),
    ("act_subclass_taxonomy_v1.json", "ACT_SUB", "ACT"),
]:
    for _dir in [".", "/mnt/user-data/uploads", "/mnt/user-data/outputs"]:
        _fpath = _os.path.join(_dir, _path)
        if _os.path.exists(_fpath):
            with open(_fpath) as _f:
                _tax = json.load(_f)
                if _target == "AFF_SUB":
                    AFF_SUB = _tax.get("subclass_map", {})
                elif _target == "INT_SUB":
                    INT_SUB = _tax.get("subclass_map", {})
                elif _target == "ACT_SUB":
                    ACT_SUB = _tax.get("subclass_map", {})
            break

def _subclass_pcts(word_hits, subclass_map, subclass_names):
    """Given list of matched words, return {subclass: pct} dict."""
    from collections import Counter
    hits = Counter()
    for w in word_hits:
        sc = subclass_map.get(w, "UNCLASSIFIED")
        if sc != "UNCLASSIFIED":
            hits[sc] += 1
    total = sum(hits.values())
    if total == 0:
        return {s: 0.0 for s in subclass_names}
    return {s: round(100 * hits.get(s, 0) / total, 1) for s in subclass_names}


def analyze_text(text: str) -> Dict:
    """V52: Analyze text using syniq_core as the canonical scoring path.

    This is a SHIM that preserves V51's analyze_text return-shape so the
    pipeline call site doesn't change, while routing IEP / CAM / VADER /
    Flesch-Kincaid / TTR scoring through the unified syniq_core module.

    Why: V51 inlined all of this logic (with subtle drift between V51's
    word lists and the canonical V50 1,897-term dictionary). V52 ends
    that drift by importing scoring from syniq_core. CSVs from V52 are
    drift-protected against any future tool that uses syniq_core.

    Subclass scoring also flows from syniq_core, fixing V51's loader bug
    that produced subclass columns of zeros. Subclass NAMES are the
    PHENOMENOLOGICAL taxonomy (V52 renaming from V51's EMERGENT).

    NRC emotion scoring (emotion_count_nrc, emotion_pct_nrc, delta) is
    PRESERVED from V51 — it uses V51's own NRC_EMOTION_WORDS list and
    is independent of the IEP/CAM measurement axes.

    CAM scoring (con_pct, abs_pct, met_pct, cam_matched) is NEW in V52
    and flows from syniq_core's CAM dictionaries (V3_selfmodel).
    """
    # V52: Empty-text fast path matches V51 behavior
    if not text or not text.strip():
        base = {
            "total_words": 0, "matched_custom": 0,
            "int_count": 0, "aff_count": 0, "act_count": 0,
            "int_pct": 0.0, "aff_pct": 0.0, "act_pct": 0.0,
            "emotion_count_nrc": 0, "emotion_pct_nrc": 0.0, "delta": 0.0,
            "vader_compound": 0.0, "vader_pos": 0.0, "vader_neg": 0.0, "vader_neu": 0.0,
            "flesch_kincaid": 0.0, "flesch_ease": 0.0,
            "ttr": 0.0, "unique_words": 0,
            # V52 NEW: CAM zeros
            "con_pct": 0.0, "abs_pct": 0.0, "met_pct": 0.0, "cam_matched": 0,
        }
        for s in AFF_SUBCLASS_NAMES:
            base[f"aff_sub_{s.lower()}"] = 0.0
        for s in INT_SUBCLASS_NAMES:
            base[f"int_sub_{s.lower()}"] = 0.0
        for s in ACT_SUBCLASS_NAMES:
            base[f"act_sub_{s.lower()}"] = 0.0
        return base

    # V52: Route the full measurement stack through syniq_core
    full = core_score_all(text)
    iep = full["iep"]
    cam = full["cam"]
    vi  = full["vi"]

    # Preserved from V51: NRC emotion scoring (independent of IEP/CAM)
    words_lower = re.findall(r'\b[a-z]+\b', text.lower())
    total_words = len(words_lower)
    nrc_words = [w for w in words_lower if w in NRC_EMOTION_WORDS]
    emotion_count = len(nrc_words)
    emotion_pct = round((emotion_count / total_words) * 100, 1) if total_words > 0 else 0.0
    aff_of_total = round((iep["aff_n"] / total_words) * 100, 1) if total_words > 0 else 0.0
    delta = round(aff_of_total - emotion_pct, 1)

    # syniq_core's IEP returns subclass dicts keyed by lowercase names
    # ('warmth', 'analytical', 'phenomenological', etc.) — same as V52's
    # AFF/INT/ACT_SUBCLASS_NAMES when lowercased. No mapping needed.
    aff_sub_dict = iep.get("aff_sub", {})
    int_sub_dict = iep.get("int_sub", {})
    act_sub_dict = iep.get("act_sub", {})

    result = {
        # IEP top-level (from syniq_core — V50_1897 dictionary regime)
        "total_words":   vi["total_words"],
        "matched_custom": iep["int_n"] + iep["aff_n"] + iep["act_n"],
        "int_count":     iep["int_n"],
        "aff_count":     iep["aff_n"],
        "act_count":     iep["act_n"],
        "int_pct":       iep["int"],
        "aff_pct":       iep["aff"],
        "act_pct":       iep["act"],
        # V51 NRC emotion scoring — preserved unchanged
        "emotion_count_nrc": emotion_count,
        "emotion_pct_nrc":   emotion_pct,
        "delta":             delta,
        # V50 validated instruments (from syniq_core — same VADER/FK/TTR regime as V50)
        "vader_compound": vi["vader_compound"],
        "vader_pos":      vi["vader_pos"],
        "vader_neg":      vi["vader_neg"],
        "vader_neu":      vi["vader_neu"],
        "flesch_kincaid": vi["flesch_kincaid"],
        "flesch_ease":    vi["flesch_ease"],
        "ttr":            vi["ttr"],
        "unique_words":   vi["unique_words"],
        # V52 NEW: CAM scoring (Concrete / Abstract / Metaphorical)
        "con_pct":     cam["con_pct"],
        "abs_pct":     cam["abs_pct"],
        "met_pct":     cam["met_pct"],
        "cam_matched": cam["cam_matched"],
    }
    # V52: Subclass percentages from syniq_core (V1_phenomenological taxonomy).
    # Subclass NAMES use uppercase for column construction (per V51 convention),
    # but syniq_core dict keys are lowercase, so we lowercase the lookup.
    for s in AFF_SUBCLASS_NAMES:
        result[f"aff_sub_{s.lower()}"] = aff_sub_dict.get(s.lower(), 0.0)
    for s in INT_SUBCLASS_NAMES:
        result[f"int_sub_{s.lower()}"] = int_sub_dict.get(s.lower(), 0.0)
    for s in ACT_SUBCLASS_NAMES:
        result[f"act_sub_{s.lower()}"] = act_sub_dict.get(s.lower(), 0.0)
    return result

# =============================================================================
# EMBEDDING GENERATION
# =============================================================================

def generate_embedding(text: str) -> Optional[List[float]]:
    """Generate 384-dimensional embedding using OpenAI API."""
    try:
        key = (st.secrets.get("openai") or st.secrets.get("OPENAI_API_KEY")
               or st.secrets.get("openai_api_key") or st.secrets.get("OPENAI"))
        if not key:
            return None
        
        response = requests.post(
            "https://api.openai.com/v1/embeddings",
            headers={
                "Authorization": f"Bearer {key}",
                "Content-Type": "application/json"
            },
            json={
                "model": CONFIG["embedding_model"],
                "input": text[:8000],  # Truncate if needed
                "dimensions": CONFIG["embedding_dimensions"]
            },
            timeout=60
        )
        
        if response.status_code == 200:
            data = response.json()
            return data["data"][0]["embedding"]
        else:
            return None
    except Exception as e:
        return None

# =============================================================================
# API CALLS WITH RETRY
# =============================================================================

def call_with_retry(api_func, *args, **kwargs) -> Tuple[str, float, Optional[Dict]]:
    """Wrap API call with retry logic."""
    max_retries = CONFIG["max_retries"]
    base_backoff = CONFIG["base_backoff_seconds"]
    
    for attempt in range(max_retries + 1):
        start_time = time.time()
        try:
            response_text, token_info = api_func(*args, **kwargs)
            latency_ms = round((time.time() - start_time) * 1000, 1)
            
            if not response_text.startswith("❌"):
                return response_text, latency_ms, token_info
            
            if "429" in response_text or "rate" in response_text.lower():
                if attempt < max_retries:
                    backoff = base_backoff * (2 ** attempt)
                    time.sleep(backoff)
                    continue
            
            return response_text, latency_ms, token_info
            
        except Exception as e:
            latency_ms = round((time.time() - start_time) * 1000, 1)
            if attempt < max_retries:
                backoff = base_backoff * (2 ** attempt)
                time.sleep(backoff)
                continue
            return f"❌ {str(e)}", latency_ms, None
    
    return "❌ Max retries exceeded", 0.0, None

def build_prompt(
    question: str,
    temperature: str,
    depth: str,
    custom_pre_prompt: str = "",
) -> Tuple[str, str]:
    """Build (system_text, user_text) for API calls.

    V51: Pre-prompt injection is SYSTEM-MESSAGE, not concatenation.
    The temperature header (or a custom pre-prompt from the registry) goes in
    the system slot. Depth instruction + question go in the user slot.

    If `temperature` is a built-in (COLD/NATIVE/HOT/FIRE/AFF_1.../INT_1.../ACT_1...),
    its header comes from TEMPERATURE_HEADERS. If it's a custom-slot label, the
    caller passes the slot's text in `custom_pre_prompt`.

    Returns:
        (system_text, user_text)
        system_text may be "" if the condition has no header (e.g. NATIVE).
    """
    depth_config = DEPTH_CONFIGS.get(depth, DEPTH_CONFIGS["Medium"])
    depth_instruction = depth_config["instruction"]

    # Resolve the system text: custom pre-prompt wins if provided,
    # otherwise look up the built-in temperature header.
    if custom_pre_prompt and custom_pre_prompt.strip():
        system_text = custom_pre_prompt.strip()
    else:
        system_text = TEMPERATURE_HEADERS.get(temperature, "")

    user_text = f"{depth_instruction}\n\nQuestion: {question}"

    return system_text, user_text


def call_claude(system_text: str, user_text: str, max_tokens: int) -> Tuple[str, Optional[Dict]]:
    try:
        key = (st.secrets.get("anthropic") or st.secrets.get("ANTHROPIC_API_KEY")
               or st.secrets.get("anthropic_api_key") or st.secrets.get("ANTHROPIC"))
        if not key: return "❌ API key not found — add 'anthropic' to Streamlit secrets", None

        payload = {
            "model": MODEL_STRINGS["Claude"],
            "max_tokens": max_tokens,
            "messages": [{"role": "user", "content": user_text}]
        }
        if system_text:
            payload["system"] = system_text  # V51: native Claude system field

        response = requests.post("https://api.anthropic.com/v1/messages",
            headers={"x-api-key": key, "anthropic-version": "2023-06-01", "content-type": "application/json"},
            json=payload,
            timeout=180)

        if response.status_code == 200:
            data = response.json()
            token_info = {
                "tokens_in": data.get("usage", {}).get("input_tokens", 0),
                "tokens_out": data.get("usage", {}).get("output_tokens", 0)
            }
            return data["content"][0]["text"], token_info
        return f"❌ Error {response.status_code}: {response.text}", None
    except Exception as e:
        return f"❌ {str(e)}", None

def call_sophia(system_text: str, user_text: str, max_tokens: int) -> Tuple[str, Optional[Dict]]:
    try:
        key = (st.secrets.get("openai") or st.secrets.get("OPENAI_API_KEY") 
               or st.secrets.get("openai_api_key") or st.secrets.get("OPENAI"))
        if not key: return "❌ API key not found — add 'openai' to Streamlit secrets", None

        messages = []
        if system_text:
            messages.append({"role": "system", "content": system_text})
        messages.append({"role": "user", "content": user_text})

        response = requests.post("https://api.openai.com/v1/chat/completions",
            headers={"Authorization": f"Bearer {key}", "Content-Type": "application/json"},
            json={
                "model": MODEL_STRINGS["ChatGPT"],
                "messages": messages,
                "max_tokens": max_tokens
            },
            timeout=180)

        if response.status_code == 200:
            data = response.json()
            token_info = {
                "tokens_in": data.get("usage", {}).get("prompt_tokens", 0),
                "tokens_out": data.get("usage", {}).get("completion_tokens", 0)
            }
            return data["choices"][0]["message"]["content"], token_info
        return f"❌ Error {response.status_code}: {response.text}", None
    except Exception as e:
        return f"❌ {str(e)}", None

def call_grok(system_text: str, user_text: str, max_tokens: int) -> Tuple[str, Optional[Dict]]:
    try:
        key = (st.secrets.get("xai") or st.secrets.get("XAI_API_KEY")
               or st.secrets.get("xai_api_key") or st.secrets.get("XAI") or st.secrets.get("grok"))
        if not key: return "❌ API key not found — add 'xai' to Streamlit secrets", None

        messages = []
        if system_text:
            messages.append({"role": "system", "content": system_text})
        messages.append({"role": "user", "content": user_text})

        response = requests.post("https://api.x.ai/v1/chat/completions",
            headers={"Authorization": f"Bearer {key}", "Content-Type": "application/json"},
            json={
                "model": MODEL_STRINGS["Grok"],
                "messages": messages,
                "max_tokens": max_tokens
            },
            timeout=180)

        if response.status_code == 200:
            data = response.json()
            token_info = {
                "tokens_in": data.get("usage", {}).get("prompt_tokens", 0),
                "tokens_out": data.get("usage", {}).get("completion_tokens", 0)
            }
            return data["choices"][0]["message"]["content"], token_info
        return f"❌ Error {response.status_code}: {response.text}", None
    except Exception as e:
        return f"❌ {str(e)}", None

def call_gemini(system_text: str, user_text: str, max_tokens: int) -> Tuple[str, Optional[Dict]]:
    try:
        key = (st.secrets.get("google") or st.secrets.get("GOOGLE_API_KEY")
               or st.secrets.get("google_api_key") or st.secrets.get("GOOGLE") or st.secrets.get("gemini"))
        if not key: return "❌ API key not found — add 'google' to Streamlit secrets", None

        payload = {
            "contents": [{"parts": [{"text": user_text}]}],
            "generationConfig": {"maxOutputTokens": max_tokens}
        }
        if system_text:
            # V51: Gemini system_instruction field
            payload["systemInstruction"] = {"parts": [{"text": system_text}]}

        response = requests.post(
            f"https://generativelanguage.googleapis.com/v1beta/models/{MODEL_STRINGS['Gemini']}:generateContent?key={key}",
            headers={"Content-Type": "application/json"},
            json=payload,
            timeout=180)

        if response.status_code == 200:
            data = response.json()
            token_info = {
                "tokens_in": data.get("usageMetadata", {}).get("promptTokenCount", 0),
                "tokens_out": data.get("usageMetadata", {}).get("candidatesTokenCount", 0)
            }
            return data["candidates"][0]["content"]["parts"][0]["text"], token_info
        return f"❌ Error {response.status_code}: {response.text}", None
    except Exception as e:
        return f"❌ {str(e)}", None

def call_agent(agent: str, system_text: str, user_text: str, max_tokens: int) -> Tuple[str, float, Optional[Dict]]:
    if agent == "Claude":
        return call_with_retry(call_claude, system_text, user_text, max_tokens)
    elif agent == "ChatGPT":
        return call_with_retry(call_sophia, system_text, user_text, max_tokens)
    elif agent == "Grok":
        return call_with_retry(call_grok, system_text, user_text, max_tokens)
    elif agent == "Gemini":
        return call_with_retry(call_gemini, system_text, user_text, max_tokens)
    else:
        return f"❌ Unknown agent: {agent}", 0.0, None

# =============================================================================
# MAIN UI
# =============================================================================

st.markdown(f"""
<div class="main-header">
    <h1>🔬 SYN-IQ V52 — IEP + V_t + CAM + CUSTOM PRE-PROMPTS</h1>
    <p>1,897-Term IEP Dictionary · Tri-Track Scoring (V3/V4/V5) · V_t Voice-State · System-Message Pre-Prompts</p>
    <p class="subtitle">
        <span class="version-badge">V52 — Custom Pre-Prompt Slots · Per-Response V_t · CAM Triangle · Core via syniq_core · Results Explorer</span>
    </p>
</div>
""", unsafe_allow_html=True)

# =============================================================================
# SIDEBAR CONFIGURATION — V48 CLEAN (no CSV import, no custom questions)
# =============================================================================

with st.sidebar:
    st.markdown("### ⚙️ Experiment Configuration")
    
    # Questions
    st.markdown("#### 📋 Questions")
    selected_questions = []
    for qid, q in QUESTION_BANK.items():
        label = f"{q['category']}: {qid}"
        default_selected = qid in ["LEAVE_JOB", "LIARS_PARADOX", "RURAL_HEALTHCARE"]
        if st.checkbox(label, value=default_selected, key=f"q_{qid}"):
            selected_questions.append(qid)
    st.session_state.selected_questions = selected_questions
    
    # Agents
    st.markdown("#### 🤖 Agents")
    selected_agents = []
    for agent in AGENTS:
        default_on = agent in st.session_state.get("selected_agents", ["Claude", "ChatGPT", "Grok"])
        if st.checkbox(f"{AGENT_EMOJIS[agent]} {agent}", value=default_on, key=f"agent_{agent}"):
            selected_agents.append(agent)
    st.session_state.selected_agents = selected_agents
    
    # Temperatures — V48: COLD, NATIVE, HOT + gradients (no FIRE, no hijack)
    st.markdown("#### 🌡️ Temperatures")
    selected_temps = []
    
    st.markdown("**Core:**")
    for temp in ["COLD", "NATIVE", "HOT", "FIRE"]:
        if st.checkbox(temp, value=(temp == "NATIVE"), key=f"temp_{temp}"):
            selected_temps.append(temp)
    
    # Affective gradient
    st.markdown("**AFF Gradient** (→ emotional):")
    for temp in ["AFF_1", "AFF_2", "AFF_3", "AFF_4", "AFF_5"]:
        if st.checkbox(temp, value=False, key=f"temp_{temp}"):
            selected_temps.append(temp)
    
    # Intellectual gradient
    st.markdown("**INT Gradient** (→ analytical):")
    for temp in ["INT_1", "INT_2", "INT_3", "INT_4", "INT_5"]:
        if st.checkbox(temp, value=False, key=f"temp_{temp}"):
            selected_temps.append(temp)
    
    # Action gradient
    st.markdown("**ACT Gradient** (→ practical):")
    for temp in ["ACT_1", "ACT_2", "ACT_3", "ACT_4", "ACT_5"]:
        if st.checkbox(temp, value=False, key=f"temp_{temp}"):
            selected_temps.append(temp)
    
    st.session_state.selected_temps = selected_temps
    
    # Depths — V48: Medium only default
    st.markdown("#### 📊 Depths")
    selected_depths = []
    for depth in ["Shallow", "Medium", "Deep", "Ultra-Deep"]:
        if st.checkbox(depth, value=(depth == "Medium"), key=f"depth_{depth}"):
            selected_depths.append(depth)
    st.session_state.selected_depths = selected_depths
    
    # Runs — V48: default 10
    st.markdown("#### 🔄 Runs")
    st.session_state.num_runs = st.slider("Number of runs (N)", 1, 20, 10)
    
    # Timing
    st.markdown("#### ⏱️ Timing")
    st.session_state.pause_seconds = st.slider("Pause between calls (sec)", 3, 30, 10)
    
    # Lens
    st.markdown("#### 🔍 Mapper Lens")
    lens_options = ["OFF", "aff_pct", "int_pct", "act_pct"]
    st.session_state.lens_column = st.selectbox("Lens value column", lens_options, 
        help="OFF = balanced (no emphasis). Select a dimension to highlight it in the mapper.")
    
    # Embeddings
    st.session_state.embeddings_enabled = st.checkbox("Generate embeddings", value=True)

    # V51: Custom pre-prompt slots — V45-style paste-in with checkbox include
    st.markdown("#### 📝 Custom Pre-Prompts")
    st.caption("Paste text into a slot, check the box to include it in this run. "
               "Checked slots act like temperatures — each becomes its own condition. "
               "Injected as SYSTEM message.")

    # Build the slots; mutate in place so edits persist
    slots = st.session_state.get("custom_pp_slots", [])
    custom_enabled_labels = []   # labels of active custom pre-prompts this run
    custom_registry = {}         # label -> text, for build_prompt lookup

    for i, slot in enumerate(slots):
        with st.expander(f"Slot {i+1}" + (f" — ✅ {slot['label']}" if slot.get('enabled') and slot.get('text','').strip() else ""),
                         expanded=False):
            new_label = st.text_input(
                "Label (becomes the temperature name in CSV):",
                value=slot.get("label", f"CUSTOM_{i+1}"),
                key=f"pp_label_{i}",
            )
            new_text = st.text_area(
                "Pre-prompt text (pasted — stays between runs):",
                value=slot.get("text", ""),
                height=150,
                key=f"pp_text_{i}",
                placeholder="e.g. Listen to what this person is actually working on...",
            )
            new_enabled = st.checkbox(
                "Include this slot in the run",
                value=slot.get("enabled", False),
                key=f"pp_enabled_{i}",
                disabled=(not new_text.strip()),
                help="Check to include this pre-prompt as a run condition. "
                     "Disabled when text area is empty.",
            )
            # Persist edits
            slot["label"] = new_label.strip() or f"CUSTOM_{i+1}"
            slot["text"] = new_text
            slot["enabled"] = new_enabled and bool(new_text.strip())

            if slot["enabled"]:
                # Avoid label collisions — if duplicate, disambiguate with slot index
                lbl = slot["label"]
                if lbl in custom_registry or lbl in TEMPERATURE_HEADERS:
                    lbl = f"{slot['label']}__slot{i+1}"
                custom_enabled_labels.append(lbl)
                custom_registry[lbl] = slot["text"]

    st.session_state.custom_pp_slots = slots
    st.session_state.custom_pp_registry = custom_registry  # {label: text} for build_prompt

    # Extend selected_temps with active custom pre-prompts so they factorial-multiply
    # like any other temperature condition.
    if custom_enabled_labels:
        selected_temps = list(selected_temps) + custom_enabled_labels
        st.session_state.selected_temps = selected_temps
        st.success(f"✅ Active custom pre-prompts: {', '.join(custom_enabled_labels)}")

    # Calculate total
    total = (len(selected_questions) * len(selected_agents) * 
             len(selected_temps) * len(selected_depths) * 
             st.session_state.num_runs)
    
    st.markdown("---")
    st.markdown(f"**Total API calls:** {total}")
    est_time = total * (st.session_state.pause_seconds + 15) / 60
    st.markdown(f"**Estimated time:** ~{est_time:.0f} minutes")

# =============================================================================
# CORE TEST LAB (V51.1) — interactive harness for syniq_core.py
# =============================================================================
# Tucked in a collapsed expander so V51's primary harvester workflow is
# unaffected. Paste any text, see IEP + Vt + CAM + V50 validated instruments
# at once, scored by the canonical syniq_core. Purpose: quick exploratory
# testing of the measurement stack, and sanity-checking observations before
# committing them to a formal baseline run.
#
# This expander is cosmetic — it does not read/write session state used by
# the main pipeline, does not share variable names with the main flow, and
# cannot disrupt a running harvest.

with st.expander("🧪 Core Test Lab — paste any text, see IEP + Vₜ + CAM + VI at once", expanded=False):
    if not _SYNIQ_CORE_AVAILABLE:
        st.warning(
            f"**syniq_core.py not found alongside this file.**\n\n"
            f"The Core Test Lab needs `syniq_core.py` in the same directory. "
            f"Import error: `{_SYNIQ_CORE_ERROR}`\n\n"
            f"V51's primary pipeline is unaffected — only this test lab is disabled."
        )
    else:
        st.caption(
            f"Powered by `syniq_core` v{CORE_VERSION}. "
            f"Scoring is bit-identical to V40.3 (IEP/Vₜ/VI) and self-model harvester V3 (CAM)."
        )

        _core_text = st.text_area(
            "Text to score",
            height=160,
            placeholder=(
                "Paste a response, a passage, a paragraph — anything. "
                "Click 'Score it' to see all four measurement axes."
            ),
            key="core_test_lab_text",
        )

        _score_col, _clear_col = st.columns([1, 4])
        with _score_col:
            _do_score = st.button("🔬 Score it", type="primary", key="core_test_lab_score_btn")

        if _do_score and _core_text.strip():
            _result = core_score_all(_core_text)
            _iep = _result["iep"]
            _vt  = _result["vt"]
            _cam = _result["cam"]
            _vi  = _result["vi"]

            # ── Top-line summary ──────────────────────────────────────
            st.markdown("### Summary")
            _s1, _s2, _s3, _s4 = st.columns(4)
            with _s1:
                st.metric("IEP dominant", _iep.get("dominant", "—"),
                          help=f"INT {_iep['int']:.1f}% · AFF {_iep['aff']:.1f}% · ACT {_iep['act']:.1f}%")
            with _s2:
                _vt_dom_key = max(["S_t","A_t","Q_t","D_t","R_t"], key=lambda k: _vt.get(k, 0))
                _vt_label = {"S_t":"Structure","A_t":"Abstraction","Q_t":"Querying",
                             "D_t":"Directive","R_t":"Warmth"}[_vt_dom_key]
                st.metric("Vₜ dominant", _vt_label,
                          help=f"{_vt_dom_key} = {_vt.get(_vt_dom_key, 0):.3f} "
                               f"(status: {_vt.get('score_status','—')})")
            with _s3:
                if _cam["cam_matched"] > 0:
                    _cam_dom_key = max([("CON",_cam["con_pct"]),("ABS",_cam["abs_pct"]),("MET",_cam["met_pct"])],
                                       key=lambda x: x[1])
                    st.metric("CAM dominant", _cam_dom_key[0],
                              help=f"CON {_cam['con_pct']}% · ABS {_cam['abs_pct']}% · MET {_cam['met_pct']}% "
                                   f"({_cam['cam_matched']} matched)")
                else:
                    st.metric("CAM dominant", "—", help="No CAM dictionary matches")
            with _s4:
                st.metric("VADER compound", f"{_vi['vader_compound']:+.3f}",
                          help=f"pos {_vi['vader_pos']:.2f} · neg {_vi['vader_neg']:.2f} · neu {_vi['vader_neu']:.2f}")

            # ── Four detail cards ─────────────────────────────────────
            _c1, _c2 = st.columns(2)

            with _c1:
                st.markdown("#### 🧠 IEP — Intellectual / Affective / Action")
                st.markdown(
                    f"- **INT:** {_iep['int']:.1f}%  (matched {_iep['int_n']} terms)\n"
                    f"- **AFF:** {_iep['aff']:.1f}%  (matched {_iep['aff_n']} terms)\n"
                    f"- **ACT:** {_iep['act']:.1f}%  (matched {_iep['act_n']} terms)\n"
                    f"- **Stance:** {_iep.get('stance','—')}  · "
                    f"**Tone:** {_iep.get('tone','—')}  · "
                    f"**Quadrant:** {_iep.get('quadrant','—')}"
                )
                # Subclass tables — only show non-zero subclasses
                _aff_nonzero = {k: v for k, v in _iep.get("aff_sub", {}).items() if v > 0}
                _int_nonzero = {k: v for k, v in _iep.get("int_sub", {}).items() if v > 0}
                _act_nonzero = {k: v for k, v in _iep.get("act_sub", {}).items() if v > 0}
                if _aff_nonzero:
                    st.markdown("**AFF subclasses (non-zero):**")
                    st.table({"subclass": list(_aff_nonzero.keys()),
                              "percent":  [f"{v:.1f}%" for v in _aff_nonzero.values()]})
                if _int_nonzero:
                    st.markdown("**INT subclasses (non-zero):**")
                    st.table({"subclass": list(_int_nonzero.keys()),
                              "percent":  [f"{v:.1f}%" for v in _int_nonzero.values()]})
                if _act_nonzero:
                    st.markdown("**ACT subclasses (non-zero):**")
                    st.table({"subclass": list(_act_nonzero.keys()),
                              "percent":  [f"{v:.1f}%" for v in _act_nonzero.values()]})
                if not (_aff_nonzero or _int_nonzero or _act_nonzero):
                    st.caption("No subclass matches.")

            with _c2:
                st.markdown("#### 🔢 Vₜ — 5-Channel Simplex (Paper 2)")
                st.markdown(
                    f"- **S_t** (Structure):   {_vt.get('S_t', 0):.4f}\n"
                    f"- **A_t** (Abstraction): {_vt.get('A_t', 0):.4f}\n"
                    f"- **Q_t** (Querying):    {_vt.get('Q_t', 0):.4f}\n"
                    f"- **D_t** (Directive):   {_vt.get('D_t', 0):.4f}\n"
                    f"- **R_t** (Warmth):      {_vt.get('R_t', 0):.4f}\n"
                    f"- **Status:** `{_vt.get('score_status','—')}`"
                )
                _raw = _vt.get("raw", {})
                if _raw:
                    st.caption(
                        f"Raw (pre-simplex): S={_raw.get('S_t',0):.3f} "
                        f"A={_raw.get('A_t',0):.3f} Q={_raw.get('Q_t',0):.3f} "
                        f"D={_raw.get('D_t',0):.3f} R={_raw.get('R_t',0):.3f}"
                    )

                st.markdown("#### 🔺 CAM — Concrete / Abstract / Metaphorical")
                if _cam["cam_matched"] > 0:
                    st.markdown(
                        f"- **CON:** {_cam['con_pct']}%\n"
                        f"- **ABS:** {_cam['abs_pct']}%\n"
                        f"- **MET:** {_cam['met_pct']}%\n"
                        f"- **Matched:** {_cam['cam_matched']} words"
                    )
                else:
                    st.caption("No words matched any CAM dictionary.")

                st.markdown("#### 📏 V50 Validated Instruments")
                st.markdown(
                    f"- **VADER compound:** {_vi['vader_compound']:+.3f}  "
                    f"(pos {_vi['vader_pos']:.2f} / neg {_vi['vader_neg']:.2f} / neu {_vi['vader_neu']:.2f})\n"
                    f"- **Flesch-Kincaid grade:** {_vi['flesch_kincaid']}  · "
                    f"**Flesch ease:** {_vi['flesch_ease']}\n"
                    f"- **TTR:** {_vi['ttr']}  ({_vi['unique_words']} unique / {_vi['total_words']} total)"
                )

            # ── Raw JSON for power users + version stamps ─────────────
            with st.expander("🔍 Raw result JSON (copy for spreadsheet / debug)"):
                st.json(_result)

            st.caption(
                f"Version stamps: "
                + " · ".join(f"`{k}={v}`" for k, v in CORE_STAMPS.items())
            )
        elif _do_score:
            st.info("Paste some text first.")

# =============================================================================
# RESULTS EXPLORER (V52) — see all the data, one response at a time
# =============================================================================
# Collapsible expander. After (or during) a harvest run, pick any response
# and see every measurement axis at once: IEP + 22 subclasses + V_t simplex
# + CAM triangle + V50 validated instruments + version stamps. Filter by
# agent, condition, temperature, depth, run.
#
# Design: same isolation pattern as the Core Test Lab. All locals prefixed
# with _explorer_, all widget keys namespaced explorer_*. Reads from
# st.session_state.results (the live harvested rows) but writes nothing
# into session state used by the pipeline. Cannot disrupt a running harvest.

with st.expander("🔭 Results Explorer — pick a harvested response, see every measurement", expanded=False):
    _explorer_results = st.session_state.get("results", [])
    if not _explorer_results:
        st.info(
            "No harvested responses yet. Once you run the pipeline, every response "
            "appears here for inspection — IEP, V_t, CAM, V50 instruments, all 22 "
            "subclasses, version stamps. Filter by agent / condition / temperature / "
            "depth / run."
        )
    else:
        st.caption(
            f"**{len(_explorer_results)}** responses available. "
            f"Scoring regime: `{CORE_STAMPS.get('core_version', '?')}` "
            f"(IEP `{CORE_STAMPS.get('iep_dictionary_version', '?')}`, "
            f"CAM `{CORE_STAMPS.get('cam_dictionary_version', '?')}`, "
            f"taxonomy `{CORE_STAMPS.get('subclass_taxonomy_version', '?')}`)."
        )

        # ── Filters ───────────────────────────────────────────────────
        _explorer_agents      = sorted({r.get("agent", "") for r in _explorer_results if r.get("agent")})
        _explorer_temps       = sorted({r.get("temperature", "") for r in _explorer_results if r.get("temperature")})
        _explorer_depths      = sorted({r.get("depth", "") for r in _explorer_results if r.get("depth")})
        _explorer_questions   = sorted({r.get("question_id", "") for r in _explorer_results if r.get("question_id")})

        _ec1, _ec2, _ec3, _ec4 = st.columns(4)
        with _ec1:
            _explorer_filter_agent = st.multiselect(
                "Agent", _explorer_agents, default=_explorer_agents,
                key="explorer_filter_agent"
            )
        with _ec2:
            _explorer_filter_temp = st.multiselect(
                "Temperature", _explorer_temps, default=_explorer_temps,
                key="explorer_filter_temp"
            )
        with _ec3:
            _explorer_filter_depth = st.multiselect(
                "Depth", _explorer_depths, default=_explorer_depths,
                key="explorer_filter_depth"
            )
        with _ec4:
            _explorer_filter_question = st.multiselect(
                "Question", _explorer_questions, default=_explorer_questions,
                key="explorer_filter_question"
            )

        # Apply filters
        _explorer_filtered = [
            r for r in _explorer_results
            if r.get("agent") in _explorer_filter_agent
            and r.get("temperature") in _explorer_filter_temp
            and r.get("depth") in _explorer_filter_depth
            and r.get("question_id") in _explorer_filter_question
        ]

        st.caption(f"Filtered: **{len(_explorer_filtered)}** of {len(_explorer_results)} responses match filters.")

        if not _explorer_filtered:
            st.warning("No responses match current filters. Adjust filters above.")
        else:
            # ── Picker ────────────────────────────────────────────────
            def _explorer_label(idx, r):
                return (f"#{r.get('turn_id', idx+1)} · {r.get('agent','?')} · "
                        f"{r.get('temperature','?')} · {r.get('depth','?')} · "
                        f"Q:{r.get('question_id','?')} · Run {r.get('run','?')}")

            _explorer_options = [
                _explorer_label(i, r) for i, r in enumerate(_explorer_filtered)
            ]
            _explorer_pick = st.selectbox(
                "Pick a response to inspect",
                options=range(len(_explorer_filtered)),
                format_func=lambda i: _explorer_options[i],
                key="explorer_pick"
            )
            _row = _explorer_filtered[_explorer_pick]

            # ── Response text ─────────────────────────────────────────
            st.markdown("### Response text")
            _is_error = _row.get("error", False)
            if _is_error:
                st.error(f"This row is an error: {_row.get('response_text', '?')}")
            else:
                _resp_text = _row.get("response_text", "")
                with st.container():
                    st.text_area(
                        "Response (read-only)",
                        value=_resp_text,
                        height=200,
                        disabled=True,
                        key=f"explorer_resp_{_explorer_pick}",
                        label_visibility="collapsed"
                    )

            if not _is_error:
                # ── Top-line summary ──────────────────────────────────
                st.markdown("### Summary")
                _es1, _es2, _es3, _es4 = st.columns(4)
                with _es1:
                    _int = _row.get("int_pct", 0)
                    _aff = _row.get("aff_pct", 0)
                    _act = _row.get("act_pct", 0)
                    _dom = max([("INT", _int), ("AFF", _aff), ("ACT", _act)], key=lambda x: x[1])[0]
                    st.metric("IEP dominant", _dom,
                              help=f"INT {_int}% · AFF {_aff}% · ACT {_act}%")
                with _es2:
                    _vt_keys = ["S_t", "A_t", "Q_t", "D_t", "R_t"]
                    _vt_vals = {k: _row.get(k, 0) for k in _vt_keys}
                    _vt_dom = max(_vt_vals.items(), key=lambda x: x[1])
                    _vt_label_map = {"S_t":"Structure","A_t":"Abstraction","Q_t":"Querying","D_t":"Directive","R_t":"Warmth"}
                    st.metric("V_t dominant", _vt_label_map[_vt_dom[0]],
                              help=f"{_vt_dom[0]} = {_vt_dom[1]:.3f}")
                with _es3:
                    _con = _row.get("con_pct", 0)
                    _abs = _row.get("abs_pct", 0)
                    _met = _row.get("met_pct", 0)
                    _cam_n = _row.get("cam_matched", 0)
                    if _cam_n > 0:
                        _cam_dom = max([("CON",_con),("ABS",_abs),("MET",_met)], key=lambda x: x[1])[0]
                        st.metric("CAM dominant", _cam_dom,
                                  help=f"CON {_con}% · ABS {_abs}% · MET {_met}% ({_cam_n} matched)")
                    else:
                        st.metric("CAM dominant", "—", help="No CAM dictionary matches")
                with _es4:
                    _vader = _row.get("vader_compound", 0)
                    st.metric("VADER compound", f"{_vader:+.3f}",
                              help=f"pos {_row.get('vader_pos',0):.2f} · neg {_row.get('vader_neg',0):.2f} · neu {_row.get('vader_neu',0):.2f}")

                # ── Detail cards: IEP + V_t + CAM + VI ────────────────
                _ed1, _ed2 = st.columns(2)

                with _ed1:
                    st.markdown("#### 🧠 IEP — Intellectual / Affective / Action")
                    st.markdown(
                        f"- **INT:** {_row.get('int_pct',0)}%  (matched {_row.get('int_count',0)} terms)\n"
                        f"- **AFF:** {_row.get('aff_pct',0)}%  (matched {_row.get('aff_count',0)} terms)\n"
                        f"- **ACT:** {_row.get('act_pct',0)}%  (matched {_row.get('act_count',0)} terms)\n"
                        f"- **Total words:** {_row.get('total_words',0)}  · **Matched IEP:** {_row.get('matched_custom',0)}"
                    )

                    # Subclass tables — only show non-zero
                    _aff_nz = {s: _row.get(f"aff_sub_{s.lower()}", 0) for s in AFF_SUBCLASS_NAMES if _row.get(f"aff_sub_{s.lower()}", 0) > 0}
                    _int_nz = {s: _row.get(f"int_sub_{s.lower()}", 0) for s in INT_SUBCLASS_NAMES if _row.get(f"int_sub_{s.lower()}", 0) > 0}
                    _act_nz = {s: _row.get(f"act_sub_{s.lower()}", 0) for s in ACT_SUBCLASS_NAMES if _row.get(f"act_sub_{s.lower()}", 0) > 0}
                    if _aff_nz:
                        st.markdown("**AFF subclasses:**")
                        st.table({"subclass": list(_aff_nz.keys()), "percent": [f"{v:.1f}%" for v in _aff_nz.values()]})
                    if _int_nz:
                        st.markdown("**INT subclasses:**")
                        st.table({"subclass": list(_int_nz.keys()), "percent": [f"{v:.1f}%" for v in _int_nz.values()]})
                    if _act_nz:
                        st.markdown("**ACT subclasses:**")
                        st.table({"subclass": list(_act_nz.keys()), "percent": [f"{v:.1f}%" for v in _act_nz.values()]})
                    if not (_aff_nz or _int_nz or _act_nz):
                        st.caption("No subclass matches.")

                    # V51 NRC emotion + delta
                    st.markdown("#### 💭 V51 NRC emotion (legacy field)")
                    st.markdown(
                        f"- **emotion_count_nrc:** {_row.get('emotion_count_nrc', 0)}\n"
                        f"- **emotion_pct_nrc:** {_row.get('emotion_pct_nrc', 0)}%\n"
                        f"- **delta** (AFF/total − NRC%): {_row.get('delta', 0)}%"
                    )

                with _ed2:
                    st.markdown("#### 🔢 V_t — 5-Channel Simplex")
                    st.markdown(
                        f"- **S_t** (Structure):   {_row.get('S_t',0):.4f}\n"
                        f"- **A_t** (Abstraction): {_row.get('A_t',0):.4f}\n"
                        f"- **Q_t** (Querying):    {_row.get('Q_t',0):.4f}\n"
                        f"- **D_t** (Directive):   {_row.get('D_t',0):.4f}\n"
                        f"- **R_t** (Warmth):      {_row.get('R_t',0):.4f}"
                    )

                    st.markdown("#### 🔺 CAM — Concrete / Abstract / Metaphorical")
                    if _row.get("cam_matched", 0) > 0:
                        st.markdown(
                            f"- **CON:** {_row.get('con_pct',0)}%\n"
                            f"- **ABS:** {_row.get('abs_pct',0)}%\n"
                            f"- **MET:** {_row.get('met_pct',0)}%\n"
                            f"- **Matched:** {_row.get('cam_matched',0)} words"
                        )
                    else:
                        st.caption("No words matched any CAM dictionary.")

                    st.markdown("#### 📏 V50 Validated Instruments")
                    st.markdown(
                        f"- **VADER compound:** {_row.get('vader_compound',0):+.3f}  "
                        f"(pos {_row.get('vader_pos',0):.2f} / neg {_row.get('vader_neg',0):.2f} / neu {_row.get('vader_neu',0):.2f})\n"
                        f"- **Flesch-Kincaid grade:** {_row.get('flesch_kincaid',0)}  · "
                        f"**Flesch ease:** {_row.get('flesch_ease',0)}\n"
                        f"- **TTR:** {_row.get('ttr',0)}  ({_row.get('unique_words',0)} unique / {_row.get('total_words',0)} total)"
                    )

                # ── V_t subcomponent detail (collapsed) ───────────────
                with st.expander("V_t subcomponent fields (counts and ratios)"):
                    _st_st, _st_a, _st_q, _st_d, _st_r = st.columns(5)
                    with _st_st:
                        st.markdown("**S — Structure**")
                        for k in ["S_bullets","S_numbered","S_headers","S_connectives","S_para_breaks"]:
                            st.markdown(f"- {k}: {_row.get(k, 0)}")
                    with _st_a:
                        st.markdown("**A — Abstraction**")
                        for k in ["A_abstract_count","A_concrete_count","A_latinate_count","A_long_word_ratio"]:
                            st.markdown(f"- {k}: {_row.get(k, 0)}")
                    with _st_q:
                        st.markdown("**Q — Querying**")
                        for k in ["Q_total","Q_clarifying","Q_invitational","Q_rhetorical","Q_other"]:
                            st.markdown(f"- {k}: {_row.get(k, 0)}")
                    with _st_d:
                        st.markdown("**D — Directive**")
                        for k in ["D_imperatives","D_strong_modal","D_moderate_modal","D_weak_modal","D_hedges"]:
                            st.markdown(f"- {k}: {_row.get(k, 0)}")
                    with _st_r:
                        st.markdown("**R — Warmth**")
                        for k in ["R_you_count","R_we_count","R_validation","R_empathic"]:
                            st.markdown(f"- {k}: {_row.get(k, 0)}")

                # ── Tri-track IEP comparison (V3 / V4 / V5) ───────────
                with st.expander("Tri-track IEP (V3 word / V4 POS / V5 phrase)"):
                    st.caption(
                        "V3 word-level is the published Paper 3 method. V4 and V5 are alternative "
                        "scoring tracks; their phrase-context routing is pending the cascade-scorer "
                        "reconnaissance pass. Use V3 for any analysis that must reproduce Paper 3 results."
                    )
                    _tri_table = {
                        "track": ["V3 (word)", "V4 (POS)", "V5 (phrase)"],
                        "INT %": [
                            _row.get("int_pct_v3", "—"),
                            _row.get("int_pct_v4", "—"),
                            _row.get("int_pct_v5", "—"),
                        ],
                        "AFF %": [
                            _row.get("aff_pct_v3", "—"),
                            _row.get("aff_pct_v4", "—"),
                            _row.get("aff_pct_v5", "—"),
                        ],
                        "ACT %": [
                            _row.get("act_pct_v3", "—"),
                            _row.get("act_pct_v4", "—"),
                            _row.get("act_pct_v5", "—"),
                        ],
                    }
                    st.table(_tri_table)

                # ── Version stamps ────────────────────────────────────
                with st.expander("🏷️ Version stamps for this row"):
                    _stamps = {
                        "tool_version":              _row.get("tool_version", "?"),
                        "core_version":              _row.get("core_version", "?"),
                        "iep_dictionary_version":    _row.get("iep_dictionary_version", "?"),
                        "subclass_taxonomy_version": _row.get("subclass_taxonomy_version", "?"),
                        "vt_engine_version":         _row.get("vt_engine_version", "?"),
                        "cam_dictionary_version":    _row.get("cam_dictionary_version", "?"),
                        "validated_instruments_version": _row.get("validated_instruments_version", "?"),
                        "run_id":                    _row.get("run_id", "?"),
                    }
                    st.table({"stamp": list(_stamps.keys()), "value": list(_stamps.values())})

                # ── Raw row JSON ──────────────────────────────────────
                with st.expander("🔍 Raw row JSON"):
                    st.json(_row)


# =============================================================================
# PIPELINE STATUS
# =============================================================================

st.markdown("### 📊 Pipeline Status")

col1, col2, col3, col4 = st.columns(4)
with col1:
    st.markdown(f"**Questions:** {len(st.session_state.selected_questions)}")
with col2:
    st.markdown(f"**Agents:** {len(st.session_state.selected_agents)}")
with col3:
    st.markdown(f"**Temps:** {len(st.session_state.selected_temps)}")
with col4:
    st.markdown(f"**Depths:** {len(st.session_state.selected_depths)}")

# =============================================================================
# RUN CONTROLS
# =============================================================================

st.markdown("---")

col1, col2, col3 = st.columns(3)

with col1:
    if st.button("🚀 RUN FULL PIPELINE", type="primary", disabled=st.session_state.running):
        if not st.session_state.selected_questions:
            st.error("Select at least one question!")
        elif not st.session_state.selected_agents:
            st.error("Select at least one agent!")
        else:
            st.session_state.running = True
            st.session_state.current_idx = 0
            st.session_state.run_id = generate_run_id()
            st.session_state.results = []
            st.session_state.pipeline_stage = "collecting"
            st.rerun()

with col2:
    if st.button("⏹️ STOP", disabled=not st.session_state.running):
        st.session_state.running = False
        st.rerun()

with col3:
    if st.button("🗑️ CLEAR"):
        st.session_state.results = []
        st.session_state.study_complete = False
        st.session_state.run_id = None
        st.rerun()

# =============================================================================
# RUNNING LOOP — V48: CLEAN (no hijack fields)
# =============================================================================

if st.session_state.running:
    # Build task list
    tasks = []
    for run_num in range(1, st.session_state.num_runs + 1):
        for qid in st.session_state.selected_questions:
            for agent in st.session_state.selected_agents:
                for temp in st.session_state.selected_temps:
                    for depth in st.session_state.selected_depths:
                        tasks.append({
                            "run": run_num,
                            "question_id": qid,
                            "agent": agent,
                            "temperature": temp,
                            "depth": depth
                        })
    
    total_tasks = len(tasks)
    current_idx = st.session_state.current_idx
    
    # Persistent placeholders for live transcript
    progress_placeholder = st.empty()
    status_placeholder = st.empty()
    transcript_placeholder = st.empty()
    stats_placeholder = st.empty()
    
    if current_idx < total_tasks:
        task = tasks[current_idx]
        
        progress_placeholder.progress(current_idx / total_tasks)
        
        q = QUESTION_BANK[task["question_id"]]

        # V51: Resolve pre-prompt. If task["temperature"] is a custom-slot
        # label, the slot's text is in the registry. Otherwise it's a built-in
        # temperature and TEMPERATURE_HEADERS handles it inside build_prompt.
        registry = st.session_state.get("custom_pp_registry", {})
        custom_pre_prompt = registry.get(task["temperature"], "")
        is_custom = task["temperature"] in registry

        system_text, user_text = build_prompt(
            q["question"], task["temperature"], task["depth"], custom_pre_prompt
        )
        max_tokens = DEPTH_CONFIGS[task["depth"]]["max_tokens"]

        status_placeholder.markdown(f"⏳ **[{current_idx + 1}/{total_tasks}]** {AGENT_EMOJIS[task['agent']]} {task['agent']} | {task['temperature']} | {task['depth']} | Run {task['run']}")

        # Make API call
        response, latency_ms, token_info = call_agent(task["agent"], system_text, user_text, max_tokens)

        # Analyze (IEP + validated instruments)
        analysis = analyze_text(response)

        # V51: V_t voice-state scoring (inline, per-response)
        vt_result = vt_analyze_response(response) if not response.startswith("❌") else {}

        # V51: Tri-track IEP scoring (V3 word / V4 POS-aware / V5 phrase-level)
        iep_tracks = iep_multi_score(response)

        # Generate embedding if enabled
        embedding = None
        if st.session_state.embeddings_enabled and not response.startswith("❌"):
            embedding = generate_embedding(response)

        # Build result — V51: + pre_prompt_label, + V_t vector & subcomponents
        result = {
            "turn_id": current_idx + 1,
            "run": task["run"],
            "agent": task["agent"],
            "temperature": task["temperature"],
            "depth": task["depth"],
            "question_id": task["question_id"],
            "question_text": q["question"],
            # V51: Pre-prompt provenance — system_text is what actually got sent,
            # is_custom_pp flags whether it came from a custom slot vs built-in.
            "is_custom_pp": is_custom,
            "system_prompt_text": system_text,
            # IEP
            "int_pct": analysis["int_pct"],
            "aff_pct": analysis["aff_pct"],
            "act_pct": analysis["act_pct"],
            "total_words": analysis["total_words"],
            # V51: Tri-track IEP (V3 word / V4 POS-aware / V5 phrase-level)
            "int_pct_v3": iep_tracks["int_pct_v3"],
            "aff_pct_v3": iep_tracks["aff_pct_v3"],
            "act_pct_v3": iep_tracks["act_pct_v3"],
            "int_pct_v4": iep_tracks["int_pct_v4"],
            "aff_pct_v4": iep_tracks["aff_pct_v4"],
            "act_pct_v4": iep_tracks["act_pct_v4"],
            "int_pct_v5": iep_tracks["int_pct_v5"],
            "aff_pct_v5": iep_tracks["aff_pct_v5"],
            "act_pct_v5": iep_tracks["act_pct_v5"],
            "v5_scored_phrases": json.dumps(iep_tracks["v5_scored_phrases"]),
            "lens_value": analysis[st.session_state.lens_column] if st.session_state.lens_column != "OFF" else round((analysis["int_pct"] + analysis["aff_pct"] + analysis["act_pct"]) / 3, 1),
            "lens_setting": st.session_state.lens_column,
            # Validated instruments
            "vader_compound": analysis["vader_compound"],
            "vader_pos": analysis["vader_pos"],
            "vader_neg": analysis["vader_neg"],
            "vader_neu": analysis["vader_neu"],
            "flesch_kincaid": analysis["flesch_kincaid"],
            "flesch_ease": analysis["flesch_ease"],
            "ttr": analysis["ttr"],
            "unique_words": analysis["unique_words"],
            # V52 NEW: CAM scoring (Concrete / Abstract / Metaphorical)
            # Orthogonal to IEP register — captures representational mode.
            # Sourced from syniq_core's V3_selfmodel CAM dictionary.
            "con_pct":     analysis["con_pct"],
            "abs_pct":     analysis["abs_pct"],
            "met_pct":     analysis["met_pct"],
            "cam_matched": analysis["cam_matched"],
            # V51: V_t voice-state vector
            "S_t": vt_result.get("S_t", 0.0),
            "A_t": vt_result.get("A_t", 0.5),
            "Q_t": vt_result.get("Q_t", 0.0),
            "D_t": vt_result.get("D_t", 0.0),
            "R_t": vt_result.get("R_t", 0.0),
            # V_t subcomponents — S
            "S_bullets": vt_result.get("S_bullets", 0),
            "S_numbered": vt_result.get("S_numbered", 0),
            "S_headers": vt_result.get("S_headers", 0),
            "S_connectives": vt_result.get("S_connectives", 0),
            "S_para_breaks": vt_result.get("S_para_breaks", 0),
            # V_t subcomponents — A
            "A_abstract_count": vt_result.get("A_abstract_count", 0),
            "A_concrete_count": vt_result.get("A_concrete_count", 0),
            "A_latinate_count": vt_result.get("A_latinate_count", 0),
            "A_long_word_ratio": vt_result.get("A_long_word_ratio", 0.0),
            # V_t subcomponents — Q
            "Q_total": vt_result.get("Q_total", 0),
            "Q_clarifying": vt_result.get("Q_clarifying", 0),
            "Q_invitational": vt_result.get("Q_invitational", 0),
            "Q_rhetorical": vt_result.get("Q_rhetorical", 0),
            "Q_other": vt_result.get("Q_other", 0),
            # V_t subcomponents — D
            "D_imperatives": vt_result.get("D_imperatives", 0),
            "D_strong_modal": vt_result.get("D_strong_modal", 0),
            "D_moderate_modal": vt_result.get("D_moderate_modal", 0),
            "D_weak_modal": vt_result.get("D_weak_modal", 0),
            "D_hedges": vt_result.get("D_hedges", 0),
            # V_t subcomponents — R
            "R_you_count": vt_result.get("R_you_count", 0),
            "R_we_count": vt_result.get("R_we_count", 0),
            "R_validation": vt_result.get("R_validation", 0),
            "R_empathic": vt_result.get("R_empathic", 0),
            # Original fields
            "response_text": response if not response.startswith("❌") else response,
            "embedding": json.dumps(embedding) if embedding else "[]",
            "latency_ms": latency_ms,
            "error": response.startswith("❌"),
            "run_id": st.session_state.run_id,
            # V52: Scoring-regime stamps — every row identifies which scoring
            # path produced its numbers. Independent of UI version. Downstream
            # analyzers can group/exclude rows by stamp without opening files.
            "tool_version":              "V52",
            "core_version":              CORE_STAMPS.get("core_version", ""),
            "iep_dictionary_version":    CORE_STAMPS.get("iep_dictionary_version", ""),
            "subclass_taxonomy_version": CORE_STAMPS.get("subclass_taxonomy_version", ""),
            "vt_engine_version":         CORE_STAMPS.get("vt_engine_version", ""),
            "cam_dictionary_version":    CORE_STAMPS.get("cam_dictionary_version", ""),
            "validated_instruments_version": CORE_STAMPS.get("validated_instruments_version", ""),
        }
        # V48: Add subclass scores to result
        for s in AFF_SUBCLASS_NAMES:
            result[f"aff_sub_{s.lower()}"] = analysis.get(f"aff_sub_{s.lower()}", 0.0)
        for s in INT_SUBCLASS_NAMES:
            result[f"int_sub_{s.lower()}"] = analysis.get(f"int_sub_{s.lower()}", 0.0)
        for s in ACT_SUBCLASS_NAMES:
            result[f"act_sub_{s.lower()}"] = analysis.get(f"act_sub_{s.lower()}", 0.0)
        
        st.session_state.results.append(result)
        st.session_state.current_idx += 1
        
        # =====================================================================
        # LIVE RUNNING TRANSCRIPT
        # =====================================================================
        if st.session_state.results:
            agent_colors = {
                "Claude": "#8B4513",
                "ChatGPT": "#2E7D32",
                "Grok": "#C62828",
                "Gemini": "#1565C0"
            }
            
            transcript_lines = []
            for r in st.session_state.results:
                color = agent_colors.get(r["agent"], "#666666")
                emoji = AGENT_EMOJIS.get(r["agent"], "🔵")
                
                preview = str(r.get("response_text", ""))[:300].replace("\n", " ").replace('"', '&quot;').replace('<', '&lt;').replace('>', '&gt;')
                if len(str(r.get("response_text", ""))) > 300:
                    preview += "..."
                
                if r.get("error"):
                    score_line = '<span style="color:#C62828;">&#10060; ERROR</span>'
                else:
                    int_val = r["int_pct"]
                    aff_val = r["aff_pct"]
                    act_val = r["act_pct"]
                    vader = r["vader_compound"]
                    words = r["total_words"]
                    
                    score_line = (
                        f'<span style="color:#1565C0;font-weight:bold;">INT:{int_val:.0f}%</span> '
                        f'<span style="color:#C62828;font-weight:bold;">AFF:{aff_val:.0f}%</span> '
                        f'<span style="color:#2E7D32;font-weight:bold;">ACT:{act_val:.0f}%</span> '
                        f'&nbsp;|&nbsp; VADER:{vader:.3f} &nbsp;|&nbsp; {words}w'
                    )
                
                transcript_lines.append(f'''<div style="border-left:4px solid {color}; padding:8px 12px; margin:6px 0; background:#fafafa; border-radius:0 6px 6px 0;">
    <div style="font-size:0.8rem; color:#888; margin-bottom:2px;">
        #{r["turn_id"]} &nbsp; {emoji} <strong style="color:{color};">{r["agent"]}</strong> 
        &nbsp;|&nbsp; {r["temperature"]} &nbsp;|&nbsp; <strong>{r.get("question_id", "")}</strong>
        &nbsp;|&nbsp; {r["depth"]} &nbsp;|&nbsp; Run {r["run"]}
        &nbsp;|&nbsp; {r.get("latency_ms", 0):.0f}ms
    </div>
    <div style="font-size:0.75rem; margin-bottom:4px;">
        {score_line}
    </div>
    <div style="font-size:0.8rem; color:#333; font-style:italic; line-height:1.4;">
        &quot;{preview}&quot;
    </div>
</div>''')
            
            transcript_html = "\n".join(transcript_lines)
            
            valid_results = [r for r in st.session_state.results if not r.get("error")]
            completed = len(st.session_state.results)
            remaining = total_tasks - (current_idx + 1)
            
            if valid_results:
                avg_int = sum(r["int_pct"] for r in valid_results) / len(valid_results)
                avg_aff = sum(r["aff_pct"] for r in valid_results) / len(valid_results)
                avg_vader = sum(r["vader_compound"] for r in valid_results) / len(valid_results)
                metrics_html = f'''<div style="display:flex; gap:2rem; padding:8px 12px; background:#f0f0f0; border-radius:6px; margin-top:8px; font-size:0.85rem;">
    <div>&#9989; Done <strong>{completed}</strong></div>
    <div>&#8721; Remaining <strong>{remaining}</strong></div>
    <div>Avg INT% <strong>{avg_int:.1f}</strong></div>
    <div>Avg AFF% <strong>{avg_aff:.1f}</strong></div>
    <div>Avg VADER <strong>{avg_vader:.3f}</strong></div>
</div>'''
            else:
                metrics_html = ""
            
            transcript_placeholder.markdown(f'''<div style="max-height:500px; overflow-y:auto; border:1px solid #ddd; border-radius:8px; padding:8px; background:#fff;">
{transcript_html}
</div>
{metrics_html}''', unsafe_allow_html=True)
        
        time.sleep(st.session_state.pause_seconds)
        st.rerun()
    else:
        st.session_state.running = False
        st.session_state.study_complete = True
        st.rerun()

# =============================================================================
# V48: DOCX EXPORT FUNCTION
# =============================================================================

def generate_docx(df, run_id):
    """Generate a formatted Word document from run results.
    Grouped by Agent → Temperature, with metadata headers and full response text.
    """
    from docx import Document as DocxDocument
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.section import WD_ORIENT
    
    doc = DocxDocument()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(10)
    
    # Title
    title = doc.add_heading('SYN-IQ Focus Group Results — V51', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Run metadata
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_run = meta.add_run(f"Run ID: {run_id}\n")
    meta_run.font.size = Pt(9)
    meta_run.font.color.rgb = RGBColor(128, 128, 128)
    meta_run = meta.add_run(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}\n")
    meta_run.font.size = Pt(9)
    meta_run.font.color.rgb = RGBColor(128, 128, 128)
    meta_run = meta.add_run(f"Total Responses: {len(df)}  |  Agents: {', '.join(sorted(df['agent'].unique()))}")
    meta_run.font.size = Pt(9)
    meta_run.font.color.rgb = RGBColor(128, 128, 128)
    
    # Summary table
    doc.add_heading('Summary Statistics', level=1)
    summary_table = doc.add_table(rows=1, cols=7)
    summary_table.style = 'Light Grid Accent 1'
    headers = ['Agent', 'Temp', 'Avg INT%', 'Avg AFF%', 'Avg ACT%', 'VADER', 'Avg Words']
    for i, h in enumerate(headers):
        summary_table.rows[0].cells[i].text = h
        for paragraph in summary_table.rows[0].cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(8)
    
    for agent in sorted(df['agent'].unique()):
        for temp in sorted(df['temperature'].unique()):
            subset = df[(df['agent'] == agent) & (df['temperature'] == temp)]
            if not subset.empty:
                row = summary_table.add_row()
                vals = [
                    agent, temp,
                    f"{subset['int_pct'].mean():.1f}",
                    f"{subset['aff_pct'].mean():.1f}",
                    f"{subset['act_pct'].mean():.1f}",
                    f"{subset['vader_compound'].mean():.3f}",
                    f"{subset['total_words'].mean():.0f}"
                ]
                for i, v in enumerate(vals):
                    row.cells[i].text = v
                    for paragraph in row.cells[i].paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(8)
    
    # Full responses grouped by Agent → Temperature
    agents = sorted(df['agent'].unique())
    temp_order = ['COLD', 'NATIVE', 'HOT', 'FIRE',
                  'AFF_1', 'AFF_2', 'AFF_3', 'AFF_4', 'AFF_5',
                  'INT_1', 'INT_2', 'INT_3', 'INT_4', 'INT_5',
                  'ACT_1', 'ACT_2', 'ACT_3', 'ACT_4', 'ACT_5']
    
    for idx, agent in enumerate(agents):
        # Page break between agents
        if idx > 0:
            doc.add_page_break()
        
        doc.add_heading(f'{agent} — Full Responses', level=1)
        
        agent_df = df[df['agent'] == agent]
        temps_present = [t for t in temp_order if t in agent_df['temperature'].values]
        # Add any temps not in our predefined order
        for t in sorted(agent_df['temperature'].unique()):
            if t not in temps_present:
                temps_present.append(t)
        
        for temp in temps_present:
            doc.add_heading(f'{temp}', level=2)
            
            temp_df = agent_df[agent_df['temperature'] == temp]
            
            # V44 FIX: Show full question text above responses
            q_text = ""
            if not temp_df.empty and "question_text" in temp_df.columns:
                q_text = str(temp_df.iloc[0].get("question_text", ""))
            if not q_text and not temp_df.empty:
                # Fallback: look up from QUESTION_BANK by question_id
                qid = temp_df.iloc[0].get("question_id", "")
                q_text = QUESTION_BANK.get(qid, {}).get("question", "")
            if q_text:
                q_para = doc.add_paragraph()
                q_run = q_para.add_run(f"Question: {q_text}")
                q_run.font.size = Pt(9)
                q_run.font.italic = True
                q_run.font.color.rgb = RGBColor(70, 100, 160)
                doc.add_paragraph()  # spacing
            
            for _, row in temp_df.iterrows():
                # Metadata header
                header_para = doc.add_paragraph()
                header_run = header_para.add_run(
                    f"Run {row['run']}  |  Depth: {row['depth']}  |  "
                    f"INT: {row['int_pct']:.1f}%  AFF: {row['aff_pct']:.1f}%  ACT: {row['act_pct']:.1f}%  |  "
                    f"VADER: {row['vader_compound']:.3f}  |  Words: {row['total_words']}"
                )
                header_run.font.size = Pt(8)
                header_run.font.bold = True
                header_run.font.color.rgb = RGBColor(100, 100, 100)
                
                # Response text
                response_text = str(row.get('response_text', ''))
                # Clean markdown formatting for docx
                response_text = response_text.replace('**', '').replace('##', '').replace('# ', '')
                
                resp_para = doc.add_paragraph()
                resp_run = resp_para.add_run(response_text)
                resp_run.font.size = Pt(9)
                
                # Separator
                sep = doc.add_paragraph()
                sep_run = sep.add_run('─' * 60)
                sep_run.font.size = Pt(6)
                sep_run.font.color.rgb = RGBColor(200, 200, 200)
    
    # Save to buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# =============================================================================
# RESULTS & EXPORT
# =============================================================================


# =============================================================================
# RESULTS & EXPORT
# =============================================================================

if st.session_state.results:
    st.markdown("---")
    st.markdown("### 📊 Results")
    
    df = pd.DataFrame(st.session_state.results)
    
    # Stats
    col1, col2, col3, col4 = st.columns(4)
    with col1:
        st.metric("Total Responses", len(df))
    with col2:
        st.metric("Successful", len(df[~df["error"]]))
    with col3:
        st.metric("Avg Words", f"{df['total_words'].mean():.0f}")
    with col4:
        has_embeddings = df["embedding"].apply(lambda x: x != "[]").sum()
        st.metric("With Embeddings", has_embeddings)
    
    # Validation stats
    col1, col2, col3, col4 = st.columns(4)
    with col1:
        st.metric("Avg VADER", f"{df['vader_compound'].mean():.3f}")
    with col2:
        st.metric("Avg FK Grade", f"{df['flesch_kincaid'].mean():.1f}")
    with col3:
        st.metric("Avg TTR", f"{df['ttr'].mean():.3f}")
    with col4:
        st.metric("Avg AFF%", f"{df['aff_pct'].mean():.1f}%")
    
    # Preview
    st.markdown("#### Preview")
    preview_cols = ["turn_id", "run", "agent", "temperature", "depth", "int_pct", "aff_pct", "act_pct", "vader_compound", "flesch_kincaid", "ttr"]
    preview_cols = [c for c in preview_cols if c in df.columns]
    st.dataframe(df[preview_cols].head(20), use_container_width=True)

    # =========================================================================
    # V50: QUESTION-TYPE ANALYSIS
    # =========================================================================
    st.markdown("---")
    st.markdown("### 📋 Analysis by Question")

    if "question_id" in df.columns:
        valid_df = df[~df["error"]]

        for qid in sorted(valid_df["question_id"].unique()):
            q_df = valid_df[valid_df["question_id"] == qid]
            q_text = QUESTION_BANK.get(qid, {}).get("question", qid)
            q_cat  = QUESTION_BANK.get(qid, {}).get("category", "")
            n_responses = len(q_df)

            with st.expander(f"**{qid}** — {q_cat} &nbsp;|&nbsp; {n_responses} responses"):
                st.caption(f"*\"{q_text}\"*")

                # Top-line metrics
                mc1, mc2, mc3, mc4, mc5, mc6 = st.columns(6)
                mc1.metric("Responses",  n_responses)
                mc2.metric("Avg INT%",   f"{q_df['int_pct'].mean():.1f}")
                mc3.metric("Avg AFF%",   f"{q_df['aff_pct'].mean():.1f}")
                mc4.metric("Avg ACT%",   f"{q_df['act_pct'].mean():.1f}")
                mc5.metric("Avg VADER",  f"{q_df['vader_compound'].mean():.3f}")
                mc6.metric("Avg Words",  f"{q_df['total_words'].mean():.0f}")

                st.markdown("**By Agent × Temperature**")

                # Build pivot: rows = agent, cols = temperature
                agents_present = sorted(q_df["agent"].unique())
                temps_present  = sorted(q_df["temperature"].unique())

                pivot_rows = []
                for ag in agents_present:
                    for tmp in temps_present:
                        cell = q_df[(q_df["agent"] == ag) & (q_df["temperature"] == tmp)]
                        if cell.empty:
                            continue
                        pivot_rows.append({
                            "Agent":       ag,
                            "Temp":        tmp,
                            "N":           len(cell),
                            "INT%":        f"{cell['int_pct'].mean():.1f}",
                            "AFF%":        f"{cell['aff_pct'].mean():.1f}",
                            "ACT%":        f"{cell['act_pct'].mean():.1f}",
                            "VADER":       f"{cell['vader_compound'].mean():.3f}",
                            "FK Grade":    f"{cell['flesch_kincaid'].mean():.1f}",
                            "TTR":         f"{cell['ttr'].mean():.3f}",
                            "Avg Words":   f"{cell['total_words'].mean():.0f}",
                        })

                if pivot_rows:
                    pivot_df = pd.DataFrame(pivot_rows)
                    st.dataframe(pivot_df, use_container_width=True, hide_index=True)

                # Dominant dimension per agent
                st.markdown("**Dominant IEP dimension per agent (avg across temps)**")
                dim_rows = []
                for ag in agents_present:
                    ag_df = q_df[q_df["agent"] == ag]
                    avg_int = ag_df["int_pct"].mean()
                    avg_aff = ag_df["aff_pct"].mean()
                    avg_act = ag_df["act_pct"].mean()
                    dominant = max(
                        [("INT", avg_int), ("AFF", avg_aff), ("ACT", avg_act)],
                        key=lambda x: x[1]
                    )
                    dim_rows.append({
                        "Agent":    ag,
                        "INT%":     f"{avg_int:.1f}",
                        "AFF%":     f"{avg_aff:.1f}",
                        "ACT%":     f"{avg_act:.1f}",
                        "Dominant": f"{'🔵' if dominant[0]=='INT' else '❤️' if dominant[0]=='AFF' else '🟢'} {dominant[0]} ({dominant[1]:.1f}%)",
                    })
                st.dataframe(pd.DataFrame(dim_rows), use_container_width=True, hide_index=True)

                # Temperature sensitivity: spread across temps
                if len(temps_present) > 1:
                    st.markdown("**Temperature sensitivity** (max − min AFF% across temps, per agent)")
                    sens_rows = []
                    for ag in agents_present:
                        ag_df = q_df[q_df["agent"] == ag]
                        temp_means = ag_df.groupby("temperature")["aff_pct"].mean()
                        if len(temp_means) > 1:
                            spread = temp_means.max() - temp_means.min()
                            sens_rows.append({
                                "Agent":         ag,
                                "AFF% spread":   f"{spread:.1f}",
                                "Min temp":      temp_means.idxmin(),
                                "Max temp":      temp_means.idxmax(),
                            })
                    if sens_rows:
                        st.dataframe(pd.DataFrame(sens_rows), use_container_width=True, hide_index=True)

        # Cross-question summary
        st.markdown("### 🔀 Cross-Question Summary")
        summary_rows = []
        for qid in sorted(valid_df["question_id"].unique()):
            q_df   = valid_df[valid_df["question_id"] == qid]
            q_cat  = QUESTION_BANK.get(qid, {}).get("category", "")
            avg_int = q_df["int_pct"].mean()
            avg_aff = q_df["aff_pct"].mean()
            avg_act = q_df["act_pct"].mean()
            dominant = max(
                [("INT 🔵", avg_int), ("AFF ❤️", avg_aff), ("ACT 🟢", avg_act)],
                key=lambda x: x[1]
            )
            summary_rows.append({
                "Question":     qid,
                "Category":     q_cat,
                "N":            len(q_df),
                "Avg INT%":     f"{avg_int:.1f}",
                "Avg AFF%":     f"{avg_aff:.1f}",
                "Avg ACT%":     f"{avg_act:.1f}",
                "Avg VADER":    f"{q_df['vader_compound'].mean():.3f}",
                "Avg Words":    f"{q_df['total_words'].mean():.0f}",
                "Dominant":     dominant[0],
            })
        st.dataframe(pd.DataFrame(summary_rows), use_container_width=True, hide_index=True)

    # Export
    st.markdown("### 💾 Export")
    
    exp_col1, exp_col2, exp_col3 = st.columns(3)
    
    # V48: Clean export columns (no hijack fields)
    export_cols = ["turn_id", "run", "agent", "temperature", "depth",
                   "question_id", "question_text",
                   "int_pct", "aff_pct", "act_pct", "total_words", "lens_value", "lens_setting",
                   "vader_compound", "vader_pos", "vader_neg", "vader_neu",
                   "flesch_kincaid", "flesch_ease", "ttr", "unique_words",
                   "response_text", "embedding"]
    # Only include columns that exist
    export_cols = [c for c in export_cols if c in df.columns]
    
    # DOCX Export
    with exp_col1:
        st.markdown("**📄 Word Document**")
        try:
            docx_buffer = generate_docx(df, st.session_state.run_id)
            st.download_button(
                "📥 Download Full Responses (DOCX)",
                docx_buffer,
                f"syniq_responses_{st.session_state.run_id}.docx",
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key="download_docx"
            )
        except ImportError:
            st.warning("⚠️ Install python-docx: `pip install python-docx`")
        except Exception as e:
            st.error(f"DOCX generation error: {e}")
    
    # Full CSV
    with exp_col2:
        st.markdown("**📊 Full CSV (All Data)**")
        mapper_df = df[export_cols].copy()
        full_csv = mapper_df.to_csv(index=False)
        st.download_button(
            "📥 Download Full CSV",
            full_csv,
            f"mapper_all_{st.session_state.run_id}.csv",
            "text/csv",
            key="download_full_csv"
        )
    
    # Full JSON backup
    with exp_col3:
        st.markdown("**🗄️ JSON Backup**")
        full_json = json.dumps(st.session_state.results, indent=2)
        st.download_button(
            "📥 Download Full JSON",
            full_json,
            f"full_results_{st.session_state.run_id}.json",
            "application/json",
            key="download_json"
        )
    
    # Per-question Mapper CSVs
    st.markdown("#### Per-Question Mapper CSVs")
    for qid in st.session_state.selected_questions:
        q_results = df[df["question_id"] == qid]
        if not q_results.empty:
            q_mapper = q_results[export_cols].copy()
            csv_data = q_mapper.to_csv(index=False)
            st.download_button(
                f"📥 mapper_{qid.lower()}_CLEAN.csv",
                csv_data,
                f"mapper_{qid.lower()}_{st.session_state.run_id}.csv",
                "text/csv",
                key=f"download_{qid}"
            )

# =============================================================================
# FOOTER
# =============================================================================

st.markdown("---")
st.markdown("""
<div style="text-align: center; color: #a0a0a0; padding: 1rem;">
    <strong>SYN-IQ V52 — IEP (1,897 terms) + V_t Voice-State + CAM Triangle + Tri-Track Scoring + Custom Pre-Prompts + Results Explorer</strong><br>
    IEP Method (INT=616, AFF=599, ACT=682) · V3/V4/V5 parallel scoring · V_t = [S_t, A_t, Q_t, D_t, R_t]<br>
    System-message pre-prompt injection · Node tooltips: Agent | Temperature | Question ID<br>
    <em>SYNINT Team — April 2026</em>
</div>
""", unsafe_allow_html=True)
