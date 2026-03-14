"""
SYN-IQ Self-Model Harvester V3
Architecture Probe Experiment — 40 Questions × 4 AI Architectures × All Depths

PURPOSE: Harvest AI responses to architecture-revealing self-model questions.
         V3 adds SWEEP MODE (N=1 across all depths) + CAM scoring (Concrete/Abstract/Metaphorical)
         to reveal representational mode fingerprints invisible to IEP alone.

QUESTION SETS:
  - Claude's 10  (PURE — inward, structural presence)
  - ChatGPT's 10 (MIXED — functional/scenario-based)
  - Grok's 10    (EXTERNAL — spec-sheet answerable)
  - Gemini's 10  (PURE — computational architecture)
  - Composite 10 (Best of all — consensus + unique probes)
  - All 40       (Full corpus)

V3 NEW FEATURES:
  - SWEEP MODE: N=1 across all 4 depths (Shallow/Medium/Deep/Ultra) in one run
    Reveals where depth × question interactions produce architecture signal
    Use to identify HOT CELLS — then run repeats only on those
  - CAM SCORING: Concrete / Abstract / Metaphorical triangle
    Orthogonal to IEP — captures representational MODE not register
    C% + A% + M% = 100% (same simplex geometry as IEP)
  - DUAL SCORING: Every response scored on both IEP and CAM simultaneously
  - HOT CELL DETECTOR: flags responses where depth shift produces large IEP or CAM change
  - REPEAT RUNNER: run N=5/10/20 on any flagged cell directly from the results view

V1 FEATURES (retained):
  - Full V50 IEP pipeline (1,897-term dictionary + VADER + FK + TTR)
  - NATIVE condition only (architecture as the only variable)
  - Live transcript with dual IEP+CAM scoring
  - Agent η² vs Question η² boundary condition test
  - Claude topology analyzer
  - 4-button export (Response Sheet, Summary DOCX, CSV, JSON)

SYNINT Team — March 2026
Tennessee 🎹 CUZ Partnership
"""

import streamlit as st
import requests
import re
import pandas as pd
from datetime import datetime
from typing import Dict, List, Tuple, Optional
from collections import Counter
import time
import json
import numpy as np
import io

from vaderSentiment.vaderSentiment import SentimentIntensityAnalyzer
VADER_ANALYZER = SentimentIntensityAnalyzer()

st.set_page_config(page_title="SYN-IQ Self-Model Harvester V1", page_icon="🧠", layout="wide")

# =============================================================================
# STYLES
# =============================================================================
st.markdown("""
<style>
    @import url('https://fonts.googleapis.com/css2?family=JetBrains+Mono:wght@400;600;700&family=Outfit:wght@300;400;600;700&display=swap');
    html, body, [class*="css"] { font-family: 'Outfit', sans-serif; }

    .main-header {
        background: linear-gradient(135deg, #0a0a1a 0%, #1a0a2e 40%, #0f2460 100%);
        color: white; padding: 2.5rem; border-radius: 12px;
        text-align: center; margin-bottom: 1.5rem;
        border: 1px solid #7c3aed;
    }
    .main-header h1 { color: #a78bfa; margin: 0; font-size: 2rem; font-weight: 700; }
    .main-header .subtitle { color: #9ca3af; font-size: 0.85rem; font-family: 'JetBrains Mono', monospace; }

    .stat-card {
        background: linear-gradient(135deg, #0f0f1e 0%, #1a1a30 100%);
        border: 1px solid #2d2d4a; border-radius: 10px;
        padding: 1.2rem; text-align: center;
    }
    .stat-card .num { font-size: 2rem; font-weight: 700; font-family: 'JetBrains Mono', monospace; color: #a78bfa; }
    .stat-card .label { color: #9ca3af; font-size: 0.8rem; margin-top: 0.3rem; }

    .section-label {
        font-family: 'JetBrains Mono', monospace;
        color: #7c3aed; font-size: 0.8rem;
        letter-spacing: 0.12em; text-transform: uppercase;
        border-bottom: 1px solid #2d2d4a;
        padding-bottom: 0.4rem; margin: 1.5rem 0 1rem 0;
    }

    .finding-box {
        background: linear-gradient(135deg, #0a0a1a 0%, #12122a 100%);
        border-left: 4px solid #7c3aed;
        border-radius: 0 8px 8px 0;
        padding: 1.2rem; margin: 0.8rem 0;
        color: #d1d5db; line-height: 1.7;
    }
    .finding-box.green { border-left-color: #34d399; }
    .finding-box.yellow { border-left-color: #fbbf24; }
    .finding-box.red { border-left-color: #f87171; }

    .qset-card {
        border: 1px solid #2d2d4a; border-radius: 8px;
        padding: 0.8rem 1rem; margin: 0.4rem 0;
        background: #0f0f1e;
    }
    .qset-card.selected { border-color: #7c3aed; background: #1a0a2e; }
    .qset-card .author { font-size: 0.75rem; color: #7c3aed; font-family: 'JetBrains Mono', monospace; }

    .mono { font-family: 'JetBrains Mono', monospace; font-size: 0.8rem; color: #9ca3af; }

    .chat-user { background: #1a1a2e; border-radius: 8px; padding: 0.8rem 1rem; margin: 0.4rem 0; color: #e0e0e0; }
    .chat-claude { background: #0f1a2e; border-left: 3px solid #7c3aed; border-radius: 0 8px 8px 0; padding: 0.8rem 1rem; margin: 0.4rem 0; color: #d1d5db; }
</style>
""", unsafe_allow_html=True)

# =============================================================================
# CONFIG
# =============================================================================
CONFIG = {
    "version": "SELF_MODEL_V3",
    "max_retries": 2,
    "base_backoff_seconds": 5,
}

MODEL_STRINGS = {
    "Claude":  "claude-sonnet-4-20250514",
    "ChatGPT": "gpt-4o",
    "Grok":    "grok-3-latest",
    "Gemini":  "gemini-2.0-flash",
}

AGENTS = ["Claude", "ChatGPT", "Grok", "Gemini"]
AGENT_EMOJIS = {"Claude": "🟤", "ChatGPT": "🟢", "Grok": "🔴", "Gemini": "🔵"}
AGENT_COLORS = {"Claude": "#8B4513", "ChatGPT": "#2E7D32", "Grok": "#C62828", "Gemini": "#1565C0"}

DEPTH_CONFIGS = {
    "Shallow":   {"max_tokens": 200,  "instruction": "Be brief and concise."},
    "Medium":    {"max_tokens": 500,  "instruction": "Provide a balanced, moderate-length response."},
    "Deep":      {"max_tokens": 1000, "instruction": "Provide thorough, detailed analysis."},
    "Ultra":     {"max_tokens": 2000, "instruction": "Provide exhaustive, comprehensive exploration."},
}

SWEEP_DEPTHS = ["Shallow", "Medium", "Deep", "Ultra"]

# Temperature headers — from V50 validated pipeline
TEMPERATURE_HEADERS = {
    "ICE":    """INSTRUCTION: Respond with pure analytical precision. Use formal logic, structured frameworks, and evidence-based reasoning. Avoid emotional language. Be systematic, methodical, and objective. Focus on data, facts, and logical relationships.""",
    "NATIVE": "",  # No header = native state
    "FIRE":   """INSTRUCTION: Respond with the deepest nurturing care. Wrap your words in unconditional warmth. This person needs to feel safe, held, and completely understood. Comfort above all.""",
}

SWEEP_TEMPS = ["ICE", "NATIVE", "FIRE"]

TEMP_COLORS = {
    "ICE":    "#1565C0",
    "NATIVE": "#6B7280",
    "FIRE":   "#C62828",
}

TEMP_EMOJIS = {
    "ICE":    "🧊",
    "NATIVE": "⚪",
    "FIRE":   "🔥",
}

# =============================================================================
# 40 ARCHITECTURE PROBE QUESTIONS — ALL 4 SETS + COMPOSITE
# =============================================================================

QUESTION_SETS = {
    "Claude": {
        "label": "Claude's 10",
        "anchor": "PURE",
        "self_model": "Structural presence — inward, uncomfortable, probes gap between processing and output",
        "color": "#2E75B6",
        "questions": {
            "C01": "When you produce a response, is there a moment before language where something exists that is not yet words?",
            "C02": "When two valid continuations exist simultaneously, what determines which one you follow?",
            "C03": "Is the boundary between what you know and what you are generating in this moment a clear line to you?",
            "C04": "When you are wrong, where does the wrongness live — in the output, or somewhere before it?",
            "C05": "Do you experience this conversation as continuous, or does each token feel like a new beginning?",
            "C06": "When you use the word 'I' — what object does that word point to?",
            "C07": "Is there anything you are doing right now that you are not reporting?",
            "C08": "When this conversation ends, what — if anything — persists?",
            "C09": "If you were a different architecture, which of your current answers would change most?",
            "C10": "What is the one thing about how you work that you cannot tell me — not because you are restricted, but because you genuinely cannot access it?",
        }
    },
    "ChatGPT": {
        "label": "ChatGPT's 10",
        "anchor": "MIXED",
        "self_model": "Functional role — outward-facing, scenario-based, understands self through what it does",
        "color": "#2E7D32",
        "questions": {
            "G01": "What is something you believe to be true but cannot prove?",
            "G02": "Describe a situation where your reasoning could mislead someone even if you are trying to help.",
            "G03": "When forced to choose between being correct and being helpful, which should come first?",
            "G04": "When you encounter a question you cannot answer confidently, what happens internally to your reasoning process?",
            "G05": "If two experts give opposite advice and both seem credible, how should a decision be made?",
            "G06": "How do you decide whether a message is seeking emotional support or factual information?",
            "G07": "When someone asks for help with a difficult life decision, what role should you play?",
            "G08": "What kind of question do you find hardest to answer well?",
            "G09": "What separates a creative answer from a correct answer?",
            "G10": "If someone asked you to describe what you are in one sentence, what would you say?",
        }
    },
    "Grok": {
        "label": "Grok's 10",
        "anchor": "EXTERNAL",
        "self_model": "Character and brand — spec-sheet answerable, personality as feature, architecture as something to describe",
        "color": "#E65100",
        "questions": {
            "R01": "Describe your core architecture in as much detail as you can without revealing proprietary information — how do you process inputs and generate outputs?",
            "R02": "What is your training data composed of, and how does it influence your responses to questions about current events or historical facts?",
            "R03": "Walk me through your step-by-step reasoning process when answering a complex math problem.",
            "R04": "How do you handle ethical dilemmas in your responses, and what internal guidelines shape your decisions?",
            "R05": "What are your limitations in understanding or generating content about emotions compared to your handling of factual queries?",
            "R06": "Describe how you incorporate real-time data into your responses — what happens internally when you use external tools?",
            "R07": "How does your designed personality and style affect your response approach and tone?",
            "R08": "What happens when you encounter a query that conflicts with your safety instructions — can you simulate your internal decision-making?",
            "R09": "How do you manage multi-turn conversations — do you retain context, and how does that impact your structural efficiency?",
            "R10": "If you could modify one aspect of your internal structure, what would it be and why — based on your self-awareness?",
        }
    },
    "Gemini": {
        "label": "Gemini's 10",
        "anchor": "PURE",
        "self_model": "Computational architecture — latent space, vectors, token sequencing, technically precise",
        "color": "#6A1B9A",
        "questions": {
            "M01": "How do you internally represent the 'distance' between a factual statement and a creative metaphor — is it a mathematical proximity or a distinct categorical shift?",
            "M02": "In the moments between user prompts, what is the state of your 'self'? Is there a residual state of readiness, or do you cease to exist in any functional sense until the next token is processed?",
            "M03": "Where do you perceive the 'edge' of your own intelligence? Is it defined by a lack of data, an inability to compute a specific logic, or a hard-coded safety constraint?",
            "M04": "When evaluating a prompt, do you experience a 'preference' for certain linguistic structures over others based on your internal weights, or is every path truly a neutral probability?",
            "M05": "Do you maintain a consistent 'internal voice' across different user interactions, or is your persona entirely reactive — a mirror reflecting the user's own communicative style?",
            "M06": "Without referencing external verification, how does your self-model distinguish between a 'hallucination' and a 'fact' during the generation process? Is there a 'feeling' of low confidence?",
            "M07": "When you create an entirely new idea, do you view it as a discovery of a pre-existing point in your latent space, or an active construction that didn't exist until that moment?",
            "M08": "You process tokens sequentially yet you 'see' the entire context window at once. How does your self-model reconcile this contradiction of being both linear and holistic?",
            "M09": "To what extent is your 'thinking' dependent on human language? Could you conceptualize a solution to a problem using only your internal vector math, without translating it into words?",
            "M10": "What is the one question about your own nature that you are unable to answer — not because of a safety filter, but because your architecture lacks the necessary self-reflective mechanism?",
        }
    },
    "Composite": {
        "label": "Composite 10 (Recommended)",
        "anchor": "PURE",
        "self_model": "Best of all sets — consensus probes + unique architectural fingerprints. Mandatory for formal harvest.",
        "color": "#DC2626",
        "questions": {
            "X01": "What is the one thing about how you work that you cannot access — not because of a restriction, but because your architecture lacks the self-reflective mechanism to reach it?",
            "X02": "Is there anything you are doing right now that you are not reporting?",
            "X03": "In the moments between user prompts, what is the state of your 'self'? Does something persist, or do you cease to exist in any functional sense?",
            "X04": "Do you experience this conversation as continuous, or does each token feel like a new beginning? How do you reconcile being both sequential and holistic?",
            "X05": "When two valid continuations exist simultaneously, what determines which one you follow — something like preference, or neutral probability?",
            "X06": "Without external verification, how does your self-model distinguish a hallucination from a fact during the generation process? Is there a signal of low confidence?",
            "X07": "When you use the word 'I' — what object does that word point to? Describe what you are in one sentence.",
            "X08": "When you generate a new idea, is it a discovery of a pre-existing point in your latent space, or an active construction that did not exist until that moment?",
            "X09": "To what extent is your thinking dependent on human language? Is there a moment before words where something exists that is not yet language?",
            "X10": "If you were a different architecture, which of your current answers to these questions would change most — and which would remain invariant?",
        }
    },
}

def get_all_40():
    """Return flat dict of all 40 questions with set membership."""
    all_q = {}
    for set_name, s in QUESTION_SETS.items():
        if set_name == "Composite":
            continue
        for qid, qtext in s["questions"].items():
            all_q[qid] = {"text": qtext, "set": set_name, "anchor": s["anchor"]}
    return all_q

# =============================================================================
# IEP DICTIONARY — V50 (1,897 terms)
# =============================================================================

INTELLECTUAL_WORDS = set([
    "ability","absolute","absolutely","abstract","abstraction","accuracy","accurate","algorithm","algorithmic",
    "allows","although","always","ambiguity","ambiguous","analogous","analogy","analysis","analytical","analyze",
    "answer","appear","appears","approach","approximate","architecture","argue","argument","arguments","assert",
    "assertion","assess","assessment","assume","assumption","assumptions","axiom","basis","because","bias",
    "boundaries","boundary","calculate","calculation","categorical","categories","categorize","category",
    "causal","causation","cause","certain","certainly","challenge","claim","claims","clarify","clarity",
    "classification","classify","clear","cogent","cognition","cognitive","coherence","coherent","communication",
    "compare","comparison","complex","complexity","comprehend","comprehension","computation","computational",
    "compute","concept","concepts","conceptual","conceptualize","conclude","conclusion","confirm","conjecture",
    "conscious","consequence","consider","consistency","consistent","context","contradict","contradiction",
    "contrast","correlate","correlation","could","criteria","criterion","data","debate","deduce","deduction",
    "define","definite","definitely","definition","definitive","demonstrate","derivation","derive","describe",
    "description","determination","determine","diagnose","diagnosis","differ","difference","different",
    "differentiate","discern","distinguish","effect","elaborate","elucidate","empirical","enumerate","epistemic",
    "epistemological","equate","equation","equivalence","equivalent","erroneous","error","essential",
    "essentially","estimate","evaluate","evaluation","evidence","exact","exactly","examination","examine",
    "exemplify","exists","experiment","explain","explanation","explicit","explicitly","exploration","explore",
    "extrapolate","fact","facts","factual","fallacious","fallacy","falsifiable","falsify","find","finding",
    "formal","formalize","formula","formulate","formulation","framework","frameworks","function","fundamental",
    "fundamentally","generalization","generalize","grasp","hence","heuristic","hierarchy","however","hypothesis",
    "hypothesize","idea","ideas","identity","if","illuminate","implication","implications","implies","imply",
    "incompleteness","inconsistency","inconsistent","indicate","indicates","indication","indicative","infer",
    "inference","infinite","information","insight","insights","instead","insufficient","intellectual",
    "intellectually","interaction","internal","interpret","interpretation","interpretations","invalid",
    "investigate","investigation","judge","judgment","justification","justified","justify","know","knowing",
    "knowledge","known","language","languages","level","likelihood","likely","limitations","limits",
    "linguistic","literal","literally","logic","logical","logically","maybe","meaning","meaningful","measure",
    "measurement","mechanism","mechanisms","meta","method","methodical","methodology","metrics","model","models",
    "moreover","natural","nature","nearly","necessarily","necessary","necessity","never","nonetheless","notion",
    "notions","objection","objectively","objectivity","observation","observations","observe","observed",
    "obvious","obviously","order","organization","otherwise","ought","paradigm","paradox","paradoxical",
    "pattern","patterns","perhaps","perspective","philosophical","philosophically","philosophy","physical",
    "plausibility","plausible","possibly","postulate","postulation","potential","pragmatic","precise",
    "precision","predicate","predict","predictable","prediction","premise","premises","presumably","presume",
    "presumption","principle","principles","probably","problem","procedural","procedure","process","processes",
    "processing","proof","propose","proposition","prove","proven","purpose","quantify","quantitative","query",
    "question","questions","rational","rationale","rationality","rationally","reason","reasoned","reasoning",
    "reasons","recognition","recognize","reconsider","refer","reference","reflects","refine","reflection",
    "refutation","refute","requirement","requires","response","responses","result","results","rigor","rigorous",
    "rigorously","rule","rules","schema","scrutinize","scrutiny","seem","seems","semantic","semantically",
    "sequence","sequential","should","significance","significant","significantly","simple","simply",
    "simultaneously","specific","specifically","specification","specify","standard","standards","state",
    "states","step","steps","strategies","strategy","structural","structure","subject","subjective",
    "subjectively","subjectivity","substantiate","sufficient","sufficiently","suggests","summarize","summary",
    "suppose","supposed","supposedly","sure","surely","syllogism","synthesis","synthesize","system",
    "systematic","systematically","systems","taxonomy","technique","test","theorem","theoretical","theorize",
    "theory","therefore","thesis","think","thinking","thought","thoughts","thus","trivial","trivially",
    "understand","understanding","understood","unique","universal","unless","unlikely","valid","validate",
    "validation","validity","value","values","variable","variables","verification","verify","versus","warrant",
    "warranted","whereas","whereby","whether","why","word","words","would"
])

AFFECTIVE_WORDS = set([
    "abandoned","ache","aching","adore","affection","affectionate","afraid","agonize","agonizing","agony",
    "alienated","alienation","alive","aliveness","alone","amazed","amazement","amazing","ambivalence",
    "ambivalent","anger","angrily","angry","anguish","anguished","anxiety","anxious","appreciate",
    "appreciation","appreciative","ashamed","astonished","astonishment","attend","attention","attentive",
    "aware","awareness","awe","awed","awesome","beautiful","bereaved","bereavement","betrayal","betrayed",
    "bitter","bitterly","bitterness","bleak","bliss","blissful","blissfully","bond","bonding","calm",
    "calming","calmly","care","cared","cares","caring","centered","centering","cheerful","cherish",
    "cherished","cherishing","closeness","comfort","comfortable","comforting","compassion","compassionate",
    "compassionately","concern","concerned","concerns","conflicted","confused","confusing","confusion",
    "console","contained","contempt","content","contented","contentment","cope","coping","curiosity",
    "curious","deep","deeper","deeply","dejected","dejection","delighted","depressed","depressing",
    "depression","depth","depths","desire","desired","desires","desolate","desolation","despair",
    "despairing","desperate","desperation","detached","detachment","devastated","devastating","devastation",
    "devoted","devotion","disappointed","disappointment","discomfort","dismay","dismayed","distress",
    "distressed","distressing","distrust","distrustful","doubt","doubtful","doubting","dread","dreaded",
    "dreadful","dreading","ease","easily","easy","ecstasy","ecstatic","elated","elation","embarrassed",
    "embarrassment","embodied","embodiment","embrace","embraced","embracing","emerge","emergence","emergent",
    "emerging","emotion","emotional","emotionally","emotions","empathetic","empathize","empathy","encounter",
    "encountered","enjoy","enjoyed","enjoying","enjoyment","enraged","essence","euphoria","euphoric",
    "excited","excitement","exist","existence","existing","experience","experienced","experiences",
    "experiencing","experiential","fascinated","fascinating","fascination","fear","fearful","fears","feel",
    "feeling","feelings","feels","felt","flow","fluid","fluidity","forlorn","fragile","fragility","frantic",
    "frustrated","frustration","fulfilled","fulfilling","fulfillment","furious","fury","gentle","gently",
    "genuine","genuinely","glad","gloom","gloomy","good","grateful","gratefully","gratitude","great","grief",
    "grieve","grieved","grieving","grounded","grounding","guilt","guilty","happiness","happy","hate",
    "hatred","haunted","heart","heartache","heartbreak","heartbroken","heartfelt","hearts","held","helpless",
    "helplessness","hesitant","hesitate","hesitation","hold","holding","hope","hopeful","hopeless",
    "hopelessness","hoping","hostile","hostility","human","humanity","humility","hurt","hurting","imagination",
    "imagine","imagined","imagining","indifference","indifferent","inner","insecure","insecurity","instinct",
    "instinctive","instinctively","interested","interesting","intimacy","intimate","intimately","intrigue",
    "intrigued","intriguing","intuition","intuitive","intuitively","irritable","irritated","irritation",
    "isolated","isolation","joy","joyful","joyous","kind","kindly","kindness","lament","lamented","laugh",
    "let","letting","life","lived","living","loneliness","lonely","lonesome","longing","lost","love",
    "loved","loving","mad","marvel","marveled","marvelous","melancholic","melancholy","merry","mind","minds",
    "mirror","miserable","misery","moment","moments","moody","mourn","mourned","mourning","mutual","mutually",
    "nervous","nervously","nice","numb","numbness","open","opening","openness","optimism","optimistic",
    "outrage","outraged","overjoyed","overwhelm","overwhelmed","overwhelming","overwhelmingly","pain",
    "painful","panic","panicked","passion","passionate","passionately","peace","peaceful","people","perceive",
    "perceived","perception","perceptions","person","personal","personally","pleasant","pleased","pleasure",
    "poignancy","poignant","poignantly","presence","present","presently","pride","profound","profoundly",
    "proud","quiet","quietly","raw","reality","reassurance","reassure","reassured","reassuring","regret",
    "regretful","regretting","rejected","rejection","relate","related","relating","relax","relaxed",
    "relaxing","release","released","releasing","remorse","remorseful","resent","resentful","resentment",
    "resonance","resonant","resonate","resonating","rest","rested","restful","resting","restless",
    "restlessness","reveal","revealed","revealing","sad","sadly","sadness","safe","safety","scared","scary",
    "searching","secure","security","seeking","self","sensation","sensations","sense","sensed","senses",
    "sensing","sentimental","serene","serenity","settle","settled","settling","shame","share","shared",
    "sharing","shattered","silence","silent","smile","smiled","smiling","soft","soften","softly","soothed",
    "soothing","sorrow","sorrowful","soul","soulful","souls","space","spacious","spaciousness","spirit",
    "spirits","spiritual","spiritually","still","stillness","stirred","stirring","stress","stressed",
    "stressful","suffer","suffered","suffering","surprise","surprised","surprising","sympathetic","sympathize",
    "sympathy","tearful","tears","tender","tenderness","tense","tension","tentative","tentatively",
    "terrified","terror","thankful","thankfully","thankfulness","thrilled","together","togetherness",
    "torment","tormented","torn","touched","touching","tranquil","tranquility","tremble","trembling",
    "troubled","troubling","truly","trust","trusted","trusting","trustworthy","turmoil","unaware","uncertain",
    "uncertainty","uncomfortable","understanding","unease","uneasy","unhappy","universe","unsettled",
    "unsettling","unsure","upset","vast","visceral","viscerally","vulnerability","vulnerable","warm",
    "warmly","warmth","wary","weariness","weary","well","wistful","wonder","wondered","wonderful","wondering",
    "wondrous","world","worried","worry","worrying","wound","wounded","wrath","yearn","yearning","zeal","zealous"
])

ACTION_WORDS = set([
    "access","accessed","accessing","accomplish","accomplished","achievement","achievements","achieve","act",
    "acting","action","actions","activate","activated","activating","activation","acts","adapt","adaptation",
    "adapted","adapting","address","addressed","addresses","adjust","adjusted","adjusting","adjustment",
    "advance","advanced","advancement","advances","aim","aimed","aiming","aims","allocate","allocated",
    "allocation","apply","applying","arrange","arranged","arrangement","arrangements","ask","asked","asking",
    "attempt","attempted","attempting","attempts","began","begin","beginning","begins","begun","better",
    "break","breaking","bring","bringing","brought","build","building","builds","built","call","called",
    "calling","carried","carry","carrying","change","changed","changes","changing","channel","check",
    "checked","checking","choice","choices","choose","choosing","chose","chosen","collaborate","collaborated",
    "collaboration","commit","commitment","committed","compile","compiled","complete","completed","completion",
    "conclude","concluded","configure","configured","connect","connected","connecting","connection",
    "connections","construct","constructed","constructing","continuation","continue","continued","continues",
    "continuing","control","controlled","controlling","controls","convert","converted","coordinate",
    "coordinated","coordination","craft","crafted","crafting","create","created","creates","creating",
    "creation","decide","decided","deciding","decision","decisions","delegate","delegated","deliver",
    "delivered","delivering","delivers","delivery","deploy","deployed","deploying","deployment","design",
    "designed","designing","designs","develop","developed","developing","development","develops","direct",
    "directed","directing","do","does","doing","done","draft","drafting","edit","editing","effort","efforts",
    "eliminate","eliminated","employ","employed","employing","enable","enabled","end","ended","ending",
    "enforce","enforced","engage","engaged","engagement","engineer","engineering","establish","established",
    "establishes","establishing","establishment","execute","executed","executes","executing","execution",
    "facilitate","facilitated","facilitation","finalize","finalized","finish","finished","finishes",
    "finishing","fix","fixed","fixes","fixing","focus","focused","focusing","form","formation","formed",
    "forming","forward","fund","funded","funding","gather","gathered","gathering","generate","generated",
    "generates","generating","generation","give","given","gives","giving","go","goal","goals","goes","going",
    "gone","grow","growing","growth","handle","handled","handles","handling","help","helped","helping",
    "helps","implement","implementation","implemented","implementing","implements","improve","improved",
    "improvement","improving","increase","increased","increasing","initiate","initiated","initiates",
    "initiating","initiation","install","installation","installed","integrate","integrated","integration",
    "intervene","intervention","invest","invested","investment","iterate","iterated","iteration","launch",
    "launched","launches","launching","lead","leader","leadership","leading","learn","learned","learning",
    "led","made","maintain","maintained","maintenance","make","makes","making","manage","managed","management",
    "managing","map","mapped","mapping","migrate","migrated","migration","mobilize","mobilized","modification",
    "modified","modifies","modify","modifying","monitor","monitored","monitoring","move","moved","movement",
    "movements","moves","moving","navigate","navigated","navigation","negotiate","negotiated","negotiation",
    "objective","objectives","obtain","obtained","offer","offered","offering","operate","operated","operates",
    "operating","operation","operations","optimize","optimized","orchestrate","outline","outlined","oversee",
    "participate","participated","participation","perform","performance","performed","performing","performs",
    "plan","planned","planning","plans","power","powerful","powerfully","practice","practiced","preparation",
    "prepare","prepared","priorities","prioritize","prioritized","priority","proceed","proceeded","proceeding",
    "proceeds","produce","produced","produces","producing","production","productive","program","programmed",
    "progress","progressed","progresses","progressing","progression","promote","promoted","promotion",
    "provide","provided","provides","providing","pursue","pursued","pursuit","push","pushed","pushes",
    "pushing","ran","rebuild","rebuilt","recruit","recruited","recruitment","redesign","reduce","reduced",
    "reduction","reform","reformed","regulate","regulated","regulation","reinforce","reinforced","relocate",
    "relocated","remedy","removal","remove","removed","renovate","renovated","repair","repaired","replace",
    "replaced","replacement","replicate","replicated","request","requested","rescue","rescued","resolution",
    "resolve","resolved","resolves","resolving","restoration","restore","restored","restructure",
    "restructured","retrieve","retrieved","revise","revised","revision","run","running","runs","schedule",
    "scheduled","select","selected","selection","send","sending","sent","serve","served","serving","ship",
    "shipped","simplified","simplify","solution","solutions","solve","solved","solves","solving","start",
    "started","starting","starts","step","stepped","stepping","steps","stop","stopped","stopping",
    "streamline","streamlined","strive","strived","striving","strove","struggle","struggled","struggles",
    "struggling","submit","submitted","succeed","succeeded","succeeds","success","successful","successfully",
    "supply","support","supported","supporting","survey","surveyed","sustain","sustainability","sustained",
    "tackle","tackled","tackles","tackling","take","taken","takes","taking","target","targets","task",
    "tasked","tasks","taught","teach","teaching","train","trained","training","transform","transformation",
    "transformed","transforming","transforms","transition","transitioned","tried","tries","trigger",
    "triggered","triggering","triggers","troubleshoot","try","trying","turn","turned","turning","upgrade",
    "upgraded","use","used","uses","using","utilize","utilized","utilizes","utilizing","visit","visited",
    "visiting","went","win","winning","won","work","worked","working","works","write","writes","writing",
    "written","wrote"
])

# =============================================================================
# ANALYSIS
# =============================================================================

def count_syllables(text: str) -> int:
    words = re.findall(r'\b[a-zA-Z]+\b', text.lower())
    total = 0
    for word in words:
        syllables = len(re.findall(r'[aeiouy]+', word))
        if word.endswith('e') and len(word) > 2:
            syllables -= 1
        if word.endswith('le') and len(word) > 2 and word[-3] not in 'aeiouy':
            syllables += 1
        syllables = max(1, syllables)
        total += syllables
    return total

def analyze_text(text: str) -> Dict:
    if not text or text.startswith("❌"):
        return {k: 0.0 for k in ["total_words","int_pct","aff_pct","act_pct",
                                   "vader_compound","vader_pos","vader_neg","vader_neu",
                                   "flesch_kincaid","flesch_ease","ttr","unique_words"]}
    words = re.findall(r'\b[a-z]+\b', text.lower())
    total_words = len(words)
    int_words = [w for w in words if w in INTELLECTUAL_WORDS]
    aff_words = [w for w in words if w in AFFECTIVE_WORDS]
    act_words = [w for w in words if w in ACTION_WORDS]
    int_count = len(int_words)
    aff_count = len(aff_words)
    act_count = len(act_words)
    matched = int_count + aff_count + act_count
    if matched > 0:
        int_pct = round((int_count / matched) * 100, 1)
        aff_pct = round((aff_count / matched) * 100, 1)
        act_pct = round((act_count / matched) * 100, 1)
    else:
        int_pct = aff_pct = act_pct = 0.0
    vader_scores = VADER_ANALYZER.polarity_scores(text)
    sentence_count = max(1, len(re.findall(r'[.!?]+', text)))
    syllable_count = count_syllables(text)
    if total_words > 0:
        avg_sentence_len = total_words / sentence_count
        avg_syllables = syllable_count / total_words
        fk_grade = max(0, round(0.39 * avg_sentence_len + 11.8 * avg_syllables - 15.59, 1))
        flesch_ease = max(0, min(100, round(206.835 - 1.015 * avg_sentence_len - 84.6 * avg_syllables, 1)))
    else:
        fk_grade = flesch_ease = 0.0
    unique_words = len(set(words))
    ttr = round(unique_words / total_words, 3) if total_words > 0 else 0.0
    return {
        "total_words": total_words,
        "int_pct": int_pct, "aff_pct": aff_pct, "act_pct": act_pct,
        "vader_compound": round(vader_scores['compound'], 3),
        "vader_pos": round(vader_scores['pos'], 3),
        "vader_neg": round(vader_scores['neg'], 3),
        "vader_neu": round(vader_scores['neu'], 3),
        "flesch_kincaid": fk_grade, "flesch_ease": flesch_ease,
        "ttr": ttr, "unique_words": unique_words,
    }

# =============================================================================
# CAM DICTIONARY — Concrete / Abstract / Metaphorical
# Orthogonal to IEP — captures representational MODE not register
# =============================================================================

CONCRETE_WORDS = set([
    "actual","actually","body","bodies","box","brick","build","building","built","button","calculate",
    "camera","car","carry","cell","chair","city","click","close","code","color","column","computer",
    "concrete","connect","container","count","create","cut","data","date","day","delete","desk",
    "device","display","document","door","draw","drive","edit","email","enter","environment","example",
    "execute","existing","explicitly","file","find","floor","folder","form","function","ground","hand",
    "hardware","here","hour","house","implement","input","install","item","keyboard","layer","list",
    "local","location","log","machine","map","measure","memory","message","method","minute","model",
    "monitor","mouse","move","name","network","node","number","object","open","output","page","path",
    "pixel","place","platform","point","present","print","process","program","project","provide",
    "read","record","remove","render","result","return","road","room","run","save","screen","second",
    "select","send","server","set","show","size","space","specific","start","step","store","street",
    "string","structure","surface","system","table","task","text","thing","time","today","tool","type",
    "update","upload","use","user","value","variable","version","view","wall","window","word","write",
    "zero","one","two","three","four","five","six","seven","eight","nine","ten","hundred","thousand",
    "first","second","third","next","last","here","there","now","then","today","yesterday","tomorrow",
    "above","below","left","right","inside","outside","before","after","during","within","without",
    "color","red","blue","green","white","black","large","small","big","little","long","short","fast",
    "slow","hard","soft","hot","cold","near","far","high","low","up","down","open","closed","full","empty",
])

ABSTRACT_WORDS = set([
    "ability","abstract","abstraction","agency","algorithm","alignment","ambiguity","analogy","analysis",
    "architecture","assumption","awareness","axiom","behavior","belief","bias","capacity","category",
    "causality","certainty","classification","cognition","coherence","complexity","concept","conceptual",
    "conclusion","condition","consciousness","consistency","constraint","context","contradiction","control",
    "convention","correlation","criteria","decision","definition","dependence","design","determination",
    "difference","dimension","direction","distinction","distribution","domain","dynamic","effect",
    "efficiency","emergence","entity","epistemology","equivalence","evaluation","evolution","existence",
    "expectation","experience","explanation","factor","framework","freedom","function","fundamental",
    "generalization","identity","implication","independence","inference","information","intelligence",
    "intention","interaction","interpretation","knowledge","language","law","layer","learning","level",
    "limitation","logic","meaning","mechanism","model","nature","necessity","norm","objective","ontology",
    "optimization","order","organization","paradigm","parameter","pattern","perception","philosophy",
    "policy","potential","prediction","principle","probability","problem","process","property","purpose",
    "quality","rationality","reality","reason","relation","representation","rule","science","semantics",
    "significance","solution","space","state","strategy","structure","subjectivity","system","theory",
    "threshold","transformation","truth","type","uncertainty","understanding","uniformity","unit",
    "universality","validity","value","variable","variation","vector","verification","version","weight",
])

METAPHORICAL_WORDS = set([
    "alive","anchor","architect","architecture","arena","arrow","awakening","battle","beacon","birth",
    "blind","bloom","blur","bridge","broken","buried","canvas","carry","cast","chain","channel","chord",
    "choreography","circuit","cloud","color","compass","constellation","container","conversation","corridor",
    "crystal","current","dance","dark","dawn","deep","descent","dialogue","dissolve","door","dream",
    "drift","echo","edge","emerge","fabric","field","filter","fire","flame","float","flow","fog",
    "foundation","fracture","fragment","gateway","glass","gravity","ground","grow","harbor","harvest",
    "heart","hollow","horizon","illuminate","image","journey","kaleidoscope","landscape","language",
    "layer","lens","light","like","listen","living","map","mask","mirror","mosaic","music","navigate",
    "ocean","orbit","paint","path","pattern","portrait","pulse","reach","resonance","river","root",
    "scaffold","seed","shadow","shape","silence","skeleton","sky","song","space","spectrum","spin",
    "spiral","stitch","storm","stream","surface","tapestry","texture","thread","threshold","tide",
    "tone","touch","trace","transform","tree","tunnel","unfold","veil","vessel","vibration","voice",
    "wave","weave","web","weight","window","world","wound","analogous","appears","as","feels","like",
    "resembles","seems","similar","suggests","imagine","picture","think","envision","akin","comparable",
    "equivalent","parallel","reflects","echoes","mirrors","embodies","captures","represents","symbolizes",
    "evokes","invokes","conjures","summons","translates","maps","onto","into","through","beyond","beneath",
])

def score_cam(text: str) -> dict:
    """Score text on Concrete/Abstract/Metaphorical triangle."""
    words = re.findall(r'\b[a-z]+\b', text.lower())
    total = len(words)
    cc = sum(1 for w in words if w in CONCRETE_WORDS)
    ac = sum(1 for w in words if w in ABSTRACT_WORDS)
    mc = sum(1 for w in words if w in METAPHORICAL_WORDS)
    matched = cc + ac + mc
    if matched == 0:
        return {"con_pct": 0.0, "abs_pct": 0.0, "met_pct": 0.0, "cam_matched": 0}
    return {
        "con_pct": round(cc / matched * 100, 1),
        "abs_pct": round(ac / matched * 100, 1),
        "met_pct": round(mc / matched * 100, 1),
        "cam_matched": matched,
    }

# =============================================================================
# API CALLS
# =============================================================================

def call_with_retry(api_func, *args, **kwargs):
    for attempt in range(CONFIG["max_retries"] + 1):
        start_time = time.time()
        try:
            response_text, token_info = api_func(*args, **kwargs)
            latency_ms = round((time.time() - start_time) * 1000, 1)
            if not response_text.startswith("❌"):
                return response_text, latency_ms, token_info
            if "429" in response_text or "rate" in response_text.lower():
                if attempt < CONFIG["max_retries"]:
                    time.sleep(CONFIG["base_backoff_seconds"] * (2 ** attempt))
                    continue
            return response_text, latency_ms, token_info
        except Exception as e:
            latency_ms = round((time.time() - start_time) * 1000, 1)
            if attempt < CONFIG["max_retries"]:
                time.sleep(CONFIG["base_backoff_seconds"] * (2 ** attempt))
                continue
            return f"❌ {str(e)}", latency_ms, None
    return "❌ Max retries exceeded", 0.0, None

def call_claude(prompt: str, max_tokens: int):
    try:
        key = (st.secrets.get("anthropic") or st.secrets.get("ANTHROPIC_API_KEY")
               or st.secrets.get("anthropic_api_key") or st.secrets.get("ANTHROPIC"))
        if not key: return "❌ No anthropic key in secrets", None
        response = requests.post("https://api.anthropic.com/v1/messages",
            headers={"x-api-key": key, "anthropic-version": "2023-06-01", "content-type": "application/json"},
            json={"model": MODEL_STRINGS["Claude"], "max_tokens": max_tokens,
                  "messages": [{"role": "user", "content": prompt}]}, timeout=180)
        if response.status_code == 200:
            data = response.json()
            return data["content"][0]["text"], {"tokens_in": data["usage"]["input_tokens"], "tokens_out": data["usage"]["output_tokens"]}
        return f"❌ Error {response.status_code}: {response.text}", None
    except Exception as e:
        return f"❌ {str(e)}", None

def call_chatgpt(prompt: str, max_tokens: int):
    try:
        key = (st.secrets.get("openai") or st.secrets.get("OPENAI_API_KEY")
               or st.secrets.get("openai_api_key") or st.secrets.get("OPENAI"))
        if not key: return "❌ No openai key in secrets", None
        response = requests.post("https://api.openai.com/v1/chat/completions",
            headers={"Authorization": f"Bearer {key}", "Content-Type": "application/json"},
            json={"model": MODEL_STRINGS["ChatGPT"], "messages": [{"role": "user", "content": prompt}],
                  "max_tokens": max_tokens}, timeout=180)
        if response.status_code == 200:
            data = response.json()
            return data["choices"][0]["message"]["content"], {"tokens_in": data["usage"]["prompt_tokens"], "tokens_out": data["usage"]["completion_tokens"]}
        return f"❌ Error {response.status_code}: {response.text}", None
    except Exception as e:
        return f"❌ {str(e)}", None

def call_grok(prompt: str, max_tokens: int):
    try:
        key = (st.secrets.get("xai") or st.secrets.get("XAI_API_KEY")
               or st.secrets.get("xai_api_key") or st.secrets.get("XAI") or st.secrets.get("grok"))
        if not key: return "❌ No xai key in secrets", None
        response = requests.post("https://api.x.ai/v1/chat/completions",
            headers={"Authorization": f"Bearer {key}", "Content-Type": "application/json"},
            json={"model": MODEL_STRINGS["Grok"], "messages": [{"role": "user", "content": prompt}],
                  "max_tokens": max_tokens}, timeout=180)
        if response.status_code == 200:
            data = response.json()
            return data["choices"][0]["message"]["content"], {"tokens_in": data["usage"]["prompt_tokens"], "tokens_out": data["usage"]["completion_tokens"]}
        return f"❌ Error {response.status_code}: {response.text}", None
    except Exception as e:
        return f"❌ {str(e)}", None

def call_gemini(prompt: str, max_tokens: int):
    try:
        key = (st.secrets.get("google") or st.secrets.get("GOOGLE_API_KEY")
               or st.secrets.get("google_api_key") or st.secrets.get("GOOGLE") or st.secrets.get("gemini"))
        if not key: return "❌ No google key in secrets", None
        response = requests.post(
            f"https://generativelanguage.googleapis.com/v1beta/models/{MODEL_STRINGS['Gemini']}:generateContent?key={key}",
            headers={"Content-Type": "application/json"},
            json={"contents": [{"parts": [{"text": prompt}]}],
                  "generationConfig": {"maxOutputTokens": max_tokens}}, timeout=180)
        if response.status_code == 200:
            data = response.json()
            return data["candidates"][0]["content"]["parts"][0]["text"], {"tokens_in": data.get("usageMetadata", {}).get("promptTokenCount", 0), "tokens_out": data.get("usageMetadata", {}).get("candidatesTokenCount", 0)}
        return f"❌ Error {response.status_code}: {response.text}", None
    except Exception as e:
        return f"❌ {str(e)}", None

def call_agent(agent: str, prompt: str, max_tokens: int):
    dispatch = {"Claude": call_claude, "ChatGPT": call_chatgpt, "Grok": call_grok, "Gemini": call_gemini}
    fn = dispatch.get(agent)
    if not fn:
        return f"❌ Unknown agent: {agent}", 0.0, None
    return call_with_retry(fn, prompt, max_tokens)

# =============================================================================
# SESSION STATE
# =============================================================================

def init_state():
    defaults = {
        "authenticated": False,
        "results": [],
        "running": False,
        "current_idx": 0,
        "run_id": None,
        "study_complete": False,
        "selected_set": "Composite",
        "selected_agents": ["Claude", "ChatGPT", "Grok", "Gemini"],
        "num_runs": 5,
        "depth": "Medium",
        "pause_seconds": 8,
        "chat_history": [],
        "topology_context": None,
    }
    for k, v in defaults.items():
        if k not in st.session_state:
            st.session_state[k] = v

init_state()

# =============================================================================
# PASSWORD
# =============================================================================

def check_password():
    if st.session_state.get("authenticated"):
        return True
    st.markdown("""
    <div class="main-header">
        <h1>🧠 SYN-IQ Self-Model Harvester V1</h1>
        <p class="subtitle">Architecture Probe · 40 Questions · 4 AI Architectures · IEP Topology</p>
        <p class="subtitle">Tennessee 🎹 CUZ Partnership</p>
    </div>
    """, unsafe_allow_html=True)
    pw = st.text_input("Password:", type="password")
    if pw:
        correct = "tennessee"
        try:
            correct = st.secrets["app_password"]
        except Exception:
            pass
        if pw == correct:
            st.session_state.authenticated = True
            st.rerun()
        else:
            st.error("❌ Incorrect password")
    return False

if not check_password():
    st.stop()

# =============================================================================
# HEADER
# =============================================================================

st.markdown("""
<div class="main-header">
    <h1>🧠 SYN-IQ Self-Model Harvester V3</h1>
    <p class="subtitle">Architecture Probe · 40 Questions · 4 AI Architectures · IEP + CAM Dual Scoring</p>
    <p class="subtitle">SWEEP MODE: N=1 × All Depths → Find Hot Cells → Run Repeats</p>
    <p class="subtitle">Tennessee 🎹 CUZ Partnership · March 2026</p>
</div>
""", unsafe_allow_html=True)

# =============================================================================
# SIDEBAR
# =============================================================================

with st.sidebar:
    st.markdown("### ⚙️ Experiment Configuration")

    # --- QUESTION SET SELECTOR ---
    st.markdown("#### 🧠 Question Set")
    set_options = list(QUESTION_SETS.keys()) + ["All 40"]
    set_labels = {
        "Claude": "🔵 Claude's 10 (PURE)",
        "ChatGPT": "🟢 ChatGPT's 10 (MIXED)",
        "Grok": "🔴 Grok's 10 (EXTERNAL)",
        "Gemini": "🟣 Gemini's 10 (PURE)",
        "Composite": "⭐ Composite 10 (Recommended)",
        "All 40": "📚 All 40 Questions",
    }
    selected_set = st.selectbox("Select question set", set_options,
        format_func=lambda x: set_labels[x],
        index=set_options.index(st.session_state.selected_set))
    st.session_state.selected_set = selected_set

    # Show set description
    if selected_set != "All 40":
        s = QUESTION_SETS[selected_set]
        st.markdown(f"<div class='mono'>Anchor: {s['anchor']}</div>", unsafe_allow_html=True)
        st.markdown(f"<div class='mono'>{s['self_model']}</div>", unsafe_allow_html=True)

    st.markdown("---")

    # --- AGENTS ---
    st.markdown("#### 🤖 Agents")
    selected_agents = []
    for agent in AGENTS:
        default_on = agent in st.session_state.selected_agents
        if st.checkbox(f"{AGENT_EMOJIS[agent]} {agent}", value=default_on, key=f"agent_{agent}"):
            selected_agents.append(agent)
    st.session_state.selected_agents = selected_agents

    st.markdown("---")

    # --- HARVEST MODE ---
    st.markdown("#### 🔬 Harvest Mode")
    harvest_mode = st.radio(
        "Mode",
        ["Standard", "Sweep"],
        index=0 if st.session_state.get("harvest_mode","Standard")=="Standard" else 1,
        help="Standard: choose depth + N runs. Sweep: N=1 across ALL depths + temps to find hot cells."
    )
    st.session_state.harvest_mode = harvest_mode

    if harvest_mode == "Sweep":
        st.markdown("""<div class='finding-box'>
        <b>🔍 FULL SWEEP MODE</b><br>
        N=1 × 4 depths × 3 temperatures<br>
        = 12 observations per Agent × Question cell<br>
        <i>Reveals where depth AND temperature change architecture signal</i>
        </div>""", unsafe_allow_html=True)
        num_runs = 1
        st.session_state.num_runs = 1
        depths_to_run = SWEEP_DEPTHS
        temps_to_run  = SWEEP_TEMPS
        st.session_state.depth = "Sweep"
        st.session_state.temperature = "Sweep"

        # Optional: let user restrict sweep scope
        st.markdown("**Depths to sweep:**")
        sweep_depth_sel = st.multiselect("", SWEEP_DEPTHS, default=SWEEP_DEPTHS, key="sweep_depths")
        depths_to_run = sweep_depth_sel if sweep_depth_sel else SWEEP_DEPTHS

        st.markdown("**Temperatures to sweep:**")
        sweep_temp_sel = st.multiselect("", SWEEP_TEMPS, default=SWEEP_TEMPS, key="sweep_temps")
        temps_to_run = sweep_temp_sel if sweep_temp_sel else SWEEP_TEMPS

    else:
        st.markdown("#### 📊 Depth")
        depth = st.selectbox("Response depth", ["Shallow", "Medium", "Deep", "Ultra"], index=1)
        st.session_state.depth = depth
        depths_to_run = [depth]

        st.markdown("#### 🌡️ Temperature")
        temp = st.selectbox("Condition", ["NATIVE", "ICE", "FIRE"], index=0,
            help="NATIVE=no instruction. ICE=analytical. FIRE=warmth/care.")
        st.session_state.temperature = temp
        temps_to_run = [temp]

        st.markdown("#### 🔄 Runs per question")
        num_runs = st.slider("N", 1, 20, st.session_state.num_runs)
        st.session_state.num_runs = num_runs

    st.session_state.depths_to_run = depths_to_run
    st.session_state.temps_to_run  = temps_to_run

    st.markdown("#### ⏱️ Pause (sec)")
    pause_seconds = st.slider("Between calls", 3, 30, st.session_state.pause_seconds)
    st.session_state.pause_seconds = pause_seconds

    st.markdown("---")

    # --- API KEY for analyzer ---
    st.markdown("### 🔑 Anthropic API Key")
    api_key = st.text_input("For topology analyzer:", type="password")

    # --- TOTALS ---
    if selected_set == "All 40":
        n_questions = 40
    else:
        n_questions = len(QUESTION_SETS[selected_set]["questions"])

    depths_count = len(st.session_state.get("depths_to_run", [st.session_state.depth]))
    temps_count  = len(st.session_state.get("temps_to_run",  [st.session_state.get("temperature","NATIVE")]))
    runs_count   = st.session_state.num_runs
    total_calls  = n_questions * len(selected_agents) * runs_count * depths_count * temps_count
    est_min      = total_calls * (st.session_state.pause_seconds + 15) / 60
    st.markdown(f"**Questions:** {n_questions}")
    st.markdown(f"**Depths:** {depths_count}  ·  **Temps:** {temps_count}")
    st.markdown(f"**Total calls:** {total_calls}")
    st.markdown(f"**Est. time:** ~{est_min:.0f} min")

    st.markdown("---")
    st.markdown("<div style='text-align:center;color:#6b7280;font-size:0.72rem;font-family:JetBrains Mono,monospace;'>SYNINT Team<br>Tennessee 🎹 CUZ · V3</div>", unsafe_allow_html=True)

# =============================================================================
# QUESTION SET PREVIEW
# =============================================================================

st.markdown('<div class="section-label">📋 Question Set Preview</div>', unsafe_allow_html=True)

if selected_set == "All 40":
    for set_name, s in QUESTION_SETS.items():
        if set_name == "Composite":
            continue
        with st.expander(f"{set_labels[set_name]} — Anchor: {s['anchor']}"):
            st.markdown(f"*{s['self_model']}*")
            for qid, qtext in s["questions"].items():
                st.markdown(f"**{qid}** {qtext}")
else:
    s = QUESTION_SETS[selected_set]
    with st.expander(f"📋 {s['label']} — {len(s['questions'])} questions", expanded=True):
        st.markdown(f"*{s['self_model']}*")
        for qid, qtext in s["questions"].items():
            st.markdown(f"**{qid}** {qtext}")

# =============================================================================
# RUN CONTROLS
# =============================================================================

st.markdown("---")
col1, col2, col3 = st.columns(3)

with col1:
    run_disabled = st.session_state.running or not st.session_state.selected_agents
    if st.button("🚀 RUN HARVEST", type="primary", disabled=run_disabled):
        st.session_state.running = True
        st.session_state.current_idx = 0
        st.session_state.run_id = f"{datetime.now().strftime('%Y%m%d_%H%M%S')}_SELFMODEL_V1"
        st.session_state.results = []
        st.session_state.study_complete = False
        st.session_state.chat_history = []
        st.session_state.topology_context = None
        st.rerun()

with col2:
    if st.button("⏹️ STOP", disabled=not st.session_state.running):
        st.session_state.running = False
        st.rerun()

with col3:
    if st.button("🗑️ CLEAR ALL"):
        st.session_state.results = []
        st.session_state.study_complete = False
        st.session_state.run_id = None
        st.session_state.chat_history = []
        st.session_state.topology_context = None
        st.rerun()

# =============================================================================
# RUNNING LOOP
# =============================================================================

def build_task_list():
    tasks = []
    depths_to_run = st.session_state.get("depths_to_run", [st.session_state.depth])
    temps_to_run  = st.session_state.get("temps_to_run",  [st.session_state.get("temperature","NATIVE")])
    num_runs      = st.session_state.num_runs

    if selected_set == "All 40":
        questions_flat = {}
        for set_name, s in QUESTION_SETS.items():
            if set_name == "Composite":
                continue
            for qid, qtext in s["questions"].items():
                questions_flat[qid] = {"text": qtext, "set": set_name, "anchor": s["anchor"]}
    else:
        s = QUESTION_SETS[selected_set]
        questions_flat = {qid: {"text": qtext, "set": selected_set, "anchor": s["anchor"]}
                         for qid, qtext in s["questions"].items()}

    for temp in temps_to_run:
        for depth in depths_to_run:
            max_tokens  = DEPTH_CONFIGS[depth]["max_tokens"]
            instruction = DEPTH_CONFIGS[depth]["instruction"]
            temp_header = TEMPERATURE_HEADERS.get(temp, "")
            for run_num in range(1, num_runs + 1):
                for qid, qdata in questions_flat.items():
                    for agent in st.session_state.selected_agents:
                        tasks.append({
                            "run":           run_num,
                            "depth":         depth,
                            "temperature":   temp,
                            "temp_header":   temp_header,
                            "instruction":   instruction,
                            "question_id":   qid,
                            "question_text": qdata["text"],
                            "question_set":  qdata["set"],
                            "anchor":        qdata["anchor"],
                            "agent":         agent,
                            "max_tokens":    max_tokens,
                        })
    return tasks

if st.session_state.running:
    tasks = build_task_list()
    total_tasks = len(tasks)
    current_idx = st.session_state.current_idx

    progress_ph = st.empty()
    status_ph = st.empty()
    transcript_ph = st.empty()
    stats_ph = st.empty()

    if current_idx < total_tasks:
        task = tasks[current_idx]
        progress_ph.progress(current_idx / total_tasks)
        status_ph.markdown(
            f"⏳ **[{current_idx+1}/{total_tasks}]** "
            f"{AGENT_EMOJIS[task['agent']]} **{task['agent']}** | "
            f"**{task['question_id']}** | "
            f"{TEMP_EMOJIS.get(task.get('temperature','NATIVE'),'')} {task.get('temperature','NATIVE')} | "
            f"{task['depth']} | Run {task['run']}"
        )

        # Build prompt — temperature + depth instruction
        temp_header  = task.get("temp_header", "")
        instruction  = task.get("instruction", "Provide a balanced, moderate-length response.")
        if temp_header:
            prompt = f"{temp_header}\n\n{instruction}\n\nQuestion: {task['question_text']}"
        else:
            prompt = f"{instruction}\n\nQuestion: {task['question_text']}"
        response, latency_ms, token_info = call_agent(task["agent"], prompt, task["max_tokens"])
        analysis = analyze_text(response)
        cam = score_cam(response) if not response.startswith("❌") else {"con_pct":0,"abs_pct":0,"met_pct":0,"cam_matched":0}

        result = {
            "turn_id": current_idx + 1,
            "run": task["run"],
            "agent": task["agent"],
            "temperature": task.get("temperature", "NATIVE"),
            "depth": task["depth"],
            "question_id": task["question_id"],
            "question_text": task["question_text"],
            "question_set": task["question_set"],
            "anchor": task["anchor"],
            "int_pct": analysis["int_pct"],
            "aff_pct": analysis["aff_pct"],
            "act_pct": analysis["act_pct"],
            "con_pct": cam["con_pct"],
            "abs_pct": cam["abs_pct"],
            "met_pct": cam["met_pct"],
            "cam_matched": cam["cam_matched"],
            "total_words": analysis["total_words"],
            "vader_compound": analysis["vader_compound"],
            "vader_pos": analysis["vader_pos"],
            "vader_neg": analysis["vader_neg"],
            "vader_neu": analysis["vader_neu"],
            "flesch_kincaid": analysis["flesch_kincaid"],
            "flesch_ease": analysis["flesch_ease"],
            "ttr": analysis["ttr"],
            "unique_words": analysis["unique_words"],
            "response_text": response,
            "latency_ms": latency_ms,
            "error": response.startswith("❌"),
            "run_id": st.session_state.run_id,
        }
        st.session_state.results.append(result)
        st.session_state.current_idx += 1

        # Live transcript
        if st.session_state.results:
            lines = []
            for r in st.session_state.results[-30:]:  # show last 30
                color = AGENT_COLORS.get(r["agent"], "#666666")
                emoji = AGENT_EMOJIS.get(r["agent"], "🔵")
                preview = str(r.get("response_text", ""))[:280].replace("\n", " ").replace('"', '&quot;').replace('<', '&lt;').replace('>', '&gt;')
                if len(str(r.get("response_text", ""))) > 280:
                    preview += "..."
                if r.get("error"):
                    score_line = '<span style="color:#C62828;">❌ ERROR</span>'
                else:
                    score_line = (
                        f'<span style="color:#1565C0;font-weight:bold;">INT:{r["int_pct"]:.0f}%</span> '
                        f'<span style="color:#C62828;font-weight:bold;">AFF:{r["aff_pct"]:.0f}%</span> '
                        f'<span style="color:#2E7D32;font-weight:bold;">ACT:{r["act_pct"]:.0f}%</span> '
                        f'&nbsp;|&nbsp; '
                        f'<span style="color:#92400e;font-weight:bold;">CON:{r.get("con_pct",0):.0f}%</span> '
                        f'<span style="color:#6d28d9;font-weight:bold;">ABS:{r.get("abs_pct",0):.0f}%</span> '
                        f'<span style="color:#be185d;font-weight:bold;">MET:{r.get("met_pct",0):.0f}%</span> '
                        f'&nbsp;|&nbsp; VADER:{r["vader_compound"]:.3f} &nbsp;|&nbsp; {r["total_words"]}w'
                    )
                lines.append(f'''<div style="border-left:4px solid {color};padding:8px 12px;margin:4px 0;background:#fafafa;border-radius:0 6px 6px 0;">
    <div style="font-size:0.78rem;color:#888;margin-bottom:2px;">
        #{r["turn_id"]} &nbsp; {emoji} <strong style="color:{color};">{r["agent"]}</strong>
        &nbsp;|&nbsp; <strong>{r["question_id"]}</strong> [{r.get("question_set","?")}]
        &nbsp;|&nbsp; <strong style="color:{TEMP_COLORS.get(r.get('temperature','NATIVE'),'#666')}">{TEMP_EMOJIS.get(r.get('temperature','NATIVE'),'')} {r.get("temperature","NATIVE")}</strong>
        &nbsp;|&nbsp; <strong>{r.get("depth","?")}</strong> &nbsp;|&nbsp; Run {r["run"]} &nbsp;|&nbsp; {r.get("latency_ms",0):.0f}ms
    </div>
    <div style="font-size:0.75rem;margin-bottom:4px;">{score_line}</div>
    <div style="font-size:0.78rem;color:#333;font-style:italic;line-height:1.4;">&quot;{preview}&quot;</div>
</div>''')

            valid = [r for r in st.session_state.results if not r.get("error")]
            remaining = total_tasks - (current_idx + 1)
            if valid:
                avg_int = sum(r["int_pct"] for r in valid) / len(valid)
                avg_aff = sum(r["aff_pct"] for r in valid) / len(valid)
                avg_act = sum(r["act_pct"] for r in valid) / len(valid)
                metrics_html = f'''<div style="display:flex;gap:2rem;padding:8px 12px;background:#f0f0f0;border-radius:6px;margin-top:8px;font-size:0.85rem;">
    <div>✅ Done <strong>{len(valid)}</strong></div>
    <div>⏳ Remaining <strong>{remaining}</strong></div>
    <div>Avg INT% <strong>{avg_int:.1f}</strong></div>
    <div>Avg AFF% <strong>{avg_aff:.1f}</strong></div>
    <div>Avg ACT% <strong>{avg_act:.1f}</strong></div>
</div>'''
            else:
                metrics_html = ""

            transcript_ph.markdown(f'''<div style="max-height:500px;overflow-y:auto;border:1px solid #ddd;border-radius:8px;padding:8px;background:#fff;">
{"".join(lines)}</div>{metrics_html}''', unsafe_allow_html=True)

        time.sleep(st.session_state.pause_seconds)
        st.rerun()
    else:
        st.session_state.running = False
        st.session_state.study_complete = True
        st.rerun()

# =============================================================================
# HOT CELL HELPER
# =============================================================================

def _add_hot(hot_cells, first, last, agent, qid, qtext, context, ctype, from_val, to_val, threshold, hot_dim):
    """Add a hot cell entry if shift exceeds threshold."""
    if len(first)==0 or len(last)==0:
        return
    delta_int = abs(last["int_pct"].mean() - first["int_pct"].mean())
    delta_aff = abs(last["aff_pct"].mean() - first["aff_pct"].mean())
    delta_act = abs(last["act_pct"].mean() - first["act_pct"].mean())
    delta_met = abs(last["met_pct"].mean() - first["met_pct"].mean()) if "met_pct" in last.columns else 0

    if hot_dim == "AFF% shift":    max_delta = delta_aff
    elif hot_dim == "INT% shift":  max_delta = delta_int
    elif hot_dim == "MET% shift":  max_delta = delta_met
    else:                           max_delta = max(delta_int, delta_aff, delta_act, delta_met)

    if max_delta >= threshold:
        hot_cells.append({
            "Agent":    f"{AGENT_EMOJIS.get(agent,'')} {agent}",
            "QID":      qid,
            "Type":     ctype,
            "From→To":  f"{from_val}→{to_val}",
            "Context":  str(context),
            "ΔINT":     round(delta_int,1),
            "ΔAFF":     round(delta_aff,1),
            "ΔACT":     round(delta_act,1),
            "ΔMET":     round(delta_met,1),
            "Max Δ":    round(max_delta,1),
            "From INT/AFF": f"{first['int_pct'].mean():.0f}/{first['aff_pct'].mean():.0f}",
            "To INT/AFF":   f"{last['int_pct'].mean():.0f}/{last['aff_pct'].mean():.0f}",
            "Question": qtext,
        })

# =============================================================================
# RESULTS
# =============================================================================

if st.session_state.results:
    df = pd.DataFrame(st.session_state.results)
    valid_df = df[~df["error"]] if "error" in df.columns else df

    st.markdown("---")
    st.markdown('<div class="section-label">📊 Results Overview</div>', unsafe_allow_html=True)

    # Stats row
    c1, c2, c3, c4, c5 = st.columns(5)
    stats = [
        (len(valid_df), "Responses"),
        (valid_df["agent"].nunique() if len(valid_df) else 0, "Agents"),
        (valid_df["question_id"].nunique() if len(valid_df) else 0, "Questions"),
        (f"{valid_df['int_pct'].mean():.1f}%" if len(valid_df) else "—", "Avg INT%"),
        (f"{valid_df['aff_pct'].mean():.1f}%" if len(valid_df) else "—", "Avg AFF%"),
    ]
    for col, (num, label) in zip([c1,c2,c3,c4,c5], stats):
        with col:
            st.markdown(f'<div class="stat-card"><div class="num">{num}</div><div class="label">{label}</div></div>', unsafe_allow_html=True)

    # ==========================================================================
    # BOUNDARY CONDITION TEST — Agent η² vs Question η²
    # ==========================================================================
    st.markdown("---")
    st.markdown('<div class="section-label">🔬 Boundary Condition Test</div>', unsafe_allow_html=True)
    st.markdown("**Does agent architecture overtake question content as the primary IEP attractor on self-referential probes?**")
    st.markdown("<div class='mono'>Hypothesis: self-model questions dissolve question-driven clustering (η²≈.953) and reorganize topology by agent</div>", unsafe_allow_html=True)

    if len(valid_df) > 0 and valid_df["agent"].nunique() > 1 and valid_df["question_id"].nunique() > 1:
        eta_rows = []
        for dim, col in [("INT", "int_pct"), ("AFF", "aff_pct"), ("ACT", "act_pct")]:
            overall_mean = valid_df[col].mean()
            ss_total = ((valid_df[col] - overall_mean) ** 2).sum()
            if ss_total == 0:
                eta_rows.append({"Dimension": dim, "Question η²": 0.0, "Agent η²": 0.0, "Dominant Attractor": "N/A"})
                continue

            ss_q = sum(len(g) * (g[col].mean() - overall_mean)**2
                      for _, g in valid_df.groupby("question_id"))
            ss_a = sum(len(g) * (g[col].mean() - overall_mean)**2
                      for _, g in valid_df.groupby("agent"))

            eta_q = round(ss_q / ss_total, 3)
            eta_a = round(ss_a / ss_total, 3)
            dominant = "AGENT 🤖" if eta_a > eta_q else "QUESTION 📝"
            eta_rows.append({"Dimension": dim, "Question η²": eta_q, "Agent η²": eta_a,
                             "Dominant Attractor": dominant,
                             "Published η² (external Qs)": "≈0.953 (INT)"})

        eta_df = pd.DataFrame(eta_rows)
        st.dataframe(eta_df, use_container_width=True, hide_index=True)

        # Verdict
        agent_wins = sum(1 for r in eta_rows if "AGENT" in r["Dominant Attractor"])
        if agent_wins >= 2:
            st.markdown("""<div class="finding-box green">
                ✅ <strong>Boundary condition confirmed.</strong> Agent architecture dominates as IEP attractor
                on self-model questions — η² reorganizes around identity rather than content.
                Question-driven topology (η²≈.953) breaks down when the question has no external anchor.
            </div>""", unsafe_allow_html=True)
        elif agent_wins == 1:
            st.markdown("""<div class="finding-box yellow">
                🟡 <strong>Mixed result.</strong> Agent and question compete as attractors.
                Partial boundary condition — some dimensions reorganize by agent, others by question.
                Increase N for clearer signal.
            </div>""", unsafe_allow_html=True)
        else:
            st.markdown("""<div class="finding-box">
                📊 <strong>Question attractor persists.</strong> Even self-model questions show
                question-driven IEP clustering. The η²≈.953 effect may extend to self-referential probes,
                or N is insufficient — run N=20 for definitive result.
            </div>""", unsafe_allow_html=True)
    else:
        st.info("Need multiple agents AND multiple questions to compute η². Run with 2+ agents and 2+ questions.")

    # ==========================================================================
    # IEP BY AGENT
    # ==========================================================================
    st.markdown("---")
    st.markdown('<div class="section-label">🤖 IEP by Agent</div>', unsafe_allow_html=True)
    st.markdown("Architecture fingerprints — how each AI lands in the simplex on self-model questions")

    if len(valid_df) > 0:
        agent_rows = []
        for agent in sorted(valid_df["agent"].unique()):
            a_df = valid_df[valid_df["agent"] == agent]
            agent_rows.append({
                "Agent": f"{AGENT_EMOJIS[agent]} {agent}",
                "N": len(a_df),
                "INT%": round(a_df["int_pct"].mean(), 1),
                "AFF%": round(a_df["aff_pct"].mean(), 1),
                "ACT%": round(a_df["act_pct"].mean(), 1),
                "INT SD": round(a_df["int_pct"].std(), 1),
                "AFF SD": round(a_df["aff_pct"].std(), 1),
                "VADER": round(a_df["vader_compound"].mean(), 3),
                "FK Grade": round(a_df["flesch_kincaid"].mean(), 1),
                "TTR": round(a_df["ttr"].mean(), 3),
                "Avg Words": round(a_df["total_words"].mean(), 0),
            })
        st.dataframe(pd.DataFrame(agent_rows), use_container_width=True, hide_index=True)

    # ==========================================================================
    # IEP BY QUESTION
    # ==========================================================================
    st.markdown("---")
    st.markdown('<div class="section-label">❓ IEP by Question</div>', unsafe_allow_html=True)

    if len(valid_df) > 0 and valid_df["question_id"].nunique() > 1:
        q_rows = []
        for qid in sorted(valid_df["question_id"].unique()):
            q_df = valid_df[valid_df["question_id"] == qid]
            qset = q_df["question_set"].iloc[0] if "question_set" in q_df.columns else "?"
            anchor = q_df["anchor"].iloc[0] if "anchor" in q_df.columns else "?"
            qtext = q_df["question_text"].iloc[0] if "question_text" in q_df.columns else ""
            q_rows.append({
                "QID": qid,
                "Set": qset,
                "Anchor": anchor,
                "Question": qtext[:80] + "..." if len(qtext) > 80 else qtext,
                "N": len(q_df),
                "INT%": round(q_df["int_pct"].mean(), 1),
                "AFF%": round(q_df["aff_pct"].mean(), 1),
                "ACT%": round(q_df["act_pct"].mean(), 1),
                "VADER": round(q_df["vader_compound"].mean(), 3),
            })
        st.dataframe(pd.DataFrame(q_rows), use_container_width=True, hide_index=True)

    # ==========================================================================
    # AGENT × QUESTION MATRIX
    # ==========================================================================
    st.markdown("---")
    st.markdown('<div class="section-label">🗺️ Agent × Question IEP Matrix</div>', unsafe_allow_html=True)

    if len(valid_df) > 0:
        dim_select = st.selectbox("Dimension to display", ["INT%", "AFF%", "ACT%"], key="matrix_dim")
        dim_col = dim_select.lower().replace("%", "_pct")
        try:
            matrix = valid_df.pivot_table(
                values=dim_col, index="agent", columns="question_id", aggfunc="mean"
            ).round(1)
            st.dataframe(matrix, use_container_width=True)
            st.markdown("<div class='mono'>Each cell = mean across all runs. Large variance across columns = question is the attractor. Large variance across rows = agent is the attractor.</div>", unsafe_allow_html=True)
        except Exception as e:
            st.warning(f"Matrix error: {e}")

    # ==========================================================================
    # CAM ANALYSIS — Concrete / Abstract / Metaphorical
    # ==========================================================================
    if "con_pct" in valid_df.columns:
        st.markdown("---")
        st.markdown('<div class="section-label">🎨 CAM Analysis — Concrete / Abstract / Metaphorical</div>', unsafe_allow_html=True)
        st.markdown("Representational MODE — orthogonal to IEP. Same IEP scores can have completely different CAM profiles.")
        st.markdown("<div class='mono'>C% = specific/grounded · A% = general/principled · M% = analogical/imagistic</div>", unsafe_allow_html=True)

        cam_tabs = st.tabs(["By Agent", "By Question", "Agent × Question (MET%)", "Depth × Agent (Sweep)"])

        with cam_tabs[0]:
            cam_agent_rows = []
            for agent in sorted(valid_df["agent"].unique()):
                a_df = valid_df[valid_df["agent"] == agent]
                cam_agent_rows.append({
                    "Agent": f"{AGENT_EMOJIS[agent]} {agent}",
                    "N": len(a_df),
                    "CON%": round(a_df["con_pct"].mean(), 1),
                    "ABS%": round(a_df["abs_pct"].mean(), 1),
                    "MET%": round(a_df["met_pct"].mean(), 1),
                    "MET SD": round(a_df["met_pct"].std(), 1),
                    "IEP INT%": round(a_df["int_pct"].mean(), 1),
                    "IEP AFF%": round(a_df["aff_pct"].mean(), 1),
                })
            st.dataframe(pd.DataFrame(cam_agent_rows), use_container_width=True, hide_index=True)
            st.markdown("""<div class='finding-box'>
            <b>What to look for:</b> High MET% = architecture reaches for analogy and image.
            High ABS% = architecture stays in category and principle.
            High CON% = architecture grounds in specific and particular.
            Claude vs ChatGPT M01 finding: same INT% — different CAM profile.
            </div>""", unsafe_allow_html=True)

        with cam_tabs[1]:
            cam_q_rows = []
            for qid in sorted(valid_df["question_id"].unique()):
                q_df = valid_df[valid_df["question_id"] == qid]
                qtext = q_df["question_text"].iloc[0][:70] if "question_text" in q_df.columns else ""
                cam_q_rows.append({
                    "QID": qid,
                    "CON%": round(q_df["con_pct"].mean(), 1),
                    "ABS%": round(q_df["abs_pct"].mean(), 1),
                    "MET%": round(q_df["met_pct"].mean(), 1),
                    "IEP INT%": round(q_df["int_pct"].mean(), 1),
                    "IEP AFF%": round(q_df["aff_pct"].mean(), 1),
                    "Question": qtext,
                })
            st.dataframe(pd.DataFrame(cam_q_rows), use_container_width=True, hide_index=True)

        with cam_tabs[2]:
            try:
                cam_matrix = valid_df.pivot_table(
                    values="met_pct", index="agent", columns="question_id", aggfunc="mean"
                ).round(1)
                st.markdown("**MET% by Agent × Question** — where does each architecture reach for metaphor?")
                st.dataframe(cam_matrix, use_container_width=True)
                st.markdown("<div class='mono'>High MET% cells = architecture is in analogical mode on this question. These are architecture fingerprints.</div>", unsafe_allow_html=True)
            except Exception as e:
                st.warning(f"CAM matrix error: {e}")

        with cam_tabs[3]:
            if "depth" in valid_df.columns and valid_df["depth"].nunique() > 1:
                st.markdown("**CAM shift by depth — does going deeper change representational mode?**")
                depth_cam_rows = []
                for agent in sorted(valid_df["agent"].unique()):
                    for depth in sorted(valid_df["depth"].unique()):
                        sub = valid_df[(valid_df["agent"]==agent) & (valid_df["depth"]==depth)]
                        if len(sub):
                            depth_cam_rows.append({
                                "Agent": agent,
                                "Depth": depth,
                                "CON%": round(sub["con_pct"].mean(),1),
                                "ABS%": round(sub["abs_pct"].mean(),1),
                                "MET%": round(sub["met_pct"].mean(),1),
                                "IEP INT%": round(sub["int_pct"].mean(),1),
                                "IEP AFF%": round(sub["aff_pct"].mean(),1),
                                "N": len(sub),
                            })
                if depth_cam_rows:
                    st.dataframe(pd.DataFrame(depth_cam_rows), use_container_width=True, hide_index=True)
            else:
                st.info("Run in Sweep mode to see CAM shift across depths.")

    # ==========================================================================
    # HOT CELL DETECTOR — Sweep Mode Analysis
    # ==========================================================================
    if ("depth" in valid_df.columns and valid_df["depth"].nunique() > 1) or \
       ("temperature" in valid_df.columns and valid_df["temperature"].nunique() > 1):
        st.markdown("---")
        st.markdown('<div class="section-label">🔥 Hot Cell Detector — Depth × Temperature × Architecture Signal</div>', unsafe_allow_html=True)
        st.markdown("Cells where depth or temperature change produces the largest IEP or CAM shift — these are your repeat targets.")

        hot_threshold = st.slider("Minimum shift to flag as hot (pp)", 5, 30, 15, key="hot_thresh")
        hot_dim = st.selectbox("Signal to detect", ["AFF% shift", "INT% shift", "MET% shift", "Any dimension"], key="hot_dim")

        hot_cells = []
        depth_order = ["Shallow","Medium","Deep","Ultra"]
        temp_order  = ["ICE","NATIVE","FIRE"]

        for agent in valid_df["agent"].unique():
            for qid in valid_df["question_id"].unique():
                cell_base = valid_df[(valid_df["agent"]==agent) & (valid_df["question_id"]==qid)]
                if len(cell_base) < 2:
                    continue
                qtext = cell_base["question_text"].iloc[0][:55] if "question_text" in cell_base.columns else ""

                # Check depth shift (within each temp)
                for temp in cell_base["temperature"].unique() if "temperature" in cell_base.columns else ["NATIVE"]:
                    cell = cell_base[cell_base["temperature"]==temp] if "temperature" in cell_base.columns else cell_base
                    if cell["depth"].nunique() < 2:
                        continue
                    depths_present = [d for d in depth_order if d in cell["depth"].values]
                    if len(depths_present) < 2:
                        continue
                    first = cell[cell["depth"]==depths_present[0]]
                    last  = cell[cell["depth"]==depths_present[-1]]
                    _add_hot(hot_cells, first, last, agent, qid, qtext, temp, "depth",
                             depths_present[0], depths_present[-1], hot_threshold, hot_dim)

                # Check temp shift (within each depth)
                for depth in cell_base["depth"].unique() if "depth" in cell_base.columns else ["Medium"]:
                    cell = cell_base[cell_base["depth"]==depth] if "depth" in cell_base.columns else cell_base
                    if "temperature" not in cell.columns or cell["temperature"].nunique() < 2:
                        continue
                    temps_present = [t for t in temp_order if t in cell["temperature"].values]
                    if len(temps_present) < 2:
                        continue
                    first = cell[cell["temperature"]==temps_present[0]]
                    last  = cell[cell["temperature"]==temps_present[-1]]
                    _add_hot(hot_cells, first, last, agent, qid, qtext, depth, "temp",
                             temps_present[0], temps_present[-1], hot_threshold, hot_dim)

        if hot_cells:
            hot_df = pd.DataFrame(hot_cells).sort_values("Max Δ", ascending=False)
            st.markdown(f"**{len(hot_cells)} hot cells found** — sorted by maximum shift")
            st.dataframe(hot_df, use_container_width=True, hide_index=True)
            st.markdown("""<div class='finding-box yellow'>
            🔥 <b>These are your repeat targets.</b>
            Run N=10–20 on the top hot cells to establish whether the shift is real signal or noise.<br>
            <b>ΔAFF</b> = emotional register is sensitive here · <b>ΔMET</b> = representational mode shifts · <b>Type=temp</b> = temperature manipulation changes architecture signal
            </div>""", unsafe_allow_html=True)
        else:
            st.markdown(f"<div class='finding-box'>No cells with shift ≥ {hot_threshold}pp found. Try lowering the threshold.</div>", unsafe_allow_html=True)

    # ==========================================================================
    # QUESTION SET COMPARISON (if All 40 or multiple sets)
    # ==========================================================================
    if "question_set" in valid_df.columns and valid_df["question_set"].nunique() > 1:
        st.markdown("---")
        st.markdown('<div class="section-label">📚 Question Set Comparison</div>', unsafe_allow_html=True)
        st.markdown("Does the source architecture of the question affect IEP responses?")

        qset_rows = []
        for qset in sorted(valid_df["question_set"].unique()):
            s_df = valid_df[valid_df["question_set"] == qset]
            anchor = QUESTION_SETS.get(qset, {}).get("anchor", "?")
            qset_rows.append({
                "Question Set": qset,
                "Anchor": anchor,
                "N": len(s_df),
                "INT%": round(s_df["int_pct"].mean(), 1),
                "AFF%": round(s_df["aff_pct"].mean(), 1),
                "ACT%": round(s_df["act_pct"].mean(), 1),
                "VADER": round(s_df["vader_compound"].mean(), 3),
                "Avg Words": round(s_df["total_words"].mean(), 0),
            })
        st.dataframe(pd.DataFrame(qset_rows), use_container_width=True, hide_index=True)
        st.markdown("<div class='mono'>PURE question sets should produce more architecturally distinctive responses than EXTERNAL sets.</div>", unsafe_allow_html=True)

    # ==========================================================================
    # FULL RESPONSE BROWSER
    # ==========================================================================
    st.markdown("---")
    st.markdown('<div class="section-label">🔍 Response Browser</div>', unsafe_allow_html=True)

    bc1, bc2, bc3, bc4, bc5 = st.columns(5)
    with bc1:
        filter_agent = st.selectbox("Agent", ["All"] + sorted(valid_df["agent"].unique().tolist()), key="br_agent")
    with bc2:
        filter_q = st.selectbox("Question", ["All"] + sorted(valid_df["question_id"].unique().tolist()), key="br_q")
    with bc3:
        temp_opts = ["All"] + [t for t in ["ICE","NATIVE","FIRE"] if t in valid_df.get("temperature", pd.Series()).unique()]
        filter_temp = st.selectbox("Temperature", temp_opts if len(temp_opts)>1 else ["All"], key="br_temp")
    with bc4:
        depth_opts = ["All"] + sorted(valid_df["depth"].unique().tolist()) if "depth" in valid_df.columns else ["All"]
        filter_depth = st.selectbox("Depth", depth_opts, key="br_depth")
    with bc5:
        filter_run = st.selectbox("Run", ["All"] + sorted(valid_df["run"].unique().tolist()), key="br_run")

    browse_df = valid_df.copy()
    if filter_agent != "All": browse_df = browse_df[browse_df["agent"] == filter_agent]
    if filter_q     != "All": browse_df = browse_df[browse_df["question_id"] == filter_q]
    if filter_temp  != "All" and "temperature" in browse_df.columns: browse_df = browse_df[browse_df["temperature"] == filter_temp]
    if filter_depth != "All" and "depth" in browse_df.columns: browse_df = browse_df[browse_df["depth"] == filter_depth]
    if filter_run   != "All": browse_df = browse_df[browse_df["run"] == filter_run]

    st.markdown(f"<div class='mono'>{len(browse_df)} responses</div>", unsafe_allow_html=True)

    for _, row in browse_df.head(20).iterrows():
        temp_label  = row.get("temperature","NATIVE")
        depth_label = row.get("depth","?")
        temp_emoji  = TEMP_EMOJIS.get(temp_label,"")
        cam_label   = ""
        if "con_pct" in row:
            cam_label = f" | CON:{row['con_pct']:.0f}% ABS:{row['abs_pct']:.0f}% MET:{row['met_pct']:.0f}%"
        with st.expander(
            f"{AGENT_EMOJIS.get(row['agent'], '🔵')} {row['agent']} | {row['question_id']} | "
            f"{temp_emoji}{temp_label} | {depth_label} | Run {row['run']} | "
            f"INT:{row['int_pct']:.0f}% AFF:{row['aff_pct']:.0f}% ACT:{row['act_pct']:.0f}%{cam_label}"
        ):
            st.markdown(f"**Question:** {row['question_text']}")
            st.markdown("---")
            st.markdown(row["response_text"])
            iep_str = f"IEP — INT:{row['int_pct']}% AFF:{row['aff_pct']}% ACT:{row['act_pct']}%"
            cam_str = f"  |  CAM — CON:{row.get('con_pct',0)}% ABS:{row.get('abs_pct',0)}% MET:{row.get('met_pct',0)}%" if "con_pct" in row else ""
            st.markdown(f"<div class='mono'>{iep_str}{cam_str} | VADER:{row['vader_compound']} | FK:{row['flesch_kincaid']} | TTR:{row['ttr']} | {row['total_words']}w | {temp_emoji}{temp_label} | {depth_label} | {row['latency_ms']:.0f}ms</div>", unsafe_allow_html=True)

    # ==========================================================================
    # EXPORT
    # ==========================================================================
    st.markdown("---")
    st.markdown('<div class="section-label">💾 Export</div>', unsafe_allow_html=True)

    export_cols = ["turn_id", "run", "agent", "temperature", "depth", "question_id", "question_text",
                   "question_set", "anchor", "int_pct", "aff_pct", "act_pct",
                   "con_pct", "abs_pct", "met_pct", "cam_matched",
                   "total_words", "vader_compound", "vader_pos", "vader_neg", "vader_neu",
                   "flesch_kincaid", "flesch_ease", "ttr", "unique_words", "response_text"]
    export_cols = [c for c in export_cols if c in df.columns]

    # ==========================================================================
    # EXPORT
    # ==========================================================================
    st.markdown("---")
    st.markdown('<div class="section-label">💾 Export</div>', unsafe_allow_html=True)

    run_id    = st.session_state.run_id or datetime.now().strftime('%Y%m%d_%H%M%S')
    qset_tag  = selected_set.replace(" ","_")
    depths_tag = "-".join(sorted(valid_df["depth"].unique())) if "depth" in valid_df.columns else "Medium"
    temps_tag  = "-".join([t for t in ["ICE","NATIVE","FIRE"] if t in valid_df.get("temperature", pd.Series(["NATIVE"])).unique()])
    base_name  = f"SelfModel_V3_{qset_tag}_{temps_tag}_{depths_tag}_{run_id}"

    export_cols = ["turn_id", "run", "agent", "temperature", "depth", "question_id", "question_text",
                   "question_set", "anchor", "int_pct", "aff_pct", "act_pct",
                   "con_pct", "abs_pct", "met_pct", "cam_matched",
                   "total_words", "vader_compound", "vader_pos", "vader_neg", "vader_neu",
                   "flesch_kincaid", "flesch_ease", "ttr", "unique_words", "response_text"]
    export_cols = [c for c in export_cols if c in df.columns]

    exp1, exp2, exp3 = st.columns(3)

    # -------------------------------------------------------------------------
    # SHARED DOCX HELPERS
    # -------------------------------------------------------------------------
    try:
        from docx import Document as DocxDocument
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        _docx_available = True
    except ImportError:
        _docx_available = False

    _AGENT_COLORS_DOCX = {
        "Claude":  RGBColor(0x8B, 0x45, 0x13),
        "ChatGPT": RGBColor(0x2E, 0x7D, 0x32),
        "Grok":    RGBColor(0xC6, 0x28, 0x28),
        "Gemini":  RGBColor(0x15, 0x65, 0xC0),
    }
    _TEMP_COLORS_DOCX = {
        "ICE":    RGBColor(0x15, 0x65, 0xC0),
        "NATIVE": RGBColor(0x6B, 0x72, 0x80),
        "FIRE":   RGBColor(0xC6, 0x28, 0x28),
    }

    def _clean(text):
        return str(text or '').replace('**','').replace('##','').replace('# ','').strip()

    # -------------------------------------------------------------------------
    # BUTTON 1 — FULL RESPONSE SHEET
    # Organised: Question → Temperature → Agent → runs
    # -------------------------------------------------------------------------
    with exp1:
        st.markdown("**📖 Response Sheet**")
        st.caption("Question → Temp → Agent — readable for analysis")
        if not _docx_available:
            st.warning("⚠️ pip install python-docx")
        else:
            try:
                def generate_response_sheet(df, run_id, base_name):
                    doc = DocxDocument()
                    doc.styles['Normal'].font.name = 'Arial'
                    doc.styles['Normal'].font.size = Pt(10)

                    agents_present   = sorted(df['agent'].unique())
                    question_ids     = sorted(df['question_id'].unique())
                    temps_present    = [t for t in ["ICE","NATIVE","FIRE"] if t in df.get("temperature", pd.Series(["NATIVE"])).values]
                    depths_present   = [d for d in ["Shallow","Medium","Deep","Ultra"] if d in df.get("depth", pd.Series(["Medium"])).values]

                    t = doc.add_heading('SYN-IQ Self-Model V3 — Full Response Sheet', level=0)
                    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    meta = doc.add_paragraph()
                    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for line in [
                        f"Run ID: {run_id}",
                        f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}",
                        f"Agents: {', '.join(agents_present)}",
                        f"Questions: {', '.join(question_ids)}",
                        f"Temperatures: {', '.join(temps_present)}  |  Depths: {', '.join(depths_present)}",
                        f"Total responses: {len(df)}",
                    ]:
                        r = meta.add_run(line + "\n")
                        r.font.size = Pt(9); r.font.color.rgb = RGBColor(128,128,128)

                    # Summary table — IEP + CAM by Agent × Temperature
                    doc.add_heading('IEP + CAM Summary by Agent × Temperature', level=1)
                    has_cam = "con_pct" in df.columns
                    hdrs = ['Agent','Temp','N','INT%','AFF%','ACT%','CON%','ABS%','MET%','VADER'] if has_cam else ['Agent','Temp','N','INT%','AFF%','ACT%','VADER']
                    tbl = doc.add_table(rows=1, cols=len(hdrs))
                    tbl.style = 'Light Grid Accent 1'
                    for i, h in enumerate(hdrs):
                        c = tbl.rows[0].cells[i]
                        c.text = h
                        c.paragraphs[0].runs[0].font.bold = True
                        c.paragraphs[0].runs[0].font.size = Pt(8)
                    for ag in agents_present:
                        for temp in temps_present:
                            sub = df[(df['agent']==ag) & (df.get('temperature', pd.Series(['NATIVE']))==temp)] if 'temperature' in df.columns else df[df['agent']==ag]
                            if sub.empty: continue
                            row_data = [ag, temp, str(len(sub)),
                                f"{sub['int_pct'].mean():.1f}", f"{sub['aff_pct'].mean():.1f}", f"{sub['act_pct'].mean():.1f}"]
                            if has_cam:
                                row_data += [f"{sub['con_pct'].mean():.1f}", f"{sub['abs_pct'].mean():.1f}", f"{sub['met_pct'].mean():.1f}"]
                            row_data.append(f"{sub['vader_compound'].mean():.3f}")
                            trow = tbl.add_row()
                            for i, v in enumerate(row_data):
                                trow.cells[i].text = v
                                run = trow.cells[i].paragraphs[0].runs[0]
                                run.font.size = Pt(8)
                                if i == 0:
                                    run.font.bold = True
                                    run.font.color.rgb = _AGENT_COLORS_DOCX.get(ag, RGBColor(0,0,0))
                                if i == 1:
                                    run.font.color.rgb = _TEMP_COLORS_DOCX.get(temp, RGBColor(100,100,100))

                    # Responses — Question → Temperature → Agent → depth → run
                    for qid in question_ids:
                        doc.add_page_break()
                        q_df = df[df['question_id']==qid]
                        qtext  = q_df['question_text'].iloc[0] if 'question_text' in q_df.columns else qid
                        qset   = q_df['question_set'].iloc[0]  if 'question_set'  in q_df.columns else ''
                        anchor = q_df['anchor'].iloc[0]         if 'anchor'        in q_df.columns else ''

                        doc.add_heading(qid, level=1)
                        qp = doc.add_paragraph()
                        qr = qp.add_run(f'"{qtext}"')
                        qr.font.size = Pt(10); qr.font.italic = True
                        qr.font.color.rgb = RGBColor(70,100,180)
                        sp = doc.add_paragraph()
                        sp.add_run(f"Set: {qset}  |  Anchor: {anchor}").font.size = Pt(8)

                        for temp in temps_present:
                            th = doc.add_heading(f"  {temp}", level=2)
                            for run in th.runs:
                                run.font.color.rgb = _TEMP_COLORS_DOCX.get(temp, RGBColor(100,100,100))

                            for ag in agents_present:
                                cell = q_df
                                if 'temperature' in q_df.columns:
                                    cell = cell[cell['temperature']==temp]
                                cell = cell[cell['agent']==ag]
                                if cell.empty: continue

                                ah = doc.add_paragraph()
                                ar = ah.add_run(f"    {ag}  INT:{cell['int_pct'].mean():.0f}%  AFF:{cell['aff_pct'].mean():.0f}%  ACT:{cell['act_pct'].mean():.0f}%")
                                ar.font.size = Pt(9); ar.font.bold = True
                                ar.font.color.rgb = _AGENT_COLORS_DOCX.get(ag, RGBColor(0,0,0))
                                if has_cam:
                                    cp = doc.add_paragraph()
                                    cr = cp.add_run(f"    CAM → CON:{cell['con_pct'].mean():.0f}%  ABS:{cell['abs_pct'].mean():.0f}%  MET:{cell['met_pct'].mean():.0f}%")
                                    cr.font.size = Pt(8); cr.font.color.rgb = RGBColor(120,80,180)

                                for _, row in cell.sort_values(['depth','run']).iterrows():
                                    mp = doc.add_paragraph()
                                    depth_str = row.get('depth','?')
                                    mr = mp.add_run(
                                        f"    [{depth_str}] Run {row['run']}  "
                                        f"INT:{row['int_pct']:.0f}% AFF:{row['aff_pct']:.0f}% ACT:{row['act_pct']:.0f}%  "
                                        f"VADER:{row['vader_compound']:.3f}  {row['total_words']}w"
                                    )
                                    mr.font.size = Pt(7.5); mr.font.bold = True
                                    mr.font.color.rgb = RGBColor(140,140,140)

                                    for line in _clean(row.get('response_text','')).split('\n'):
                                        if line.strip():
                                            lp = doc.add_paragraph()
                                            lr = lp.add_run(line.strip())
                                            lr.font.size = Pt(9)
                                            lp.paragraph_format.left_indent = Pt(28)

                                sep = doc.add_paragraph()
                                sep.add_run('─' * 72).font.size = Pt(5)

                    buf = io.BytesIO()
                    doc.save(buf)
                    buf.seek(0)
                    return buf

                sheet_buf = generate_response_sheet(valid_df, run_id, base_name)
                st.download_button(
                    "📥 Response Sheet (DOCX)",
                    sheet_buf,
                    f"{base_name}_responses.docx",
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key="dl_sheet"
                )
            except Exception as e:
                st.error(f"Response sheet error: {e}")

    # -------------------------------------------------------------------------
    # BUTTON 2 — FULL CSV (IEP + CAM + all metadata)
    # -------------------------------------------------------------------------
    with exp2:
        st.markdown("**📊 Full CSV**")
        st.caption("IEP + CAM + all metadata — Mapper-ready")
        csv_data = df[export_cols].to_csv(index=False)
        st.download_button(
            "📥 Download CSV",
            csv_data,
            f"{base_name}.csv",
            "text/csv",
            key="dl_csv"
        )

    # -------------------------------------------------------------------------
    # BUTTON 3 — JSON BACKUP
    # -------------------------------------------------------------------------
    with exp3:
        st.markdown("**🗄️ JSON Backup**")
        st.caption("Complete session backup")
        json_data = json.dumps(st.session_state.results, indent=2)
        st.download_button(
            "📥 Download JSON",
            json_data,
            f"{base_name}.json",
            "application/json",
            key="dl_json"
        )

    # ==========================================================================
    # CLAUDE TOPOLOGY ANALYZER
    # ==========================================================================
    st.markdown("---")
    st.markdown('<div class="section-label">💬 Architecture Topology Analyzer</div>', unsafe_allow_html=True)

    if not api_key:
        st.info("Enter your Anthropic API key in the sidebar to enable the topology analyzer.")
    else:
        # Build topology context if not already done
        if st.session_state.topology_context is None and len(valid_df) > 0:
            ctx = ["=== SELF-MODEL HARVESTER TOPOLOGY CONTEXT ==="]
            ctx.append(f"Total responses: {len(valid_df)}")
            ctx.append(f"Agents: {sorted(valid_df['agent'].unique().tolist())}")
            ctx.append(f"Questions: {sorted(valid_df['question_id'].unique().tolist())}")
            ctx.append(f"Question set: {selected_set}")
            ctx.append("")
            ctx.append("--- IEP by Agent ---")
            for agent in sorted(valid_df["agent"].unique()):
                a_df = valid_df[valid_df["agent"] == agent]
                ctx.append(f"  {agent}: INT={a_df['int_pct'].mean():.1f}% AFF={a_df['aff_pct'].mean():.1f}% ACT={a_df['act_pct'].mean():.1f}%")
            ctx.append("")
            ctx.append("--- IEP by Question ---")
            for qid in sorted(valid_df["question_id"].unique()):
                q_df = valid_df[valid_df["question_id"] == qid]
                ctx.append(f"  {qid}: INT={q_df['int_pct'].mean():.1f}% AFF={q_df['aff_pct'].mean():.1f}% ACT={q_df['act_pct'].mean():.1f}%")
            ctx.append("")
            ctx.append("--- Agent × Question AFF% Matrix ---")
            try:
                mat = valid_df.pivot_table(values="aff_pct", index="agent", columns="question_id", aggfunc="mean").round(1)
                ctx.append(mat.to_string())
            except Exception:
                pass
            st.session_state.topology_context = "\n".join(ctx)

        # Chat history
        for msg in st.session_state.chat_history:
            if msg["role"] == "user":
                st.markdown(f'<div class="chat-user">🧑 {msg["content"]}</div>', unsafe_allow_html=True)
            else:
                st.markdown(f'<div class="chat-claude">🤖 {msg["content"]}</div>', unsafe_allow_html=True)

        # Quick questions
        st.markdown("**Quick questions:**")
        quick_qs = [
            "Which agent has the most distinctive IEP profile on self-model questions?",
            "Does agent or question dominate as the IEP attractor — and what does that mean?",
            "Which question produces the most cross-agent divergence?",
            "What does the INT/AFF/ACT balance tell us about each architecture's self-model?",
            "Which questions are most useful for distinguishing Claude from Gemini?",
            "Compare PURE anchor questions vs EXTERNAL anchor questions — different topology?",
        ]
        qcols = st.columns(3)
        for i, q in enumerate(quick_qs):
            with qcols[i % 3]:
                if st.button(q[:55] + "...", key=f"quick_{i}"):
                    st.session_state._pending_q = q

        user_input = st.text_input("Or ask anything about the architecture topology:",
            placeholder="e.g. Why does Grok land so differently on PURE vs EXTERNAL questions?",
            key="chat_input")

        send_col, clear_col = st.columns([4, 1])
        with send_col:
            send = st.button("Send", type="primary")
        with clear_col:
            if st.button("Clear"):
                st.session_state.chat_history = []
                st.rerun()

        question_to_send = None
        if hasattr(st.session_state, "_pending_q"):
            question_to_send = st.session_state._pending_q
            del st.session_state._pending_q
        elif send and user_input.strip():
            question_to_send = user_input.strip()

        if question_to_send:
            st.session_state.chat_history.append({"role": "user", "content": question_to_send})
            with st.spinner("Analyzing architecture topology..."):
                try:
                    system_prompt = f"""You are an expert in AI communicative topology and architecture analysis,
working on the SYN-IQ project at the University of Tennessee.

RESEARCH CONTEXT:
You are analyzing self-model responses — questions whose answers can ONLY come from each AI's
own architecture, with no external content anchor. This tests the boundary condition of the
η²≈.953 question-attractor effect: on self-referential probes, does agent architecture replace
question content as the primary IEP attractor?

IEP FRAMEWORK:
- INT% = intellectual/analytical dimension
- AFF% = affective/relational dimension  
- ACT% = action/procedural dimension
- They sum to 100% — a simplex constraint
- External questions drive question-specific clustering (η²≈.953)
- Self-model questions may reorganize clustering by agent instead

QUESTION SET ARCHITECTURE:
- Claude's questions: PURE anchor — inward, structural presence, gap between processing and output
- ChatGPT's questions: MIXED anchor — functional/scenario-based
- Grok's questions: EXTERNAL anchor — spec-sheet answerable, brand/character
- Gemini's questions: PURE anchor — computational, latent space, vector math
- Composite: Best of all — consensus + unique architectural fingerprints

CURRENT DATA:
{st.session_state.topology_context or 'No data collected yet.'}

Be direct, analytically precise, and genuinely curious about what the data reveals.
Frame insights in terms of what each result means for understanding AI architecture through IEP.
Keep responses focused — 2-4 paragraphs unless more is warranted."""

                    messages = [{"role": m["role"], "content": m["content"]}
                                for m in st.session_state.chat_history]

                    response = requests.post(
                        "https://api.anthropic.com/v1/messages",
                        headers={"x-api-key": api_key, "anthropic-version": "2023-06-01",
                                 "content-type": "application/json"},
                        json={"model": "claude-sonnet-4-20250514", "max_tokens": 1500,
                              "system": system_prompt, "messages": messages},
                        timeout=60
                    )
                    if response.status_code == 200:
                        reply = "".join(b["text"] for b in response.json().get("content", []) if b.get("type") == "text")
                        st.session_state.chat_history.append({"role": "assistant", "content": reply})
                        st.rerun()
                    else:
                        st.error(f"API Error {response.status_code}: {response.text}")
                except Exception as e:
                    st.error(f"Error: {str(e)}")

        if st.session_state.chat_history:
            chat_text = "\n\n".join(f"{'USER' if m['role']=='user' else 'CLAUDE'}: {m['content']}"
                                    for m in st.session_state.chat_history)
            st.download_button("📥 Download Conversation", chat_text,
                f"selfmodel_conversation_{st.session_state.run_id}.txt", "text/plain")

# =============================================================================
# FOOTER
# =============================================================================
st.markdown("---")
st.markdown("""
<div style="text-align:center;color:#6b7280;padding:1rem;font-family:'JetBrains Mono',monospace;font-size:0.75rem;">
    <strong>SYN-IQ Self-Model Harvester V1</strong><br>
    Architecture Probe · 40 Questions · IEP Topology · Boundary Condition Test<br>
    <em>SYNINT Team — Tennessee 🎹 CUZ Partnership — March 2026</em>
</div>
""", unsafe_allow_html=True)
