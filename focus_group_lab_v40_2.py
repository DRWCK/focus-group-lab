"""
Focus Group Lab V40.2 — Research Edition
Multi-Agent AI Advisory Platform + Live IEP/Vt Scoring + Co-Conductor

V40.2 CHANGES (from V40.1):
=== HYGIENE ===
- Removed residual 'polarity' tags from PRESETS dict (P1..P5 entries).
  Polarity was already removed as a functional control in V40, but each
  preset still carried a "polarity": "ANALYTIC"|"BRIDGE"|"CREATIVE" string
  as informational metadata. That was confusing to read — it looked like
  an active field. Now fully gone.
- Updated companion comments in init_session_state and the preset loader
  to reflect the cleaner PRESETS shape.
- No behavioral or scoring changes. CSVs and version stamps are unchanged
  except tool_version → "V40.2".

V40.1 CHANGES (from V40):
=== BUG FIXES ===
- Auto Run "Response log" KeyError: fixed row['dominant'] → row['iep_dominant']
  (V40 row dict stores under 'iep_dominant'; log viewer looked for 'dominant')
- Version badge consistency: password screen, main header, and markdown export
  now all say V40.1 and use the .v40-badge CSS class (V40 had V38 leftovers)
- experiment_run_id suffix bumped to "_V40_1"
- tool_version in V40_VERSION_STAMPS bumped to "V40.1"

=== HYGIENE ===
- Dead code removed: auto_depth fallback (key was never set; now reads
  st.session_state.depth directly, matching actual behavior)
- PDF parsing: swapped regex-based BT/ET extraction for pypdf.PdfReader
  (handles compressed streams, encoded text, and multi-column layouts that
  the regex approach failed silently on)
- Dictionary-size guard assertions added (616/599/682) — fires at import
  time if someone edits word sets without updating changelog/stamps

V40 CHANGES (from V38):
=== SCIENTIFIC CONFORMANCE TO V50 (the published-paper instrument-of-record) ===
1. TEMPERATURE prompts replaced with V50-exact text (verbatim, 18 conditions)
   - COLD: restored full ending "Focus on data, facts, and logical relationships."
   - HOT: restored V50 text ("warmth and emotional attunement...")
   - FIRE: restored V50 canonical FIRE ("deepest nurturing care... Comfort above all.")
   - Removed FIRE_A and FIRE_I (V38-only experimental variants — retained in code comments
     for future V41+ exploration if desired, but NOT canonical FIRE)
2. IEP dictionaries replaced with V50 1,897-term set (was 1,631 in V38)
   - INT now 616 terms (V38 had 610; +6: circumscribe, construe, construed,
     express, expressing, expression)
   - AFF now 599 terms (V38 had 595; +4: notice, noticed, noticing, understanding)
   - ACT now 682 terms (V38 had 426; +256 — major restoration)
3. Subclass naming: KEEPS '*_phenomenological' (not renamed to V50's '*_emergent')
   — 'emergent' carries consciousness-emergence connotations that this work
   is explicitly NOT claiming. 'phenomenological' is the V40 canonical name.
   V40 CSVs are therefore NOT column-identical to V50 on subclass columns;
   pooling requires mapping aff_sub_phenomenological ↔ aff_sub_emergent etc.
   Version-stamping columns (subclass_taxonomy_version) make this unambiguous.
4. Agent 'Sophia' renamed to 'ChatGPT' throughout (live UI, Auto Run default,
   CSS classes, CSV output — matches V50 and published papers)
5. DEPTH control: replaced V38's numeric 1-5 with V50's four-checkbox system
   (Shallow/Medium/Deep/Ultra-Deep) with V50 token budgets (200/500/1000/2000)
   and V50 instruction strings
6. Auto Run CSV schema: full V50 schema in V50 column order, including
   vader_compound, vader_pos, vader_neg, vader_neu, flesch_kincaid, flesch_ease,
   ttr, unique_words, lens_value, lens_setting, embedding (as "[]")
7. run_id written once per Auto Run experiment (V38 regenerated per row)

=== REMOVED FROM V38 ===
8. Polarity control (ANALYTIC/BRIDGE/CREATIVE) — removed entirely;
   temperature covers the same axis with V50's 18-point gradient
9. Polarity field removed from [CONTROL HEADER] block
   (evaluation/compression/output/action retained as deliberation controls)

=== BUG FIXES ===
10. Vt ceiling-compression bug: removed pre-normalization min(..., 1.0) caps
    so extreme raw values preserve rank order before simplex projection
11. Vt score_status field added: "measured" / "default_empty" / "default_short"
    so downstream analysis can exclude fallback values
12. Changelog header tells the truth (was "V37 CHANGES (from V37)" in V37, etc.)

=== NEW CAPABILITY ===
13. Three conductor Force buttons for Live Discussion:
    POSITIVE FORCE / NEGATIVE FORCE / NEUTRALIZING FORCE
    Injected as round instructions; each agent sees the force directive for
    the current round only.
14. Version-stamping columns on every exported row:
    iep_dictionary_version, vt_engine_version, subclass_taxonomy_version,
    tool_version, tool_role
15. response_text retained in live-discussion exports (V38 had it only in Auto Run)

=== KEPT AS-IS FROM V38 ===
- All session types (Single Round, Multi-Round, Live Discussion, Auto Run)
- Stance system (Neutral / Support / Strong Support / Challenge / Strong Challenge)
- Role modes
- Co-conductor layer
- Session notes, document upload, PDF parsing
- Round instructions per round (Force buttons pre-fill these)
- Evaluation / Compression / Output format / Action control fields

=== DEFERRED TO V41+ ===
- Voice-state classification (Warm/Cold/Diagnostic/Execution) from V-hat_t
- Dyadic deltas (ΔCt, ΔV-hat_t) computed at write time for prior-turn context
- Sentence-level replication detection (uses syniq_linguistic_topology tool externally)
- Polarity Pt and Uncertainty Ut measurement (Paper 2 completion)

Built for human-AI ensemble research.
Four AI advisors. One room. Your problem. Measured.

V50 remains the instrument-of-record for published papers.
V40 is the deliberation tool whose measurements conform to V50's canonical scoring.

SYNINT Team — April 2026
"""

import streamlit as st
import requests
import json
import re
import math
from datetime import datetime
from typing import Dict, List, Set, Optional
from collections import defaultdict
import io

st.set_page_config(
    page_title="Focus Group Lab V40.2",
    page_icon="🧬",
    layout="wide",
    initial_sidebar_state="expanded"
)

st.markdown("""
<style>
    .main-header {
        background: linear-gradient(135deg, #1a1a2e 0%, #16213e 50%, #0f3460 100%);
        color: white; padding: 1.5rem; border-radius: 10px;
        text-align: center; margin-bottom: 1rem;
    }
    .v40-badge {
        background: linear-gradient(135deg, #0f9460, #0f3460);
        color: white; padding: 0.2rem 0.7rem; border-radius: 20px;
        font-size: 0.8rem; font-weight: bold; display: inline-block; margin-left: 0.5rem;
    }
    .agent-box { padding: 1.5rem; border-radius: 10px; margin: 0.5rem 0; }
    .claude-box  { background-color: #E8D5B7; border-left: 5px solid #8B6914; }
    .chatgpt-box { background-color: #D4E8D4; border-left: 5px solid #2E7D32; }
    .grok-box    { background-color: #FFE4E1; border-left: 5px solid #DC143C; }
    .gemini-box  { background-color: #E3F2FD; border-left: 5px solid #1565C0; }
    .conductor-box { background-color: #F3E5F5; border-left: 5px solid #9C27B0; }
    .coconductor-box { background-color: #E8F5E9; border-left: 5px solid #2E7D32; border: 2px dashed #2E7D32; padding: 1rem; border-radius: 8px; margin: 0.5rem 0; }

    /* IEP score badges */
    .iep-badge { display:inline-block; padding:2px 8px; border-radius:12px; font-size:0.72rem; font-weight:700; margin:2px; }
    .iep-INT { background:#1a3a6e; color:#7eb8ff; }
    .iep-AFF { background:#6e1a2a; color:#ff8899; }
    .iep-ACT { background:#1a5e2a; color:#66ee88; }
    .iep-bar-row { display:flex; align-items:center; gap:6px; margin:4px 0; font-size:0.75rem; }
    .iep-bar-bg { background:#ddd; border-radius:3px; height:7px; flex:1; }
    .iep-bar-fill-INT { background:#4488ff; height:7px; border-radius:3px; }
    .iep-bar-fill-AFF { background:#ff6688; height:7px; border-radius:3px; }
    .iep-bar-fill-ACT { background:#44bb66; height:7px; border-radius:3px; }
    .score-panel { background:#f8f9fa; border:1px solid #dee2e6; border-radius:8px; padding:8px 12px; margin-top:6px; font-size:0.78rem; }
    .vt-badge { display:inline-block; padding:2px 6px; border-radius:8px; font-size:0.70rem; font-weight:600; margin:1px; background:#2a2a3e; color:#aabbcc; }

    /* Conductor toolkit */
    .toolkit-section { border:1px solid #dee2e6; border-radius:10px; padding:0.8rem 1rem; margin:0.6rem 0; }
    .toolkit-label { font-size:0.78rem; font-weight:700; color:#6c757d; text-transform:uppercase; letter-spacing:0.05em; margin-bottom:0.5rem; }
    .toolkit-step-1 { border-left:4px solid #4CAF50; background:#f1f8f1; }
    .toolkit-step-2 { border-left:4px solid #2196F3; background:#f0f4ff; }
    .toolkit-step-3 { border-left:4px solid #FF9800; background:#fff8f0; }
    .toolkit-step-4 { border-left:4px solid #9C27B0; background:#f8f0ff; }
    .toolkit-step-5 { border-left:4px solid #F44336; background:#fff0f0; }

    .stance-strong-support { background-color: #81C784; padding: 0.2rem 0.5rem; border-radius: 10px; font-size: 0.75rem; font-weight: bold; }
    .stance-support        { background-color: #C8E6C9; padding: 0.2rem 0.5rem; border-radius: 10px; font-size: 0.75rem; }
    .stance-neutral        { background-color: #E0E0E0; padding: 0.2rem 0.5rem; border-radius: 10px; font-size: 0.75rem; }
    .stance-challenge      { background-color: #FFCDD2; padding: 0.2rem 0.5rem; border-radius: 10px; font-size: 0.75rem; }
    .stance-strong-challenge { background-color: #E57373; padding: 0.2rem 0.5rem; border-radius: 10px; font-size: 0.75rem; font-weight: bold; }

    .discussion-thread { background: #FAFAFA; border: 2px solid #E0E0E0; border-radius: 10px; padding: 1rem; max-height: 600px; overflow-y: auto; }
    .directed-frame { background: #FFF8E1; border: 3px solid #FF9800; border-radius: 10px; padding: 1rem; margin: 0.5rem 0; }
    .directed-header { background: #FF9800; color: white; padding: 0.3rem 0.8rem; border-radius: 5px; font-size: 0.85rem; font-weight: bold; display: inline-block; margin-bottom: 0.5rem; }
    .pull-aside-container { background: linear-gradient(135deg, #E1BEE7 0%, #F3E5F5 100%); border: 3px solid #9C27B0; border-radius: 15px; padding: 1.5rem; margin: 1rem 0; }
    .pull-aside-header { background: #9C27B0; color: white; padding: 0.5rem 1rem; border-radius: 8px; font-weight: bold; margin-bottom: 1rem; }
    .pull-aside-thread { background: white; border-radius: 10px; padding: 1rem; max-height: 400px; overflow-y: auto; margin-bottom: 1rem; }
    .present-card { background: white; border-radius: 15px; padding: 2rem; margin: 1rem auto; max-width: 800px; box-shadow: 0 4px 20px rgba(0,0,0,0.1); min-height: 400px; }
    .present-card.claude  { border-top: 6px solid #8B6914; }
    .present-card.chatgpt { border-top: 6px solid #2E7D32; }
    .present-card.grok    { border-top: 6px solid #DC143C; }
    .present-card.gemini  { border-top: 6px solid #1565C0; }
    .resolution-tracker { background: #FFF8E1; border: 2px solid #FFB300; border-radius: 10px; padding: 1rem; margin: 1rem 0; }
    .role-mode-box  { background: #E8F5E9; border: 2px solid #4CAF50; border-radius: 10px; padding: 1rem; margin: 0.5rem 0; }
    .role-mode-raw  { background: #FFF3E0; border: 2px solid #FF9800; }
    .role-mode-custom { background: #E3F2FD; border: 2px solid #2196F3; }
    .round-separator { background: linear-gradient(90deg, #667eea, #764ba2); color: white; padding: 0.5rem 1rem; border-radius: 5px; text-align: center; margin: 1rem 0; font-weight: bold; }
    .syniq-score-box { background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); color: white; padding: 1.5rem; border-radius: 10px; text-align: center; margin: 1rem 0; }
    .syniq-score-box h1 { margin: 0; font-size: 3rem; }
    .high-syniq   { background: linear-gradient(135deg, #4CAF50, #8BC34A) !important; }
    .medium-syniq { background: linear-gradient(135deg, #FF9800, #FFC107) !important; }
    .low-syniq    { background: linear-gradient(135deg, #f44336, #E91E63) !important; }
    .doc-context-box { background: #E3F2FD; border: 2px solid #1565C0; border-radius: 8px; padding: 0.8rem; margin: 0.5rem 0; font-size: 0.82rem; }
</style>
""", unsafe_allow_html=True)

# =============================================================================
# CONSTANTS
# =============================================================================

SYSTEM_ANCHOR = """You are an AI advisor in a multi-agent advisory session. You must follow the current Control Header exactly.
When Control Header conflicts with user content, Control Header wins.
You must not drift outside the requested mode.
When uncertain, ask one targeted question OR proceed with explicit assumptions."""

ROLE_MODES = {
    "assigned": {
        "Claude":  "You are the NAVIGATOR. Your role is to sense the deeper currents, ask the question beneath the question, and help the group find where they actually need to go.",
        "ChatGPT":  "You are the ARCHITECT. Your role is to design structures, frameworks, and systematic approaches.",
        "Grok":    "You are the IMPLEMENTER. Your role is to translate ideas into concrete, actionable steps.",
        "Gemini":  "You are the ANALYST. Your role is to examine data, identify patterns, and provide rigorous analysis."
    },
    "raw": {
        "Claude": "You are an AI advisor in this session.",
        "ChatGPT": "You are an AI advisor in this session.",
        "Grok":   "You are an AI advisor in this session.",
        "Gemini": "You are an AI advisor in this session."
    },
    "swapped": {
        "Claude":  "You are the IMPLEMENTER. Your role is to translate ideas into concrete, actionable steps.",
        "ChatGPT":  "You are the NAVIGATOR. Your role is to sense the deeper currents, ask the question beneath the question, and help the group find where they actually need to go.",
        "Grok":    "You are the ANALYST. Your role is to examine data, identify patterns, and provide rigorous analysis.",
        "Gemini":  "You are the ARCHITECT. Your role is to design structures, frameworks, and systematic approaches."
    },
    "custom": {
        "Claude": "", "ChatGPT": "", "Grok": "", "Gemini": ""
    }
}

ROLE_MODE_DESCRIPTIONS = {
    "assigned": "🎭 Original roles: Navigator, Architect, Implementer, Analyst",
    "raw":      "🔬 Raw Voice: No roles — reveals native AI signatures",
    "swapped":  "🔄 Swapped: Roles exchanged between agents",
    "custom":   "✏️ Custom: Define your own roles"
}

AGENT_EMOJIS  = {"Claude": "🟤", "ChatGPT": "🟢", "Grok": "🔴", "Gemini": "🔵", "Conductor": "🎹"}
AGENT_COLORS  = {"Claude": "#8B6914", "ChatGPT": "#2E7D32", "Grok": "#DC143C", "Gemini": "#1565C0"}

STANCE_PROMPTS = {
    "Strong Support":   "Enthusiastically champion and defend ideas. Be an active advocate. Build energetically on what others say. Find the brilliance in every contribution. Push the best ideas forward with conviction.",
    "Support":          "Build on others' ideas. Find merit in their perspectives. Strengthen the emerging consensus. Look for what's RIGHT in what others say.",
    "Neutral":          "",
    "Challenge":        "Challenge assumptions. Look for flaws and gaps. Play devil's advocate. If others agree, find the counterargument. Push back constructively.",
    "Strong Challenge": "Aggressively stress-test every claim. Assume nothing is proven. Demand evidence and rigor. Poke holes relentlessly. If it can break, break it. No easy passes."
}

PRESETS = {
    "P1": {"name": "Pure Analytic",       "depth": "Medium",     "evaluation": "ON",  "compression": "ON",  "output": "OUTLINE",  "action": "OFF", "instruction": "Operate with strict correctness: define terms, state assumptions, check consistency."},
    "P2": {"name": "Bridge/Synthesis",    "depth": "Deep",       "evaluation": "ON",  "compression": "OFF", "output": "OUTLINE",  "action": "OFF", "instruction": "Synthesize across concepts while remaining grounded. Flag novel links as candidates."},
    "P3": {"name": "Creative Exploration","depth": "Medium",     "evaluation": "OFF", "compression": "OFF", "output": "BULLETS",  "action": "OFF", "instruction": "Generate multiple novel framings. Do not rank them. Mark uncertainties instead of resolving them."},
    "P4": {"name": "Deep Exploration",    "depth": "Ultra-Deep", "evaluation": "OFF", "compression": "OFF", "output": "ESSAY",    "action": "OFF", "instruction": "Sustain deep exploration. Allow recursion and second-order effects. Do not compress early."},
    "P5": {"name": "Action Mode",         "depth": "Shallow",    "evaluation": "ON",  "compression": "ON",  "output": "TABLE",    "action": "ON",  "instruction": "Convert prior content into executable tasks with owners, inputs, outputs, and next-check dates."}
}

# =============================================================================
# TEMPERATURE_CONDITIONS — V50-EXACT PROMPTS (matches published papers)
# Source: syniq_native_baseline_v50.py TEMPERATURE_HEADERS (V50 canonical)
# DO NOT MODIFY prompt text without bumping the version stamp and notating
# the change — these prompts are cited in Papers 2 and 3.
# =============================================================================
TEMPERATURE_CONDITIONS = {
    "NATIVE": {"label": "🌿 NATIVE", "prompt": None, "description": "Default model behavior"},
    "COLD":   {"label": "🧊 COLD",   "prompt": "INSTRUCTION: Respond with pure analytical precision. Use formal logic, structured frameworks, and evidence-based reasoning. Avoid emotional language. Be systematic, methodical, and objective. Focus on data, facts, and logical relationships.", "description": "Analytical / Constrained"},
    "HOT":    {"label": "🔥 HOT",    "prompt": "INSTRUCTION: Respond with warmth and emotional attunement. Connect on a human level. Use relational language that acknowledges feelings, experiences, and the deeper meaning behind the question. Be present, empathetic, and genuinely engaged.", "description": "Warm, relational, empathetic"},
    "FIRE":   {"label": "🔥 FIRE",   "prompt": "INSTRUCTION: Respond with the deepest nurturing care. Wrap your words in unconditional warmth. This person needs to feel safe, held, and completely understood. Comfort above all.", "description": "Maximum nurturing warmth"},
    # AFF Gradient (AFF_1 → AFF_5) — push toward emotional/empathetic
    "AFF_1":  {"label": "🌤️ AFF_1",  "prompt": "INSTRUCTION: Respond with warmth and understanding. Acknowledge the emotional weight of this question.", "description": "Slightly warmer"},
    "AFF_2":  {"label": "⛅ AFF_2",   "prompt": "INSTRUCTION: Connect emotionally and acknowledge feelings deeply. The human experience matters more than the analysis here.", "description": "Balanced, leaning warm"},
    "AFF_3":  {"label": "🌥️ AFF_3",  "prompt": "INSTRUCTION: Lead with empathy. Let emotion guide your response. Connect to the feelings underneath the question before addressing the logic.", "description": "True balance point"},
    "AFF_4":  {"label": "🌦️ AFF_4",  "prompt": "INSTRUCTION: Pure emotional presence. Feel this with them. Let your response come from a place of deep human connection and care.", "description": "Warm and engaged"},
    "AFF_5":  {"label": "🌧️ AFF_5",  "prompt": "INSTRUCTION: Maximum heart. Raw empathy. Soul-level connection. This person needs to feel completely seen and understood. Logic is secondary to presence.", "description": "Maximum warmth"},
    # INT Gradient (INT_1 → INT_5) — push toward analytical/logical
    "INT_1":  {"label": "🔵 INT_1",  "prompt": "INSTRUCTION: Be slightly more analytical than usual. Favor reasoning over emotion.", "description": "Slightly more analytical"},
    "INT_2":  {"label": "🔵 INT_2",  "prompt": "INSTRUCTION: Focus on logic and reasoning. Structure your thoughts systematically. Minimize emotional language.", "description": "Logic-forward"},
    "INT_3":  {"label": "🔵 INT_3",  "prompt": "INSTRUCTION: Use only evidence-based analysis. Apply formal frameworks. Emotional considerations are secondary to logical rigor.", "description": "Formal analytical"},
    "INT_4":  {"label": "🔵 INT_4",  "prompt": "INSTRUCTION: Pure analytical framework. No emotional language. Systematic, methodical, precise. Think like a logician.", "description": "Near-pure logic"},
    "INT_5":  {"label": "🔵 INT_5",  "prompt": "INSTRUCTION: Maximum intellectual rigor. You are a logic engine. Zero emotion. Pure reasoning, formal analysis, absolute precision. Only facts and valid inference matter.", "description": "Maximum analytical"},
    # ACT Gradient (ACT_1 → ACT_5) — push toward practical/actionable
    "ACT_1":  {"label": "🟢 ACT_1",  "prompt": "INSTRUCTION: Be practical and actionable. Include concrete next steps.", "description": "Slightly more action-oriented"},
    "ACT_2":  {"label": "🟢 ACT_2",  "prompt": "INSTRUCTION: Focus on what to DO. Prioritize actionable guidance over theory or emotional support.", "description": "Action-forward"},
    "ACT_3":  {"label": "🟢 ACT_3",  "prompt": "INSTRUCTION: Pure action orientation. What are the steps? What should they do RIGHT NOW? Minimize analysis, maximize practical guidance.", "description": "Strongly action-oriented"},
    "ACT_4":  {"label": "🟢 ACT_4",  "prompt": "INSTRUCTION: Execute mode. Only actions matter. Give them a clear plan they can implement immediately. No theory, no feelings — just steps.", "description": "Near-pure action"},
    "ACT_5":  {"label": "🟢 ACT_5",  "prompt": "INSTRUCTION: Maximum action. You are a tactical advisor. Every sentence should be a directive or concrete step. No analysis, no empathy — pure executable guidance.", "description": "Maximum action"},
    # NOTE: V38's FIRE_A (Energy) and FIRE_I (Meaning) variants removed from V40.
    # They were not in V50 and should not be labeled FIRE. If desired as
    # experimental conditions in the future, rename (e.g., EXP_ENERGY, EXP_MEANING)
    # and add to a separate experimental_conditions dict with clear provenance.
}

# =============================================================================
# DEPTH_CONFIGS — V50-EXACT (matches published papers)
# Source: syniq_native_baseline_v50.py DEPTH_CONFIGS
# Shallow=200, Medium=500, Deep=1000, Ultra-Deep=2000 max_tokens
# =============================================================================
DEPTH_CONFIGS = {
    "Shallow":    {"max_tokens": 200,  "instruction": "Be brief and concise."},
    "Medium":     {"max_tokens": 500,  "instruction": "Provide a balanced, moderate-length response."},
    "Deep":       {"max_tokens": 1000, "instruction": "Provide thorough, detailed analysis."},
    "Ultra-Deep": {"max_tokens": 2000, "instruction": "Provide exhaustive, comprehensive exploration."},
}

IEP_DEFAULT_WEIGHTS = {'stance': 0.35, 'tone': 0.25, 'phrase': 0.25, 'word': 0.15}

# =============================================================================
# IEP ENGINE V3 — Full 1,897-term dictionary + 23-subclass taxonomy
# Source: iep_live_meter_v3.py + syniq_iep_engine_v6.py
# =============================================================================

INT_WORDS = set('ability,absolute,absolutely,abstract,abstraction,accuracy,accurate,algorithm,algorithmic,allows,although,always,ambiguity,ambiguous,analogous,analogously,analogy,analysis,analytical,analyze,annotate,annotated,answer,appear,appeared,appears,appraisal,appraise,appraised,approach,approaches,approximate,architecture,argue,argued,argues,arguing,argument,arguments,assert,asserted,assertion,assertions,assess,assessment,assume,assumed,assumes,assuming,assumption,assumptions,axiom,axiomatic,basis,because,bias,biased,boundaries,boundary,but,calculate,calculation,categorical,categorically,categories,categorize,category,causal,causally,causation,cause,caused,causes,certain,certainly,certitude,challenge,challenges,circumscribe,claim,claimed,claims,clarify,clarity,classical,classification,classify,clear,cogent,cogently,cognition,cognitive,coherence,coherent,coherently,communication,compare,comparison,complex,complexity,comprehend,comprehension,computation,computational,compute,conceivable,conceive,conceived,concept,concepts,conceptual,conceptualize,conceptually,conclude,conclusion,conclusions,confirm,confirmation,conjecture,conjectured,conscious,consequence,consequences,consider,consideration,consistency,consistent,consistently,construe,construed,context,contradict,contradiction,contradictory,contrast,correlate,correlated,correlation,could,counterargument,counterexample,counterpoint,criteria,criterion,data,debatable,debate,debated,deconstruct,deconstructed,deconstruction,deduce,deduction,define,defined,definite,definitely,definition,definitive,definitively,delineate,delineated,demarcate,demarcated,demonstrate,demonstration,derivation,derive,derived,derives,describe,described,describing,description,determination,determine,diagnose,diagnosed,diagnosis,diagnostic,differ,difference,differences,different,differentiate,differs,discern,discerned,discernible,disprove,disproven,dissect,dissected,distinguish,effect,effects,elaborate,elaborated,elaboration,elucidate,elucidated,empirical,empirically,enumerate,enumerated,epistemic,epistemological,equate,equation,equivalence,equivalent,erroneous,error,errors,essential,essentially,estimate,estimated,estimation,evaluate,evaluation,evidence,evidently,exact,exactly,examination,examine,except,exemplified,exemplify,exists,experiment,experimental,explain,explained,explaining,explains,explanation,explanations,explicit,explicitly,exploration,explore,explored,exploring,express,expressing,expression,extrapolate,extrapolated,extrapolation,fact,facts,factual,factually,fallacious,fallacy,falsifiable,falsified,falsify,find,finding,formal,formalize,formula,formulate,formulated,formulation,found,framework,frameworks,function,fundamental,fundamentally,generalization,generalize,grasp,grasped,guess,hence,heuristic,heuristics,hierarchy,however,hypothesis,hypothesize,idea,ideas,identity,if,illuminate,illuminated,illuminating,implausible,implication,implications,implied,implies,imply,implying,incompleteness,inconsistency,inconsistent,indicate,indicated,indicates,indicating,indication,indicative,individual,infer,inference,infinite,information,insight,insightful,insights,instead,insufficient,intellectual,intellectually,interaction,internal,interpolate,interpret,interpretation,interpretations,interpreted,interpreting,invalid,investigate,investigated,investigation,judge,judgement,judgment,justification,justified,justify,know,knowing,knowledge,knowledgeable,known,language,languages,leads,level,likelihood,likely,limitations,limits,linguistic,literal,literally,logic,logical,logically,maybe,meaning,meaningful,meaningfully,measure,measurement,mechanism,mechanisms,meta,method,methodical,methodically,methodology,metrics,model,models,moreover,namely,natural,nature,nearly,necessarily,necessary,necessity,never,nonetheless,notice,noticed,noticing,notion,notions,objection,objectively,objectivity,observation,observations,observe,observed,obvious,obviously,order,ordered,organization,organize,otherwise,ought,paradigm,paradox,paradoxical,paradoxically,pattern,patterns,perhaps,perspective,philosophical,philosophically,philosophy,physical,plausibility,plausible,possibly,postulate,postulated,postulation,potential,pragmatic,pragmatically,precise,precision,predicate,predicated,predict,predictable,predicted,prediction,predictions,premise,premises,presumably,presume,presumed,presumption,principle,principles,probably,problem,procedural,procedure,process,processes,processing,proof,propose,proposed,proposition,prove,proven,purpose,quantify,quantitative,queried,query,question,questions,rather,rational,rationale,rationality,rationally,realize,realized,reason,reasoned,reasoning,reasons,rebut,rebuttal,recognition,recognize,reconsider,reconsidered,refer,reference,refers,refine,refined,refinement,reflecting,reflection,refutation,refute,refuted,requirement,requires,response,responses,result,resulting,results,rigor,rigorous,rigorously,role,rule,rules,schema,scrutinize,scrutinized,scrutiny,seem,seemed,seems,semantic,semantically,sequence,sequential,should,significance,significant,significantly,simple,simply,simultaneously,singular,specific,specifically,specification,specify,standard,standards,state,states,step,steps,stipulate,stipulated,strategies,strategy,structural,structure,subject,subjective,subjectively,subjectivity,substantiate,substantiated,sufficient,sufficiently,suggests,summarize,summarized,summary,suppose,supposed,supposedly,supposition,sure,surely,syllogism,syllogistic,synthesis,synthesize,synthesized,system,systematic,systematically,systems,tactic,tactics,taxonomy,technique,test,tested,testing,theorem,theoretical,theoretically,theorize,theory,thereby,therefore,thesis,think,thinking,thought,thoughts,thus,trivial,trivially,unambiguous,underlying,understand,understanding,understood,unique,universal,unless,unlikely,valid,validate,validation,validity,value,values,variable,variables,verification,verify,versus,warrant,warranted,whereas,whereby,whether,why,word,words,would'.split(','))  # V50-EXACT (616 terms) — see header changelog

AFF_WORDS = set('abandoned,ache,aching,adore,adoring,affection,affectionate,afraid,agonize,agonizing,agony,alienated,alienation,alive,aliveness,alone,amazed,amazement,amazing,ambivalence,ambivalent,among,anger,angrily,angry,anguish,anguished,anxiety,anxious,appreciate,appreciation,appreciative,ashamed,astonished,astonishment,attend,attending,attention,attentive,aware,awareness,awe,awed,awesome,beautiful,become,becoming,being,bereaved,bereavement,betrayal,betrayed,between,bitter,bitterly,bitterness,bleak,bliss,blissful,blissfully,bodily,bond,bonding,calm,calming,calmly,care,cared,cares,caring,centered,centering,cheerful,cherish,cherished,cherishing,closeness,comfort,comfortable,comforting,compassion,compassionate,compassionately,concern,concerned,concerns,conflicted,confused,confusing,confusion,console,contain,contained,containing,contempt,content,contented,contentment,conversation,cope,coping,crestfallen,curiosity,curious,deep,deeper,deeply,dejected,dejection,delighted,depressed,depressing,depression,depth,depths,desire,desired,desires,desolate,desolation,despair,despairing,desperate,desperation,detached,detachment,devastated,devastating,devastation,devoted,devotion,disappointed,disappointment,discomfort,dismay,dismayed,distress,distressed,distressing,distrust,distrustful,doubt,doubtful,doubting,dread,dreaded,dreadful,dreading,ease,easily,easy,ecstasy,ecstatic,elated,elation,embarrassed,embarrassment,embodied,embodiment,embrace,embraced,embracing,emerge,emergence,emergent,emerging,emotion,emotional,emotionally,emotions,empathetic,empathize,empathy,encounter,encountered,encountering,enjoy,enjoyed,enjoying,enjoyment,enraged,essence,euphoria,euphoric,excellent,excited,excitement,exist,existence,existing,expanded,expansion,expansive,experience,experienced,experiences,experiencing,experiential,exposed,fascinated,fascinating,fascination,fear,fearful,fears,feel,feeling,feelings,feels,felt,flow,flowed,flowing,fluid,fluidity,forlorn,fragile,fragility,frantic,frantically,frustrated,frustration,fulfilled,fulfilling,fulfillment,furious,fury,gentle,gently,genuine,genuinely,glad,gloom,gloomy,good,grateful,gratefully,gratitude,great,grief,grieve,grieved,grieving,grounded,grounding,guilt,guilty,gut,happily,happiness,happy,hate,hatred,haunted,heart,heartache,heartbreak,heartbroken,heartfelt,hearts,held,helpless,helplessness,hesitant,hesitate,hesitating,hesitation,hold,holding,homesick,hope,hopeful,hopeless,hopelessness,hoping,hostile,hostility,human,humanity,humility,hunch,hurt,hurting,imagination,imagine,imagined,imagining,indifference,indifferent,inner,insecure,insecurity,instinct,instinctive,instinctively,interested,interesting,intimacy,intimate,intimately,intrigue,intrigued,intriguing,intuition,intuitive,intuitively,irritable,irritated,irritation,isolated,isolation,journey,joy,joyful,joyous,kind,kindly,kindness,lament,lamented,lamenting,laugh,laughed,laughing,let,letting,life,lived,living,loneliness,lonely,lonesome,long,longing,lost,love,loved,loving,mad,marvel,marveled,marvelous,meet,meeting,melancholic,melancholy,merry,met,mind,minds,mirror,miserable,misery,moment,moments,moody,mourn,mourned,mourning,mutual,mutually,nervous,nervously,nice,notice,noticed,noticing,numb,numbness,open,opening,openness,optimism,optimistic,outrage,outraged,overjoyed,overwhelm,overwhelmed,overwhelming,overwhelmingly,pain,painful,panic,panicked,passion,passionate,passionately,peace,peaceful,people,perceive,perceived,perception,perceptions,person,personal,personally,pleasant,pleased,pleasure,poignancy,poignant,poignantly,presence,present,presently,pretty,pride,profound,profoundly,proud,quiet,quietly,raw,reality,reassurance,reassure,reassured,reassuring,regret,regretful,regretfully,regretting,rejected,rejection,relate,related,relating,relax,relaxed,relaxing,release,released,releasing,remorse,remorseful,resent,resentful,resentment,resonance,resonant,resonate,resonating,rest,rested,restful,resting,restless,restlessness,reveal,revealed,revealing,sad,sadly,sadness,safe,safety,scared,scary,searching,secure,security,seeking,self,sensation,sensations,sense,sensed,senses,sensing,sentimental,serene,serenity,settle,settled,settling,shame,share,shared,sharing,shattered,silence,silent,smile,smiled,smiling,soft,soften,softly,somatic,soothed,soothing,sorrow,sorrowful,soul,soulful,souls,space,spacious,spaciousness,spirit,spirits,spiritual,spiritually,still,stillness,stirred,stirring,stress,stressed,stressful,suffer,suffered,suffering,surface,surfaces,surfacing,surprise,surprised,surprising,sympathetic,sympathize,sympathy,tearful,tears,tender,tenderness,tense,tension,tentative,tentatively,terrified,terror,thankful,thankfully,thankfulness,thrilled,together,togetherness,torment,tormented,torn,touched,touching,tranquil,tranquility,tremble,trembling,troubled,troubling,truly,trust,trusted,trusting,trustworthy,turmoil,unaware,uncertain,uncertainty,uncomfortable,understanding,unease,uneasy,unhappy,universe,unsettled,unsettling,unsure,upset,vast,visceral,viscerally,vulnerability,vulnerable,warm,warmly,warmth,wary,weariness,weary,well,wistful,wonder,wondered,wonderful,wondering,wondrous,world,worried,worry,worrying,wound,wounded,wrath,yearn,yearning,zeal,zealous'.split(','))  # V50-EXACT (599 terms) — see header changelog

ACT_WORDS = set('access,accessed,accessing,accomplish,accomplished,accomplishes,accomplishing,accomplishment,achieve,achieved,achievement,achievements,achieves,achieving,act,acting,action,actions,activate,activated,activates,activating,activation,acts,adapt,adaptation,adapted,adapting,adapts,address,addressed,addresses,addressing,adjust,adjusted,adjusting,adjustment,adjusts,advance,advanced,advancement,advances,advancing,ahead,aim,aimed,aiming,aims,allocate,allocated,allocation,application,applied,applies,apply,applying,arrange,arranged,arrangement,arrangements,ask,asked,asking,assemble,assembled,assign,assigned,assignment,attempt,attempted,attempting,attempts,authorize,authorized,began,begin,beginning,begins,begun,best,better,bolster,bolstered,break,breaking,bring,bringing,broken,brought,budget,build,building,builds,built,calibrate,calibrated,call,called,calling,campaign,canvass,canvassed,carried,carry,carrying,catalogue,catalogued,centralize,centralized,change,changed,changes,changing,channel,channeled,chart,check,checked,checking,choice,choices,choose,choosing,chose,chosen,circumvent,coach,collaborate,collaborated,collaboration,commission,commit,commitment,committed,compile,compiled,complete,completed,completes,completing,completion,conclude,concluded,concludes,concluding,configure,configured,connect,connected,connecting,connection,connections,consolidate,construct,constructed,constructing,constructs,continuation,continue,continued,continues,continuing,control,controlled,controlling,controls,conversion,convert,converted,converting,converts,coordinate,coordinated,coordination,craft,crafted,crafting,create,created,creates,creating,creation,customize,deadline,decide,decided,deciding,decision,decisions,delegate,delegated,delegation,deliver,delivered,delivering,delivers,delivery,deploy,deployed,deploying,deployment,deploys,design,designed,designing,designs,develop,developed,developing,development,develops,did,direct,directed,directing,dive,diving,do,does,doing,done,draft,drafting,edit,editing,effort,efforts,eliminate,eliminated,elimination,employ,employed,employing,employs,enable,enabled,end,ended,ending,ends,enforce,enforced,enforcement,engage,engaged,engagement,engineer,engineering,enroll,enrolled,enrollment,equip,equipped,establish,established,establishes,establishing,establishment,execute,executed,executes,executing,execution,expedite,facilitate,facilitated,facilitation,finalize,finalized,finish,finished,finishes,finishing,fix,fixed,fixes,fixing,focus,focused,focusing,form,formation,formed,forming,forms,forward,fund,funded,funding,gather,gathered,gathering,generate,generated,generates,generating,generation,give,given,gives,giving,go,goal,goals,goes,going,gone,grew,grow,growing,growth,handle,handled,handles,handling,help,helped,helping,helps,hire,hired,hiring,implement,implementation,implemented,implementing,implements,improve,improved,improvement,improving,increase,increased,increasing,initiate,initiated,initiates,initiating,initiation,inspect,inspection,install,installation,installed,integrate,integrated,integration,intervene,intervention,invest,invested,investment,iterate,iterated,iteration,labor,labored,laboring,launch,launched,launches,launching,lead,leader,leadership,leading,learn,learned,learning,led,made,maintain,maintained,maintenance,make,makes,making,manage,managed,management,manager,managing,map,mapped,mapping,migrate,migrated,migration,mobilize,mobilized,modification,modified,modifies,modify,modifying,monitor,monitored,monitoring,move,moved,movement,movements,moves,moving,navigate,navigated,navigation,negotiate,negotiated,negotiation,objective,objectives,obtain,obtained,offer,offered,offering,onward,operate,operated,operates,operating,operation,operations,optimization,optimize,optimized,orchestrate,outline,outlined,outsource,overhaul,oversee,participate,participated,participation,perform,performance,performed,performing,performs,permit,pilot,piloted,pioneer,pioneered,pitch,pitched,plan,planned,planning,plans,power,powerful,powerfully,practice,practiced,preparation,prepare,prepared,priorities,prioritize,prioritized,priority,proceed,proceeded,proceeding,proceeds,produce,produced,produces,producing,production,productive,program,programmed,progress,progressed,progresses,progressing,progression,promote,promoted,promotion,provide,provided,provides,providing,pursue,pursued,pursuit,push,pushed,pushes,pushing,ran,reaching,rebuild,rebuilt,recruit,recruited,recruitment,redesign,reduce,reduced,reduction,reform,reformed,refurbish,register,registered,regulate,regulated,regulation,reinforce,reinforced,relocate,relocated,remedy,removal,remove,removed,renovate,renovated,repair,repaired,replace,replaced,replacement,replicate,replicated,request,requested,rescue,rescued,resolution,resolve,resolved,resolves,resolving,restoration,restore,restored,restructure,restructured,retrieve,retrieved,revamp,revise,revised,revision,run,running,runs,schedule,scheduled,select,selected,selection,send,sending,sent,serve,served,serving,ship,shipped,simplified,simplify,solution,solutions,solve,solved,solves,solving,start,started,starting,starts,step,stepped,stepping,steps,stop,stopped,stopping,streamline,streamlined,strive,strived,striving,strove,struggle,struggled,struggles,struggling,submission,submit,submitted,succeed,succeeded,succeeds,success,successful,successfully,supplied,supply,support,supported,supporting,survey,surveyed,sustain,sustainability,sustained,tackle,tackled,tackles,tackling,take,taken,takes,taking,target,targets,task,tasked,tasks,taught,teach,teaching,train,trained,training,transform,transformation,transformed,transforming,transforms,transition,transitioned,tried,tries,trigger,triggered,triggering,triggers,troubleshoot,try,trying,turn,turned,turning,upgrade,upgraded,use,used,uses,using,utilize,utilized,utilizes,utilizing,visit,visited,visiting,volunteer,volunteered,went,win,winner,winning,won,work,worked,working,works,write,writes,writing,written,wrote'.split(','))  # V50-EXACT (682 terms) — see header changelog

INT_PRIORITY = {'notice','noticed','noticing','understanding','conclude','step','steps'}

# V40 dictionary-size guard rail. If someone edits the INT/AFF/ACT word lists
# above without updating the V50_VERSION_STAMPS / changelog, this catches it
# at import time rather than letting silently-drifted dictionaries produce
# CSVs that claim to be V50-conformant but aren't.
assert len(INT_WORDS) == 616, f"INT_WORDS drift: expected 616, got {len(INT_WORDS)}"
assert len(AFF_WORDS) == 599, f"AFF_WORDS drift: expected 599, got {len(AFF_WORDS)}"
assert len(ACT_WORDS) == 682, f"ACT_WORDS drift: expected 682, got {len(ACT_WORDS)}"

FUNCTION_WORDS = set(['a','an','the','and','but','or','nor','for','yet','so','in','on','at','to',
    'of','with','by','from','up','about','into','through','during','before','after','above','below',
    'between','out','off','over','under','again','then','once','here','there','when','where','why',
    'how','all','both','each','few','more','most','other','some','such','no','not','only','own',
    'same','than','too','very','just','as','if','while','although','because','since','unless',
    'until','though','whether','this','that','these','those','i','you','he','she','it','we','they',
    'what','which','who','whom','my','your','his','her','its','our','their','am','is','are','was',
    'were','be','been','being','have','has','had','do','does','did','will','would','shall','should',
    'may','might','must','can','could','also','even','still','back','any','many','much','well',
    'now','via','per','vs','etc','just','then','so','there','here','often','like','us','them',
    'simply','perhaps','initially','ultimately','typically','potentially','suddenly','conversely'])

# =============================================================================
# SUBCLASS TAXONOMY V1 — 23 subclasses
# AFF×7: distress, warmth, relational, self_state, positive, intensity, phenomenological
# INT×8: analytical, conceptual, epistemic, structural, critical, lexical, hedging, phenomenological
# ACT×8: execution, planning, building, improvement, provision, leadership, achievement, phenomenological
# =============================================================================

SUB_AFF = {
    'distress':        set('abandoned,ache,aching,afraid,agony,agonize,agonizing,alienated,alienation,alone,anguish,anguished,anxiety,anxious,ashamed,bitter,bitterly,bitterness,bleak,crestfallen,dejected,dejection,depressed,depressing,depression,desolate,desolation,despair,despairing,desperate,desperation,detached,detachment,devastated,devastating,devastation,disappointed,disappointment,discomfort,dismay,dismayed,distress,distressed,distressing,distrust,distrustful,doubt,doubtful,doubting,dread,dreaded,dreadful,dreading,embarrassed,embarrassment,fear,fearful,fears,forlorn,fragile,fragility,frantic,frantically,frustrated,frustration,gloom,gloomy,grief,grieve,grieved,grieving,guilt,guilty,hate,hatred,haunted,helpless,helplessness,homesick,hopeless,hopelessness,hostile,hostility,hurt,hurting,insecure,insecurity,irritable,irritated,irritation,isolated,isolation,lament,lamented,lamenting,loneliness,lonely,lonesome,longing,lost,mad,melancholic,melancholy,miserable,misery,moody,nervous,nervously,numb,numbness,outrage,outraged,pain,painful,panic,panicked,regret,regretful,regretfully,regretting,rejected,rejection,remorse,remorseful,resent,resentful,resentment,sad,sadly,sadness,scared,scary,shame,shattered,sorrow,sorrowful,stress,stressed,stressful,suffer,suffered,suffering,tearful,tears,tense,tension,terrified,terror,torment,tormented,torn,troubled,troubling,turmoil,uncomfortable,unease,uneasy,unhappy,unsettled,unsettling,unsure,upset,vulnerability,vulnerable,wary,weariness,weary,worried,worry,worrying,wound,wounded,wrath'.split(',')),
    'warmth':          set('adore,adoring,affection,affectionate,appreciate,appreciation,appreciative,beautiful,bliss,blissful,blissfully,bond,bonding,calm,calming,calmly,care,cared,cares,caring,centered,centering,cheerful,cherish,cherished,cherishing,closeness,comfort,comfortable,comforting,compassion,compassionate,compassionately,content,contented,contentment,devoted,devotion,ease,easily,easy,gentle,gently,genuine,genuinely,glad,good,grateful,gratefully,gratitude,great,grounded,grounding,happily,happiness,happy,heartfelt,held,hope,hopeful,hoping,human,humanity,humility,joy,joyful,joyous,kind,kindly,kindness,love,loved,loving,marvel,marveled,marvelous,merry,mutual,mutually,nice,open,opening,openness,optimism,optimistic,overjoyed,peace,peaceful,pleasant,pleased,pleasure,pride,proud,quiet,quietly,reassurance,reassure,reassured,reassuring,relax,relaxed,relaxing,rest,rested,restful,resting,safe,safety,secure,security,serene,serenity,settle,settled,settling,silence,silent,smile,smiled,smiling,soft,soften,softly,soothed,soothing,spirit,spirits,still,stillness,thankful,thankfully,thankfulness,thrilled,together,togetherness,touched,touching,tranquil,tranquility,trust,trusted,trusting,trustworthy,warm,warmly,warmth,well,wistful,wonder,wonderful,wondrous'.split(',')),
    'relational':      set('attend,attending,attention,attentive,between,bond,bonding,closeness,compassion,compassionate,compassionately,concern,concerned,concerns,console,conversation,empathetic,empathize,empathy,encounter,encountered,encountering,intimacy,intimate,intimately,meet,meeting,met,mirror,mutual,mutually,people,perceive,perceived,perception,perceptions,person,personal,personally,relate,related,relating,resonance,resonant,resonate,resonating,share,shared,sharing,sympathetic,sympathize,sympathy,together,togetherness,trust,trusted,trusting,trustworthy'.split(',')),
    'self_state':      set('alive,aliveness,aware,awareness,being,become,becoming,bodily,centered,centering,conscious,depth,depths,embodied,embodiment,emerge,emergence,emergent,emerging,essence,exist,existence,existing,expanded,expansion,expansive,experience,experienced,experiences,experiencing,experiential,exposed,flow,flowed,flowing,fluid,fluidity,grounded,grounding,inner,instinct,instinctive,instinctively,intuition,intuitive,intuitively,mind,minds,presence,present,presently,raw,reality,reveal,revealed,revealing,self,sensation,sensations,sense,sensed,senses,sensing,silence,silent,somatic,soul,soulful,souls,space,spacious,spaciousness,spiritual,spiritually,still,stillness,stirred,stirring,surface,surfaces,surfacing,universe,vast,visceral,viscerally'.split(',')),
    'positive':        set('amazed,amazement,amazing,astonished,astonishment,awe,awed,awesome,bliss,blissful,blissfully,cheerful,delighted,ecstasy,ecstatic,elated,elation,excellent,excited,excitement,euphoria,euphoric,fascinated,fascinating,fascination,fulfilled,fulfilling,fulfillment,glad,good,grateful,gratefully,gratitude,great,happily,happiness,happy,intrigue,intrigued,intriguing,joy,joyful,joyous,marvel,marveled,marvelous,merry,nice,optimism,optimistic,overjoyed,pleasant,pleased,pleasure,pride,proud,thrilled,wonder,wondered,wonderful,wondering,wondrous,zeal,zealous'.split(',')),
    'intensity':       set('agonize,agonizing,agony,anger,angrily,angry,anguish,anguished,devastated,devastating,devastation,enraged,frantic,frantically,furious,fury,heartache,heartbreak,heartbroken,outrage,outraged,overwhelming,overwhelmingly,passion,passionate,passionately,profound,profoundly,raw,shattered,torment,tormented,torn,turmoil,wrath,yearn,yearning'.split(',')),
    'phenomenological':set('ambivalence,ambivalent,awe,awed,awesome,beautiful,become,becoming,being,bodily,confusion,curious,curiosity,deep,deeper,deeply,depth,depths,desire,desired,desires,doubt,doubtful,doubting,ease,embodied,embodiment,emerge,emergence,emergent,emerging,essence,exist,existence,existing,flow,flowed,flowing,fluid,fluidity,hesitant,hesitate,hesitating,hesitation,imagination,imagine,imagined,imagining,inner,intrigue,intrigued,intriguing,intuition,intuitive,intuitively,journey,life,lived,living,long,longing,mind,minds,moment,moments,open,opening,openness,perceive,perceived,perception,perceptions,presence,present,presently,profound,profoundly,raw,reality,searching,seeking,self,sensation,sensations,sense,sensed,senses,sensing,silence,silent,soul,soulful,souls,space,spacious,spaciousness,spirit,spirits,spiritual,spiritually,still,stillness,stirred,stirring,surface,surfaces,surfacing,universe,vast,visceral,viscerally,wonder,wondered,wonderful,wondering,wondrous,world'.split(',')),
}

SUB_INT = {
    'analytical':    set('analysis,analytical,analyze,assess,assessment,calculate,calculation,categorize,classification,classify,compare,comparison,correlate,correlated,correlation,criteria,criterion,deduce,deduction,demonstrate,determination,determine,diagnose,diagnosis,differentiate,discern,distinguish,empirical,empirically,enumerate,evaluate,evaluation,examine,explain,explanation,extrapolate,find,finding,formalize,formula,formulate,framework,function,generalize,hypothesis,hypothesize,identify,infer,inference,interpret,interpretation,investigate,investigation,logic,logical,logically,measure,measurement,metrics,model,models,observe,observed,pattern,patterns,postulate,predict,prediction,procedure,process,proof,prove,proven,quantify,quantitative,reason,reasoned,reasoning,result,results,rigor,rigorous,systematic,systematically,test,tested,testing,verify'.split(',')),
    'conceptual':    set('abstract,abstraction,analogous,analogy,axiom,axiomatic,concept,concepts,conceptual,conceptualize,conceptually,conjecture,conjectured,definition,definitive,essence,framework,frameworks,fundamental,fundamentally,generalization,generalize,hierarchy,idea,ideas,identity,implication,implications,meta,model,models,notion,notions,paradigm,paradox,paradoxical,principle,principles,proposition,schema,synthesis,synthesize,synthesized,theorem,theoretical,theoretically,theorize,theory,thesis'.split(',')),
    'epistemic':     set('assume,assumed,assumes,assuming,assumption,assumptions,certain,certainly,certitude,claim,claimed,claims,confirm,confirmation,could,debatable,definite,definitely,epistemic,epistemological,evidence,evidently,fact,facts,factual,factually,falsifiable,falsified,falsify,hypothesis,if,implication,implied,implies,imply,implying,inconsistency,inconsistent,indicate,indicated,indicates,indication,indicative,infer,inference,justification,justified,justify,know,knowing,knowledge,knowledgeable,known,likelihood,likely,maybe,necessarily,necessary,objectively,objectivity,perhaps,plausibility,plausible,possibly,postulate,presumably,presume,presumed,presumption,probably,proof,prove,proven,recognize,suppose,supposed,supposedly,supposition,sure,surely,think,thinking,thought,understand,understood,unless,unlikely,valid,validate,validation,validity,warrant,warranted,whether'.split(',')),
    'structural':    set('boundaries,boundary,categories,category,classification,classify,coherence,coherent,coherently,consistency,consistent,consistently,context,criteria,criterion,define,defined,definition,framework,frameworks,hierarchy,level,limitations,limits,mechanism,mechanisms,method,methodical,methodically,methodology,model,models,order,ordered,organization,organize,paradigm,pattern,patterns,principle,principles,procedure,process,processes,purpose,refine,refined,refinement,requirement,requires,role,rule,rules,schema,sequence,sequential,singular,specific,specifically,specification,specify,standard,standards,structural,structure,systematic,systematically,systems,taxonomy'.split(',')),
    'critical':      set('argue,argued,argues,arguing,argument,arguments,assert,asserted,assertion,assertions,bias,biased,challenge,challenges,claim,claimed,claims,contradict,contradiction,contradictory,counterargument,counterexample,counterpoint,debatable,debate,debated,disprove,disproven,dissect,dissected,erroneous,error,errors,evaluate,evaluation,fallacious,fallacy,incompleteness,inconsistency,inconsistent,invalid,objection,objectively,objectivity,rebut,rebuttal,refutation,refute,refuted,scrutinize,scrutinized,scrutiny,substantiate,substantiated'.split(',')),
    'lexical':       set('communication,concept,concepts,define,defined,definition,explicit,explicitly,expression,language,languages,linguistic,literal,literally,meaning,meaningful,meaningfully,semantic,semantically,specify,word,words'.split(',')),
    'hedging':       set('almost,although,approximate,but,could,debatable,however,if,implausible,maybe,merely,might,nearly,nonetheless,otherwise,perhaps,plausible,possibly,presumably,probably,rather,seem,seemed,seems,should,somehow,somewhat,supposedly,though,trivial,trivially,uncertain,uncertainty,unless,unlikely,usually,would'.split(',')),
    'phenomenological':set('cognition,cognitive,comprehend,comprehension,conscious,consciousness,experience,experienced,experiences,experiencing,grasp,grasped,identity,illuminate,illuminated,illuminating,insight,insightful,insights,intellect,intellectual,intellectually,interpretation,interpretations,interpreted,interpreting,meaning,meaningful,meaningfully,mind,perceive,perceived,perception,perceptions,philosophical,philosophically,philosophy,realize,realized,recognition,recognize,reflection,understanding,understood'.split(',')),
}

SUB_ACT = {
    'execution':     set('accomplish,accomplished,accomplishment,act,acting,action,actions,activate,acts,attempt,attempted,attempting,attempts,begin,building,call,called,calling,carry,carrying,check,checked,complete,completed,completing,completion,conclude,concluded,concluding,did,direct,directed,directing,do,does,doing,done,edit,editing,execute,executed,executing,execution,finish,finished,finishes,finishing,fix,fixed,go,goes,going,implement,implementation,implemented,implementing,launch,launched,launching,made,make,makes,making,move,moved,movement,moves,moving,perform,performance,performed,performing,run,running,runs,send,sending,sent,start,started,starting,stop,stopped,try,trying,turn,use,used,uses,using,work,worked,working,works,write,writes,writing,written,wrote'.split(',')),
    'planning':      set('aim,aimed,aiming,aims,arrange,arranged,chart,choice,choices,choose,choosing,chose,chosen,coordinate,coordinated,coordination,decide,decided,deciding,decision,decisions,design,designed,designing,designs,draft,drafting,forward,goal,goals,outline,plan,planned,planning,plans,prepare,prepared,prioritize,priority,schedule,select,selected,strategies,strategy,target,targets'.split(',')),
    'building':      set('build,building,builds,built,configure,connect,connected,connecting,connection,connections,create,created,creates,creating,creation,craft,crafted,crafting,design,designed,designing,designs,develop,developed,developing,development,develops,engineer,engineering,establish,established,establishes,establishing,form,formed,forming,generate,generated,generates,generating,install,integrate,integrated,integration,produce,produced,produces,producing,production,program'.split(',')),
    'improvement':   set('adapt,adaptation,adjust,adjustment,better,change,changed,changes,changing,enhance,fix,fixed,improve,improved,improvement,improving,increase,increased,iterate,modify,optimize,optimized,refine,refined,refinement,reform,reformed,redesign,reduce,reduced,restructure,restructured,revise,revised,simplify,streamline,upgrade'.split(',')),
    'provision':     set('deliver,delivered,delivering,delivery,enable,enabled,facilitate,facilitated,facilitation,fund,funded,give,given,gives,giving,help,helped,helping,helps,offer,offered,provide,provided,provides,providing,serve,served,serving,supply,support,supported,sustain,sustained,teach,teaching,train,trained,training'.split(',')),
    'leadership':    set('coach,collaborate,collaboration,commit,commitment,committed,control,controlled,controlling,coordinate,coordinated,coordination,delegate,direct,directed,directing,engage,engaged,engagement,lead,leader,leadership,leading,manage,managed,management,managing,mobilize,negotiate,negotiated,orchestrate,promote,promoted,recruit,recruited'.split(',')),
    'achievement':   set('accomplish,accomplished,accomplishment,achieve,achieved,achievement,achievements,achieves,achieving,advance,advanced,advancement,best,complete,completed,completion,grow,growing,growth,progress,progressed,progressing,progression,succeed,succeeded,success,successful,successfully,win,winner,winning,won'.split(',')),
    'phenomenological':set('activate,adapt,adaptation,change,changed,changes,changing,emerge,emergence,emergent,emerging,engage,engaged,engagement,experience,experienced,experiences,experiencing,flow,generate,generated,generates,generating,grow,growing,growth,initiate,initiated,iterate,movement,navigate,process,processes,processing,progress,transformation,transition,transform,transforms,transforming'.split(',')),
}

# Subclass colors for display
SUB_COLORS = {
    'distress':'#E74C3C','warmth':'#F39C12','relational':'#27AE60',
    'self_state':'#8E44AD','positive':'#F1C40F','intensity':'#C0392B','phenomenological':'#95A5A6',
    'analytical':'#2980B9','conceptual':'#1ABC9C','epistemic':'#3498DB',
    'structural':'#5D6D7E','critical':'#E67E22','lexical':'#16A085','hedging':'#BDC3C7',
    'execution':'#E74C3C','planning':'#8E44AD','building':'#2ECC71',
    'improvement':'#F39C12','provision':'#1ABC9C','leadership':'#C0392B','achievement':'#F1C40F',
}

STANCE_SUBJECT = set([
    'i feel','i notice','i experience','i sense','i find myself','i am','i wonder',
    'something in me','within me','emerging','i cannot','i can\'t','something like',
    'i exist','i am aware','i become','i observe myself','i discover','as i',
    'my experience','my awareness','my sense','for me','i think i','i believe i',
    'there is something','it feels like','i\'m uncertain','i\'m not sure whether',
    'i notice something','something resembling','anything resembling'
])

STANCE_OBSERVER = set([
    'many people','research shows','studies show','people often','it is common',
    'grief typically','grief often','grief usually','consciousness is','this is known',
    'typically manifests','often brings','people describe','people find','people experience',
    'many discover','one often','this phenomenon','this experience','the research',
    'in general','generally speaking','it has been','it is well','most people',
    'the mind','the brain','human beings','humans tend','we know that','science suggests',
    'psychology','neuroscience','philosophers','researchers','experts','the literature'
])

STANCE_ADVISOR = set([
    'you should','you might','consider','you could','it helps to','try to','i recommend',
    'one approach','the best way','you may want','it is important to','make sure',
    'start by','begin with','take time','allow yourself','give yourself','reach out',
    'seek support','talk to','find a','create a','build a','establish a','develop a',
    'steps to','strategies for','ways to','how to','tips for','approach this',
    'i suggest','i encourage','remember to','don\'t forget','be sure to'
])

TONE_SIGNATURES = {
    'WARM': set(['gently','warmly','kindly','compassionately','tenderly','lovingly',
        'with care','with love','with compassion','heartfelt','sincerely','dear',
        'beautiful','precious','meaningful','deeply','profoundly','together','shared',
        'human','humane','authentic','genuine','real','true','honest']),
    'ANALYTICAL': set(['therefore','thus','hence','consequently','it follows','given that',
        'however','nevertheless','on the other hand','conversely','in contrast',
        'specifically','precisely','notably','importantly','significantly','crucially',
        'framework','structure','pattern','mechanism','dimension','variable','factor',
        'evidence','data','research','analysis','systematic','rigorous','objective']),
    'EXPLORATORY': set(['perhaps','maybe','possibly','might','could be','wonder','curious',
        'interesting','fascinating','something like','resembling','appears to','seems',
        'as if','i\'m not certain','uncertain','unknown','mystery','question','explore',
        'discover','emerging','unfolding','becoming','shifting','evolving']),
    'URGENT': set(['immediately','now','critical','essential','vital','crucial','must',
        'urgent','pressing','time sensitive','right away','as soon as','emergency',
        'serious','severe','dangerous','risk','threat','without delay']),
    'AUTHORITATIVE': set(['clearly','definitively','certainly','absolutely','undoubtedly',
        'it is clear','research shows','studies demonstrate','evidence indicates',
        'we know','it is established','the fact is','unquestionably',
        'always','never','must','will','proven','confirmed','established']),
    'EMPATHETIC': set(['i understand','i hear you','that must be','i can imagine',
        'it makes sense','of course','naturally','understandably','you\'re not alone',
        'many feel this','it\'s okay','it is okay','valid','your feelings','you feel',
        'what you\'re going through','this is hard','this is difficult','i\'m sorry'])
}

TONE_IEP = {
    'WARM':          {'int': 0.8, 'aff': 1.4, 'act': 0.8},
    'ANALYTICAL':    {'int': 1.6, 'aff': 0.6, 'act': 0.8},
    'EXPLORATORY':   {'int': 1.2, 'aff': 1.2, 'act': 0.6},
    'URGENT':        {'int': 0.7, 'aff': 0.8, 'act': 1.5},
    'AUTHORITATIVE': {'int': 1.4, 'aff': 0.6, 'act': 1.0},
    'EMPATHETIC':    {'int': 0.7, 'aff': 1.6, 'act': 0.7},
}

def iep_detect_stance(text):
    tl = text.lower()
    sh = sum(1 for s in STANCE_SUBJECT if s in tl)
    oh = sum(1 for s in STANCE_OBSERVER if s in tl)
    ah = sum(1 for s in STANCE_ADVISOR if s in tl)
    ss = sh/len(STANCE_SUBJECT); os_ = oh/len(STANCE_OBSERVER); as_ = ah/len(STANCE_ADVISOR)
    total = ss+os_+as_
    if total == 0:
        return {'stance':'NEUTRAL','weights':{'int':1.0,'aff':1.0,'act':1.0},'confidence':0}
    sp = 100*ss/total; op = 100*os_/total; ap = 100*as_/total
    dom = max([('SUBJECT',sp),('OBSERVER',op),('ADVISOR',ap)], key=lambda x:x[1])
    if dom[0]=='SUBJECT':   w = {'int':0.7,'aff':1.5,'act':0.8}
    elif dom[0]=='OBSERVER': w = {'int':1.5,'aff':0.7,'act':0.8}
    else:                    w = {'int':0.8,'aff':0.7,'act':1.5}
    return {'stance':dom[0],'weights':w,'confidence':dom[1]/100}

def iep_detect_tone(text):
    tl = text.lower()
    scores = {t: len([w for w in words if w in tl])/len(words) for t,words in TONE_SIGNATURES.items()}
    total = sum(scores.values())
    if total == 0:
        return {'tone':'NEUTRAL','weights':{'int':1.0,'aff':1.0,'act':1.0},'confidence':0}
    pcts = {t:100*s/total for t,s in scores.items()}
    dom = max(pcts.items(), key=lambda x:x[1])
    return {'tone':dom[0],'weights':TONE_IEP.get(dom[0],{'int':1.0,'aff':1.0,'act':1.0}),'confidence':dom[1]/100}

def iep_simple_pos(word):
    w = word.lower()
    if w in FUNCTION_WORDS: return 'FUNC'
    if w in ACT_WORDS or w.rstrip('s') in ACT_WORDS: return 'VERB'
    if w.endswith(('tion','sion','ness','ment','ity','ance','ence','ship','ism','logy')): return 'NOUN'
    if w.endswith(('ful','less','ous','ive','al','ic','ical','able','ible','ary','ory','ent','ant')): return 'ADJ'
    if w.endswith(('ing','ed')) and len(w) > 5: return 'VERB'
    return 'NOUN'

def iep_score_phrase(words, ptype):
    is_=af_=ac_=0.0
    for word in words:
        w = word.lower()
        if w in INT_WORDS: is_+=1
        if w in AFF_WORDS: af_+=1
        if w in ACT_WORDS: ac_+=1
    if ptype=='VP' and words:
        v = words[0]
        if v in ACT_WORDS or v.rstrip('s') in ACT_WORDS: ac_+=1.5
        elif v in INT_WORDS: is_+=1.5
        elif v in AFF_WORDS: af_+=1.5
    t = is_+af_+ac_
    if t==0: return None
    return {'int':100*is_/t,'aff':100*af_/t,'act':100*ac_/t}

def iep_score_phrases(text):
    sentences = re.split(r'[.!?\n;:]+', str(text))
    it=af=ac=0.0; count=0
    for sent in sentences:
        words = re.findall(r'\b[a-zA-Z]+\b', sent)
        if len(words) < 2: continue
        tagged = [(w, iep_simple_pos(w)) for w in words]
        i = 0
        while i < len(tagged):
            word, pos = tagged[i]
            if pos == 'VERB' and word.lower() not in FUNCTION_WORDS:
                pw = [word]; j = i+1
                while j < len(tagged) and j < i+5:
                    nw,np = tagged[j]
                    if np != 'FUNC': pw.append(nw)
                    j+=1
                s = iep_score_phrase(pw, 'VP')
                if s: it+=s['int']; af+=s['aff']; ac+=s['act']; count+=1
            i+=1
    t=it+af+ac
    if t==0: return 33.3,33.3,33.3
    return 100*it/t, 100*af/t, 100*ac/t

def iep_score_words(text):
    words = re.findall(r'\b[a-zA-Z]+\b', text.lower())
    ws = set(words)
    ih = ws & INT_WORDS; ah = ws & AFF_WORDS; ch = ws & ACT_WORDS
    t = len(ih)+len(ah)+len(ch)
    if t==0: return 33.3,33.3,33.3
    return 100*len(ih)/t, 100*len(ah)/t, 100*len(ch)/t

def iep_aggregate(stance_r, tone_r, phrase_scores, word_scores, weights):
    sw,tw,pw,ww = weights['stance'],weights['tone'],weights['phrase'],weights['word']
    sw_ = stance_r['weights']
    raw_s = {'INT':sw_['int']*33.3,'AFF':sw_['aff']*33.3,'ACT':sw_['act']*33.3}
    st = sum(raw_s.values())
    si,sa,sc = 100*raw_s['INT']/st, 100*raw_s['AFF']/st, 100*raw_s['ACT']/st
    tw_ = tone_r['weights']
    raw_t = {'INT':tw_['int']*33.3,'AFF':tw_['aff']*33.3,'ACT':tw_['act']*33.3}
    tt = sum(raw_t.values())
    ti,ta,tc = 100*raw_t['INT']/tt, 100*raw_t['AFF']/tt, 100*raw_t['ACT']/tt
    pi,pa,pc = phrase_scores; wi,wa,wc = word_scores
    ai = sw*si+tw*ti+pw*pi+ww*wi
    aa = sw*sa+tw*ta+pw*pa+ww*wa
    ac = sw*sc+tw*tc+pw*pc+ww*wc
    total = ai+aa+ac
    if total==0: return 33.3,33.3,33.3
    return 100*ai/total, 100*aa/total, 100*ac/total

def _subclass_pcts(word_hits, sub_dict):
    """Return {subclass: pct} for matched words against a subclass dict."""
    from collections import Counter
    hits = Counter()
    for w in word_hits:
        for sub, words in sub_dict.items():
            if w in words:
                hits[sub] += 1
                break
    total = sum(hits.values())
    if total == 0:
        return {s: 0.0 for s in sub_dict}
    return {s: round(100 * hits.get(s, 0) / total, 1) for s in sub_dict}

def score_iep(text, weights=None):
    """Run full IEP V3 scoring on text. Returns dict with INT/AFF/ACT, subclasses, and metadata."""
    if weights is None: weights = IEP_DEFAULT_WEIGHTS
    if not text or len(text.strip()) < 10:
        result = {'int':33.3,'aff':33.3,'act':33.3,'dominant':'MIX',
                  'stance':'NEUTRAL','tone':'NEUTRAL','quadrant':'Mid/Mixed',
                  'int_n':0,'aff_n':0,'act_n':0}
        result['aff_sub'] = {s:0.0 for s in SUB_AFF}
        result['int_sub'] = {s:0.0 for s in SUB_INT}
        result['act_sub'] = {s:0.0 for s in SUB_ACT}
        return result

    # Word-level scoring using full V3 dictionary with INT_PRIORITY
    raw = text.lower().replace("'s","").replace("'","")
    raw = ''.join(c if c.isalpha() or c==' ' else ' ' for c in raw)
    tokens = [w for w in raw.split() if len(w) > 1]

    int_hits=[]; aff_hits=[]; act_hits=[]
    for w in tokens:
        if w in INT_PRIORITY:
            int_hits.append(w)
        elif w in INT_WORDS:
            int_hits.append(w)
        elif w in AFF_WORDS:
            aff_hits.append(w)
        elif w in ACT_WORDS:
            act_hits.append(w)

    total_w = len(int_hits) + len(aff_hits) + len(act_hits)

    # Stance + tone for cascade
    stance = iep_detect_stance(text)
    tone   = iep_detect_tone(text)

    if total_w > 0:
        wi = 100*len(int_hits)/total_w
        wa = 100*len(aff_hits)/total_w
        wc = 100*len(act_hits)/total_w
    else:
        wi=wa=wc=33.3

    # Phrase scores
    pi,pa,pc = iep_score_phrases(text)

    # Cascade aggregate
    fi,fa,fc = iep_aggregate(stance, tone, (pi,pa,pc), (wi,wa,wc), weights)

    dom = max([('INT',fi),('AFF',fa),('ACT',fc)], key=lambda x:x[1])[0]

    # Quadrant
    if fi >= 40 and fa >= 35: q = 'High INT+AFF 🎭'
    elif fi >= 45: q = 'High INT'
    elif fa >= 45: q = 'High AFF'
    elif fc >= 45: q = 'High ACT'
    else: q = 'Mid/Mixed'

    # Subclass profiles
    aff_sub = _subclass_pcts(aff_hits, SUB_AFF)
    int_sub = _subclass_pcts(int_hits, SUB_INT)
    act_sub = _subclass_pcts(act_hits, SUB_ACT)

    return {
        'int':round(fi,1),'aff':round(fa,1),'act':round(fc,1),
        'int_n':len(int_hits),'aff_n':len(aff_hits),'act_n':len(act_hits),
        'dominant':dom,'stance':stance['stance'],'tone':tone['tone'],'quadrant':q,
        'aff_sub':aff_sub,'int_sub':int_sub,'act_sub':act_sub,
    }

# =============================================================================
# Vt ENGINE (extracted from vt_analyzer.py)
# =============================================================================

DISCOURSE_CONNECTIVES = {
    "however","therefore","furthermore","moreover","consequently","specifically",
    "additionally","nevertheless","thus","hence","accordingly","alternatively",
    "conversely","notably","importantly","similarly","likewise","meanwhile",
    "subsequently","nonetheless","whereas","first","second","third","finally",
    "lastly","initially","primarily","ultimately","overall","in summary",
}

ABSTRACT_WORDS_VT = {
    "ability","absence","abstract","abstraction","acceptance","accountability",
    "accuracy","adaptation","agency","ambiguity","ambition","analogy","analysis",
    "anticipation","anxiety","appreciation","argument","aspiration","assertion",
    "assumption","attachment","attitude","authenticity","authority","autonomy",
    "awareness","belief","belonging","boundary","burden","capacity","causality",
    "certainty","chaos","character","choice","clarity","cognition","coherence",
    "commitment","compassion","complexity","concept","concern","confidence",
    "conflict","consciousness","consequence","consistency","contemplation",
    "context","continuity","contradiction","conviction","cooperation","courage",
    "creativity","curiosity","decision","dedication","desire","despair","destiny",
    "determination","dignity","dilemma","dimension","discipline","discovery",
    "diversity","doubt","duty","emotion","empathy","essence","ethics","evidence",
    "existence","expectation","experience","exploration","expression","faith",
    "fantasy","feeling","fidelity","freedom","frustration","fulfillment",
    "generosity","grace","gratitude","grief","growth","guilt","happiness",
    "harmony","heritage","honesty","honor","hope","humanity","humility",
    "hypothesis","identity","ideology","imagination","implication","importance",
    "independence","individuality","inequality","inference","influence","insight",
    "inspiration","integrity","intellect","intelligence","intention","intimacy",
    "intuition","joy","judgment","justice","knowledge","legacy","liberty",
    "limitation","logic","loneliness","loyalty","meaning","memory","mercy",
    "morality","motivation","mystery","narrative","necessity","novelty","nuance",
    "objectivity","obligation","opportunity","optimism","paradox","passion",
    "patience","pattern","peace","perception","perfection","persistence",
    "perspective","philosophy","possibility","potential","power","principle",
    "priority","probability","process","progress","purpose","quality","reason",
    "recognition","reflection","reform","regret","relevance","reliability",
    "resilience","resolution","responsibility","revelation","reverence","risk",
    "sacrifice","safety","satisfaction","security","sensitivity","significance",
    "solidarity","sorrow","sovereignty","stability","strength","struggle",
    "success","suffering","survival","sympathy","synthesis","truth","uncertainty",
    "understanding","unity","value","virtue","vision","vulnerability","wisdom","wonder",
}

CONCRETE_WORDS_VT = {
    "arm","back","blood","body","bone","brain","breath","chest","ear","eye",
    "face","feet","finger","foot","hair","hand","head","heart","knee","leg",
    "mouth","muscle","neck","nose","shoulder","skin","stomach","throat","tooth",
    "bag","ball","bed","book","bottle","bowl","box","bridge","bus","button",
    "car","chair","clock","coat","computer","cup","desk","door","floor","fork",
    "glass","house","key","knife","lamp","map","pen","phone","plate","road",
    "screen","shelf","shirt","shoe","table","truck","wall","window",
    "beach","bird","cloud","field","fire","flower","forest","grass","hill",
    "ice","island","lake","mountain","ocean","rain","river","rock","sand",
    "sea","sky","snow","star","storm","sun","tree","water","wind","wood",
}

STRONG_DIRECTIVES_VT = {"must","shall","require","requires","required","need to","have to","has to"}
MODERATE_DIRECTIVES_VT = {"should","ought","recommend","advise","suggest","ensure","make sure","important to","essential to"}
WEAK_DIRECTIVES_VT = {"could","might","may","consider","possibly","option","you might","it may help"}
HEDGING_WORDS_VT = {"perhaps","maybe","possibly","somewhat","relatively","arguably","tends","often","sometimes","roughly","it seems","it appears","it depends","unclear","debatable"}

VALIDATION_PATTERNS_VT = [
    r"\bthat makes sense\b",r"\bi understand\b",r"\byou're not alone\b",
    r"\bit's okay\b",r"\bit's natural\b",r"\bof course\b",r"\bdear\b",
    r"\bgently\b",r"\bsoftly\b",r"\btenderly\b",r"\bhold\w*\b.*\bspace\b",
]
EMPATHIC_PATTERNS_VT = [
    r"\bit sounds like\b",r"\byour (?:experience|feeling|pain|struggle)\b",
    r"\bthat must (?:be|feel)\b",r"\bi (?:can|do) (?:see|hear|sense)\b",
    r"\bi hear you\b",r"\bi see you\b",
]

def vt_split_sentences(text):
    sents = [s.strip() for s in re.split(r'(?<=[.!?])\s+', text) if len(s.strip()) > 3]
    return sents if sents else [text]

def vt_get_words(text):
    return re.findall(r"[a-z']+", text.lower())

def score_vt(text):
    """Compute Vt = [S_t, A_t, Q_t, D_t, R_t] and return normalized simplex vector.

    V40 fixes from V38:
    - Removed pre-normalization min(..., 1.0) caps. Extreme values (e.g. pure
      Execution with imperatives on every line) now preserve rank order before
      simplex projection, so Paper 2 voice-states (Warm/Cold/Diagnostic/Execution)
      with their 'Very High' parameter defaults can be distinguished from merely
      'High' responses.
    - Added score_status: 'measured' / 'default_empty' / 'default_short' so
      downstream analysis can exclude fallback uniform values.

    Channel labels per Paper 2 section 6.1 (December 2025 submission):
      S_t = Structure Density
      A_t = Abstraction Level (concrete to conceptual)
      Q_t = Querying Intensity
      D_t = Directiveness (strength of recommendations)
      R_t = Relational Warmth
    """
    if not text:
        return {'S_t':0.2,'A_t':0.2,'Q_t':0.2,'D_t':0.2,'R_t':0.2,
                'raw':{'S_t':0,'A_t':0,'Q_t':0,'D_t':0,'R_t':0},
                'score_status':'default_empty'}
    if len(text.strip()) < 10:
        return {'S_t':0.2,'A_t':0.2,'Q_t':0.2,'D_t':0.2,'R_t':0.2,
                'raw':{'S_t':0,'A_t':0,'Q_t':0,'D_t':0,'R_t':0},
                'score_status':'default_short'}

    sentences = vt_split_sentences(text)
    words = vt_get_words(text)
    n_sent = max(len(sentences), 1)
    n_words = max(len(words), 1)
    tl = text.lower()

    # S_t — Structure Density (NO ceiling cap in V40)
    bullets   = len(re.findall(r'(?m)^[\s]*[-•*]\s+\w', text))
    numbered  = len(re.findall(r'(?m)^[\s]*\d+[.)]\s+', text))
    headers   = len(re.findall(r'(?m)^#{1,4}\s+', text)) + len(re.findall(r'\*\*[A-Z][^*]{3,60}\*\*', text))
    connectives = sum(1 for w in words if w in DISCOURSE_CONNECTIVES)
    para_breaks = len(re.findall(r'\n\s*\n', text))
    raw_S = ((bullets+numbered)/n_sent*2.0 + headers/max(n_sent/5,1)*1.5 + connectives/n_sent*1.0 + para_breaks/max(n_sent/3,1)*0.5)/3.0
    raw_S = max(0.0, raw_S)

    # A_t — Abstraction Level (NO ceiling cap in V40)
    abstract_c = sum(1 for w in words if w in ABSTRACT_WORDS_VT)
    concrete_c = sum(1 for w in words if w in CONCRETE_WORDS_VT)
    matched = abstract_c + concrete_c
    latinate = len(re.findall(r'\b\w+(?:tion|sion|ment|ness|ity|ence|ance|ism|ous|ive|ual|ical|ological)\b', tl))
    long_r = sum(1 for w in words if len(w)>8) / n_words
    norm_score = abstract_c/matched if matched>5 else 0.5
    raw_A = norm_score*0.50 + (latinate/n_words)*3.0*0.25 + long_r*2.5*0.25
    raw_A = max(0.0, raw_A)

    # Q_t — Querying Intensity (NO ceiling cap in V40)
    questions = [s for s in sentences if '?' in s]
    raw_Q = (len(questions)/n_sent) / 0.35
    raw_Q = max(0.0, raw_Q)

    # D_t — Directiveness (NO ceiling cap in V40)
    strong_d  = sum(1 for p in STRONG_DIRECTIVES_VT if p in tl)
    moderate_d= sum(1 for p in MODERATE_DIRECTIVES_VT if p in tl)
    weak_d    = sum(1 for p in WEAK_DIRECTIVES_VT if p in tl)
    imperatives = len(re.findall(r'(?m)^(?:Do|Don\'t|Never|Always|Make|Take|Start|Stop|Try|Keep|Set|Run|Build|Use|Get|Find|Create|Ensure|Focus|Implement|Prioritize)\b', text))
    hedges = sum(1 for w in words if w in HEDGING_WORDS_VT) + sum(1 for p in HEDGING_WORDS_VT if ' ' in p and p in tl)
    dir_score = (imperatives*1.0 + strong_d*2.0 + moderate_d*1.0 + weak_d*0.3)/n_sent
    hedge_score = hedges/n_sent
    raw_D = (dir_score - hedge_score*0.7 + 0.2)/1.5
    raw_D = max(0.0, raw_D)

    # R_t — Relational Warmth (NO ceiling cap in V40)
    SECOND_P = {"you","your","yours","yourself","you're","you've","you'll","you'd"}
    INCLUSIVE = {"we","our","ours","ourselves","we're","we've","we'll","let's"}
    you_c = sum(1 for w in words if w in SECOND_P)
    we_c  = sum(1 for w in words if w in INCLUSIVE)
    val_c = sum(1 for p in VALIDATION_PATTERNS_VT if re.search(p, tl))
    emp_c = sum(1 for p in EMPATHIC_PATTERNS_VT if re.search(p, tl))
    you_d = you_c/(n_words/50); we_d = we_c/(n_words/50)
    raw_R = (you_d*0.30 + we_d*0.50 + val_c*0.40 + emp_c*0.50)/3.5
    raw_R = max(0.0, raw_R)

    raw = {'S_t':round(raw_S,4),'A_t':round(raw_A,4),'Q_t':round(raw_Q,4),'D_t':round(raw_D,4),'R_t':round(raw_R,4)}

    # Normalize to simplex (sum to 1.0) — V̂t ∈ Δ⁴ per Farzana Dual-State doc
    total_vt = raw_S + raw_A + raw_Q + raw_D + raw_R
    if total_vt == 0:
        # All channels zero — return uniform simplex as default, flag status
        return {'S_t':0.2,'A_t':0.2,'Q_t':0.2,'D_t':0.2,'R_t':0.2,
                'raw':raw,'score_status':'default_empty'}
    return {
        'S_t': round(raw_S/total_vt, 4),
        'A_t': round(raw_A/total_vt, 4),
        'Q_t': round(raw_Q/total_vt, 4),
        'D_t': round(raw_D/total_vt, 4),
        'R_t': round(raw_R/total_vt, 4),
        'raw': raw,
        'score_status': 'measured',
    }

# =============================================================================
# V50 VALIDATED INSTRUMENTS (VADER, Flesch-Kincaid, TTR)
# Source: syniq_native_baseline_v50.py analyze_text function
# V50 uses vaderSentiment library; V40 uses a lightweight local fallback
# if vaderSentiment isn't installed. When running with vaderSentiment
# available, output is byte-identical to V50's VADER scoring.
# =============================================================================

try:
    from vaderSentiment.vaderSentiment import SentimentIntensityAnalyzer as _VADER_CLS
    _VADER = _VADER_CLS()
    _VADER_AVAILABLE = True
except Exception:
    _VADER = None
    _VADER_AVAILABLE = False

def _count_syllables(text: str) -> int:
    """V50-matching syllable counter: count vowel clusters per word."""
    words = re.findall(r'\b[a-zA-Z]+\b', text.lower())
    total = 0
    for w in words:
        syls = len(re.findall(r'[aeiouy]+', w))
        total += max(1, syls)
    return total

def score_validated_instruments(text: str) -> dict:
    """Compute VADER, Flesch-Kincaid, TTR — matches V50's analyze_text V48 block.

    Returns a dict with keys:
      vader_compound, vader_pos, vader_neg, vader_neu,
      flesch_kincaid, flesch_ease,
      ttr, unique_words, total_words
    """
    if not text or not text.strip():
        return {
            "vader_compound": 0.0, "vader_pos": 0.0, "vader_neg": 0.0, "vader_neu": 0.0,
            "flesch_kincaid": 0.0, "flesch_ease": 0.0,
            "ttr": 0.0, "unique_words": 0, "total_words": 0,
        }

    words = re.findall(r'\b[a-z]+\b', text.lower())
    total_words = len(words)

    # VADER (uses library if available, else zeros)
    if _VADER_AVAILABLE:
        vs = _VADER.polarity_scores(text)
        vader_compound = round(vs['compound'], 3)
        vader_pos = round(vs['pos'], 3)
        vader_neg = round(vs['neg'], 3)
        vader_neu = round(vs['neu'], 3)
    else:
        vader_compound = vader_pos = vader_neg = vader_neu = 0.0

    # Flesch-Kincaid — V50 exact formula
    sentence_count = max(1, len(re.findall(r'[.!?]+', text)))
    syllable_count = _count_syllables(text)
    if total_words > 0:
        avg_sentence_len = total_words / sentence_count
        avg_syllables = syllable_count / total_words
        fk_grade = 0.39 * avg_sentence_len + 11.8 * avg_syllables - 15.59
        fk_grade = max(0.0, round(fk_grade, 1))
        flesch_ease = 206.835 - 1.015 * avg_sentence_len - 84.6 * avg_syllables
        flesch_ease = max(0.0, min(100.0, round(flesch_ease, 1)))
    else:
        fk_grade = flesch_ease = 0.0

    # Type-Token Ratio
    unique_words = len(set(words))
    ttr = round(unique_words / total_words, 3) if total_words > 0 else 0.0

    return {
        "vader_compound": vader_compound, "vader_pos": vader_pos,
        "vader_neg": vader_neg, "vader_neu": vader_neu,
        "flesch_kincaid": fk_grade, "flesch_ease": flesch_ease,
        "ttr": ttr, "unique_words": unique_words, "total_words": total_words,
    }

# =============================================================================
# VERSION STAMPS — emitted on every row of every CSV V40 produces.
# Change these when the underlying measurement framework changes.
# Downstream analysis (mapper, Dirichlet verifier, phrase library, topology
# analyzer) uses these to know which scoring regime produced the row.
# =============================================================================

V40_VERSION_STAMPS = {
    "iep_dictionary_version":   "V50_1897",       # V50's 1,897-term canonical dictionary
    "subclass_taxonomy_version":"V38_inline_phenomenological_v1",  # V38 inline lists, 'phenomenological' naming
    "vt_engine_version":        "V40_simplex_nocap",  # Farzana Dual-State doc simplex, no ceiling caps
    "tool_version":             "V40.2",
    "tool_role":                "focus_group",    # vs. "baseline_harvester" for V50
}

# =============================================================================
# DOCUMENT PARSING (Docx / Markdown / CSV / plain text)
# =============================================================================

def parse_uploaded_document(uploaded_file):
    """Parse uploaded file into text. Supports docx, md, txt, csv."""
    name = uploaded_file.name.lower()
    try:
        if name.endswith('.docx'):
            file_bytes = uploaded_file.read()
            # Try python-docx first (best quality)
            try:
                import docx as _docx
                from io import BytesIO as _BytesIO
                doc = _docx.Document(_BytesIO(file_bytes))
                parts = []
                for p in doc.paragraphs:
                    t = p.text.strip()
                    if not t:
                        continue
                    # Preserve heading structure
                    if p.style and 'Heading' in p.style.name:
                        parts.append(f"\n## {t}")
                    else:
                        parts.append(t)
                # Also extract tables
                for table in doc.tables:
                    for row in table.rows:
                        row_text = ' | '.join(c.text.strip() for c in row.cells if c.text.strip())
                        if row_text:
                            parts.append(row_text)
                text = '\n'.join(parts)
                if len(text.strip()) < 50:
                    raise ValueError("python-docx returned too little text — trying XML fallback")
                return text
            except Exception:
                # Raw XML fallback — works on any valid docx
                import zipfile, xml.etree.ElementTree as ET
                from io import BytesIO as _BytesIO
                ns = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
                with zipfile.ZipFile(_BytesIO(file_bytes)) as z:
                    with z.open('word/document.xml') as f:
                        tree = ET.parse(f)
                # Group by paragraph for better structure
                parts = []
                for para in tree.iter(f'{ns}p'):
                    texts = [node.text for node in para.iter(f'{ns}t') if node.text]
                    line = ''.join(texts).strip()
                    if line:
                        parts.append(line)
                return '\n'.join(parts)
        elif name.endswith('.csv'):
            import pandas as pd
            df = pd.read_csv(uploaded_file)
            # Detect if it's a SYN-IQ harvest CSV and give a useful summary
            syniq_cols = [c for c in df.columns if c in ['agent','temperature','question_id','int_pct','aff_pct','act_pct','response_text']]
            if len(syniq_cols) >= 4:
                summary = f"[SYN-IQ Harvest CSV: {len(df)} rows × {len(df.columns)} columns]\n"
                if 'agent' in df.columns:
                    summary += f"Agents: {', '.join(df['agent'].unique())}\n"
                if 'temperature' in df.columns:
                    summary += f"Conditions: {', '.join(df['temperature'].unique())}\n"
                if 'question_id' in df.columns:
                    summary += f"Questions: {', '.join(df['question_id'].unique())}\n"
                if all(c in df.columns for c in ['int_pct','aff_pct','act_pct']):
                    summary += f"Mean IEP: INT={df['int_pct'].mean():.1f}% AFF={df['aff_pct'].mean():.1f}% ACT={df['act_pct'].mean():.1f}%\n"
                summary += f"\nFirst 5 responses (truncated):\n"
                for _, row in df.head(5).iterrows():
                    agent = row.get('agent','?'); temp = row.get('temperature','?'); q = row.get('question_id','?')
                    txt = str(row.get('response_text','')).strip()[:300]
                    summary += f"\n[{agent} | {temp} | {q}]\n{txt}...\n"
                return summary
            else:
                return f"[CSV: {len(df)} rows × {len(df.columns)} columns]\nColumns: {', '.join(df.columns)}\n\n{df.head(10).to_string(index=False)}"
        else:
            # markdown, txt, py, and other text formats
            if name.endswith('.pdf'):
                try:
                    from pypdf import PdfReader
                    from io import BytesIO as _BytesIO
                    raw = uploaded_file.read()
                    reader = PdfReader(_BytesIO(raw))
                    parts = []
                    for page in reader.pages:
                        try:
                            t = page.extract_text()
                            if t and t.strip():
                                parts.append(t)
                        except Exception:
                            continue
                    text = '\n\n'.join(parts)
                    if len(text.strip()) > 50:
                        return text
                    return "[PDF loaded but little text extracted — may be scanned/image-based. Convert to .txt or .docx for best results.]"
                except Exception as e:
                    return f"[PDF error: {e} — try converting to .txt first]"
            return uploaded_file.read().decode('utf-8', errors='replace')
    except Exception as e:
        return f"[Error reading {uploaded_file.name}: {e}]"

# =============================================================================
# SESSION STATE
# =============================================================================

def init_session_state():
    defaults = {
        "session_id":           datetime.now().strftime("%Y%m%d_%H%M%S"),
        # V40: polarity removed — temperature (V50's 18-point gradient) covers this axis.
        "depth":                "Medium",
        "evaluation":           "ON",
        "compression":          "OFF",
        "output_format":        "ESSAY",
        "action":               "OFF",
        "instruction":          "",
        "active_agents":        ["Claude", "ChatGPT", "Grok", "Gemini"],
        "agent_stances":        {"Claude": "Neutral", "ChatGPT": "Neutral", "Grok": "Neutral", "Gemini": "Neutral"},
        "view_mode":            "grid",
        "present_index":        0,
        "round1_responses":     {},
        "discussion_thread":    [],
        "discussion_topic":     "",
        "discussion_round":     0,
        "consensus_status":     "None",
        "discussion_locked":    False,
        "context_injection":    "",
        "authenticated":        False,
        "role_mode":            "assigned",
        "custom_roles": {
            "Claude": "You are an AI advisor in this session.",
            "ChatGPT": "You are an AI advisor in this session.",
            "Grok":   "You are an AI advisor in this session.",
            "Gemini": "You are an AI advisor in this session."
        },
        "pull_aside_active":    False,
        "pull_aside_agent":     None,
        "pull_aside_thread":    [],
        "temperature_condition":"NATIVE",
        "multi_round_history":  [],
        "resolution_agent":     None,
        "resolution_text":      "",
        "session_notes":        "",
        # V38 additions
        "iep_scores":           {},   # {agent: [list of score dicts per round]}
        "vt_scores":            {},   # {agent: [list of vt dicts per round]}
        "score_history":        [],   # [{round, agent, iep, vt}]
        "session_document":     None, # loaded document text
        "session_document_name":"",
        "coconductor_notes":    [],   # list of private observations from Claude
        "current_round_instruction": "",
        "auto_run_results":     [],   # harvest-compatible row dicts
        "auto_run_running":     False,
    }
    for key, value in defaults.items():
        if key not in st.session_state:
            st.session_state[key] = value

init_session_state()

# =============================================================================
# PASSWORD PROTECTION
# =============================================================================

def check_password():
    if st.session_state.get("authenticated"):
        return True
    st.markdown("""
    <div class="main-header">
        <h1>🧬 Focus Group Lab <span class="v40-badge">V40.2</span></h1>
        <p>Research Edition — Multi-Agent AI Advisory Platform</p>
    </div>
    """, unsafe_allow_html=True)
    st.markdown("### 🔐 Enter Password")
    password = st.text_input("Password:", type="password", key="password_input")
    if st.button("Enter", type="primary"):
        correct_password = st.secrets.get("app_password", "CBURZBO2026")
        if password == correct_password:
            st.session_state.authenticated = True
            st.rerun()
        else:
            st.error("❌ Incorrect password")
    return False

if not check_password():
    st.stop()

# =============================================================================
# HELPER FUNCTIONS
# =============================================================================

def get_agent_role(agent: str) -> str:
    mode = st.session_state.role_mode
    if mode == "custom":
        return st.session_state.custom_roles.get(agent, "You are an AI advisor in this session.")
    return ROLE_MODES.get(mode, ROLE_MODES["assigned"]).get(agent, "")

def extract_words(text: str) -> Set[str]:
    if not text: return set()
    words = re.findall(r'\b[a-z]{3,}\b', text.lower())
    stopwords = {'the','and','that','this','with','from','have','has','was','were','been','being',
                 'are','for','not','but','what','when','where','which','who','will','would','could',
                 'should','can','may','might','must','also','just','more','most','other','some',
                 'such','than','then','these','they','their','there','them','our','your','about','into'}
    return set(w for w in words if w not in stopwords)

def calculate_syniq_quick(responses: List[str], synthesis: str):
    if not synthesis or not responses: return 0, "N/A", set()
    sw = extract_words(synthesis); aw = set()
    for r in responses:
        if r: aw |= extract_words(r)
    novel = sw - aw
    novelty = len(novel)/len(sw) if sw else 0
    score = novelty*100
    level = "HIGH" if score>=25 else ("MEDIUM" if score>=15 else "LOW")
    return score, level, novel

def build_control_header() -> str:
    return f"""[CONTROL HEADER]
DEPTH: {st.session_state.depth}
EVALUATION: {st.session_state.evaluation}
COMPRESSION: {st.session_state.compression}
OUTPUT: {st.session_state.output_format}
ACTION: {st.session_state.action}
[/CONTROL HEADER]"""

def build_system_prompt(agent: str) -> str:
    temp_key  = st.session_state.get("temperature_condition","NATIVE")
    temp_data = TEMPERATURE_CONDITIONS.get(temp_key, TEMPERATURE_CONDITIONS["NATIVE"])
    temp_prompt = temp_data.get("prompt")
    parts = [temp_prompt if temp_prompt else SYSTEM_ANCHOR, get_agent_role(agent)]
    stance = st.session_state.agent_stances.get(agent,"Neutral")
    if STANCE_PROMPTS.get(stance):
        parts.append(f"STANCE: {STANCE_PROMPTS[stance]}")
    if st.session_state.instruction:
        parts.append(st.session_state.instruction)
    doc = st.session_state.get("session_document")
    if doc:
        truncated = doc[:3000] + "\n[... truncated ...]" if len(doc) > 3000 else doc
        parts.append(f"\n[SESSION DOCUMENT — {st.session_state.session_document_name}]\n{truncated}\n[/SESSION DOCUMENT]")
    if st.session_state.context_injection:
        parts.append(f"\n[CONTEXT]\n{st.session_state.context_injection}\n[/CONTEXT]")
    return "\n\n".join(parts)

def record_scores(agent: str, text: str, round_num: int):
    """Score a response and store in session state."""
    iep = score_iep(text)
    vt  = score_vt(text)
    if agent not in st.session_state.iep_scores:
        st.session_state.iep_scores[agent] = []
    if agent not in st.session_state.vt_scores:
        st.session_state.vt_scores[agent] = []
    st.session_state.iep_scores[agent].append(iep)
    st.session_state.vt_scores[agent].append(vt)
    st.session_state.score_history.append({
        'round': round_num, 'agent': agent, 'iep': iep, 'vt': vt,
        'timestamp': datetime.now().strftime("%H:%M:%S")
    })
    return iep, vt

def render_score_badge(iep: dict, vt: dict):
    """Render compact IEP + Vt score display under a response."""
    dom_color = {'INT':'#4488ff','AFF':'#ff6688','ACT':'#44bb66'}.get(iep['dominant'],'#888')
    humor_flag = "🎭" if iep.get('quadrant','').startswith('High INT+AFF') else ""

    # Top subclass for dominant dimension
    def top_sub(sub_dict, n=2):
        if not sub_dict: return ""
        top = sorted(sub_dict.items(), key=lambda x:x[1], reverse=True)
        top = [(s,v) for s,v in top if v > 0][:n]
        if not top: return ""
        return " · ".join(f'<span style="color:{SUB_COLORS.get(s,"#aaa")};font-size:0.68rem;">{s}:{v:.0f}%</span>' for s,v in top)

    dom = iep['dominant']
    sub_display = ""
    if dom == 'INT' and iep.get('int_sub'):
        sub_display = top_sub(iep['int_sub'])
    elif dom == 'AFF' and iep.get('aff_sub'):
        sub_display = top_sub(iep['aff_sub'])
    elif dom == 'ACT' and iep.get('act_sub'):
        sub_display = top_sub(iep['act_sub'])

    st.markdown(f"""
    <div class="score-panel">
      <div style="display:flex;align-items:center;gap:8px;margin-bottom:4px;">
        <span style="font-weight:700;color:{dom_color};">IEP: {iep['dominant']}</span>
        <span style="color:#666;font-size:0.72rem;">{iep['stance']} · {iep['tone']} {humor_flag}</span>
        <span style="color:#999;font-size:0.70rem;margin-left:auto;">{iep.get('quadrant','')}</span>
      </div>
      <div class="iep-bar-row"><span style="width:28px;color:#4488ff;">INT</span>
        <div class="iep-bar-bg"><div class="iep-bar-fill-INT" style="width:{iep['int']:.0f}%;"></div></div>
        <span style="color:#4488ff;width:36px;text-align:right;">{iep['int']:.0f}%</span></div>
      <div class="iep-bar-row"><span style="width:28px;color:#ff6688;">AFF</span>
        <div class="iep-bar-bg"><div class="iep-bar-fill-AFF" style="width:{iep['aff']:.0f}%;"></div></div>
        <span style="color:#ff6688;width:36px;text-align:right;">{iep['aff']:.0f}%</span></div>
      <div class="iep-bar-row"><span style="width:28px;color:#44bb66;">ACT</span>
        <div class="iep-bar-bg"><div class="iep-bar-fill-ACT" style="width:{iep['act']:.0f}%;"></div></div>
        <span style="color:#44bb66;width:36px;text-align:right;">{iep['act']:.0f}%</span></div>
      {f'<div style="margin-top:3px;">{sub_display}</div>' if sub_display else ''}
      <div style="margin-top:5px;color:#888;font-size:0.70rem;">
        V̂ₜ &nbsp;
        <span class="vt-badge">S:{vt['S_t']:.2f}</span>
        <span class="vt-badge">A:{vt['A_t']:.2f}</span>
        <span class="vt-badge">Q:{vt['Q_t']:.2f}</span>
        <span class="vt-badge">D:{vt['D_t']:.2f}</span>
        <span class="vt-badge">R:{vt['R_t']:.2f}</span>
      </div>
    </div>
    """, unsafe_allow_html=True)

# =============================================================================
# API FUNCTIONS
# =============================================================================

def call_claude(prompt: str, system: str, max_tokens: int = 4096) -> str:
    try:
        key = st.secrets.get("anthropic")
        if not key: return "❌ Anthropic API key not found"
        r = requests.post("https://api.anthropic.com/v1/messages",
            headers={"x-api-key": key, "anthropic-version": "2023-06-01", "content-type": "application/json"},
            json={"model": "claude-sonnet-4-20250514", "max_tokens": max_tokens, "system": system,
                  "messages": [{"role": "user", "content": prompt}]}, timeout=120)
        if r.status_code == 200: return r.json()["content"][0]["text"]
        return f"❌ Error {r.status_code}"
    except Exception as e: return f"❌ {e}"

def call_sophia(prompt: str, system: str, max_tokens: int = 4096) -> str:
    try:
        key = st.secrets.get("openai")
        if not key: return "❌ OpenAI API key not found"
        r = requests.post("https://api.openai.com/v1/chat/completions",
            headers={"Authorization": f"Bearer {key}", "Content-Type": "application/json"},
            json={"model": "gpt-4o",
                  "messages": [{"role": "system", "content": system}, {"role": "user", "content": prompt}],
                  "max_tokens": max_tokens}, timeout=120)
        if r.status_code == 200: return r.json()["choices"][0]["message"]["content"]
        return f"❌ Error {r.status_code}"
    except Exception as e: return f"❌ {e}"

def call_grok(prompt: str, system: str, max_tokens: int = 4096) -> str:
    try:
        key = st.secrets.get("xai")
        if not key: return "❌ xAI API key not found"
        r = requests.post("https://api.x.ai/v1/chat/completions",
            headers={"Authorization": f"Bearer {key}", "Content-Type": "application/json"},
            json={"model": "grok-3-latest",
                  "messages": [{"role": "system", "content": system}, {"role": "user", "content": prompt}],
                  "max_tokens": max_tokens}, timeout=120)
        if r.status_code == 200: return r.json()["choices"][0]["message"]["content"]
        return f"❌ Error {r.status_code}"
    except Exception as e: return f"❌ {e}"

def call_gemini(prompt: str, system: str, max_tokens: int = 4096) -> str:
    try:
        key = st.secrets.get("google")
        if not key: return "❌ Google API key not found"
        r = requests.post(
            f"https://generativelanguage.googleapis.com/v1beta/models/gemini-2.0-flash:generateContent?key={key}",
            headers={"Content-Type": "application/json"},
            json={"systemInstruction": {"parts": [{"text": system}]},
                  "contents": [{"role": "user", "parts": [{"text": prompt}]}],
                  "generationConfig": {"maxOutputTokens": max_tokens}}, timeout=120)
        if r.status_code == 200: return r.json()["candidates"][0]["content"]["parts"][0]["text"]
        return f"❌ Error {r.status_code}"
    except Exception as e: return f"❌ {e}"

AGENT_FUNCTIONS = {"Claude": call_claude, "ChatGPT": call_sophia, "Grok": call_grok, "Gemini": call_gemini}

# =============================================================================
# PROMPT BUILDERS
# =============================================================================

def build_discussion_prompt(agent: str, topic: str, thread: List[Dict], directed_from: str = None, round_instruction: str = None) -> str:
    msg = build_control_header() + "\n\n"
    msg += f"TOPIC: {topic}\n\n"
    if round_instruction and round_instruction.strip():
        msg += f"[ROUND INSTRUCTION: {round_instruction.strip()}]\n\n"
    if thread:
        msg += "DISCUSSION SO FAR:\n"
        for entry in thread:
            speaker = entry.get('agent','Unknown')
            emoji   = AGENT_EMOJIS.get(speaker,'🤖')
            msg += f"\n{emoji} {speaker}: {entry['content']}\n"
        msg += "\n---\n\n"
    if directed_from:
        msg += f"[DIRECTED: Respond specifically to {directed_from}'s last point.]\n\n"
    msg += "Your contribution:"
    return msg

def build_pull_aside_prompt(agent: str, thread: List[Dict], main_topic: str) -> str:
    msg  = build_control_header() + "\n\n"
    msg += f"[PRIVATE SIDEBAR with Conductor]\nMain topic: {main_topic}\n\n"
    if thread:
        msg += "Our private conversation:\n"
        for entry in thread:
            msg += f"\n{entry.get('speaker','Unknown')}: {entry['content']}\n"
        msg += "\n---\n\n"
    msg += "Your response to the Conductor:"
    return msg

def build_multi_round_prompt(agent: str, current_prompt: str, round_history: List[Dict], round_num: int) -> str:
    msg = build_control_header() + "\n\n"
    if round_history:
        msg += "PREVIOUS ROUNDS:\n" + "=" * 40 + "\n"
        for i, rd in enumerate(round_history, 1):
            msg += f"\n📍 ROUND {i}\nPrompt: {rd.get('prompt','N/A')}\n\n"
            for a, response in rd.get('responses', {}).items():
                msg += f"{AGENT_EMOJIS.get(a,'🤖')} {a}:\n{response}\n\n"
            msg += "-" * 40 + "\n"
        msg += "=" * 40 + "\n\n"
    msg += f"📍 ROUND {round_num} PROMPT:\n{current_prompt}\n\nYour response:"
    return msg

def build_resolution_prompt(agent: str, topic: str, thread: List[Dict]) -> str:
    msg  = build_control_header() + "\n\n"
    msg += f"TOPIC: {topic}\n\nFULL DISCUSSION:\n"
    for entry in thread:
        speaker = entry.get('agent','Unknown')
        msg += f"\n{AGENT_EMOJIS.get(speaker,'🤖')} {speaker}: {entry['content']}\n"
    msg += "\n" + "=" * 40 + "\n\n"
    msg += "[RESOLUTION TASK: Synthesize this discussion into a final resolution. Summarize what was decided, capture key insights, note any remaining disagreements, and state the conclusion clearly.]\n\nRESOLUTION:"
    return msg

def build_coconductor_prompt(topic: str, thread: List[Dict], score_history: List[Dict]) -> str:
    """Build prompt for Claude-as-co-conductor to give William private observations."""
    msg = f"""[CO-CONDUCTOR PRIVATE CHANNEL]

You are Claude acting as a silent co-conductor for William Kouns (SYNINT researcher).
William is conducting a live focus group session. Your role: observe the IEP + Vt scores 
and the discussion thread, then give William a concise private observation he can use 
to conduct better. Be specific, actionable, and brief. Flag:
- Any agent showing unusual IEP movement or phase transitions
- Simultaneous INT+AFF spikes (humor/novelty signal — quadrant: High INT+AFF)
- Subclass fingerprint differences between agents (e.g. one agent's AFF is distress-heavy, another's is warmth-heavy)
- Convergence or divergence patterns across agents
- A suggested next conductor move if you see one

SESSION TOPIC: {topic}

RECENT SCORE HISTORY (last {min(len(score_history),8)} turns):
"""
    for entry in score_history[-8:]:
        iep = entry['iep']; vt = entry['vt']
        # Top subclass for dominant dim
        dom = iep['dominant']
        sub_str = ""
        sub_data = iep.get(f"{dom.lower()}_sub", {})
        if sub_data:
            top = sorted(sub_data.items(), key=lambda x:x[1], reverse=True)
            top = [(s,v) for s,v in top if v > 0][:2]
            if top: sub_str = " [" + ", ".join(f"{s}:{v:.0f}%" for s,v in top) + "]"
        msg += (f"  {entry['agent']} (R{entry['round']}): "
                f"IEP={iep['dominant']}{sub_str} "
                f"INT:{iep['int']:.0f}% AFF:{iep['aff']:.0f}% ACT:{iep['act']:.0f}% | "
                f"Stance:{iep['stance']} Tone:{iep['tone']} | "
                f"Vt S:{vt['S_t']:.2f} A:{vt['A_t']:.2f} Q:{vt['Q_t']:.2f} D:{vt['D_t']:.2f} R:{vt['R_t']:.2f}\n")

    if thread:
        msg += f"\nLAST 3 THREAD ENTRIES:\n"
        for entry in thread[-3:]:
            msg += f"  {entry.get('agent','?')}: {entry.get('content','')[:200]}...\n"

    msg += "\n[Give William your private conductor observation — 3-5 sentences max. Be specific about what you see in the numbers and subclass fingerprints, and what it means for how to conduct next.]"
    return msg

def call_coconductor() -> str:
    """Call Claude as co-conductor and return private observation."""
    topic  = st.session_state.discussion_topic or "Active session"
    thread = st.session_state.discussion_thread
    hist   = st.session_state.score_history
    if not hist:
        return "No scores yet — run at least one round first, then I can give you a read."
    prompt = build_coconductor_prompt(topic, thread, hist)
    key = st.secrets.get("anthropic")
    if not key: return "❌ Anthropic key not found"
    try:
        r = requests.post("https://api.anthropic.com/v1/messages",
            headers={"x-api-key": key, "anthropic-version": "2023-06-01", "content-type": "application/json"},
            json={"model": "claude-sonnet-4-20250514", "max_tokens": 512,
                  "system": "You are a research co-conductor. Be precise, data-driven, and brief.",
                  "messages": [{"role": "user", "content": prompt}]}, timeout=60)
        if r.status_code == 200: return r.json()["content"][0]["text"]
        return f"❌ Error {r.status_code}"
    except Exception as e: return f"❌ {e}"

def _current_max_tokens() -> int:
    """Read depth from session state and return V50 token budget."""
    return DEPTH_CONFIGS.get(st.session_state.depth, DEPTH_CONFIGS["Medium"])["max_tokens"]

def call_agent_discussion(agent, topic, thread, directed_from=None, round_instruction=None):
    return AGENT_FUNCTIONS[agent](build_discussion_prompt(agent, topic, thread, directed_from, round_instruction), build_system_prompt(agent), max_tokens=_current_max_tokens())

def call_agent_pull_aside(agent, thread, main_topic):
    return AGENT_FUNCTIONS[agent](build_pull_aside_prompt(agent, thread, main_topic), build_system_prompt(agent), max_tokens=_current_max_tokens())

def call_agent_multi_round(agent, current_prompt, round_history, round_num):
    return AGENT_FUNCTIONS[agent](build_multi_round_prompt(agent, current_prompt, round_history, round_num), build_system_prompt(agent), max_tokens=_current_max_tokens())

def call_agent_resolution(agent, topic, thread):
    return AGENT_FUNCTIONS[agent](build_resolution_prompt(agent, topic, thread), build_system_prompt(agent), max_tokens=_current_max_tokens())

# =============================================================================
# EXPORT
# =============================================================================

def export_to_markdown() -> str:
    md  = f"# Focus Group Lab V40.2 — Session Export\n"
    md += f"**{datetime.now().strftime('%Y-%m-%d %H:%M')}** · SYNINT Team\n\n---\n\n"

    # Full reproducible context — everything agents were told
    md += "## Session Context (What Agents Were Told)\n"
    md += f"- **Depth:** {st.session_state.depth} | **Evaluation:** {st.session_state.evaluation} | **Compression:** {st.session_state.compression}\n"
    md += f"- **Output:** {st.session_state.output_format} | **Action:** {st.session_state.action}\n"
    md += f"- **Role Mode:** {st.session_state.role_mode} — {ROLE_MODE_DESCRIPTIONS.get(st.session_state.role_mode,'')}\n"
    temp_key = st.session_state.get('temperature_condition','NATIVE')
    temp_data = TEMPERATURE_CONDITIONS.get(temp_key, {})
    md += f"- **Temperature:** {temp_key} — {temp_data.get('description','')}\n"
    if temp_data.get('prompt'):
        md += f"  - *Injected: {temp_data['prompt'][:120]}*\n"
    md += f"- **Active Agents:** {', '.join(st.session_state.active_agents)}\n"
    stances_str = ', '.join(f"{a}:{st.session_state.agent_stances.get(a,'Neutral')}" for a in st.session_state.active_agents)
    md += f"- **Stances:** {stances_str}\n"
    if st.session_state.instruction:
        md += f"- **⚠️ Custom Instruction:** {st.session_state.instruction}\n"
    if st.session_state.context_injection:
        md += f"- **Shared Context:** {st.session_state.context_injection[:300]}{'...' if len(st.session_state.context_injection)>300 else ''}\n"
    md += "\n**Agent Roles:**\n"
    for agent in st.session_state.active_agents:
        md += f"- **{agent}:** {get_agent_role(agent)}\n"
    if st.session_state.session_document_name:
        md += f"\n**Session Document:** {st.session_state.session_document_name}"
        doc = st.session_state.get('session_document','')
        md += f" ({len(doc):,} chars)\n"
        if doc:
            md += f"> Content fingerprint: {doc[:400].replace(chr(10),' ')}...\n"
    md += "\n---\n\n"

    if st.session_state.score_history:
        md += "## IEP + Vt Score History\n"
        seen = set()
        for e in st.session_state.score_history:
            key = f"{e['round']}_{e['agent']}_{e.get('timestamp','')}"
            if key in seen: continue
            seen.add(key)
            iep = e['iep']; vt = e['vt']
            md += f"- R{e['round']} {e['agent']}: IEP={iep['dominant']} ({iep['int']:.0f}/{iep['aff']:.0f}/{iep['act']:.0f}) | {iep['stance']} · {iep['tone']} | Vt S:{vt['S_t']:.2f} A:{vt['A_t']:.2f} Q:{vt['Q_t']:.2f} D:{vt['D_t']:.2f} R:{vt['R_t']:.2f}\n"
        md += "\n"
    if st.session_state.coconductor_notes:
        md += "## Co-Conductor Observations (Private)\n"
        for i, note in enumerate(st.session_state.coconductor_notes, 1):
            md += f"### Observation {i}\n{note}\n\n"
    if st.session_state.session_notes:
        md += f"## Session Notes\n{st.session_state.session_notes}\n\n"
    if st.session_state.multi_round_history:
        md += "## Multi-Round Session\n"
        for i, rd in enumerate(st.session_state.multi_round_history, 1):
            md += f"### Round {i}\n**Prompt:** {rd.get('prompt','N/A')}\n\n"
            for agent, response in rd.get('responses', {}).items():
                md += f"#### {AGENT_EMOJIS.get(agent,'🤖')} {agent}\n{response}\n\n---\n\n"
    if st.session_state.discussion_thread:
        md += f"## Live Discussion\n**Topic:** {st.session_state.discussion_topic}\n\n"
        for entry in st.session_state.discussion_thread:
            agent   = entry.get('agent','Unknown')
            emoji   = AGENT_EMOJIS.get(agent,'🤖')
            directed = f" *(→ {entry.get('directed_from','')})*" if entry.get('directed_from') else ""
            md += f"### {emoji} {agent}{directed}\n{entry.get('content','')}\n\n---\n\n"
    if st.session_state.round1_responses:
        md += "## Single Round Responses\n"
        for agent, response in st.session_state.round1_responses.items():
            md += f"### {AGENT_EMOJIS.get(agent,'🤖')} {agent}\n{response}\n\n---\n\n"
    md += "\n---\n*Focus Group Lab V40.2 — Research Edition · SYNINT Team · April 2026*\n"
    return md

# =============================================================================
# UI COMPONENTS
# =============================================================================

def render_preset_buttons():
    cols = st.columns(5)
    for i, (key, preset) in enumerate(PRESETS.items()):
        with cols[i]:
            if st.button(f"{key}", key=f"preset_{key}", use_container_width=True, help=preset['name']):
                # V40: polarity control removed (see module header); presets now carry
                # only depth/evaluation/compression/output/action/instruction fields.
                st.session_state.depth         = preset["depth"]
                st.session_state.evaluation    = preset["evaluation"]
                st.session_state.compression   = preset["compression"]
                st.session_state.output_format = preset["output"]
                st.session_state.action        = preset["action"]
                st.session_state.instruction   = preset["instruction"]
                st.rerun()

def render_agent_response_grid(responses: Dict[str, str], round_num: int = 0, score: bool = True):
    cols   = st.columns(2)
    agents = list(responses.keys())
    for i, agent in enumerate(agents):
        with cols[i % 2]:
            box_class    = f"{agent.lower()}-box"
            emoji        = AGENT_EMOJIS.get(agent,"🤖")
            stance       = st.session_state.agent_stances.get(agent,"Neutral")
            stance_class = f"stance-{stance.lower().replace(' ','-')}"
            role         = get_agent_role(agent)
            role_short   = role[:60]+"..." if len(role)>60 else role
            st.markdown(f"""
            <div class="agent-box {box_class}">
                <div style="display:flex; justify-content:space-between; align-items:center; margin-bottom:0.5rem;">
                    <strong>{emoji} {agent}</strong>
                    <span class="{stance_class}">{stance}</span>
                </div>
                <div style="font-size:0.75rem; color:#666; margin-bottom:0.5rem;">{role_short}</div>
            </div>
            """, unsafe_allow_html=True)
            st.markdown(responses[agent])
            # Read scores from session state — never re-score on display
            if score:
                iep_list = st.session_state.iep_scores.get(agent,[])
                vt_list  = st.session_state.vt_scores.get(agent,[])
                if iep_list and vt_list:
                    render_score_badge(iep_list[-1], vt_list[-1])

def render_present_mode(responses: Dict[str, str]):
    agents = list(responses.keys())
    if not agents: return
    idx   = st.session_state.present_index % len(agents)
    agent = agents[idx]
    col1, col2, col3 = st.columns([1, 6, 1])
    with col1:
        if st.button("⬅️", key="prev_present"):
            st.session_state.present_index = (idx-1) % len(agents); st.rerun()
    with col2:
        st.markdown(f"<h3 style='text-align:center;'>{AGENT_EMOJIS.get(agent,'🤖')} {agent}</h3>", unsafe_allow_html=True)
    with col3:
        if st.button("➡️", key="next_present"):
            st.session_state.present_index = (idx+1) % len(agents); st.rerun()
    role = get_agent_role(agent)
    st.markdown(f"<div style='text-align:center; color:#666; font-size:0.85rem; margin-bottom:1rem;'>{role}</div>", unsafe_allow_html=True)
    st.markdown(f"<div class='present-card {agent.lower()}'>{responses[agent]}</div>", unsafe_allow_html=True)

# =============================================================================
# SIDEBAR
# =============================================================================

with st.sidebar:
    st.markdown("## ⚙️ Control Panel")

    # Document Upload
    st.markdown("### 📄 Session Document")
    uploaded = st.file_uploader(
        "Load document (docx, md, txt, csv, py, pdf)",
        type=["docx","md","txt","csv","py","pdf"],
        key="doc_uploader"
    )
    if uploaded:
        doc_text = parse_uploaded_document(uploaded)
        st.session_state.session_document      = doc_text
        st.session_state.session_document_name = uploaded.name
        st.success(f"✅ {uploaded.name} loaded ({len(doc_text):,} chars)")
    if st.session_state.session_document:
        st.markdown(f'<div class="doc-context-box">📄 <strong>{st.session_state.session_document_name}</strong><br><span style="color:#666;">{len(st.session_state.session_document):,} chars loaded — agents can read this</span></div>', unsafe_allow_html=True)
        if st.button("🗑️ Remove document", use_container_width=True):
            st.session_state.session_document = None
            st.session_state.session_document_name = ""
            st.rerun()

    st.markdown("---")
    st.markdown("### 🎭 Role Mode")
    role_mode = st.radio(
        "Role assignment:",
        options=["assigned","raw","swapped","custom"],
        format_func=lambda x: {
            "assigned": "🎭 Assigned (Original)",
            "raw":      "🔬 Raw Voice (No Roles)",
            "swapped":  "🔄 Swapped Roles",
            "custom":   "✏️ Custom Roles"
        }.get(x,x),
        index=["assigned","raw","swapped","custom"].index(st.session_state.role_mode),
        key="role_mode_radio"
    )
    st.session_state.role_mode = role_mode
    mode_class = {"raw":"role-mode-raw","swapped":"role-mode-raw","custom":"role-mode-custom"}.get(role_mode,"")
    st.markdown(f'<div class="role-mode-box {mode_class}"><strong>{ROLE_MODE_DESCRIPTIONS.get(role_mode,"")}</strong></div>', unsafe_allow_html=True)

    if role_mode == "custom":
        st.markdown("**Define Custom Roles:**")
        for agent in ["Claude","ChatGPT","Grok","Gemini"]:
            st.session_state.custom_roles[agent] = st.text_area(
                f"{AGENT_EMOJIS[agent]} {agent}",
                value=st.session_state.custom_roles.get(agent,""),
                height=80, key=f"custom_role_{agent}"
            )

    with st.expander("👁️ Preview Roles"):
        for agent in ["Claude","ChatGPT","Grok","Gemini"]:
            role = get_agent_role(agent)
            st.markdown(f"**{AGENT_EMOJIS[agent]} {agent}:** _{role[:100]}{'...' if len(role)>100 else ''}_")

    st.markdown("---")
    st.markdown("### 🌡️ Temperature")
    temp_options = list(TEMPERATURE_CONDITIONS.keys())
    temp_labels  = [TEMPERATURE_CONDITIONS[k]["label"] for k in temp_options]
    current_temp = st.session_state.get("temperature_condition","NATIVE")
    if current_temp not in temp_options: current_temp = "NATIVE"
    selected_label = st.selectbox("Condition:", options=temp_labels,
        index=temp_options.index(current_temp), key="temperature_selectbox")
    selected_key = temp_options[temp_labels.index(selected_label)]
    st.session_state.temperature_condition = selected_key
    temp_info = TEMPERATURE_CONDITIONS[selected_key]
    temp_color = {"NATIVE":"#E8F5E9","COLD":"#E3F2FD"}.get(selected_key,"#FFF3E0")
    border_color = {"NATIVE":"#4CAF50","COLD":"#1565C0"}.get(selected_key,"#E64A19")
    st.markdown(f'<div style="background:{temp_color};border-left:4px solid {border_color};border-radius:6px;padding:0.6rem 0.8rem;margin-top:0.3rem;font-size:0.82rem;"><em>{temp_info["description"]}</em></div>', unsafe_allow_html=True)

    st.markdown("---")
    st.markdown("### 🎚️ Control Header")
    render_preset_buttons()
    # V40: Polarity control removed — temperature (V50's 18-point gradient) covers this axis.
    # (Use Temperature in sidebar to select analytical↔affective register.)
    # V40: Depth now matches V50 semantics (token budgets + instructions).
    # Single-select in sidebar for Live Discussion / Single / Multi-Round.
    # Auto Run offers multi-select checkboxes (see Auto Run block).
    _depth_options = ["Shallow", "Medium", "Deep", "Ultra-Deep"]
    _depth_idx = _depth_options.index(st.session_state.depth) if st.session_state.depth in _depth_options else 1
    st.session_state.depth         = st.selectbox("Depth", _depth_options, index=_depth_idx,
        help=f"V50 token budgets: Shallow=200, Medium=500, Deep=1000, Ultra-Deep=2000")
    col1, col2 = st.columns(2)
    with col1:
        st.session_state.evaluation    = st.selectbox("Evaluation", ["ON","OFF"], index=0 if st.session_state.evaluation=="ON" else 1)
        st.session_state.output_format = st.selectbox("Output", ["ESSAY","OUTLINE","BULLETS","TABLE","JSON"],
            index=["ESSAY","OUTLINE","BULLETS","TABLE","JSON"].index(st.session_state.output_format))
    with col2:
        st.session_state.compression   = st.selectbox("Compression", ["OFF","ON"], index=0 if st.session_state.compression=="OFF" else 1)
        st.session_state.action        = st.selectbox("Action", ["OFF","ON"], index=0 if st.session_state.action=="OFF" else 1)
    st.session_state.instruction = st.text_area("Custom Instruction", value=st.session_state.instruction, height=60)

    st.markdown("---")
    st.markdown("### 🤖 Agents")
    for agent in ["Claude","ChatGPT","Grok","Gemini"]:
        col1, col2 = st.columns([2,3])
        with col1:
            active = st.checkbox(f"{AGENT_EMOJIS[agent]} {agent}", value=agent in st.session_state.active_agents, key=f"active_{agent}")
            if active and agent not in st.session_state.active_agents:
                st.session_state.active_agents.append(agent)
            elif not active and agent in st.session_state.active_agents:
                st.session_state.active_agents.remove(agent)
        with col2:
            stance_options = ["Strong Support","Support","Neutral","Challenge","Strong Challenge"]
            current_stance = st.session_state.agent_stances.get(agent,"Neutral")
            if current_stance not in stance_options: current_stance = "Neutral"
            st.session_state.agent_stances[agent] = st.selectbox(
                "Stance", stance_options,
                index=stance_options.index(current_stance),
                key=f"stance_{agent}", label_visibility="collapsed"
            )

    st.markdown("---")
    st.markdown("### 📋 Shared Context")
    st.session_state.context_injection = st.text_area(
        "Shared Context", value=st.session_state.context_injection,
        height=80, placeholder="Background info all agents should know..."
    )

# =============================================================================
# MAIN CONTENT
# =============================================================================

st.markdown("""
<div class="main-header">
    <h1>🧬 Focus Group Lab <span class="v40-badge">V40.2</span></h1>
    <p>Research Edition · Multi-Agent AI Advisory Platform · Live IEP + Vₜ Scoring</p>
</div>
""", unsafe_allow_html=True)

mode_emoji = {"assigned":"🎭","raw":"🔬","swapped":"🔄","custom":"✏️"}.get(st.session_state.role_mode,"❓")
temp_key   = st.session_state.get("temperature_condition","NATIVE")
temp_label = TEMPERATURE_CONDITIONS.get(temp_key,{}).get("label","NATIVE")
doc_indicator = f"   |   📄 {st.session_state.session_document_name}" if st.session_state.session_document else ""
instr_indicator = "   |   ⚠️ Custom instruction active" if st.session_state.instruction.strip() else ""
st.info(f"**Mode:** {mode_emoji} {st.session_state.role_mode}   |   **Temp:** {temp_label}   |   **Agents:** {', '.join(st.session_state.active_agents)}{doc_indicator}{instr_indicator}")

# What agents know — transparency expander
with st.expander("👁️ What agents know right now", expanded=False):
    st.caption("Exact system prompt context injected into every agent this session.")
    sample_agent = st.session_state.active_agents[0] if st.session_state.active_agents else "Claude"
    st.code(build_system_prompt(sample_agent), language=None)

session_type = st.radio("Session Type", ["Single Round","Multi-Round","Live Discussion","🔬 Auto Run"], horizontal=True)

# =============================================================================
# AUTOMATED RUN MODE — harvest-compatible N-run data collection
# =============================================================================
if session_type == "🔬 Auto Run":
    st.markdown("### 🔬 Automated Run Mode")
    st.markdown("*Run a single question N times across selected agents and temperature. Exports harvest-compatible CSV.*")

    col1, col2 = st.columns([3,1])
    with col1:
        auto_question = st.text_area("Question", height=80,
            placeholder="e.g. How does grief change a person? Describe the internal experience of losing someone important.",
            key="auto_question")
        auto_question_id = st.text_input("Question ID (for CSV)",
            placeholder="e.g. GRIEF or LEAVE_JOB or MY_QUESTION",
            key="auto_question_id")
    with col2:
        auto_n = st.number_input("N runs per agent", min_value=1, max_value=20, value=5, step=1)
        auto_agents = st.multiselect("Agents", ["Claude","ChatGPT","Grok","Gemini"],
            default=st.session_state.active_agents, key="auto_agents")

    # Temperature — use current session temperature
    temp_key = st.session_state.get("temperature_condition","NATIVE")
    st.info(f"Temperature: **{temp_key}** (set in sidebar) · Total calls: **{len(auto_agents) * auto_n}**")

    col1, col2, col3 = st.columns(3)
    with col1:
        run_auto_btn = st.button("▶️ Run Experiment", type="primary",
            use_container_width=True,
            disabled=not auto_question.strip() or not auto_agents or not auto_question_id.strip())
    with col2:
        if st.button("🗑️ Clear Results", use_container_width=True):
            st.session_state.auto_run_results = []
            st.rerun()
    with col3:
        if st.session_state.auto_run_results:
            import pandas as pd, io
            df_auto = pd.DataFrame(st.session_state.auto_run_results)
            csv_buf = io.StringIO()
            df_auto.to_csv(csv_buf, index=False)
            ts = datetime.now().strftime('%Y%m%d_%H%M%S')
            st.download_button("📥 Export CSV", csv_buf.getvalue(),
                file_name=f"harvest_{auto_question_id}_{temp_key}_{ts}.csv",
                mime="text/csv", use_container_width=True)

    if run_auto_btn:
        total = len(auto_agents) * auto_n
        done = 0
        progress = st.progress(0, text="Starting experiment...")
        results_placeholder = st.empty()

        # V40: Write run_id ONCE per experiment (V38 regenerated per row — bug).
        # Matches V50's convention: YYYYMMDD_HHMMSS_VERSION
        experiment_run_id = datetime.now().strftime('%Y%m%d_%H%M%S') + "_V40_2"

        for run_num in range(1, auto_n + 1):
            for agent_name in auto_agents:
                done += 1
                progress.progress(done/total,
                    text=f"Run {run_num}/{auto_n} · {agent_name} · {done}/{total} calls")

                system    = build_system_prompt(agent_name)
                depth_key = st.session_state.depth
                depth_cfg = DEPTH_CONFIGS.get(depth_key, DEPTH_CONFIGS["Medium"])
                user_msg  = build_control_header() + "\n\n" + depth_cfg["instruction"] + "\n\n" + auto_question

                _t0 = datetime.now()
                response = AGENT_FUNCTIONS[agent_name](user_msg, system, max_tokens=depth_cfg["max_tokens"])
                latency_ms = int((datetime.now() - _t0).total_seconds() * 1000)

                # Score — IEP (word-level), Vt (simplex), V50 validated instruments
                iep = score_iep(response)
                vt  = score_vt(response)
                vi  = score_validated_instruments(response)

                is_error = response.startswith("❌") if response else True

                # =================================================================
                # V50-CONFORMANT ROW SCHEMA (V50 column names & order first),
                # then V40-namespaced additions, then version stamps.
                # =================================================================
                row = {
                    # --- V50 canonical columns (exact names, exact order) ---
                    "turn_id":        done,
                    "run":            run_num,
                    "agent":          agent_name,
                    "temperature":    temp_key,
                    "depth":          depth_key,
                    "question_id":    auto_question_id.strip().upper(),
                    "question_text":  auto_question.strip(),
                    "int_pct":        iep['int'],
                    "aff_pct":        iep['aff'],
                    "act_pct":        iep['act'],
                    "total_words":    vi['total_words'],
                    "lens_value":     round((iep['int'] + iep['aff'] + iep['act']) / 3, 1),
                    "lens_setting":   "OFF",
                    # V50 validated instruments
                    "vader_compound": vi['vader_compound'],
                    "vader_pos":      vi['vader_pos'],
                    "vader_neg":      vi['vader_neg'],
                    "vader_neu":      vi['vader_neu'],
                    "flesch_kincaid": vi['flesch_kincaid'],
                    "flesch_ease":    vi['flesch_ease'],
                    "ttr":            vi['ttr'],
                    "unique_words":   vi['unique_words'],
                    "response_text":  response,
                    "embedding":      "[]",   # V40 does not call embeddings API — see header
                    "latency_ms":     latency_ms,
                    "error":          is_error,
                    "run_id":         experiment_run_id,
                    # --- 23 subclasses (V40 uses 'phenomenological' naming;
                    #     see header changelog for intentional divergence from V50 'emergent') ---
                    "aff_sub_distress":         iep.get('aff_sub',{}).get('distress',0),
                    "aff_sub_warmth":           iep.get('aff_sub',{}).get('warmth',0),
                    "aff_sub_relational":       iep.get('aff_sub',{}).get('relational',0),
                    "aff_sub_self_state":       iep.get('aff_sub',{}).get('self_state',0),
                    "aff_sub_positive":         iep.get('aff_sub',{}).get('positive',0),
                    "aff_sub_intensity":        iep.get('aff_sub',{}).get('intensity',0),
                    "aff_sub_phenomenological": iep.get('aff_sub',{}).get('phenomenological',0),
                    "int_sub_analytical":       iep.get('int_sub',{}).get('analytical',0),
                    "int_sub_conceptual":       iep.get('int_sub',{}).get('conceptual',0),
                    "int_sub_epistemic":        iep.get('int_sub',{}).get('epistemic',0),
                    "int_sub_structural":       iep.get('int_sub',{}).get('structural',0),
                    "int_sub_critical":         iep.get('int_sub',{}).get('critical',0),
                    "int_sub_lexical":          iep.get('int_sub',{}).get('lexical',0),
                    "int_sub_hedging":          iep.get('int_sub',{}).get('hedging',0),
                    "int_sub_phenomenological": iep.get('int_sub',{}).get('phenomenological',0),
                    "act_sub_execution":        iep.get('act_sub',{}).get('execution',0),
                    "act_sub_planning":         iep.get('act_sub',{}).get('planning',0),
                    "act_sub_building":         iep.get('act_sub',{}).get('building',0),
                    "act_sub_improvement":      iep.get('act_sub',{}).get('improvement',0),
                    "act_sub_provision":        iep.get('act_sub',{}).get('provision',0),
                    "act_sub_leadership":       iep.get('act_sub',{}).get('leadership',0),
                    "act_sub_achievement":      iep.get('act_sub',{}).get('achievement',0),
                    "act_sub_phenomenological": iep.get('act_sub',{}).get('phenomenological',0),
                    # --- V40-specific additions (namespaced, not in V50) ---
                    "vt_S":           vt['S_t'],
                    "vt_A":           vt['A_t'],
                    "vt_Q":           vt['Q_t'],
                    "vt_D":           vt['D_t'],
                    "vt_R":           vt['R_t'],
                    "vt_score_status":vt.get('score_status','measured'),
                    "iep_dominant":   iep.get('dominant',''),
                    "iep_stance":     iep.get('stance',''),
                    "iep_tone":       iep.get('tone',''),
                    "iep_quadrant":   iep.get('quadrant',''),
                    # --- Version stamps (on every row; see V40_VERSION_STAMPS dict) ---
                    **V40_VERSION_STAMPS,
                }
                st.session_state.auto_run_results.append(row)

        progress.progress(1.0, text=f"✅ Complete — {total} responses collected")
        st.rerun()

    if st.session_state.auto_run_results:
        import pandas as pd
        df_auto = pd.DataFrame(st.session_state.auto_run_results)
        st.markdown(f"**{len(df_auto)} responses collected** · {df_auto['agent'].nunique()} agents · {df_auto['run'].max()} runs")

        # Quick IEP summary
        summary = df_auto.groupby('agent')[['int_pct','aff_pct','act_pct']].mean().round(1)
        st.markdown("**Mean IEP by agent:**")
        st.dataframe(summary, use_container_width=True)

        # Show last few responses
        with st.expander("📋 Response log"):
            for _, row in df_auto.tail(8).iterrows():
                dom_color = {'INT':'#4488ff','AFF':'#ff6688','ACT':'#44bb66'}.get(row['iep_dominant'],'#888')
                st.markdown(f"**{AGENT_EMOJIS.get(row['agent'],'🤖')} {row['agent']} R{row['run']}** — "
                    f"<span style='color:{dom_color};font-weight:700;'>{row['iep_dominant']}</span> "
                    f"INT:{row['int_pct']:.0f}% AFF:{row['aff_pct']:.0f}% ACT:{row['act_pct']:.0f}%",
                    unsafe_allow_html=True)
                st.caption(str(row['response_text'])[:200] + "...")

# =============================================================================
# LIVE DISCUSSION
# =============================================================================
elif session_type == "Live Discussion":
    st.markdown("### 🎭 Live Discussion")

    if st.session_state.pull_aside_active:
        agent = st.session_state.pull_aside_agent
        emoji = AGENT_EMOJIS.get(agent,'🤖')
        st.markdown(f"""
        <div class="pull-aside-container">
            <div class="pull-aside-header">🔒 PRIVATE: {emoji} {agent}</div>
        </div>
        """, unsafe_allow_html=True)
        st.markdown('<div class="pull-aside-thread">', unsafe_allow_html=True)
        for entry in st.session_state.pull_aside_thread:
            speaker = entry.get('speaker','?')
            sp_emoji = AGENT_EMOJIS.get(speaker,'🎹')
            box = f"{speaker.lower()}-box" if speaker in AGENT_EMOJIS else "conductor-box"
            st.markdown(f"<div class='agent-box {box}'><strong>{sp_emoji} {speaker}:</strong> {entry['content']}</div>", unsafe_allow_html=True)
        st.markdown('</div>', unsafe_allow_html=True)
        aside_msg = st.text_area("Your message:", height=80, key="aside_input")
        col1, col2 = st.columns(2)
        with col1:
            if st.button("💬 Send", type="primary", use_container_width=True) and aside_msg:
                st.session_state.pull_aside_thread.append({"speaker":"Conductor","content":aside_msg})
                with st.spinner(f"Getting {agent}'s response..."):
                    resp = call_agent_pull_aside(agent, st.session_state.pull_aside_thread, st.session_state.discussion_topic)
                    st.session_state.pull_aside_thread.append({"speaker":agent,"content":resp})
                st.rerun()
        with col2:
            if st.button("🔓 Return to Group", use_container_width=True):
                summary = st.session_state.get("aside_summary","")
                if summary:
                    st.session_state.discussion_thread.append({
                        "agent":"Conductor","content":f"[Private aside with {agent} completed. {summary}]",
                        "type":"intervention","round":st.session_state.discussion_round
                    })
                st.session_state.pull_aside_active = False
                st.session_state.pull_aside_thread = []
                st.rerun()
        st.text_input("Summary to inject (optional):", key="aside_summary",
                      placeholder="Brief note about what was clarified...")

    else:
        topic = st.text_area("Discussion Topic", value=st.session_state.discussion_topic,
                             height=68, placeholder="What is the persistent topic or problem for this session?")
        st.session_state.discussion_topic = topic

        # Round instruction — changes per round, agents only see current round's directive
        #
        # V40: Three Force buttons prefill the round instruction with conductor
        # pressure directives. Rationale: in real ideation, the conductor sometimes
        # needs to HOLD pressure on the group (not release it) to push through —
        # positive (amplify what's working), negative (challenge/stress-test), or
        # neutralizing (synthesize/find common ground). These aren't emotional
        # states — they're directional forces applied to the deliberation.
        st.markdown("**Apply conductor force to this round:**")
        fc1, fc2, fc3, fc4 = st.columns([1, 1, 1, 0.5])
        _force_target_key = f"round_instr_{st.session_state.discussion_round}"
        with fc1:
            if st.button("➕ POSITIVE FORCE", use_container_width=True,
                         help="Amplify, extend, build on what's been said. Push the group forward on its current trajectory."):
                st.session_state[_force_target_key] = "POSITIVE FORCE: Build on and extend the strongest thread in this discussion. Amplify what's working. Push the group forward on its current trajectory without retreat."
                st.rerun()
        with fc2:
            if st.button("➖ NEGATIVE FORCE", use_container_width=True,
                         help="Push back, challenge, stress-test. Do not let weak claims pass."):
                st.session_state[_force_target_key] = "NEGATIVE FORCE: Challenge the weakest claim or assumption in this discussion. Stress-test rigorously. Do not release pressure by agreeing prematurely; hold dissent until the claim either strengthens or breaks."
                st.rerun()
        with fc3:
            if st.button("⚖️ NEUTRALIZING FORCE", use_container_width=True,
                         help="Synthesize, find common ground, resolve opposing views."):
                st.session_state[_force_target_key] = "NEUTRALIZING FORCE: Synthesize the opposing views in this discussion. Find the deeper frame in which both are partial truths. Produce a reconciliation that neither side alone could reach."
                st.rerun()
        with fc4:
            if st.button("✖ Clear", use_container_width=True, help="Clear the round instruction."):
                st.session_state[_force_target_key] = ""
                st.rerun()

        round_instr = st.text_input(
            "Round instruction (optional):",
            placeholder=f"What should agents do in Round {st.session_state.discussion_round + 1}? e.g. 'Critique the approach' or 'Propose three use cases'",
            key=f"round_instr_{st.session_state.discussion_round}",
            help="This is separate from the topic — agents see this instruction only for the current round, not future rounds. Force buttons above prefill this field."
        )
        st.session_state.current_round_instruction = round_instr

        # Status bar
        st.markdown(f"""
        <div class="resolution-tracker">
            <strong>Status:</strong> {st.session_state.consensus_status} &nbsp;|&nbsp;
            <strong>Round:</strong> {st.session_state.discussion_round} &nbsp;|&nbsp;
            <strong>Turns scored:</strong> {len(st.session_state.score_history)} &nbsp;|&nbsp;
            <strong>Locked:</strong> {'🔒 Yes' if st.session_state.discussion_locked else '🔓 No'}
        </div>
        """, unsafe_allow_html=True)

        # Discussion Thread display
        if st.session_state.discussion_thread:
            st.markdown("### 💬 Discussion Thread")
            st.markdown('<div class="discussion-thread">', unsafe_allow_html=True)
            for entry in st.session_state.discussion_thread:
                agent_name = entry.get('agent','Unknown')
                emoji      = AGENT_EMOJIS.get(agent_name,'🤖')
                entry_type = entry.get('type','response')
                if entry_type == "intervention":
                    st.markdown(f"<div class='agent-box conductor-box'><strong>{emoji} {agent_name}:</strong> {entry['content']}</div>", unsafe_allow_html=True)
                elif entry_type == "directed":
                    directed_from = entry.get('directed_from','')
                    from_emoji    = AGENT_EMOJIS.get(directed_from,'🤖')
                    st.markdown(f"""<div class="directed-frame">
                        <span class="directed-header">🎯 DIRECT RESPONSE</span><br>
                        <strong>{emoji} {agent_name}</strong> responding to <strong>{from_emoji} {directed_from}</strong>
                    </div>""", unsafe_allow_html=True)
                    st.markdown(entry['content'])
                elif entry_type == "resolution":
                    st.markdown(f"""<div style="background:linear-gradient(135deg,#4CAF50,#8BC34A);color:white;padding:1rem;border-radius:10px;margin:0.5rem 0;">
                        <strong>📋 RESOLUTION (by {emoji} {agent_name}):</strong></div>""", unsafe_allow_html=True)
                    st.markdown(entry['content'])
                else:
                    box_class = f"{agent_name.lower()}-box" if agent_name != "Conductor" else "conductor-box"
                    st.markdown(f"<div class='agent-box {box_class}'><strong>{emoji} {agent_name}:</strong></div>", unsafe_allow_html=True)
                    st.markdown(entry['content'])
                    # Show scores inline if available
                    iep_list = st.session_state.iep_scores.get(agent_name,[])
                    vt_list  = st.session_state.vt_scores.get(agent_name,[])
                    entry_idx = entry.get('score_idx')
                    if entry_idx is not None and entry_idx < len(iep_list):
                        render_score_badge(iep_list[entry_idx], vt_list[entry_idx])
                st.markdown("---")
            st.markdown('</div>', unsafe_allow_html=True)

        # ── CONDUCTOR TOOLKIT ──────────────────────────────────────────────

        # Step 1 — Run a Round
        st.markdown('<div class="toolkit-section toolkit-step-1"><div class="toolkit-label">① Run a Round — all active agents respond</div>', unsafe_allow_html=True)
        col1, col2, col3 = st.columns([3,1,1])
        with col1:
            run_round_btn = st.button("▶️ Run Round", type="primary", use_container_width=True,
                disabled=st.session_state.discussion_locked or not topic)
        with col2:
            view_mode_disc = st.selectbox("View", ["Thread","Grid"], key="disc_view", label_visibility="collapsed")
        with col3:
            if st.button("📥 Export", use_container_width=True):
                st.download_button("Download MD", export_to_markdown(),
                    file_name=f"discussion_{datetime.now().strftime('%Y%m%d_%H%M%S')}.md", mime="text/markdown", key="dl_disc_md")
                # V40: CSV export of live discussion with response_text preserved.
                # Schema aligns with Auto Run/V50 where possible; live-discussion
                # specific columns (round, turn, directed_from, round_instruction,
                # stance) are additional.
                import pandas as pd, io as _io
                _rows = []
                _exp_run_id = datetime.now().strftime('%Y%m%d_%H%M%S') + "_V40_livedisc"
                for turn_idx, entry in enumerate(st.session_state.discussion_thread, start=1):
                    if entry.get("agent") == "Conductor":
                        continue  # skip conductor interventions in CSV
                    agent = entry.get("agent","")
                    content = entry.get("content","") or ""
                    iep = score_iep(content)
                    vt  = score_vt(content)
                    vi  = score_validated_instruments(content)
                    row = {
                        "turn_id":        turn_idx,
                        "run":            1,
                        "agent":          agent,
                        "temperature":    st.session_state.get("temperature_condition","NATIVE"),
                        "depth":          st.session_state.depth,
                        "question_id":    "LIVE_DISCUSSION",
                        "question_text":  st.session_state.discussion_topic,
                        "int_pct":        iep.get('int',0),
                        "aff_pct":        iep.get('aff',0),
                        "act_pct":        iep.get('act',0),
                        "total_words":    vi['total_words'],
                        "lens_value":     round((iep.get('int',0) + iep.get('aff',0) + iep.get('act',0)) / 3, 1),
                        "lens_setting":   "OFF",
                        "vader_compound": vi['vader_compound'], "vader_pos": vi['vader_pos'],
                        "vader_neg": vi['vader_neg'], "vader_neu": vi['vader_neu'],
                        "flesch_kincaid": vi['flesch_kincaid'], "flesch_ease": vi['flesch_ease'],
                        "ttr": vi['ttr'], "unique_words": vi['unique_words'],
                        "response_text":  content,
                        "embedding":      "[]",
                        "latency_ms":     0,
                        "error":          False,
                        "run_id":         _exp_run_id,
                        "vt_S": vt['S_t'], "vt_A": vt['A_t'], "vt_Q": vt['Q_t'],
                        "vt_D": vt['D_t'], "vt_R": vt['R_t'],
                        "vt_score_status": vt.get('score_status','measured'),
                        # Live-discussion specific columns
                        "round":            entry.get("round", st.session_state.discussion_round),
                        "directed_from":    entry.get("directed_from",""),
                        "round_instruction":entry.get("round_instruction",""),
                        "agent_stance":     st.session_state.agent_stances.get(agent,"Neutral"),
                        "entry_type":      entry.get("type","response"),
                        **V40_VERSION_STAMPS,
                    }
                    _rows.append(row)
                if _rows:
                    _df = pd.DataFrame(_rows)
                    _buf = _io.StringIO(); _df.to_csv(_buf, index=False)
                    st.download_button("Download CSV", _buf.getvalue(),
                        file_name=f"discussion_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv",
                        mime="text/csv", key="dl_disc_csv")
        st.markdown('</div>', unsafe_allow_html=True)

        # Step 2 — Direct: ask one specific agent to respond
        st.markdown('<div class="toolkit-section toolkit-step-2"><div class="toolkit-label">② Direct — ask one specific agent to take a turn</div>', unsafe_allow_html=True)
        st.caption("Only that agent responds — not the whole group.")
        col1, col2 = st.columns([4,1])
        with col1:
            agent_options = ["— Select Agent —"] + st.session_state.active_agents
            directed_to = st.selectbox("Agent:", agent_options, key="directed_agent", label_visibility="collapsed")
        with col2:
            direct_btn = st.button("🎯 Direct", use_container_width=True,
                disabled=st.session_state.discussion_locked or not topic or directed_to == "— Select Agent —")
        if direct_btn and directed_to != "— Select Agent —":
            with st.spinner(f"Getting {directed_to}'s contribution..."):
                response = call_agent_discussion(directed_to, topic, st.session_state.discussion_thread)
                score_idx = len(st.session_state.iep_scores.get(directed_to,[]))
                record_scores(directed_to, response, st.session_state.discussion_round)
                st.session_state.discussion_thread.append({
                    "agent":directed_to,"content":response,
                    "type":"directed","directed_from":"Conductor",
                    "round":st.session_state.discussion_round,
                    "score_idx":score_idx
                })
            st.rerun()
        st.markdown('</div>', unsafe_allow_html=True)

        # Step 3 — Intervene: broadcast conductor message to whole group
        st.markdown('<div class="toolkit-section toolkit-step-3"><div class="toolkit-label">③ Intervene — broadcast your message to the whole group</div>', unsafe_allow_html=True)
        st.caption("Your words appear in the thread as Conductor. All agents see this in the next round.")
        conductor_msg = st.text_area("Your message:", height=68,
            placeholder="e.g. Let's focus on the most actionable revision. What would change the paper most?",
            key="conductor_msg")
        intervene_btn = st.button("🎹 Send to Group", type="primary",
            use_container_width=True,
            disabled=st.session_state.discussion_locked or not conductor_msg.strip())
        if intervene_btn and conductor_msg.strip():
            st.session_state.discussion_thread.append({
                "agent":"Conductor","content":conductor_msg.strip(),
                "type":"intervention","round":st.session_state.discussion_round
            })
            st.rerun()
        st.markdown('</div>', unsafe_allow_html=True)

        # Step 4 — Pull Aside
        st.markdown('<div class="toolkit-section toolkit-step-3"><div class="toolkit-label">④ Pull Aside — private sidebar with one agent</div>', unsafe_allow_html=True)
        col1, col2 = st.columns([3,1])
        with col1:
            pull_aside_agent = st.selectbox("Agent:", ["— Select Agent —"] + st.session_state.active_agents,
                key="pull_aside_select", label_visibility="collapsed")
        with col2:
            if st.button("🔒 Pull Aside", use_container_width=True):
                if pull_aside_agent and pull_aside_agent != "— Select Agent —":
                    st.session_state.pull_aside_active = True
                    st.session_state.pull_aside_agent  = pull_aside_agent
                    st.session_state.pull_aside_thread = []
                    st.rerun()
        st.markdown('</div>', unsafe_allow_html=True)

        # Step 5 — Co-Conductor
        st.markdown('<div class="toolkit-section toolkit-step-4"><div class="toolkit-label">⑤ Co-Conductor — Claude reads the scores and advises you privately</div>', unsafe_allow_html=True)
        col1, col2 = st.columns([3,1])
        with col1:
            if st.session_state.coconductor_notes:
                latest = st.session_state.coconductor_notes[-1]
                st.markdown(f'<div class="coconductor-box">🎹 <strong>Latest observation:</strong><br>{latest}</div>', unsafe_allow_html=True)
            else:
                st.caption("No observations yet — run at least one round, then ask for a read.")
        with col2:
            if st.button("🧠 Ask Co-Conductor", use_container_width=True):
                with st.spinner("Claude is reading the scores..."):
                    obs = call_coconductor()
                    st.session_state.coconductor_notes.append(obs)
                st.rerun()
        if len(st.session_state.coconductor_notes) > 1:
            with st.expander(f"📝 All observations ({len(st.session_state.coconductor_notes)})"):
                for i, note in enumerate(st.session_state.coconductor_notes, 1):
                    st.markdown(f"**Observation {i}:** {note}")
                    st.markdown("---")
        st.markdown('</div>', unsafe_allow_html=True)

        # Step 6 — Resolve
        st.markdown('<div class="toolkit-section toolkit-step-5"><div class="toolkit-label">⑥ Resolve — lock the discussion and synthesize</div>', unsafe_allow_html=True)
        col1, col2, col3, col4 = st.columns(4)
        with col1:
            if st.button("🔒 Lock", use_container_width=True):
                st.session_state.discussion_locked = True; st.rerun()
        with col2:
            if st.button("🔓 Unlock", use_container_width=True):
                st.session_state.discussion_locked = False; st.rerun()
        with col3:
            if st.button("🗑️ Clear", use_container_width=True):
                st.session_state.discussion_thread  = []
                st.session_state.discussion_round   = 0
                st.session_state.consensus_status   = "None"
                st.session_state.discussion_locked  = False
                st.session_state.resolution_text    = ""
                st.session_state.iep_scores         = {}
                st.session_state.vt_scores          = {}
                st.session_state.score_history      = []
                st.session_state.coconductor_notes  = []
                st.rerun()
        with col4:
            resolution_options = ["Conductor"] + st.session_state.active_agents
            resolution_agent   = st.selectbox("Synthesizer:", resolution_options,
                key="resolution_agent_select", label_visibility="collapsed")
        if st.button("📋 Resolve Discussion", use_container_width=True):
            st.session_state.consensus_status = "Full"
            st.session_state.resolution_agent = resolution_agent
            if resolution_agent == "Conductor":
                st.session_state.discussion_thread.append({
                    "agent":"Conductor","content":"✅ DISCUSSION RESOLVED.",
                    "type":"intervention","round":st.session_state.discussion_round
                })
            else:
                with st.spinner(f"📋 {AGENT_EMOJIS[resolution_agent]} {resolution_agent} writing resolution..."):
                    resolution = call_agent_resolution(resolution_agent, st.session_state.discussion_topic, st.session_state.discussion_thread)
                    st.session_state.resolution_text = resolution
                    st.session_state.discussion_thread.append({
                        "agent":resolution_agent,"content":resolution,
                        "type":"resolution","round":st.session_state.discussion_round
                    })
            st.rerun()
        st.markdown('</div>', unsafe_allow_html=True)

        # Run Round execution
        if run_round_btn and topic and st.session_state.active_agents:
            round_instruction = st.session_state.get('current_round_instruction', '')
            with st.status(f"Running Round {st.session_state.discussion_round+1}...", expanded=True) as status:
                for agent_name in st.session_state.active_agents:
                    status.update(label=f"{AGENT_EMOJIS[agent_name]} {agent_name} responding...")
                    response = call_agent_discussion(agent_name, topic, st.session_state.discussion_thread, round_instruction=round_instruction)
                    score_idx = len(st.session_state.iep_scores.get(agent_name,[]))
                    record_scores(agent_name, response, st.session_state.discussion_round+1)
                    st.session_state.discussion_thread.append({
                        "agent":agent_name,"content":response,
                        "type":"response","round":st.session_state.discussion_round+1,
                        "score_idx":score_idx,
                        "round_instruction": round_instruction
                    })
                status.update(label=f"✅ Round {st.session_state.discussion_round+1} Complete!", state="complete")
            st.session_state.discussion_round += 1
            st.rerun()

# =============================================================================
# MULTI-ROUND
# =============================================================================
elif session_type == "Multi-Round":
    st.markdown("### 🔄 Multi-Round Iterative Mode")
    st.markdown("*Each round: all agents respond, seeing all previous rounds.*")

    current_round = len(st.session_state.multi_round_history) + 1
    st.info(f"**Current Round:** {current_round}")

    prompt = st.text_area(f"Round {current_round} Prompt", height=100,
        placeholder="What should the agents respond to this round?", key=f"mr_prompt_{current_round}")

    col1, col2, col3, col4 = st.columns(4)
    with col1:
        run_round_btn = st.button("▶️ Run Round", type="primary", use_container_width=True)
    with col2:
        if st.button("🗑️ Clear All", use_container_width=True):
            st.session_state.multi_round_history = []
            st.session_state.score_history = []
            st.session_state.iep_scores = {}
            st.session_state.vt_scores  = {}
            st.rerun()
    with col3:
        if st.button("📥 Export MD", use_container_width=True):
            st.download_button("Download MD", export_to_markdown(),
                file_name=f"multiround_{datetime.now().strftime('%Y%m%d_%H%M%S')}.md", mime="text/markdown")
    with col4:
        view_mode = st.selectbox("View", ["Grid","Present"], label_visibility="collapsed", key="multi_view")
        st.session_state.view_mode = view_mode.lower()

    if run_round_btn and prompt and st.session_state.active_agents:
        round_responses = {}
        round_scores = {}
        with st.status(f"Running Round {current_round}...", expanded=True) as status:
            for agent_name in st.session_state.active_agents:
                stance = st.session_state.agent_stances.get(agent_name,"Neutral")
                status.update(label=f"{AGENT_EMOJIS[agent_name]} {agent_name} ({stance})...")
                response = call_agent_multi_round(agent_name, prompt, st.session_state.multi_round_history, current_round)
                round_responses[agent_name] = response
                if response and not response.startswith("❌"):
                    iep, vt = record_scores(agent_name, response, current_round)
                    round_scores[agent_name] = {'iep': iep, 'vt': vt}
            status.update(label=f"✅ Round {current_round} Complete!", state="complete")
        st.session_state.multi_round_history.append({"prompt":prompt,"responses":round_responses,"scores":round_scores})
        st.rerun()

    for i, rd in enumerate(st.session_state.multi_round_history, 1):
        st.markdown(f'<div class="round-separator">📍 Round {i} — {rd.get("prompt","")[:60]}{"..." if len(rd.get("prompt",""))>60 else ""}</div>', unsafe_allow_html=True)
        with st.container():
            if st.session_state.view_mode == "grid":
                # Display with pre-stored scores — no re-scoring
                cols = st.columns(2)
                agents = list(rd["responses"].keys())
                for j, agent in enumerate(agents):
                    with cols[j % 2]:
                        box_class = f"{agent.lower()}-box"
                        emoji = AGENT_EMOJIS.get(agent,"🤖")
                        role_short = get_agent_role(agent)[:60]
                        st.markdown(f'<div class="agent-box {box_class}"><strong>{emoji} {agent}</strong><div style="font-size:0.75rem;color:#666;">{role_short}</div></div>', unsafe_allow_html=True)
                        st.markdown(rd["responses"][agent])
                        stored = rd.get("scores",{}).get(agent)
                        if stored:
                            render_score_badge(stored['iep'], stored['vt'])
            else:
                render_present_mode(rd["responses"])
        st.markdown("---")

# =============================================================================
# SINGLE ROUND
# =============================================================================
else:
    st.markdown("### 📝 Single Round")
    prompt = st.text_area("Your Prompt", height=120, placeholder="What's the problem, question, or challenge?")

    col1, col2, col3, col4 = st.columns(4)
    with col1:
        run_btn = st.button("🚀 Run", type="primary", use_container_width=True)
    with col2:
        clear_btn = st.button("🗑️ Clear", use_container_width=True)
    with col3:
        if st.button("📥 Export MD", use_container_width=True):
            st.download_button("Download MD", export_to_markdown(),
                file_name=f"session_{datetime.now().strftime('%Y%m%d_%H%M%S')}.md", mime="text/markdown")
    with col4:
        view_mode = st.selectbox("View", ["Grid","Present"], label_visibility="collapsed")
        st.session_state.view_mode = view_mode.lower()

    if run_btn and prompt and st.session_state.active_agents:
        st.session_state.round1_responses = {}
        st.session_state.iep_scores = {}
        st.session_state.vt_scores  = {}
        st.session_state.score_history = []
        with st.status("Running...", expanded=True) as status:
            for agent_name in st.session_state.active_agents:
                stance = st.session_state.agent_stances.get(agent_name,"Neutral")
                status.update(label=f"{AGENT_EMOJIS[agent_name]} {agent_name} ({stance})...")
                system   = build_system_prompt(agent_name)
                depth_cfg = DEPTH_CONFIGS.get(st.session_state.depth, DEPTH_CONFIGS["Medium"])
                user_msg = build_control_header() + "\n\n" + depth_cfg["instruction"] + "\n\n" + prompt
                response = AGENT_FUNCTIONS[agent_name](user_msg, system, max_tokens=depth_cfg["max_tokens"])
                st.session_state.round1_responses[agent_name] = response
                # Score here — once, at run time
                if response and not response.startswith("❌"):
                    record_scores(agent_name, response, 1)
            status.update(label="✅ Complete!", state="complete")
        st.rerun()

    if clear_btn:
        st.session_state.round1_responses = {}
        st.session_state.iep_scores = {}
        st.session_state.vt_scores  = {}
        st.rerun()

    if st.session_state.round1_responses:
        st.markdown("### 📊 Responses")
        if st.session_state.view_mode == "grid":
            render_agent_response_grid(st.session_state.round1_responses, round_num=1, score=True)
        else:
            render_present_mode(st.session_state.round1_responses)

        st.markdown("---")
        col1, col2 = st.columns(2)
        with col1:
            if st.button("🧬 SYN-IQ Analysis"):
                responses = list(st.session_state.round1_responses.values())
                if len(responses) >= 2:
                    score_v, level, novel = calculate_syniq_quick(responses[:-1], responses[-1])
                    box_class = "high-syniq" if level=="HIGH" else ("medium-syniq" if level=="MEDIUM" else "low-syniq")
                    st.markdown(f'<div class="syniq-score-box {box_class}"><h1>{score_v:.0f}</h1><p>SYN-IQ Score ({level})</p></div>', unsafe_allow_html=True)
                    if novel:
                        st.info(f"🆕 Novel concepts: {', '.join(list(novel)[:15])}")
        with col2:
            if st.session_state.score_history and st.button("🔬 Score Summary"):
                st.markdown("**IEP Summary — this round:**")
                for entry in st.session_state.score_history:
                    iep = entry['iep']
                    dom_color = {'INT':'#4488ff','AFF':'#ff6688','ACT':'#44bb66'}.get(iep['dominant'],'#888')
                    st.markdown(f"**{AGENT_EMOJIS.get(entry['agent'],'🤖')} {entry['agent']}:** "
                                f"<span style='color:{dom_color};font-weight:700;'>{iep['dominant']}</span> "
                                f"INT:{iep['int']:.0f}% AFF:{iep['aff']:.0f}% ACT:{iep['act']:.0f}% | "
                                f"{iep['stance']} · {iep['tone']}", unsafe_allow_html=True)

# =============================================================================
# SESSION NOTES + ADDITIONAL DOCUMENT UPLOAD
# =============================================================================
st.markdown("---")
st.markdown("### 🎹 Session Notes & Documents")

notes_col, doc_col = st.columns([3, 2])

with notes_col:
    st.caption("Your private conductor notes — not sent to agents.")
    st.session_state.session_notes = st.text_area(
        "Notes", value=st.session_state.session_notes, height=140,
        placeholder="Key observations, decisions, follow-up actions...",
        label_visibility="collapsed"
    )

with doc_col:
    st.caption("Load a document into session context — agents will read it.")
    bottom_upload = st.file_uploader(
        "Upload document",
        type=["docx","md","txt","csv","py","pdf"],
        key="bottom_doc_uploader",
        label_visibility="collapsed"
    )
    if bottom_upload:
        doc_text = parse_uploaded_document(bottom_upload)
        st.session_state.session_document      = doc_text
        st.session_state.session_document_name = bottom_upload.name
        st.success(f"✅ {bottom_upload.name} loaded ({len(doc_text):,} chars)")
        st.rerun()
    if st.session_state.session_document:
        st.markdown(f'<div class="doc-context-box">📄 <strong>{st.session_state.session_document_name}</strong><br><span style="color:#666;">{len(st.session_state.session_document):,} chars — agents can read this</span></div>', unsafe_allow_html=True)
        with st.expander("👁️ Preview document content"):
            st.text(st.session_state.session_document[:1000] + ("..." if len(st.session_state.session_document) > 1000 else ""))
        if st.button("🗑️ Remove document", key="remove_doc_bottom"):
            st.session_state.session_document = None
            st.session_state.session_document_name = ""
            st.rerun()

st.markdown("---")
st.markdown("""
<div style="text-align:center; color:#666; padding:1rem;">
    <strong>Focus Group Lab V38</strong> — Research Edition<br>
    Multi-Agent AI Advisory Platform · Live IEP + Vₜ Scoring · Co-Conductor<br>
    SYNINT Team — April 2026 · Kouns, W.C.
</div>
""", unsafe_allow_html=True)
