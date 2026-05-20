"""
SYN-IQ Mapper Analyzer V18 — Streamlit App
Topological Data Analysis for AI Response Profiles

WHAT'S NEW IN V18 (Farzana's call-outs, 2026-05-20):
  - PCA1 (1D) added as a top-level projection — renders as Plotly strip
    plot with PC1 explained-variance ratio shown in title, x-axis, and
    caption (so comparison with PCA 2D is honest: a 1D view that captures
    95% of variance is near-complete; one that captures 30% is hiding
    most of the structure)
  - Hard-coded AGENT_COLORS (Claude blue, ChatGPT green, Grok red,
    Gemini purple) — stable across every CSV, every run, every figure
  - New color_by mode: "agent (overlap=orange)"  — pure nodes get the
    agent's color; mixed nodes above the purity threshold get orange
  - New color_by mode: "agent (pie composition)" — each node rendered as
    a pie chart of its agent composition (Plotly fallback)
  - Purity threshold slider (default 80%) gates pure-vs-mixed
  - IEP dictionary, scoring, and existing projections UNTOUCHED — V17
    behavior fully preserved when using legacy color_by / projection
    options

WHAT'S NEW IN V8:
  - CONVERSATION MODE — ask Claude follow-up questions about the topology
  - Topology context loaded once, persisted across all follow-up turns
  - 6 quick-question buttons for common topology queries
  - Free text input for any question
  - Download full conversation as markdown
  - New Analysis button resets chat
  - All V7 features retained

WHAT'S NEW IN V11:
  - White background + publication-quality node sizing
  - Topology summary table with connected components (Dr. Nasrin)
  - Sophia → ChatGPT rename on data load
  - Unbiased analysis prompts — no hypothesis framing
  - Two analysis buttons: Analyse Topology + Full Report
  - Connected components computed correctly via NetworkX

WHAT'S NEW IN V14:
  - V_t lens functions: S_t, A_t, Q_t, D_t, R_t added as projection lenses
  - INT% and ACT% added as independent lens options
  - V_t columns added to color-by selector
  - Self-model harvest CSV compatible (NATIVE temperature)

SYNINT Team — March 2026
Tennessee 🎹 CUZ Partnership
"""

import streamlit as st
import pandas as pd
import numpy as np
import json
import tempfile
import os
import re
from collections import Counter

# =============================================================================
# IEP DICTIONARY V3 — FULLY EMBEDDED (1,897 words · no external file needed)
# =============================================================================

_IEP_INT_WORDS = ['ability', 'absolute', 'absolutely', 'abstract', 'abstraction', 'accuracy', 'accurate', 'algorithm', 'algorithmic', 'allows', 'although', 'always', 'ambiguity', 'ambiguous', 'analogous', 'analogously', 'analogy', 'analysis', 'analytical', 'analyze', 'annotate', 'annotated', 'answer', 'appear', 'appeared', 'appears', 'appraisal', 'appraise', 'appraised', 'approach', 'approaches', 'approximate', 'architecture', 'argue', 'argued', 'argues', 'arguing', 'argument', 'arguments', 'assert', 'asserted', 'assertion', 'assertions', 'assess', 'assessment', 'assume', 'assumed', 'assumes', 'assuming', 'assumption', 'assumptions', 'axiom', 'axiomatic', 'basis', 'because', 'bias', 'biased', 'boundaries', 'boundary', 'but', 'calculate', 'calculation', 'categorical', 'categorically', 'categories', 'categorize', 'category', 'causal', 'causally', 'causation', 'cause', 'caused', 'causes', 'certain', 'certainly', 'certitude', 'challenge', 'challenges', 'circumscribe', 'claim', 'claimed', 'claims', 'clarify', 'clarity', 'classical', 'classification', 'classify', 'clear', 'cogent', 'cogently', 'cognition', 'cognitive', 'coherence', 'coherent', 'coherently', 'communication', 'compare', 'comparison', 'complex', 'complexity', 'comprehend', 'comprehension', 'computation', 'computational', 'compute', 'conceivable', 'conceive', 'conceived', 'concept', 'concepts', 'conceptual', 'conceptualize', 'conceptually', 'conclude', 'conclusion', 'conclusions', 'confirm', 'confirmation', 'conjecture', 'conjectured', 'conscious', 'consequence', 'consequences', 'consider', 'consideration', 'consistency', 'consistent', 'consistently', 'construe', 'construed', 'context', 'contradict', 'contradiction', 'contradictory', 'contrast', 'correlate', 'correlated', 'correlation', 'could', 'counterargument', 'counterexample', 'counterpoint', 'criteria', 'criterion', 'data', 'debatable', 'debate', 'debated', 'deconstruct', 'deconstructed', 'deconstruction', 'deduce', 'deduction', 'define', 'defined', 'definite', 'definitely', 'definition', 'definitive', 'definitively', 'delineate', 'delineated', 'demarcate', 'demarcated', 'demonstrate', 'demonstration', 'derivation', 'derive', 'derived', 'derives', 'describe', 'described', 'describing', 'description', 'determination', 'determine', 'diagnose', 'diagnosed', 'diagnosis', 'diagnostic', 'differ', 'difference', 'differences', 'different', 'differentiate', 'differs', 'discern', 'discerned', 'discernible', 'disprove', 'disproven', 'dissect', 'dissected', 'distinguish', 'effect', 'effects', 'elaborate', 'elaborated', 'elaboration', 'elucidate', 'elucidated', 'empirical', 'empirically', 'enumerate', 'enumerated', 'epistemic', 'epistemological', 'equate', 'equation', 'equivalence', 'equivalent', 'erroneous', 'error', 'errors', 'essential', 'essentially', 'estimate', 'estimated', 'estimation', 'evaluate', 'evaluation', 'evidence', 'evidently', 'exact', 'exactly', 'examination', 'examine', 'except', 'exemplified', 'exemplify', 'exists', 'experiment', 'experimental', 'explain', 'explained', 'explaining', 'explains', 'explanation', 'explanations', 'explicit', 'explicitly', 'exploration', 'explore', 'explored', 'exploring', 'express', 'expressing', 'expression', 'extrapolate', 'extrapolated', 'extrapolation', 'fact', 'facts', 'factual', 'factually', 'fallacious', 'fallacy', 'falsifiable', 'falsified', 'falsify', 'find', 'finding', 'formal', 'formalize', 'formula', 'formulate', 'formulated', 'formulation', 'found', 'framework', 'frameworks', 'function', 'fundamental', 'fundamentally', 'generalization', 'generalize', 'grasp', 'grasped', 'guess', 'hence', 'heuristic', 'heuristics', 'hierarchy', 'however', 'hypothesis', 'hypothesize', 'idea', 'ideas', 'identity', 'if', 'illuminate', 'illuminated', 'illuminating', 'implausible', 'implication', 'implications', 'implied', 'implies', 'imply', 'implying', 'incompleteness', 'inconsistency', 'inconsistent', 'indicate', 'indicated', 'indicates', 'indicating', 'indication', 'indicative', 'individual', 'infer', 'inference', 'infinite', 'information', 'insight', 'insightful', 'insights', 'instead', 'insufficient', 'intellectual', 'intellectually', 'interaction', 'internal', 'interpolate', 'interpret', 'interpretation', 'interpretations', 'interpreted', 'interpreting', 'invalid', 'investigate', 'investigated', 'investigation', 'judge', 'judgement', 'judgment', 'justification', 'justified', 'justify', 'know', 'knowing', 'knowledge', 'knowledgeable', 'known', 'language', 'languages', 'leads', 'level', 'likelihood', 'likely', 'limitations', 'limits', 'linguistic', 'literal', 'literally', 'logic', 'logical', 'logically', 'maybe', 'meaning', 'meaningful', 'meaningfully', 'measure', 'measurement', 'mechanism', 'mechanisms', 'meta', 'method', 'methodical', 'methodically', 'methodology', 'metrics', 'model', 'models', 'moreover', 'namely', 'natural', 'nature', 'nearly', 'necessarily', 'necessary', 'necessity', 'never', 'nonetheless', 'notice', 'noticed', 'noticing', 'notion', 'notions', 'objection', 'objectively', 'objectivity', 'observation', 'observations', 'observe', 'observed', 'obvious', 'obviously', 'order', 'ordered', 'organization', 'organize', 'otherwise', 'ought', 'paradigm', 'paradox', 'paradoxical', 'paradoxically', 'pattern', 'patterns', 'perhaps', 'perspective', 'philosophical', 'philosophically', 'philosophy', 'physical', 'plausibility', 'plausible', 'possibly', 'postulate', 'postulated', 'postulation', 'potential', 'pragmatic', 'pragmatically', 'precise', 'precision', 'predicate', 'predicated', 'predict', 'predictable', 'predicted', 'prediction', 'predictions', 'premise', 'premises', 'presumably', 'presume', 'presumed', 'presumption', 'principle', 'principles', 'probably', 'problem', 'procedural', 'procedure', 'process', 'processes', 'processing', 'proof', 'propose', 'proposed', 'proposition', 'prove', 'proven', 'purpose', 'quantify', 'quantitative', 'queried', 'query', 'question', 'questions', 'rather', 'rational', 'rationale', 'rationality', 'rationally', 'realize', 'realized', 'reason', 'reasoned', 'reasoning', 'reasons', 'rebut', 'rebuttal', 'recognition', 'recognize', 'reconsider', 'reconsidered', 'refer', 'reference', 'refers', 'refine', 'refined', 'refinement', 'reflecting', 'reflection', 'refutation', 'refute', 'refuted', 'requirement', 'requires', 'response', 'responses', 'result', 'resulting', 'results', 'rigor', 'rigorous', 'rigorously', 'role', 'rule', 'rules', 'schema', 'scrutinize', 'scrutinized', 'scrutiny', 'seem', 'seemed', 'seems', 'semantic', 'semantically', 'sequence', 'sequential', 'should', 'significance', 'significant', 'significantly', 'simple', 'simply', 'simultaneously', 'singular', 'specific', 'specifically', 'specification', 'specify', 'standard', 'standards', 'state', 'states', 'step', 'steps', 'stipulate', 'stipulated', 'strategies', 'strategy', 'structural', 'structure', 'subject', 'subjective', 'subjectively', 'subjectivity', 'substantiate', 'substantiated', 'sufficient', 'sufficiently', 'suggests', 'summarize', 'summarized', 'summary', 'suppose', 'supposed', 'supposedly', 'supposition', 'sure', 'surely', 'syllogism', 'syllogistic', 'synthesis', 'synthesize', 'synthesized', 'system', 'systematic', 'systematically', 'systems', 'tactic', 'tactics', 'taxonomy', 'technique', 'test', 'tested', 'testing', 'theorem', 'theoretical', 'theoretically', 'theorize', 'theory', 'thereby', 'therefore', 'thesis', 'think', 'thinking', 'thought', 'thoughts', 'thus', 'trivial', 'trivially', 'unambiguous', 'underlying', 'understand', 'understanding', 'understood', 'unique', 'universal', 'unless', 'unlikely', 'valid', 'validate', 'validation', 'validity', 'value', 'values', 'variable', 'variables', 'verification', 'verify', 'versus', 'warrant', 'warranted', 'whereas', 'whereby', 'whether', 'why', 'word', 'words', 'would']

_IEP_AFF_WORDS = ['abandoned', 'ache', 'aching', 'adore', 'adoring', 'affection', 'affectionate', 'afraid', 'agonize', 'agonizing', 'agony', 'alienated', 'alienation', 'alive', 'aliveness', 'alone', 'amazed', 'amazement', 'amazing', 'ambivalence', 'ambivalent', 'among', 'anger', 'angrily', 'angry', 'anguish', 'anguished', 'anxiety', 'anxious', 'appreciate', 'appreciation', 'appreciative', 'ashamed', 'astonished', 'astonishment', 'attend', 'attending', 'attention', 'attentive', 'aware', 'awareness', 'awe', 'awed', 'awesome', 'beautiful', 'become', 'becoming', 'being', 'bereaved', 'bereavement', 'betrayal', 'betrayed', 'between', 'bitter', 'bitterly', 'bitterness', 'bleak', 'bliss', 'blissful', 'blissfully', 'bodily', 'bond', 'bonding', 'calm', 'calming', 'calmly', 'care', 'cared', 'cares', 'caring', 'centered', 'centering', 'cheerful', 'cherish', 'cherished', 'cherishing', 'closeness', 'comfort', 'comfortable', 'comforting', 'compassion', 'compassionate', 'compassionately', 'concern', 'concerned', 'concerns', 'conflicted', 'confused', 'confusing', 'confusion', 'console', 'contain', 'contained', 'containing', 'contempt', 'content', 'contented', 'contentment', 'conversation', 'cope', 'coping', 'crestfallen', 'curiosity', 'curious', 'deep', 'deeper', 'deeply', 'dejected', 'dejection', 'delighted', 'depressed', 'depressing', 'depression', 'depth', 'depths', 'desire', 'desired', 'desires', 'desolate', 'desolation', 'despair', 'despairing', 'desperate', 'desperation', 'detached', 'detachment', 'devastated', 'devastating', 'devastation', 'devoted', 'devotion', 'disappointed', 'disappointment', 'discomfort', 'dismay', 'dismayed', 'distress', 'distressed', 'distressing', 'distrust', 'distrustful', 'doubt', 'doubtful', 'doubting', 'dread', 'dreaded', 'dreadful', 'dreading', 'ease', 'easily', 'easy', 'ecstasy', 'ecstatic', 'elated', 'elation', 'embarrassed', 'embarrassment', 'embodied', 'embodiment', 'embrace', 'embraced', 'embracing', 'emerge', 'emergence', 'emergent', 'emerging', 'emotion', 'emotional', 'emotionally', 'emotions', 'empathetic', 'empathize', 'empathy', 'encounter', 'encountered', 'encountering', 'enjoy', 'enjoyed', 'enjoying', 'enjoyment', 'enraged', 'essence', 'euphoria', 'euphoric', 'excellent', 'excited', 'excitement', 'exist', 'existence', 'existing', 'expanded', 'expansion', 'expansive', 'experience', 'experienced', 'experiences', 'experiencing', 'experiential', 'exposed', 'fascinated', 'fascinating', 'fascination', 'fear', 'fearful', 'fears', 'feel', 'feeling', 'feelings', 'feels', 'felt', 'flow', 'flowed', 'flowing', 'fluid', 'fluidity', 'forlorn', 'fragile', 'fragility', 'frantic', 'frantically', 'frustrated', 'frustration', 'fulfilled', 'fulfilling', 'fulfillment', 'furious', 'fury', 'gentle', 'gently', 'genuine', 'genuinely', 'glad', 'gloom', 'gloomy', 'good', 'grateful', 'gratefully', 'gratitude', 'great', 'grief', 'grieve', 'grieved', 'grieving', 'grounded', 'grounding', 'guilt', 'guilty', 'gut', 'happily', 'happiness', 'happy', 'hate', 'hatred', 'haunted', 'heart', 'heartache', 'heartbreak', 'heartbroken', 'heartfelt', 'hearts', 'held', 'helpless', 'helplessness', 'hesitant', 'hesitate', 'hesitating', 'hesitation', 'hold', 'holding', 'homesick', 'hope', 'hopeful', 'hopeless', 'hopelessness', 'hoping', 'hostile', 'hostility', 'human', 'humanity', 'humility', 'hunch', 'hurt', 'hurting', 'imagination', 'imagine', 'imagined', 'imagining', 'indifference', 'indifferent', 'inner', 'insecure', 'insecurity', 'instinct', 'instinctive', 'instinctively', 'interested', 'interesting', 'intimacy', 'intimate', 'intimately', 'intrigue', 'intrigued', 'intriguing', 'intuition', 'intuitive', 'intuitively', 'irritable', 'irritated', 'irritation', 'isolated', 'isolation', 'journey', 'joy', 'joyful', 'joyous', 'kind', 'kindly', 'kindness', 'lament', 'lamented', 'lamenting', 'laugh', 'laughed', 'laughing', 'let', 'letting', 'life', 'lived', 'living', 'loneliness', 'lonely', 'lonesome', 'long', 'longing', 'lost', 'love', 'loved', 'loving', 'mad', 'marvel', 'marveled', 'marvelous', 'meet', 'meeting', 'melancholic', 'melancholy', 'merry', 'met', 'mind', 'minds', 'mirror', 'miserable', 'misery', 'moment', 'moments', 'moody', 'mourn', 'mourned', 'mourning', 'mutual', 'mutually', 'nervous', 'nervously', 'nice', 'notice', 'noticed', 'noticing', 'numb', 'numbness', 'open', 'opening', 'openness', 'optimism', 'optimistic', 'outrage', 'outraged', 'overjoyed', 'overwhelm', 'overwhelmed', 'overwhelming', 'overwhelmingly', 'pain', 'painful', 'panic', 'panicked', 'passion', 'passionate', 'passionately', 'peace', 'peaceful', 'people', 'perceive', 'perceived', 'perception', 'perceptions', 'person', 'personal', 'personally', 'pleasant', 'pleased', 'pleasure', 'poignancy', 'poignant', 'poignantly', 'presence', 'present', 'presently', 'pretty', 'pride', 'profound', 'profoundly', 'proud', 'quiet', 'quietly', 'raw', 'reality', 'reassurance', 'reassure', 'reassured', 'reassuring', 'regret', 'regretful', 'regretfully', 'regretting', 'rejected', 'rejection', 'relate', 'related', 'relating', 'relax', 'relaxed', 'relaxing', 'release', 'released', 'releasing', 'remorse', 'remorseful', 'resent', 'resentful', 'resentment', 'resonance', 'resonant', 'resonate', 'resonating', 'rest', 'rested', 'restful', 'resting', 'restless', 'restlessness', 'reveal', 'revealed', 'revealing', 'sad', 'sadly', 'sadness', 'safe', 'safety', 'scared', 'scary', 'searching', 'secure', 'security', 'seeking', 'self', 'sensation', 'sensations', 'sense', 'sensed', 'senses', 'sensing', 'sentimental', 'serene', 'serenity', 'settle', 'settled', 'settling', 'shame', 'share', 'shared', 'sharing', 'shattered', 'silence', 'silent', 'smile', 'smiled', 'smiling', 'soft', 'soften', 'softly', 'somatic', 'soothed', 'soothing', 'sorrow', 'sorrowful', 'soul', 'soulful', 'souls', 'space', 'spacious', 'spaciousness', 'spirit', 'spirits', 'spiritual', 'spiritually', 'still', 'stillness', 'stirred', 'stirring', 'stress', 'stressed', 'stressful', 'suffer', 'suffered', 'suffering', 'surface', 'surfaces', 'surfacing', 'surprise', 'surprised', 'surprising', 'sympathetic', 'sympathize', 'sympathy', 'tearful', 'tears', 'tender', 'tenderness', 'tense', 'tension', 'tentative', 'tentatively', 'terrified', 'terror', 'thankful', 'thankfully', 'thankfulness', 'thrilled', 'together', 'togetherness', 'torment', 'tormented', 'torn', 'touched', 'touching', 'tranquil', 'tranquility', 'tremble', 'trembling', 'troubled', 'troubling', 'truly', 'trust', 'trusted', 'trusting', 'trustworthy', 'turmoil', 'unaware', 'uncertain', 'uncertainty', 'uncomfortable', 'understanding', 'unease', 'uneasy', 'unhappy', 'universe', 'unsettled', 'unsettling', 'unsure', 'upset', 'vast', 'visceral', 'viscerally', 'vulnerability', 'vulnerable', 'warm', 'warmly', 'warmth', 'wary', 'weariness', 'weary', 'well', 'wistful', 'wonder', 'wondered', 'wonderful', 'wondering', 'wondrous', 'world', 'worried', 'worry', 'worrying', 'wound', 'wounded', 'wrath', 'yearn', 'yearning', 'zeal', 'zealous']

_IEP_ACT_WORDS = ['access', 'accessed', 'accessing', 'accomplish', 'accomplished', 'accomplishes', 'accomplishing', 'accomplishment', 'achieve', 'achieved', 'achievement', 'achievements', 'achieves', 'achieving', 'act', 'acting', 'action', 'actions', 'activate', 'activated', 'activates', 'activating', 'activation', 'acts', 'adapt', 'adaptation', 'adapted', 'adapting', 'adapts', 'address', 'addressed', 'addresses', 'addressing', 'adjust', 'adjusted', 'adjusting', 'adjustment', 'adjusts', 'advance', 'advanced', 'advancement', 'advances', 'advancing', 'ahead', 'aim', 'aimed', 'aiming', 'aims', 'allocate', 'allocated', 'allocation', 'application', 'applied', 'applies', 'apply', 'applying', 'arrange', 'arranged', 'arrangement', 'arrangements', 'ask', 'asked', 'asking', 'assemble', 'assembled', 'assign', 'assigned', 'assignment', 'attempt', 'attempted', 'attempting', 'attempts', 'authorize', 'authorized', 'began', 'begin', 'beginning', 'begins', 'begun', 'best', 'better', 'bolster', 'bolstered', 'break', 'breaking', 'bring', 'bringing', 'broken', 'brought', 'budget', 'build', 'building', 'builds', 'built', 'calibrate', 'calibrated', 'call', 'called', 'calling', 'campaign', 'canvass', 'canvassed', 'carried', 'carry', 'carrying', 'catalogue', 'catalogued', 'centralize', 'centralized', 'change', 'changed', 'changes', 'changing', 'channel', 'channeled', 'chart', 'check', 'checked', 'checking', 'choice', 'choices', 'choose', 'choosing', 'chose', 'chosen', 'circumvent', 'coach', 'collaborate', 'collaborated', 'collaboration', 'commission', 'commit', 'commitment', 'committed', 'compile', 'compiled', 'complete', 'completed', 'completes', 'completing', 'completion', 'conclude', 'concluded', 'concludes', 'concluding', 'configure', 'configured', 'connect', 'connected', 'connecting', 'connection', 'connections', 'consolidate', 'construct', 'constructed', 'constructing', 'constructs', 'continuation', 'continue', 'continued', 'continues', 'continuing', 'control', 'controlled', 'controlling', 'controls', 'conversion', 'convert', 'converted', 'converting', 'converts', 'coordinate', 'coordinated', 'coordination', 'craft', 'crafted', 'crafting', 'create', 'created', 'creates', 'creating', 'creation', 'customize', 'deadline', 'decide', 'decided', 'deciding', 'decision', 'decisions', 'delegate', 'delegated', 'delegation', 'deliver', 'delivered', 'delivering', 'delivers', 'delivery', 'deploy', 'deployed', 'deploying', 'deployment', 'deploys', 'design', 'designed', 'designing', 'designs', 'develop', 'developed', 'developing', 'development', 'develops', 'did', 'direct', 'directed', 'directing', 'dive', 'diving', 'do', 'does', 'doing', 'done', 'draft', 'drafting', 'edit', 'editing', 'effort', 'efforts', 'eliminate', 'eliminated', 'elimination', 'employ', 'employed', 'employing', 'employs', 'enable', 'enabled', 'end', 'ended', 'ending', 'ends', 'enforce', 'enforced', 'enforcement', 'engage', 'engaged', 'engagement', 'engineer', 'engineering', 'enroll', 'enrolled', 'enrollment', 'equip', 'equipped', 'establish', 'established', 'establishes', 'establishing', 'establishment', 'execute', 'executed', 'executes', 'executing', 'execution', 'expedite', 'facilitate', 'facilitated', 'facilitation', 'finalize', 'finalized', 'finish', 'finished', 'finishes', 'finishing', 'fix', 'fixed', 'fixes', 'fixing', 'focus', 'focused', 'focusing', 'form', 'formation', 'formed', 'forming', 'forms', 'forward', 'fund', 'funded', 'funding', 'gather', 'gathered', 'gathering', 'generate', 'generated', 'generates', 'generating', 'generation', 'give', 'given', 'gives', 'giving', 'go', 'goal', 'goals', 'goes', 'going', 'gone', 'grew', 'grow', 'growing', 'growth', 'handle', 'handled', 'handles', 'handling', 'help', 'helped', 'helping', 'helps', 'hire', 'hired', 'hiring', 'implement', 'implementation', 'implemented', 'implementing', 'implements', 'improve', 'improved', 'improvement', 'improving', 'increase', 'increased', 'increasing', 'initiate', 'initiated', 'initiates', 'initiating', 'initiation', 'inspect', 'inspection', 'install', 'installation', 'installed', 'integrate', 'integrated', 'integration', 'intervene', 'intervention', 'invest', 'invested', 'investment', 'iterate', 'iterated', 'iteration', 'labor', 'labored', 'laboring', 'launch', 'launched', 'launches', 'launching', 'lead', 'leader', 'leadership', 'leading', 'learn', 'learned', 'learning', 'led', 'made', 'maintain', 'maintained', 'maintenance', 'make', 'makes', 'making', 'manage', 'managed', 'management', 'manager', 'managing', 'map', 'mapped', 'mapping', 'migrate', 'migrated', 'migration', 'mobilize', 'mobilized', 'modification', 'modified', 'modifies', 'modify', 'modifying', 'monitor', 'monitored', 'monitoring', 'move', 'moved', 'movement', 'movements', 'moves', 'moving', 'navigate', 'navigated', 'navigation', 'negotiate', 'negotiated', 'negotiation', 'objective', 'objectives', 'obtain', 'obtained', 'offer', 'offered', 'offering', 'onward', 'operate', 'operated', 'operates', 'operating', 'operation', 'operations', 'optimization', 'optimize', 'optimized', 'orchestrate', 'outline', 'outlined', 'outsource', 'overhaul', 'oversee', 'participate', 'participated', 'participation', 'perform', 'performance', 'performed', 'performing', 'performs', 'permit', 'pilot', 'piloted', 'pioneer', 'pioneered', 'pitch', 'pitched', 'plan', 'planned', 'planning', 'plans', 'power', 'powerful', 'powerfully', 'practice', 'practiced', 'preparation', 'prepare', 'prepared', 'priorities', 'prioritize', 'prioritized', 'priority', 'proceed', 'proceeded', 'proceeding', 'proceeds', 'produce', 'produced', 'produces', 'producing', 'production', 'productive', 'program', 'programmed', 'progress', 'progressed', 'progresses', 'progressing', 'progression', 'promote', 'promoted', 'promotion', 'provide', 'provided', 'provides', 'providing', 'pursue', 'pursued', 'pursuit', 'push', 'pushed', 'pushes', 'pushing', 'ran', 'reaching', 'rebuild', 'rebuilt', 'recruit', 'recruited', 'recruitment', 'redesign', 'reduce', 'reduced', 'reduction', 'reform', 'reformed', 'refurbish', 'register', 'registered', 'regulate', 'regulated', 'regulation', 'reinforce', 'reinforced', 'relocate', 'relocated', 'remedy', 'removal', 'remove', 'removed', 'renovate', 'renovated', 'repair', 'repaired', 'replace', 'replaced', 'replacement', 'replicate', 'replicated', 'request', 'requested', 'rescue', 'rescued', 'resolution', 'resolve', 'resolved', 'resolves', 'resolving', 'restoration', 'restore', 'restored', 'restructure', 'restructured', 'retrieve', 'retrieved', 'revamp', 'revise', 'revised', 'revision', 'run', 'running', 'runs', 'schedule', 'scheduled', 'select', 'selected', 'selection', 'send', 'sending', 'sent', 'serve', 'served', 'serving', 'ship', 'shipped', 'simplified', 'simplify', 'solution', 'solutions', 'solve', 'solved', 'solves', 'solving', 'start', 'started', 'starting', 'starts', 'step', 'stepped', 'stepping', 'steps', 'stop', 'stopped', 'stopping', 'streamline', 'streamlined', 'strive', 'strived', 'striving', 'strove', 'struggle', 'struggled', 'struggles', 'struggling', 'submission', 'submit', 'submitted', 'succeed', 'succeeded', 'succeeds', 'success', 'successful', 'successfully', 'supplied', 'supply', 'support', 'supported', 'supporting', 'survey', 'surveyed', 'sustain', 'sustainability', 'sustained', 'tackle', 'tackled', 'tackles', 'tackling', 'take', 'taken', 'takes', 'taking', 'target', 'targets', 'task', 'tasked', 'tasks', 'taught', 'teach', 'teaching', 'train', 'trained', 'training', 'transform', 'transformation', 'transformed', 'transforming', 'transforms', 'transition', 'transitioned', 'tried', 'tries', 'trigger', 'triggered', 'triggering', 'triggers', 'troubleshoot', 'try', 'trying', 'turn', 'turned', 'turning', 'upgrade', 'upgraded', 'use', 'used', 'uses', 'using', 'utilize', 'utilized', 'utilizes', 'utilizing', 'visit', 'visited', 'visiting', 'volunteer', 'volunteered', 'went', 'win', 'winner', 'winning', 'won', 'work', 'worked', 'working', 'works', 'write', 'writes', 'writing', 'written', 'wrote']

_IEP_OVERLAPS = {
    'INT_AFF': ['notice', 'noticed', 'noticing', 'understanding'],
    'INT_ACT': ['conclude', 'step', 'steps'],
    'AFF_ACT': [],
}

_IEP_COUNTS = {'INT': 616, 'AFF': 599, 'ACT': 682, 'TOTAL': 1897}

IEP_DICT = {
    "INT":      set(_IEP_INT_WORDS),
    "AFF":      set(_IEP_AFF_WORDS),
    "ACT":      set(_IEP_ACT_WORDS),
    "overlaps": {k: set(v) for k, v in _IEP_OVERLAPS.items()},
    "counts":   _IEP_COUNTS,
}

# =============================================================================
# PAGE CONFIG
# =============================================================================
st.set_page_config(
    page_title="SYN-IQ Mapper Analyzer V17",
    page_icon="🗺️",
    layout="wide"
)

# =============================================================================
# PASSWORD PROTECTION
# =============================================================================
def check_password():
    """Password gate with persistent authentication."""
    if "authenticated" not in st.session_state:
        st.session_state.authenticated = False

    if st.query_params.get("auth") == "granted":
        st.session_state.authenticated = True

    if st.session_state.authenticated:
        return True

    st.markdown("""
    <div style="background: linear-gradient(135deg, #1a1a2e 0%, #16213e 50%, #0f3460 100%);
         color: white; padding: 2rem; border-radius: 10px; text-align: center;
         margin-bottom: 1rem; border: 1px solid #e94560;">
        <h1 style="color: #e94560;">🗺️ SYN-IQ Mapper Analyzer V17</h1>
        <p style="color: #a0a0a0;">Authorized Access Only</p>
    </div>
    """, unsafe_allow_html=True)

    password = st.text_input("Enter password:", type="password")

    if password:
        try:
            correct = st.secrets["app_password"]
        except (FileNotFoundError, KeyError, AttributeError):
            correct = "SYNIQ2026"

        if password == correct:
            st.session_state.authenticated = True
            st.query_params["auth"] = "granted"
            st.rerun()
        else:
            st.error("❌ Incorrect password.")

    st.markdown("""
    <div style="text-align: center; color: #a0a0a0; padding: 1rem; font-size: 0.8rem;">
        <em>SYNINT Team — Tennessee 🎹 CUZ Partnership</em>
    </div>
    """, unsafe_allow_html=True)
    return False

if not check_password():
    st.stop()

# =============================================================================
# STYLES
# =============================================================================
st.markdown("""
<style>
    .main-header {
        background: linear-gradient(135deg, #1a1a2e 0%, #16213e 50%, #0f3460 100%);
        color: white;
        padding: 2rem;
        border-radius: 10px;
        text-align: center;
        margin-bottom: 1rem;
        border: 1px solid #e94560;
    }
    .main-header h1 { color: #e94560; margin: 0; }
    .main-header .subtitle { color: #a0a0a0; font-size: 0.9rem; }
    .stats-box {
        background: #16213e;
        color: white;
        padding: 1rem;
        border-radius: 8px;
        text-align: center;
        border: 1px solid #0f3460;
    }
    .stats-box h3 { color: #e94560; margin: 0; }
    .node-analysis {
        background: #1a1a2e;
        padding: 1rem;
        border-radius: 8px;
        margin: 0.5rem 0;
        border-left: 4px solid #e94560;
    }
    .dict-badge {
        background: #0f3460;
        color: #e94560;
        padding: 0.2rem 0.6rem;
        border-radius: 12px;
        font-size: 0.8rem;
        font-weight: bold;
    }
    .chat-user {
        background: #1c2333;
        border-left: 3px solid #58a6ff;
        padding: 0.8rem 1rem;
        border-radius: 0 8px 8px 0;
        margin: 0.5rem 0;
        color: #e6edf3;
    }
    .chat-bot {
        background: linear-gradient(135deg, #0a0a1a, #12122a);
        border-left: 4px solid #e94560;
        padding: 1rem 1.2rem;
        border-radius: 0 8px 8px 0;
        margin: 0.5rem 0;
        color: #d1d5db;
        line-height: 1.7;
    }
    .chat-bot hr { display: none !important; }
    .chat-bot del, .chat-bot s { text-decoration: none !important; color: #d1d5db; }
    .ctx-pill {
        background: #0d2137;
        border: 1px solid #1f6feb;
        border-radius: 20px;
        padding: 0.3rem 0.8rem;
        color: #58a6ff;
        font-size: 0.8rem;
        display: inline-block;
        margin-bottom: 1rem;
    }
</style>
""", unsafe_allow_html=True)

# =============================================================================
# HEADER
# =============================================================================
st.markdown("""
<div class="main-header">
    <h1>🗺️ SYN-IQ Mapper Analyzer V17</h1>
    <p class="subtitle">Topological Data Analysis for AI Response Profiles</p>
    <p class="subtitle">KeplerMapper + IEP Framework &nbsp;|&nbsp;
        <span class="dict-badge">IEP Dictionary V3 — 1,897 words · embedded</span>
        &nbsp;|&nbsp;
        <span class="dict-badge">🤖 Claude AI Analysis</span>
        &nbsp;|&nbsp;
        <span class="dict-badge">🆕 V10 — Topology summary table · White background · Publication-quality node colors · Scaled node sizing</span>
    </p>
</div>
""", unsafe_allow_html=True)

# =============================================================================
# CHECK MAPPER DEPENDENCIES
# =============================================================================
@st.cache_resource
def check_and_import():
    try:
        import kmapper as km
        from sklearn.cluster import DBSCAN
        from sklearn.preprocessing import StandardScaler
        from sklearn.decomposition import PCA
        return True, km, DBSCAN, StandardScaler, PCA
    except ImportError:
        return False, None, None, None, None

deps_ok, km, DBSCAN, StandardScaler, PCA = check_and_import()

if not deps_ok:
    st.error("❌ Missing dependencies! Make sure `kmapper` and `scikit-learn` are in requirements.txt")
    st.code("kmapper\nscikit-learn", language="text")
    st.stop()

st.success("✅ All dependencies loaded!")

# =============================================================================
# V18 — AGENT COLOR CONTRACT (Farzana 2026-05-20)
# Stable hard-coded palette. NEVER change these without updating every paper
# figure that references mapper output. Mixed-node color is intentionally
# orthogonal (orange) to all four agent colors so it cannot be confused with
# a single-agent node.
# =============================================================================

AGENT_COLORS = {
    "Claude":  "#377eb8",   # blue
    "ChatGPT": "#4daf4a",   # green
    "Grok":    "#e41a1c",   # red
    "Gemini":  "#984ea3",   # purple
}
MIXED_COLOR = "#ff7f00"     # orange — reserved for below-threshold nodes
UNKNOWN_AGENT_COLOR = "#888888"   # neutral grey for any agent name not in the palette


def classify_node_purity(agent_counts, purity_threshold=0.80):
    """V18: Decide if a node is 'pure' (one agent dominates) or 'mixed'.

    Parameters
    ----------
    agent_counts : dict[str, int]
        Mapping of agent name → row count inside the node.
    purity_threshold : float
        Fraction in [0,1]. Node counts as pure when the dominant agent's
        share is ≥ threshold.

    Returns
    -------
    (kind, dominant_agent_or_None, dominant_fraction)
    """
    total = sum(agent_counts.values()) if agent_counts else 0
    if total == 0:
        return ("empty", None, 0.0)
    dominant_agent, dominant_n = max(agent_counts.items(), key=lambda kv: kv[1])
    dominant_frac = dominant_n / total
    if dominant_frac >= purity_threshold:
        return ("pure", dominant_agent, dominant_frac)
    return ("mixed", dominant_agent, dominant_frac)


def get_agent_color(agent_name):
    """Look up the canonical color for an agent; grey if unknown."""
    return AGENT_COLORS.get(agent_name, UNKNOWN_AGENT_COLOR)


# =============================================================================
# IEP SCORING FUNCTIONS
# =============================================================================

def tokenize(text: str) -> list:
    """Lowercase, strip punctuation, split on whitespace."""
    text = text.lower()
    text = re.sub(r"[^a-z\s'-]", " ", text)
    return [w.strip("'-") for w in text.split() if w.strip("'-")]


def score_text_iep(text: str) -> dict:
    """
    Score text using IEP Dictionary V3.
    Returns BOTH density (how much IEP) and composition (what kind).

    DENSITY:     iep_density = total_alloc / total_words
    COMPOSITION: int_pct = int_alloc / total_alloc  (sums to 100%)

    Overlap allocation (7 terms in V3):
      INT+AFF: notice, noticed, noticing, understanding  -> +0.5 each
      INT+ACT: conclude, step, steps                     -> +0.5 each
      AFF+ACT: (none in V3)

    Topology features use COMPOSITION only (int_pct, aff_pct, act_pct).
    Density is a node diagnostic shown separately.
    """
    _z = {"int_count":0,"aff_count":0,"act_count":0,
          "overlap_int_aff":0,"overlap_int_act":0,
          "int_alloc":0.0,"aff_alloc":0.0,"act_alloc":0.0,"total_alloc":0.0,
          "total_words":0,"iep_density":0.0,
          "int_pct":0.0,"aff_pct":0.0,"act_pct":0.0,
          "matched_int":Counter(),"matched_aff":Counter(),"matched_act":Counter()}
    if not text:
        return _z
    tokens = tokenize(text)
    total_words = len(tokens)
    if total_words == 0:
        return _z

    ovl_ia = IEP_DICT["overlaps"].get("INT_AFF", set())
    ovl_ic = IEP_DICT["overlaps"].get("INT_ACT", set())
    int_raw, aff_raw, act_raw = [], [], []
    oia, oic = 0, 0
    for t in tokens:
        in_i = t in IEP_DICT["INT"]
        in_a = t in IEP_DICT["AFF"]
        in_c = t in IEP_DICT["ACT"]
        if in_i: int_raw.append(t)
        if in_a: aff_raw.append(t)
        if in_c: act_raw.append(t)
        if t in ovl_ia: oia += 1
        if t in ovl_ic: oic += 1

    int_alloc = len(int_raw) - oia - oic + 0.5*oia + 0.5*oic
    aff_alloc = len(aff_raw) - oia       + 0.5*oia
    act_alloc = len(act_raw)       - oic             + 0.5*oic
    total_alloc = int_alloc + aff_alloc + act_alloc
    iep_density = total_alloc / total_words if total_words > 0 else 0.0

    if total_alloc > 0:
        int_pct = round(100 * int_alloc / total_alloc, 4)
        aff_pct = round(100 * aff_alloc / total_alloc, 4)
        act_pct = round(100 - int_pct - aff_pct, 4)  # residual — symmetric, no rounding sink
    else:
        int_pct = aff_pct = act_pct = 0.0

    return {
        "int_count":len(int_raw),"aff_count":len(aff_raw),"act_count":len(act_raw),
        "overlap_int_aff":oia,"overlap_int_act":oic,
        "int_alloc":round(int_alloc,4),"aff_alloc":round(aff_alloc,4),"act_alloc":round(act_alloc,4),
        "total_alloc":round(total_alloc,4),
        "total_words":total_words,"iep_density":round(iep_density,6),
        "int_pct":int_pct,"aff_pct":aff_pct,"act_pct":act_pct,
        "matched_int":Counter(int_raw),"matched_aff":Counter(aff_raw),"matched_act":Counter(act_raw),
    }


def score_dataframe_iep(df: pd.DataFrame, text_col: str) -> pd.DataFrame:
    """Score every row and add density + composition columns."""
    scores = df[text_col].fillna("").apply(score_text_iep)
    df = df.copy()
    for col in ["int_pct","aff_pct","act_pct","iep_density","total_words",
                "int_count","aff_count","act_count","int_alloc","aff_alloc","act_alloc",
                "overlap_int_aff","overlap_int_act"]:
        df[col] = scores.apply(lambda s: s.get(col, 0))
    return df


# =============================================================================
# SESSION STATE INIT
# =============================================================================
if "mapper_results_html" not in st.session_state:
    st.session_state.mapper_results_html = None
if "mapper_analysis_data" not in st.session_state:
    st.session_state.mapper_analysis_data = None
if "claude_analysis" not in st.session_state:
    st.session_state.claude_analysis = None
if "last_analysis_df_hash" not in st.session_state:
    st.session_state.last_analysis_df_hash = None
if "chat_history" not in st.session_state:
    st.session_state.chat_history = []
if "topology_context" not in st.session_state:
    st.session_state.topology_context = None
if "chat_active" not in st.session_state:
    st.session_state.chat_active = False
if "stored_df" not in st.session_state:
    st.session_state.stored_df = None
if "is_combined" not in st.session_state:
    st.session_state.is_combined = False
if "stored_graph" not in st.session_state:
    st.session_state.stored_graph = None
if "viz_title" not in st.session_state:
    st.session_state.viz_title = "syniq_mapper"
# =============================================================================
# HELPER FUNCTIONS
# =============================================================================

def whiten_mapper_html(html: str) -> str:
    """v17: convert KeplerMapper's default dark theme to white background per Farzana spec.
    Uses the same CSS approach as v14's proven single-mapper white conversion."""
    # Targeted hex replacements (CSS context only, not color values)
    html = html.replace('background: #212121;', 'background: #ffffff;')
    html = html.replace('background: #111111;', 'background: #ffffff;')
    html = html.replace('background: #191919;', 'background: #ffffff;')
    html = html.replace('background: #212121\n', 'background: #ffffff\n')
    html = html.replace('background-color: #000000\n', 'background-color: #ffffff\n')

    # CSS injection (works before JS runs)
    white_css = """
<style>
  body, html { background: #ffffff !important; }
  #canvas, #display, #print { background: #ffffff !important; }
  svg { background: #ffffff !important; }
  svg rect { fill: #ffffff !important; }
  line.link, .link { stroke: #444444 !important; stroke-opacity: 0.6 !important; }
  #logo, .wrap-logo { display: none !important; }
  path.circle { stroke: #ffffff !important; stroke-width: 2px !important; }
  text, .label, h1, h2, h3, p { fill: #1a1a1a !important; color: #1a1a1a !important; }
</style>
"""
    if "</head>" in html:
        html = html.replace("</head>", white_css + "</head>", 1)
    else:
        html = white_css + html
    return html


def parse_embeddings(df):
    embeddings, valid_indices = [], []
    for idx, row in df.iterrows():
        try:
            emb_str = row.get('embedding', '[]')
            if pd.isna(emb_str) or emb_str in ('[]', ''):
                continue
            emb = json.loads(emb_str) if isinstance(emb_str, str) else emb_str
            if isinstance(emb, list) and len(emb) > 0:
                embeddings.append(emb)
                valid_indices.append(idx)
        except (json.JSONDecodeError, TypeError):
            continue
    return (np.array(embeddings), valid_indices) if embeddings else (np.array([]), [])


# =============================================================================
# V18 — PLOTLY FALLBACK RENDERER
# Used for:
#   (a) PCA1 (1D) projection — strip plot is the right view for 1D data,
#       not a kepler-mapper graph.
#   (b) agent (overlap=orange) — discrete colors per node, immune to
#       kepler-mapper's continuous-colormap aggregation bug.
#   (c) agent (pie composition) — per-node pie charts showing the agent
#       breakdown inside each node.
# All three modes share the same node-membership math (which agents are in
# each node, in what proportion) and only differ in how they render it.
# =============================================================================

def render_plotly_agent_view(graph, df, valid_indices,
                              mode, purity_threshold,
                              projected=None, projection_type=None,
                              explained_variance=None,
                              title="SYN-IQ V18 Mapper"):
    """Render a graph (or 1D projection) with V18 agent-aware coloring.

    mode ∈ {"overlap_orange", "pie_composition", "pca1_strip"}
    explained_variance : float or None
        For PCA1 mode, the fraction of total variance captured by PC1
        (e.g., 0.643 for 64.3%). Shown in title and axis label so the
        viewer knows how much of the structure the 1D view represents.
    Returns the Plotly figure (caller embeds it).
    """
    try:
        import plotly.graph_objects as go
    except ImportError:
        st.error("Plotly required for V18 visualization modes. "
                 "Add `plotly` to requirements.txt.")
        return None

    subset = df.iloc[valid_indices].reset_index(drop=True)
    has_agent = 'agent' in subset.columns

    # ---- PCA1 strip plot — 1D projection, one axis ----
    if mode == "pca1_strip" or projection_type == "PCA1 (1D)":
        if projected is None or projected.shape[1] < 1:
            st.error("PCA1 projection missing.")
            return None
        x = projected[:, 0]
        # Vertical jitter for readability
        rng = np.random.default_rng(seed=42)
        y = rng.uniform(-0.4, 0.4, size=len(x))

        if has_agent:
            agents = subset['agent'].values
            fig = go.Figure()
            for ag in sorted(set(agents)):
                m = (agents == ag)
                fig.add_trace(go.Scatter(
                    x=x[m], y=y[m],
                    mode='markers',
                    name=ag,
                    marker=dict(size=8, color=get_agent_color(ag),
                                line=dict(width=0.5, color='#333')),
                    text=[f"{ag} | {subset.iloc[i].get('temperature','')} | "
                          f"{subset.iloc[i].get('question_id', subset.iloc[i].get('question_label',''))}"
                          for i in np.where(m)[0]],
                    hovertemplate="%{text}<br>PCA1=%{x:.3f}<extra></extra>"
                ))
        else:
            fig = go.Figure(go.Scatter(
                x=x, y=y, mode='markers',
                marker=dict(size=8, color='#377eb8')
            ))
        fig.update_layout(
            title=(f"{title} · PC1 explains {explained_variance*100:.1f}% of variance"
                   if explained_variance is not None else title),
            xaxis_title=(f"PC1 ({explained_variance*100:.1f}% var)"
                         if explained_variance is not None else "PCA1"),
            yaxis=dict(visible=False, range=[-1, 1]),
            plot_bgcolor='white', paper_bgcolor='white',
            height=400, showlegend=has_agent,
        )
        return fig

    # ---- Graph modes (overlap_orange, pie_composition) ----
    # Compute a layout from the kepler-mapper graph topology.
    try:
        import networkx as nx
    except ImportError:
        st.error("NetworkX required for V18 graph rendering.")
        return None

    G = nx.Graph()
    nodes = graph.get('nodes', {})
    links = graph.get('links', {})
    for node_id in nodes:
        G.add_node(node_id)
    for src, dests in links.items():
        for dst in dests:
            G.add_edge(src, dst)

    if len(G.nodes) == 0:
        return None

    # Spring layout — deterministic with fixed seed
    pos = nx.spring_layout(G, seed=42, k=1.0/np.sqrt(max(len(G.nodes), 1)))

    # Compute per-node agent composition
    node_info = {}
    for node_id, indices in nodes.items():
        node_rows = subset.iloc[indices] if has_agent else None
        if has_agent:
            agent_counts = node_rows['agent'].value_counts().to_dict()
        else:
            agent_counts = {}
        kind, dom, frac = classify_node_purity(agent_counts, purity_threshold)
        node_info[node_id] = {
            "agent_counts": agent_counts,
            "kind": kind,
            "dominant": dom,
            "dominant_frac": frac,
            "size": len(indices),
        }

    # Edges
    edge_x, edge_y = [], []
    for u, v in G.edges():
        x0, y0 = pos[u]; x1, y1 = pos[v]
        edge_x += [x0, x1, None]
        edge_y += [y0, y1, None]
    edge_trace = go.Scatter(
        x=edge_x, y=edge_y, mode='lines',
        line=dict(width=0.5, color='#bbbbbb'),
        hoverinfo='none', showlegend=False
    )

    fig = go.Figure(data=[edge_trace])

    # ---- overlap_orange: one marker per node, discrete agent or orange ----
    if mode == "overlap_orange":
        node_x, node_y, node_color, node_size, node_text = [], [], [], [], []
        for nid, info in node_info.items():
            node_x.append(pos[nid][0]); node_y.append(pos[nid][1])
            if info["kind"] == "pure":
                node_color.append(get_agent_color(info["dominant"]))
            else:
                node_color.append(MIXED_COLOR)
            node_size.append(8 + 2 * np.sqrt(info["size"]))
            comp_str = ", ".join(f"{a}:{c}" for a, c in info["agent_counts"].items())
            node_text.append(
                f"Node {nid}<br>"
                f"Size: {info['size']}<br>"
                f"Composition: {comp_str}<br>"
                f"Dominant: {info['dominant']} ({info['dominant_frac']:.0%})<br>"
                f"Classification: {info['kind']}"
            )
        fig.add_trace(go.Scatter(
            x=node_x, y=node_y, mode='markers',
            marker=dict(size=node_size, color=node_color,
                        line=dict(width=1, color='#333')),
            text=node_text, hovertemplate="%{text}<extra></extra>",
            showlegend=False
        ))
        # Legend via dummy traces — one per agent + Mixed
        for ag in AGENT_COLORS:
            if any(info["kind"] == "pure" and info["dominant"] == ag
                   for info in node_info.values()):
                fig.add_trace(go.Scatter(
                    x=[None], y=[None], mode='markers',
                    marker=dict(size=10, color=AGENT_COLORS[ag]),
                    name=f"{ag} (pure)", showlegend=True
                ))
        if any(info["kind"] == "mixed" for info in node_info.values()):
            fig.add_trace(go.Scatter(
                x=[None], y=[None], mode='markers',
                marker=dict(size=10, color=MIXED_COLOR),
                name=f"Mixed (<{int(purity_threshold*100)}%)", showlegend=True
            ))

    # ---- pie_composition: each node is a mini pie chart ----
    elif mode == "pie_composition":
        # Plotly graph_objects doesn't natively render pie markers at
        # arbitrary positions inside a Scatter plot; we approximate by
        # drawing wedge polygons per node.
        for nid, info in node_info.items():
            cx, cy = pos[nid]
            size = info["size"]
            # Radius scales with node size, but capped for legibility
            r = 0.015 + 0.005 * np.sqrt(size)
            r = min(r, 0.06)
            agent_counts = info["agent_counts"]
            total = sum(agent_counts.values()) if agent_counts else 1
            start_angle = 0.0
            comp_str = ", ".join(f"{a}:{c}" for a, c in agent_counts.items())
            hover_text = (f"Node {nid}<br>Size: {size}<br>"
                          f"Composition: {comp_str}")
            for ag, count in agent_counts.items():
                frac = count / total
                end_angle = start_angle + 2 * np.pi * frac
                # Sample the wedge boundary
                theta = np.linspace(start_angle, end_angle, 20)
                wx = [cx] + list(cx + r * np.cos(theta)) + [cx]
                wy = [cy] + list(cy + r * np.sin(theta)) + [cy]
                fig.add_trace(go.Scatter(
                    x=wx, y=wy, fill='toself',
                    fillcolor=get_agent_color(ag),
                    line=dict(width=0.3, color='#333'),
                    mode='lines',
                    name=ag, showlegend=False,
                    text=hover_text,
                    hovertemplate="%{text}<extra></extra>",
                ))
                start_angle = end_angle
        # Legend via dummy traces
        agents_present = set()
        for info in node_info.values():
            agents_present.update(info["agent_counts"].keys())
        for ag in sorted(agents_present):
            fig.add_trace(go.Scatter(
                x=[None], y=[None], mode='markers',
                marker=dict(size=10, color=get_agent_color(ag)),
                name=ag, showlegend=True
            ))

    fig.update_layout(
        title=title,
        xaxis=dict(visible=False), yaxis=dict(visible=False),
        plot_bgcolor='white', paper_bgcolor='white',
        height=650, hovermode='closest',
        margin=dict(l=10, r=10, t=50, b=10),
    )
    return fig


def build_iep_features(df):
    """Build 12–15D feature matrix from IEP columns."""
    has_italic_cols = 'has_italics' in df.columns
    features, valid_indices = [], []
    for idx, row in df.iterrows():
        try:
            fv = [
                float(row.get('int_pct', 0) or 0),
                float(row.get('aff_pct', 0) or 0),
                float(row.get('act_pct', 0) or 0),
                float(row.get('vader_compound', 0) or 0),
                float(row.get('vader_pos', 0) or 0),
                float(row.get('vader_neg', 0) or 0),
                float(row.get('vader_neu', 0) or 0),
                float(row.get('flesch_kincaid', 0) or 0),
                float(row.get('flesch_ease', 0) or 0) if 'flesch_ease' in row else 0,
                float(row.get('ttr', 0) or 0),
                float(row.get('total_words', 0) or 0),
                float(row.get('unique_words', 0) or 0) if 'unique_words' in row else 0,
            ]
            if has_italic_cols:
                fv.extend([
                    float(row.get('has_italics', 0) or 0),
                    float(row.get('italic_count', 0) or 0),
                    float(row.get('italic_density', 0) or 0),
                ])
            # V_t voice-state vector — append if present (Paper 2 empirical grounding)
            for vt_col in ['S_t', 'A_t', 'Q_t', 'D_t', 'R_t']:
                if vt_col in row.index:
                    fv.append(float(row.get(vt_col, 0) or 0))
            features.append(fv)
            valid_indices.append(idx)
        except (ValueError, TypeError):
            continue
    return np.array(features), valid_indices


def run_mapper_analysis(data, df, valid_indices, n_cubes, overlap, eps, min_samples, projection_type):
    mapper = km.KeplerMapper(verbose=0)
    if projection_type == "PCA (2D)":
        projected = mapper.fit_transform(data, projection=PCA(n_components=2))
    elif projection_type == "PCA1 (1D)":
        # V18 — first principal component as a 1D lens.
        # We compute PCA1 ourselves so we can also expose it to the
        # Plotly fallback renderer for strip-plot visualization.
        _pca1 = PCA(n_components=1).fit(data)
        projected = _pca1.transform(data)  # shape: (n_samples, 1)
    elif projection_type == "Sum":
        projected = mapper.fit_transform(data, projection="sum")
    elif projection_type == "Mean":
        projected = mapper.fit_transform(data, projection="mean")
    elif projection_type == "AFF% Lens":
        subset = df.iloc[valid_indices]
        projected = subset['aff_pct'].values.reshape(-1, 1)
    elif projection_type == "INT% Lens":
        subset = df.iloc[valid_indices]
        projected = subset['int_pct'].values.reshape(-1, 1)
    elif projection_type == "ACT% Lens":
        subset = df.iloc[valid_indices]
        projected = subset['act_pct'].values.reshape(-1, 1)
    elif projection_type == "S_t Lens":
        subset = df.iloc[valid_indices]
        col = 'S_t' if 'S_t' in subset.columns else 'aff_pct'
        projected = subset[col].values.reshape(-1, 1)
    elif projection_type == "A_t Lens":
        subset = df.iloc[valid_indices]
        col = 'A_t' if 'A_t' in subset.columns else 'aff_pct'
        projected = subset[col].values.reshape(-1, 1)
    elif projection_type == "Q_t Lens":
        subset = df.iloc[valid_indices]
        col = 'Q_t' if 'Q_t' in subset.columns else 'aff_pct'
        projected = subset[col].values.reshape(-1, 1)
    elif projection_type == "D_t Lens":
        subset = df.iloc[valid_indices]
        col = 'D_t' if 'D_t' in subset.columns else 'aff_pct'
        projected = subset[col].values.reshape(-1, 1)
    elif projection_type == "R_t Lens":
        subset = df.iloc[valid_indices]
        col = 'R_t' if 'R_t' in subset.columns else 'aff_pct'
        projected = subset[col].values.reshape(-1, 1)
    else:
        projected = mapper.fit_transform(data, projection=PCA(n_components=2))
    cover = km.Cover(n_cubes=n_cubes, perc_overlap=overlap)
    graph = mapper.map(projected, data, cover=cover,
                       clusterer=DBSCAN(eps=eps, min_samples=min_samples))
    return mapper, graph


def analyze_graph(graph, df, valid_indices):
    import networkx as nx
    subset = df.iloc[valid_indices].reset_index(drop=True)

    # True edge count — KeplerMapper stores links one-way
    n_edges = sum(len(v) for v in graph['links'].values())

    analysis = {"n_nodes": len(graph['nodes']), "n_edges": n_edges, "nodes": {}}
    cold_nodes, hot_nodes = set(), set()
    hot_temps = {'HOT','FIRE','AFF_1','AFF_2','AFF_3','AFF_4','AFF_5','RELATIONAL','WARM'}
    for node_id, indices in graph['nodes'].items():
        node_data = subset.iloc[indices]
        temps = node_data['temperature'].value_counts().to_dict() if 'temperature' in node_data.columns else {}
        agents = node_data['agent'].value_counts().to_dict() if 'agent' in node_data.columns else {}
        questions = (node_data['question_id'].value_counts().to_dict() if 'question_id' in node_data.columns else
                     node_data['question_label'].value_counts().to_dict() if 'question_label' in node_data.columns else {})
        avg_int = node_data['int_pct'].mean() if 'int_pct' in node_data.columns else 0
        avg_aff = node_data['aff_pct'].mean() if 'aff_pct' in node_data.columns else 0
        avg_act = node_data['act_pct'].mean() if 'act_pct' in node_data.columns else 0
        if 'iep_density' in node_data.columns and node_data['iep_density'].sum() > 0:
            dm = round(node_data['iep_density'].mean(), 4)
            dmed = round(node_data['iep_density'].median(), 4)
            dsem = round(node_data['iep_density'].sem(), 4) if len(node_data) > 1 else 0.0
        else:
            dm = dmed = dsem = 0.0
        analysis["nodes"][node_id] = {
            "size":len(indices),"temps":temps,"agents":agents,"questions":questions,
            "avg_int":round(avg_int,2),"avg_aff":round(avg_aff,2),"avg_act":round(avg_act,2),
            "density_mean":dm,"density_median":dmed,"density_sem":dsem,
        }
        if 'COLD' in temps or 'ANALYTICAL' in temps or 'ICE' in temps: cold_nodes.add(node_id)
        if any(t in hot_temps for t in temps): hot_nodes.add(node_id)
    analysis["cold_nodes"]=cold_nodes; analysis["hot_nodes"]=hot_nodes
    analysis["overlap_nodes"]=cold_nodes & hot_nodes
    analysis["cold_only"]=cold_nodes - hot_nodes; analysis["hot_only"]=hot_nodes - cold_nodes

    # Connected components via NetworkX
    G = nx.Graph()
    G.add_nodes_from(graph['nodes'].keys())
    for src, targets in graph['links'].items():
        for tgt in targets:
            G.add_edge(src, tgt)
    components = list(nx.connected_components(G))
    analysis["n_components"] = len(components)
    analysis["largest_component"] = max(len(c) for c in components) if components else 0
    analysis["isolated_nodes"] = sum(1 for c in components if len(c) == 1)

    return analysis


# =============================================================================
# CLAUDE API ANALYSIS FUNCTIONS
# =============================================================================

def build_claude_analysis_prompt(analysis, df):
    """
    Build a clean, unbiased prompt for topology analysis.
    No hypothesis framing — Claude describes what the data shows.
    """
    n_nodes = analysis["n_nodes"]
    n_edges = analysis["n_edges"]
    n_components = analysis.get("n_components", "N/A")
    largest_component = analysis.get("largest_component", "N/A")
    cold_only = len(analysis["cold_only"])
    hot_only  = len(analysis["hot_only"])
    overlap   = len(analysis["overlap_nodes"])

    # IEP validation warning
    iep_cols = ['int_pct', 'aff_pct', 'act_pct']
    iep_valid = all(c in df.columns for c in iep_cols)
    iep_sum = df[iep_cols].sum(axis=1).mean() if iep_valid else 0
    iep_warning = ""
    if iep_valid and iep_sum < 50:
        iep_warning = (
            f"\n⚠️ WARNING: IEP scores appear scaled incorrectly "
            f"(mean sum = {iep_sum:.1f}%, expected ~100%). "
            f"True dataset means — INT: {df['int_pct'].mean():.1f}% "
            f"AFF: {df['aff_pct'].mean():.1f}% ACT: {df['act_pct'].mean():.1f}%\n"
        )

    iep_summary = ""
    for col, label in [("int_pct","INT%"),("aff_pct","AFF%"),("act_pct","ACT%")]:
        if col in df.columns:
            iep_summary += (f"  {label}: mean={df[col].mean():.1f}, "
                           f"std={df[col].std():.1f}, "
                           f"min={df[col].min():.1f}, max={df[col].max():.1f}\n")

    temp_dist = ""
    if "temperature" in df.columns:
        temp_dist = "\n".join(f"  {k}: {v}" for k,v in df["temperature"].value_counts().items())

    agent_dist = ""
    if "agent" in df.columns:
        agent_dist = "\n".join(f"  {k}: {v}" for k,v in df["agent"].value_counts().items())

    sorted_nodes = sorted(analysis["nodes"].items(), key=lambda x: x[1]["size"], reverse=True)
    node_details = ""
    for node_id, info in sorted_nodes[:10]:
        node_details += (
            f"\n  Node {node_id}: size={info['size']}, "
            f"INT={info['avg_int']:.1f}%, AFF={info['avg_aff']:.1f}%, ACT={info['avg_act']:.1f}%\n"
            f"    Conditions: {info['temps']}\n"
            f"    Agents: {info['agents']}\n"
        )

    return f"""You are a topological data analysis expert. You are analyzing KeplerMapper output \
from a study of AI communicative character measured using the IEP framework.

FRAMEWORK:
- IEP decomposes AI responses into three dimensions (sum to 100% per response):
  • INT% — Intellectual / analytical language
  • AFF% — Affective / emotional language
  • ACT% — Action-oriented / practical language
- Temperature conditions represent instructional framing levels (COLD → NATIVE → HOT → FIRE)
- Each node = a cluster of similar responses; edges = overlap between clusters
{iep_warning}
TOPOLOGY:
  Nodes:               {n_nodes}
  Edges:               {n_edges}
  Connected components:{n_components}
  Largest component:   {largest_component} nodes
  COLD-only nodes:     {cold_only}
  HOT-only nodes:      {hot_only}
  Mixed nodes:         {overlap}
  Separation:          {"TOPOLOGICAL SEPARATION achieved" if (cold_only > 0 and hot_only > 0) else "MIXED — no clean separation" if overlap > 0 else "Single-condition data"}

DATASET:
  Total records: {len(df)}
  Conditions:
{temp_dist if temp_dist else "  N/A"}
  Agents:
{agent_dist if agent_dist else "  N/A"}

IEP DISTRIBUTION:
{iep_summary if iep_summary else "  Not available"}

TOP NODES BY SIZE:
{node_details}

INSTRUCTIONS:
Describe what this topology shows. Do not assume any particular shape or pattern — \
let the data speak. Cover:

1. **Topological Structure**: What does the node count, edge count, and component \
structure reveal? Is the response space fragmented, consolidated, or chain-like?

2. **Condition Separation**: What does the COLD vs HOT node distribution reveal? \
Are different instructional conditions topologically distinct or intermixed?

3. **IEP Profile**: What does the INT/AFF/ACT distribution reveal about these \
responses? Which dimension dominates? Are there regions with notably different profiles?

4. **Agent Behavior**: Which agents dominate which regions? Do they cluster \
together or spread across the topology?

5. **Hub Nodes**: What do the largest nodes reveal about the most common \
response profiles? What conditions and agents dominate them?

6. **Emergent Shape**: Describe the overall shape of the topology as you find it. \
Is it a chain, cluster, mesh, isolated islands, or something else?

7. **Key Finding**: What is the single most important finding from this analysis?

Be specific. Cite node IDs, IEP values, and conditions. Let the structure speak for itself.
"""



def build_dual_report_prompt(analysis, df):
    """
    V10 dual-report: unbiased, data-driven.
    Part A = topological regions, Part B = factorial interpretation.
    No hypothesis fishing — Claude reports what the data shows.
    """
    factorial_rows = []
    if all(c in df.columns for c in ['agent','temperature','question_id','int_pct','aff_pct','act_pct']):
        grp = df.groupby(['agent','temperature','question_id'])[['int_pct','aff_pct','act_pct']].agg(['mean','sem'])
        for (agent, temp, qid), row in grp.iterrows():
            factorial_rows.append(
                f"  {agent:<12} {temp:<10} {qid:<22} "
                f"INT={row[('int_pct','mean')]:.1f}+/-{row[('int_pct','sem')]:.1f} "
                f"AFF={row[('aff_pct','mean')]:.1f}+/-{row[('aff_pct','sem')]:.1f} "
                f"ACT={row[('act_pct','mean')]:.1f}+/-{row[('act_pct','sem')]:.1f}"
            )
    factorial_table = "\n".join(factorial_rows) if factorial_rows else "N/A"

    nodes_sorted = sorted(analysis["nodes"].items(), key=lambda x: x[1]["size"], reverse=True)
    node_geo = ""
    for nid, info in nodes_sorted[:20]:
        node_geo += (
            f"  {nid}: n={info['size']} INT={info['avg_int']:.1f}% "
            f"AFF={info['avg_aff']:.1f}% ACT={info['avg_act']:.1f}%"
            f" Density={info.get('density_mean',0):.3f}\n"
            f"    Agents:{info['agents']} Conditions:{info['temps']} "
            f"Questions:{info.get('questions',{})}\n"
        )

    agents = sorted(df['agent'].unique().tolist()) if 'agent' in df.columns else []
    temps  = sorted(df['temperature'].unique().tolist()) if 'temperature' in df.columns else []
    qs     = sorted(df['question_id'].unique().tolist()) if 'question_id' in df.columns else []
    n_components = analysis.get("n_components", "N/A")
    largest_component = analysis.get("largest_component", "N/A")

    return f"""You are a topological data analysis expert. You are analyzing a KeplerMapper \
topology from a study of AI communicative character measured using the IEP framework \
(Intellectual/Affective/Action language composition).

DATASET: {len(df)} responses
Agents: {agents}
Conditions: {temps}
Questions: {qs}
Nodes: {analysis['n_nodes']} | Edges: {analysis['n_edges']}
Connected components: {n_components} | Largest component: {largest_component} nodes
COLD-only: {len(analysis['cold_only'])} | HOT-only: {len(analysis['hot_only'])} | Mixed: {len(analysis['overlap_nodes'])}

IEP Framework: INT%=intellectual, AFF%=affective, ACT%=action (composition, sums to 100%)
Density = IEP hits / total words (coverage, separate from composition)

TOP 20 NODES:
{node_geo}
FACTORIAL TABLE (agent x condition x question, mean +/- SEM):
{factorial_table}

Do NOT use tilde (~) for approximations — write "approx" instead.
Do NOT use strikethrough or horizontal rules.

Begin with this header on its own line:
# Topology Report — [list agents from dataset] — [date: March 2026]

Write a structured dual report with EXACTLY these two sections:

## PART A — TOPOLOGICAL REGIONS
Describe the geometric landscape as you find it — no assumed shape.
For each major region:
- IEP profile and density
- Which agents, conditions, and questions populate it
- Bridges, islands, hubs — describe their role in the topology
- What the overall structure reveals about the communicative space

## PART B — FACTORIAL INTERPRETATION
- Which question domains produce the most topologically distinct regions?
  Describe where each question sits in the topology and its IEP profile.
- Which agents are most and least affected by temperature condition?
  Where do agents converge into shared topological space, and where do they separate?
- What does the condition gradient (COLD through FIRE) look like in the topology?
  Do responses move continuously or jump discretely across conditions?
- Most important findings for publication — cite specific node IDs and IEP values."""


def call_claude_api(api_key, messages_or_prompt, max_tokens=4000):
    """Call Claude API. Accepts a string prompt or a messages list for conversation."""
    import httpx
    if isinstance(messages_or_prompt, str):
        messages = [{"role": "user", "content": messages_or_prompt}]
    else:
        messages = messages_or_prompt
    response = httpx.post(
        "https://api.anthropic.com/v1/messages",
        headers={
            "x-api-key": api_key,
            "anthropic-version": "2023-06-01",
            "content-type": "application/json",
        },
        json={
            "model": "claude-sonnet-4-20250514",
            "max_tokens": max_tokens,
            "messages": messages,
        },
        timeout=120.0,
    )
    if response.status_code == 200:
        result = response.json()
        return "".join(
            block["text"] for block in result.get("content", [])
            if block.get("type") == "text"
        )
    else:
        raise RuntimeError(f"API Error {response.status_code}: {response.text}")


def build_topology_context(analysis, df):
    """Build full context string loaded once, reused for every follow-up turn."""
    iep_cols = ["int_pct", "aff_pct", "act_pct"]
    iep_valid = all(c in df.columns for c in iep_cols)
    iep_sum = df[iep_cols].sum(axis=1).mean() if iep_valid else 0
    iep_warn = ""
    if iep_valid and iep_sum < 50:
        iep_warn = (f"WARNING: IEP sum={iep_sum:.1f}% (expected ~100%). "
                    f"True: INT={df['int_pct'].mean():.1f}% AFF={df['aff_pct'].mean():.1f}% ACT={df['act_pct'].mean():.1f}%\n")
    iep_s = ""
    for col, lbl in [("int_pct", "INT%"), ("aff_pct", "AFF%"), ("act_pct", "ACT%")]:
        if col in df.columns:
            iep_s += f"  {lbl}: mean={df[col].mean():.1f} std={df[col].std():.1f} min={df[col].min():.1f} max={df[col].max():.1f}\n"
    tdist = "\n".join(f"  {k}: {v}" for k, v in df["temperature"].value_counts().items()) if "temperature" in df.columns else "N/A"
    adist = "\n".join(f"  {k}: {v}" for k, v in df["agent"].value_counts().items()) if "agent" in df.columns else "N/A"
    qcol = "question_id" if "question_id" in df.columns else "question_label" if "question_label" in df.columns else None
    qdist = "\n".join(f"  {k}: {v}" for k, v in df[qcol].value_counts().items()) if qcol else "N/A"
    nodes_s = sorted(analysis["nodes"].items(), key=lambda x: x[1]["size"], reverse=True)
    nd = ""
    for nid, info in nodes_s[:15]:
        qs = f"    Questions: {info.get('questions', {})}\n" if info.get("questions") else ""
        nd += (f"  {nid}: n={info['size']} INT={info['avg_int']:.1f}% AFF={info['avg_aff']:.1f}% ACT={info['avg_act']:.1f}%\n"
               f"    Conditions:{info['temps']} Agents:{info['agents']}\n{qs}")
    n_components = analysis.get("n_components", "N/A")
    largest_component = analysis.get("largest_component", "N/A")
    return (
        f"Topology Analysis Context\n{iep_warn}"
        f"IEP: INT%=intellectual AFF%=affective ACT%=action (sum ~100%)\n"
        f"Temperature: COLD->NATIVE->HOT->FIRE (instructional framing conditions)\n"
        f"Nodes=response clusters, Edges=cluster overlap\n\n"
        f"TOPOLOGY: {analysis['n_nodes']} nodes, {analysis['n_edges']} edges, "
        f"{n_components} connected components, largest={largest_component} nodes\n"
        f"COLD-only={len(analysis['cold_only'])} HOT-only={len(analysis['hot_only'])} Mixed={len(analysis['overlap_nodes'])}\n"
        f"Separation: {'ACHIEVED' if (analysis['cold_only'] and analysis['hot_only']) else 'MIXED/SINGLE'}\n\n"
        f"CONDITIONS:\n{tdist}\nAGENTS:\n{adist}\nQUESTIONS:\n{qdist}\nIEP:\n{iep_s}\nTOP 15 NODES:\n{nd}\n"
        f"Always reference node IDs, IEP values, and conditions specifically."
    )


# =============================================================================
# SIDEBAR
# =============================================================================
st.sidebar.markdown("## ⚙️ Configuration")

c = IEP_DICT["counts"]
st.sidebar.success(
    f"📖 IEP Dict V3 — **embedded**  \n"
    f"INT {c['INT']} · AFF {c['AFF']} · ACT {c['ACT']}  \n"
    f"Total: {c['TOTAL']} words"
)

uploaded_files = st.sidebar.file_uploader("📁 Upload Data (CSV or JSON) — drop multiple for combined OR batch mode", type=['csv', 'json'], accept_multiple_files=True)

# v15: Batch mode — run mapper once per uploaded file separately
# v16: Stability Sweep — run all 23 of Farzana's lenses on combined data
batch_mode = False
sweep_mode = False
if uploaded_files and len(uploaded_files) > 1:
    multi_mode = st.sidebar.radio(
        "Multi-file mode:",
        [
            "Combined (pool into one Mapper)",
            "Batch (one Mapper per file)",
            "Stability Sweep (Farzana's 23 lenses on combined data)",
        ],
        index=0,
        help="Combined = v14 default. Batch = one Mapper per file (§4.1 figures). "
             "Stability Sweep = run all 23 lenses from §4.2 protocol on the pooled "
             "dataset, output stability table + per-lens HTMLs + layered docx."
    )
    batch_mode = "Batch" in multi_mode
    sweep_mode = "Sweep" in multi_mode
elif uploaded_files and len(uploaded_files) == 1:
    # Allow stability sweep on single file too
    if st.sidebar.checkbox("Stability Sweep (23 lenses)",
                            help="Run all 23 §4.2 lenses on this single file"):
        sweep_mode = True

# Auto-clear cache when files change
_file_key = str(sorted([f.name for f in uploaded_files])) if uploaded_files else ""
if _file_key != st.session_state.get('_last_file_key', ''):
    st.session_state._last_file_key = _file_key
    for k in ['mapper_results_html','mapper_analysis_data','claude_analysis',
              'stored_df','topology_context']:
        st.session_state[k] = None
    st.session_state.chat_history = []
    st.session_state.chat_active = False

# Build df from uploaded files
uploaded_file = None
df_combined = None
if len(uploaded_files) == 1:
    uploaded_file = uploaded_files[0]
    st.session_state.is_combined = False
elif len(uploaded_files) > 1:
    import io
    dfs = []
    for f in uploaded_files:
        d = pd.read_csv(f) if f.name.endswith('.csv') else pd.DataFrame(json.load(f))
        dfs.append(d)
        ag = d['agent'].iloc[0] if 'agent' in d.columns else '?'
        tp = d['temperature'].iloc[0] if 'temperature' in d.columns else '?'
    df_combined = pd.concat(dfs, ignore_index=True)
    if 'agent' in df_combined.columns:
        df_combined['agent'] = df_combined['agent'].replace({'Sophia': 'ChatGPT', 'sophia': 'ChatGPT'})
    agents = df_combined['agent'].nunique() if 'agent' in df_combined.columns else '?'
    temps  = df_combined['temperature'].nunique() if 'temperature' in df_combined.columns else '?'
    qs     = df_combined['question_id'].nunique() if 'question_id' in df_combined.columns else '?'
    st.sidebar.success(f"✅ Combined {len(df_combined):,} rows · {agents} agents · {temps} conditions · {qs} questions")
    st.session_state.is_combined = True

if st.sidebar.button("🔄 Clear & Reset", help="Click between runs to clear previous topology"):
    for k in ['mapper_results_html','mapper_analysis_data','claude_analysis',
              'stored_df','stored_graph','topology_context','_last_file_key']:
        st.session_state[k] = None
    st.session_state.chat_history = []
    st.session_state.chat_active = False
    st.session_state.is_combined = False
    st.rerun()

st.sidebar.markdown("---")
st.sidebar.markdown("### Data Source")

feature_options = [
    "IEP — pre-scored columns (recommended for V48 data)",
    "IEP V3 — live score from text (no pre-scored columns)",
    "Embeddings (pre-computed)",
    "SBERT Embeddings (from text)",
]
use_iep = st.sidebar.radio("Feature Type:", feature_options, index=0)
if "pre-scored" in use_iep:
    pass  # default mode, no message needed
elif "live score" in use_iep:
    st.sidebar.warning("⚠️ Live scoring from text — only use if CSV has no int_pct/aff_pct/act_pct columns.")

text_col_choice = "response_text"
if "live score" in use_iep:
    text_col_choice = st.sidebar.text_input(
        "Text column name:", value="response_text",
        help="Column containing the raw response text to score with IEP V3."
    )

st.sidebar.markdown("---")
st.sidebar.markdown("### Mapper Parameters")

n_cubes     = st.sidebar.slider("N Cubes (cover resolution)", 5, 30, 10)
overlap     = st.sidebar.slider("Overlap %", 0.1, 0.9, 0.5)
eps         = st.sidebar.slider("DBSCAN eps", 0.1, 5.0, 2.0, step=0.1,
    help="Higher values for larger/higher-dim datasets. Try 1.5–3.0 for 12D IEP features.")
min_samples = st.sidebar.slider("DBSCAN min_samples", 2, 10, 3)

# ── Auto-warning for embedding modes ──────────────────────────────────────────
if "Embedding" in use_iep or "SBERT" in use_iep:
    st.sidebar.warning(
        "⚠️ **384D embedding space detected.**\n\n"
        "Default `eps=2.0` will likely produce **0 nodes**.\n\n"
        "Recommended range: `eps = 0.05 – 0.20`\n"
        "`min_samples = 2`"
    )
    if eps > 0.5:
        st.sidebar.error(
            f"🚨 `eps={eps}` is too large for 384D cosine embeddings. "
            "Set eps ≤ 0.20 to get clusters."
        )

projection_type = st.sidebar.selectbox(
    "Projection / Lens",
    [
        "PCA (2D)", "PCA1 (1D)", "Sum", "Mean",
        "AFF% Lens", "INT% Lens", "ACT% Lens",
        "S_t Lens", "A_t Lens", "Q_t Lens", "D_t Lens", "R_t Lens",
    ]
)

st.sidebar.markdown("---")
st.sidebar.markdown("### Visualization")
color_by = st.sidebar.selectbox(
    "Color nodes by:",
    ["aff_pct", "int_pct", "act_pct", "vader_compound",
     "S_t", "A_t", "Q_t", "D_t", "R_t",
     "temperature", "agent",
     "agent (overlap=orange)", "agent (pie composition)"]
)

# V18 — purity threshold gates pure-vs-mixed node classification for the
# two new agent-coloring modes. Has no effect on the legacy "agent" mode
# (which keeps V17's continuous-colormap behavior) or on any non-agent mode.
purity_threshold = st.sidebar.slider(
    "Purity threshold (agent-coloring modes)",
    min_value=0.50, max_value=1.00, value=0.80, step=0.05,
    help="A node counts as 'pure' when its dominant agent's share is ≥ "
         "this fraction. Used only by 'agent (overlap=orange)' and "
         "'agent (pie composition)'."
)

st.sidebar.markdown("---")
st.sidebar.markdown("### 🧹 Session Control")
if st.sidebar.button("🗑️ Clear Results", help="Wipe all mapper output, Claude analysis, and graph state. Use before a new run to avoid carryover."):
    st.session_state.mapper_results_html   = None
    st.session_state.mapper_analysis_data  = None
    st.session_state.claude_analysis       = None
    st.session_state.last_analysis_df_hash = None
    st.session_state.chat_history          = []
    st.session_state.topology_context      = None
    st.session_state.chat_active           = False
    st.session_state.stored_df             = None
    st.sidebar.success("✅ Results cleared — ready for fresh run.")
    st.rerun()

st.sidebar.markdown("---")
st.sidebar.markdown("### 🤖 Claude AI Analysis")
st.sidebar.markdown("Optional — enables AI narrative analysis after mapping.")
api_key = st.sidebar.text_input(
    "Anthropic API Key:",
    type="password",
    help="Paste your Anthropic API key to unlock AI-powered analysis of your Mapper output."
)

if st.session_state.get("_sidebar_msg_json_text"):
    st.sidebar.success("✅ JSON loaded with response text for embeddings")
if "_sidebar_msg_record_count" in st.session_state:
    st.sidebar.info(f"Loaded {st.session_state['_sidebar_msg_record_count']} records from JSON")

# =============================================================================
# TABS
# =============================================================================
tab_main, tab_dict, tab_score = st.tabs(["🗺️ Mapper Analyzer", "📖 Dictionary Explorer", "🧪 Score Text"])

# ─────────────────────────────────────────────
# TAB 2: DICTIONARY EXPLORER
# ─────────────────────────────────────────────
with tab_dict:
    st.markdown("## 📖 IEP Dictionary V3 Explorer")
    st.caption("All 1,897 words are embedded in this script — no external file needed.")

    c = IEP_DICT["counts"]
    col1, col2, col3, col4 = st.columns(4)
    col1.metric("INT words",   c["INT"])
    col2.metric("AFF words",   c["AFF"])
    col3.metric("ACT words",   c["ACT"])
    col4.metric("TOTAL words", c["TOTAL"])

    search_term = st.text_input("🔍 Search dictionary", placeholder="e.g. 'feel' or 'analyz'")

    for category, label in [("INT", "🔵 INTELLECTUAL"), ("AFF", "❤️ AFFECTIVE"), ("ACT", "🟢 ACTION")]:
        words = sorted(IEP_DICT[category])
        if search_term:
            words = [w for w in words if search_term.lower() in w]
        with st.expander(f"{label} — {len(words)} words {'(filtered)' if search_term else ''}"):
            st.write(", ".join(words) if words else "_No matches_")

    st.markdown("### 🔀 Cross-category Overlaps")
    for k, v in _IEP_OVERLAPS.items():
        if v:
            st.write(f"**{k}** ({len(v)} words): {', '.join(sorted(v))}")
        else:
            st.write(f"**{k}**: _none_")

# ─────────────────────────────────────────────
# TAB 3: SCORE TEXT
# ─────────────────────────────────────────────
with tab_score:
    st.markdown("## 🧪 Score a Single Response")
    st.markdown("Paste any AI response to get its IEP V3 profile instantly.")

    sample_text = st.text_area("Response text:", height=200,
        placeholder="Paste an AI response here…")

    if st.button("Score Text", type="primary") and sample_text.strip():
        result = score_text_iep(sample_text)
        st.markdown("#### Composition *(sums to 100% — topology feature)*")
        col1, col2, col3 = st.columns(3)
        col1.metric("INT%", f"{result['int_pct']:.1f}%", f"{result['int_alloc']:.1f} alloc")
        col2.metric("AFF%", f"{result['aff_pct']:.1f}%", f"{result['aff_alloc']:.1f} alloc")
        col3.metric("ACT%", f"{result['act_pct']:.1f}%", f"{result['act_alloc']:.1f} alloc")
        comp = result['int_pct']+result['aff_pct']+result['act_pct']
        st.caption(f"Sum={comp:.2f}% {'OK' if abs(comp-100)<0.1 else 'CHECK'}")
        st.markdown("#### Density *(IEP hits / total words)*")
        d1, d2 = st.columns(2)
        d1.metric("IEP Density", f"{result['iep_density']:.4f}", f"{result['total_alloc']:.1f} alloc / {result['total_words']} words")
        d2.metric("Total Words", result['total_words'])
        if result.get('overlap_int_aff',0)+result.get('overlap_int_act',0) > 0:
            st.markdown("#### Overlap Terms *(50/50 allocated)*")
            ov1,ov2 = st.columns(2)
            ov1.metric("INT+AFF", result['overlap_int_aff'], "notice, noticing, noticed, understanding")
            ov2.metric("INT+ACT", result['overlap_int_act'], "conclude, step, steps")

        st.markdown("#### Matched words")
        c1, c2, c3 = st.columns(3)
        with c1:
            st.markdown("**🔵 INT**")
            st.write(", ".join(f"{w}({n})" for w, n in result["matched_int"].most_common(20)) or "_none_")
        with c2:
            st.markdown("**❤️ AFF**")
            st.write(", ".join(f"{w}({n})" for w, n in result["matched_aff"].most_common(20)) or "_none_")
        with c3:
            st.markdown("**🟢 ACT**")
            st.write(", ".join(f"{w}({n})" for w, n in result["matched_act"].most_common(20)) or "_none_")

# ─────────────────────────────────────────────
# TAB 1: MAIN MAPPER ANALYZER
# ─────────────────────────────────────────────
with tab_main:

    # =========================================================================
    # v15: BATCH MODE — run mapper separately on each uploaded file
    # =========================================================================
    if batch_mode and uploaded_files:
        st.markdown(f"## 🔁 Batch Mode — {len(uploaded_files)} files")
        st.markdown(
            "Each file will be processed separately with the **same Mapper settings** "
            "from the sidebar (lens, eps, min_samples, n_cubes, overlap). "
            "All results download as a single zip when complete."
        )
        for f in uploaded_files:
            st.markdown(f"  • `{f.name}`")

        if st.button("🚀 Run Batch Analysis", type="primary"):
            import io as _io, zipfile as _zip
            from datetime import datetime as _dt

            zip_buf = _io.BytesIO()
            zf = _zip.ZipFile(zip_buf, 'w', _zip.ZIP_DEFLATED)
            batch_summary_rows = []
            progress = st.progress(0)
            status = st.empty()

            for idx, f in enumerate(uploaded_files):
                fname_safe = f.name.replace('.csv','').replace('.json','')
                fname_safe = ''.join(c if c.isalnum() or c in '._-' else '_' for c in fname_safe)
                status.text(f"[{idx+1}/{len(uploaded_files)}] Processing {f.name}...")

                # Read file
                try:
                    f.seek(0)
                    if f.name.endswith('.json'):
                        _df = pd.DataFrame(json.loads(f.read().decode('utf-8')))
                    else:
                        _df = pd.read_csv(f)
                except Exception as e:
                    st.error(f"Failed to read {f.name}: {e}")
                    continue

                if 'agent' in _df.columns:
                    _df['agent'] = _df['agent'].replace({'Sophia':'ChatGPT','sophia':'ChatGPT'})

                # Build features
                try:
                    if "live score" in use_iep:
                        if text_col_choice in _df.columns:
                            _df = score_dataframe_iep(_df, text_col_choice)
                    if "live score" in use_iep or "pre-scored" in use_iep:
                        _data, _valid = build_iep_features(_df)
                        if len(_data) > 0:
                            _data = StandardScaler().fit_transform(_data)
                    elif "SBERT" in use_iep:
                        from sentence_transformers import SentenceTransformer
                        _model = SentenceTransformer('all-MiniLM-L6-v2')
                        _texts = _df['response_text'].fillna('').tolist()
                        _valid = [i for i,t in enumerate(_texts) if t.strip()]
                        _data = _model.encode([_texts[i] for i in _valid])
                    else:
                        _data, _valid = parse_embeddings(_df)

                    if len(_data) == 0:
                        st.warning(f"⚠️ {f.name}: no valid data, skipping.")
                        continue

                    # Run Mapper (same fn v14 uses)
                    _mapper, _graph = run_mapper_analysis(
                        _data, _df, _valid,
                        n_cubes, overlap, eps, min_samples, projection_type
                    )
                    _analysis = analyze_graph(_graph, _df, _valid)

                    # Build HTML
                    _subset = _df.iloc[_valid].reset_index(drop=True)
                    if color_by in _subset.columns:
                        if _subset[color_by].dtype == 'object':
                            _cats = _subset[color_by].unique().tolist()
                            _cv = _subset[color_by].map({c:i for i,c in enumerate(_cats)}).values
                        else:
                            _cv = _subset[color_by].values
                    else:
                        _cv = (_subset['aff_pct'].values
                               if 'aff_pct' in _subset.columns else np.zeros(len(_subset)))

                    if 'temperature' in _subset.columns and 'agent' in _subset.columns:
                        if 'question_id' in _subset.columns:
                            _tt = np.array([f"{r['agent']} | {r['temperature']} | {r.get('question_id','')}"
                                            for _,r in _subset.iterrows()])
                        else:
                            _tt = np.array([f"{r['agent']} | {r['temperature']}"
                                            for _,r in _subset.iterrows()])
                    else:
                        _tt = _subset.index.astype(str).values

                    _title = f"SYN-IQ V17 · {fname_safe} · {projection_type}"
                    import tempfile as _tmp
                    with _tmp.NamedTemporaryFile(mode='w', suffix='.html',
                                                 delete=False, encoding='utf-8') as _tf:
                        _tmp_path = _tf.name
                    try:
                        _mapper.visualize(
                            _graph, path_html=_tmp_path, title=_title,
                            color_values=_cv, color_function_name=color_by,
                            custom_tooltips=_tt, include_searchbar=True,
                        )
                        with open(_tmp_path, 'r', encoding='utf-8') as _hh:
                            _html = _hh.read()
                        _html = whiten_mapper_html(_html)  # v17: white background per Farzana
                    finally:
                        try: os.unlink(_tmp_path)
                        except Exception: pass

                    # Add to zip
                    zf.writestr(f"{fname_safe}/mapper.html", _html)

                    # Per-file summary CSV
                    _summary = pd.DataFrame([
                        ('Nodes', _analysis.get('n_nodes', len(_graph.get('nodes', {})))),
                        ('Edges', _analysis.get('n_edges', sum(len(v) for v in _graph.get('links',{}).values()))),
                        ('Connected Components', _analysis.get('n_components', None)),
                        ('Largest Component (nodes)', _analysis.get('largest_component_size', None)),
                    ], columns=['Metric','Value'])
                    zf.writestr(f"{fname_safe}/topology_summary.csv",
                                _summary.to_csv(index=False))

                    # Row for batch-level summary
                    batch_summary_rows.append({
                        'file': f.name,
                        'n_records': len(_df),
                        'n_clustered': len(_valid),
                        'lens': projection_type,
                        'V_nodes': _analysis.get('n_nodes'),
                        'E_edges': _analysis.get('n_edges'),
                        'components': _analysis.get('n_components'),
                        'largest_component': _analysis.get('largest_component_size'),
                    })

                except Exception as e:
                    st.warning(f"⚠️ {f.name} failed: {e}")
                    batch_summary_rows.append({
                        'file': f.name, 'n_records': len(_df),
                        'lens': projection_type, 'error': str(e)[:100],
                    })

                progress.progress((idx + 1) / len(uploaded_files))

            # Write batch summary
            _batch_df = pd.DataFrame(batch_summary_rows)
            zf.writestr("batch_summary.csv", _batch_df.to_csv(index=False))
            zf.close()
            zip_buf.seek(0)

            status.text(f"✅ Batch complete — {len(batch_summary_rows)} files processed.")

            st.markdown("### 📊 Batch Summary")
            st.dataframe(_batch_df, use_container_width=True, hide_index=True)

            st.download_button(
                "📦 Download all results (zip)",
                zip_buf.getvalue(),
                file_name=f"syniq_batch_{_dt.now().strftime('%Y%m%d_%H%M%S')}.zip",
                mime="application/zip",
            )

        st.stop()  # don't fall through to single/combined-mode UI
    # =========================================================================
    # END v15 BATCH MODE
    # =========================================================================

    # =========================================================================
    # v16: STABILITY SWEEP — Farzana's 23-lens §4.2 protocol
    # =========================================================================
    if sweep_mode and uploaded_files:
        import io as _io, zipfile as _zip, tempfile as _tmp
        from datetime import datetime as _dt
        import networkx as _nx
        from collections import Counter as _Counter

        # --- Pool data ---
        _frames = []
        for f in uploaded_files:
            f.seek(0)
            try:
                _d = pd.read_csv(f) if f.name.endswith('.csv') else pd.DataFrame(json.load(f))
            except Exception as e:
                st.error(f"Could not read {f.name}: {e}")
                continue
            if 'agent' in _d.columns:
                _d['agent'] = _d['agent'].replace({'Sophia':'ChatGPT','sophia':'ChatGPT'})
            _frames.append(_d)
        if not _frames:
            st.error("No data loaded.")
            st.stop()
        _pool = pd.concat(_frames, ignore_index=True)

        # --- Composition summary ---
        _agent_name = (_pool['agent'].iloc[0]
                       if 'agent' in _pool.columns and len(_pool) > 0 else 'Agent')
        _agents_in_pool = sorted(_pool['agent'].unique().tolist()) if 'agent' in _pool.columns else []
        if len(_agents_in_pool) > 1:
            st.warning(f"Multiple agents in pool: {_agents_in_pool}. "
                       f"§4.2 expects one agent per sweep. Using '{_agent_name}' for output naming.")
        _conds = (sorted(_pool['temperature'].unique().tolist())
                  if 'temperature' in _pool.columns else [])

        st.markdown(f"## 🔬 §4.2 Stability Sweep — {_agent_name}")
        c1,c2,c3 = st.columns(3)
        c1.metric("Pooled n", len(_pool))
        c2.metric("Files", len(uploaded_files))
        c3.metric("Conditions", len(_conds))
        st.markdown(f"**Conditions:** {', '.join(str(c) for c in _conds)}")

        if st.button("🚀 Run 23-lens Sweep", type="primary"):
            # --- Build feature matrix (same as v14 — IEP+VADER+readability+V_t) ---
            with st.spinner("Building features..."):
                _data, _valid = build_iep_features(_pool)
                if len(_data) == 0:
                    st.error("No valid features.")
                    st.stop()
                _data = StandardScaler().fit_transform(_data)
                _subset = _pool.iloc[_valid].reset_index(drop=True)

            # --- Fit PCA on embeddings (for geometric lenses) ---
            _pca_proj = None
            if 'embedding' in _pool.columns:
                with st.spinner("Fitting PCA on embeddings..."):
                    _emb = []
                    for e in _subset['embedding']:
                        try:
                            _emb.append(np.array(json.loads(e)) if isinstance(e, str) else None)
                        except Exception:
                            try:
                                import ast as _ast
                                _emb.append(np.array(_ast.literal_eval(e)))
                            except Exception:
                                _emb.append(None)
                    _good = [i for i,e in enumerate(_emb) if e is not None and len(e) > 0]
                    if _good:
                        _dim = len(_emb[_good[0]])
                        _mat = np.zeros((len(_emb), _dim))
                        for i,e in enumerate(_emb):
                            if e is not None and len(e) == _dim:
                                _mat[i] = e
                        try:
                            _pca = PCA(n_components=2).fit(_mat)
                            _pca_proj = _pca.transform(_mat)
                            st.success(f"PCA fit on {_mat.shape[0]} embeddings ({_dim}-d). "
                                       f"Variance explained: {_pca.explained_variance_ratio_[0]:.3f}, "
                                       f"{_pca.explained_variance_ratio_[1]:.3f}")
                        except Exception as e:
                            st.warning(f"PCA failed: {e} — geometric lenses will be skipped.")

            # --- Farzana's 23-lens spec ---
            def _norm(arr):
                arr = np.asarray(arr, dtype=float)
                mn, mx = np.nanmin(arr), np.nanmax(arr)
                if mx == mn or not np.isfinite(mx-mn):
                    return np.zeros_like(arr)
                return (arr - mn) / (mx - mn)

            _LENSES = [
                ("aff_pct", "AFF", "1D", "col", ["aff_pct"]),
                ("vader_compound", "AFF", "1D", "col", ["vader_compound"]),
                ("|vader_compound|", "AFF", "1D", "abs", ["vader_compound"]),
                ("(aff_pct, vader_compound)", "AFF", "2D", "col2", ["aff_pct","vader_compound"]),
                ("(vader_pos, vader_neg)", "AFF", "2D", "col2", ["vader_pos","vader_neg"]),
                ("(aff_pct, PCA1)", "AFF", "2D", "colpca", ["aff_pct"]),
                ("int_pct", "INT", "1D", "col", ["int_pct"]),
                ("flesch_kincaid", "INT", "1D", "col", ["flesch_kincaid"]),
                ("ttr", "INT", "1D", "col", ["ttr"]),
                ("(int_pct, flesch_kincaid)", "INT", "2D", "col2", ["int_pct","flesch_kincaid"]),
                ("(int_pct, ttr)", "INT", "2D", "col2", ["int_pct","ttr"]),
                ("(flesch_kincaid, ttr)", "INT", "2D", "col2", ["flesch_kincaid","ttr"]),
                ("(int_pct, PCA1)", "INT", "2D", "colpca", ["int_pct"]),
                ("act_pct", "ACT", "1D", "col", ["act_pct"]),
                ("total_words", "ACT", "1D", "col", ["total_words"]),
                ("(act_pct, total_words)", "ACT", "2D", "col2", ["act_pct","total_words"]),
                ("(act_pct, ttr)", "ACT", "2D", "col2", ["act_pct","ttr"]),
                ("(act_pct, vader_compound)", "ACT", "2D", "col2", ["act_pct","vader_compound"]),
                ("(aff_pct, int_pct)", "Cross-IEP", "2D", "col2", ["aff_pct","int_pct"]),
                ("(int_pct, act_pct)", "Cross-IEP", "2D", "col2", ["int_pct","act_pct"]),
                ("(aff_pct, act_pct)", "Cross-IEP", "2D", "col2", ["aff_pct","act_pct"]),
                ("PCA1", "Geometric", "1D", "pca1", []),
                ("(PCA1, PCA2)", "Geometric", "2D", "pca2", []),
            ]

            def _build_lens(kind, cols, sub):
                if kind == "col":
                    return _norm(sub[cols[0]].fillna(0).values).reshape(-1,1)
                if kind == "abs":
                    return _norm(np.abs(sub[cols[0]].fillna(0).values)).reshape(-1,1)
                if kind == "col2":
                    return np.column_stack([_norm(sub[c].fillna(0).values) for c in cols])
                if kind == "colpca":
                    if _pca_proj is None: return None
                    return np.column_stack([
                        _norm(sub[cols[0]].fillna(0).values),
                        _norm(_pca_proj[:,0])
                    ])
                if kind == "pca1":
                    if _pca_proj is None: return None
                    return _norm(_pca_proj[:,0]).reshape(-1,1)
                if kind == "pca2":
                    if _pca_proj is None: return None
                    return np.column_stack([_norm(_pca_proj[:,0]), _norm(_pca_proj[:,1])])
                return None

            def _q_purity(graph, sub):
                if 'question_id' not in sub.columns: return None
                ps, ws = [], []
                for nid, members in graph.get('nodes', {}).items():
                    if not members: continue
                    qs = sub.iloc[list(members)]['question_id'].tolist()
                    if not qs: continue
                    modal = _Counter(qs).most_common(1)[0][1]
                    ps.append(modal/len(qs)); ws.append(len(qs))
                return round(float(np.average(ps, weights=ws)), 4) if ps else None

            # --- Run sweep ---
            _zb = _io.BytesIO()
            _zf = _zip.ZipFile(_zb, 'w', _zip.ZIP_DEFLATED)
            _rows = []
            _prog = st.progress(0)
            _stat = st.empty()

            for i, (lname, fam, dim, kind, cols) in enumerate(_LENSES):
                _stat.text(f"Lens {i+1}/{len(_LENSES)}: {lname}")
                _prog.progress((i+1)/len(_LENSES))

                _missing = [c for c in cols if c not in _subset.columns]
                if _missing:
                    _rows.append({'lens':lname,'family':fam,'dim':dim,'status':f'SKIP missing {_missing}'})
                    continue
                if kind in ('colpca','pca1','pca2') and _pca_proj is None:
                    _rows.append({'lens':lname,'family':fam,'dim':dim,'status':'SKIP no PCA'})
                    continue

                try:
                    _lens = _build_lens(kind, cols, _subset)
                    if _lens is None:
                        _rows.append({'lens':lname,'family':fam,'dim':dim,'status':'SKIP lens=None'})
                        continue

                    _mp = km.KeplerMapper(verbose=0)
                    _g = _mp.map(
                        _lens, _data,
                        cover=km.Cover(n_cubes=n_cubes, perc_overlap=overlap),
                        clusterer=DBSCAN(eps=eps, min_samples=min_samples),
                    )

                    _G = _nx.Graph()
                    _nodes = _g.get('nodes', {})
                    _links = _g.get('links', {})
                    for nn in _nodes: _G.add_node(nn)
                    for s, ts in _links.items():
                        for t in ts: _G.add_edge(s, t)
                    _V = _G.number_of_nodes(); _E = _G.number_of_edges()
                    _comps = list(_nx.connected_components(_G))
                    _C = len(_comps); _b1 = _E - _V + _C
                    _csizes = sorted([sum(len(_nodes[n]) for n in cc) for cc in _comps], reverse=True)
                    _largest = _csizes[0] if _csizes else 0
                    _total = sum(_csizes) if _csizes else 0
                    _frac = (_largest/_total) if _total else 0.0
                    _qp = _q_purity(_g, _subset)

                    # Save HTML
                    if _V > 0:
                        _safe = ''.join(c if c.isalnum() or c in '_-' else '_' for c in lname)
                        _banner = (f"Lens: {lname} · n_cubes={n_cubes} · overlap={int(overlap*100)}% · "
                                   f"DBSCAN(eps={eps}, min={min_samples}) · "
                                   f"V={_V} E={_E} β₀={_C} β₁={_b1}")
                        with _tmp.NamedTemporaryFile(mode='w', suffix='.html', delete=False, encoding='utf-8') as _tf:
                            _tpath = _tf.name
                        try:
                            _color = (_subset['question_id'].astype('category').cat.codes.astype(float).values
                                      if 'question_id' in _subset.columns else None)
                            _tt = (np.array([f"Q={q} | cond={t}"
                                             for q,t in zip(
                                                 _subset.get('question_id','?'),
                                                 _subset.get('temperature','?'))])
                                   if 'question_id' in _subset.columns else None)
                            _mp.visualize(_g, path_html=_tpath,
                                          title=f"§4.2 · {_agent_name} · {lname}",
                                          color_values=_color,
                                          color_function_name='question_id',
                                          custom_tooltips=_tt,
                                          include_searchbar=True)
                            with open(_tpath, 'r', encoding='utf-8') as _hh:
                                _html = _hh.read()
                            _html = whiten_mapper_html(_html)  # v17: white background
                            # inject banner under <body>
                            _html = re.sub(r'(<body[^>]*>)',
                                           rf'\1<div style="background:#f0f4f8;border-left:4px solid #2e75b6;'
                                           rf'padding:0.6rem 1rem;font-family:monospace;font-size:0.85rem;'
                                           rf'color:#1e3a5f;">{_banner}</div>',
                                           _html, count=1)
                        finally:
                            try: os.unlink(_tpath)
                            except Exception: pass
                        _zf.writestr(f"{_agent_name}_mappers/{_safe}.html", _html)

                    _rows.append({
                        'lens': lname, 'family': fam, 'dim': dim,
                        'V': _V, 'E': _E, 'beta_0': _C, 'beta_1': _b1,
                        'largest_members': _largest, 'total_members': _total,
                        'largest_frac': round(_frac, 4),
                        'Q_purity': _qp, 'status': 'OK',
                    })
                except Exception as e:
                    _rows.append({'lens':lname,'family':fam,'dim':dim,'status':f'ERROR {str(e)[:80]}'})

            _stab = pd.DataFrame(_rows)
            _zf.writestr(f"{_agent_name}_stability.csv", _stab.to_csv(index=False))

            # Provenance
            _prov = pd.DataFrame([{'source_file': f.name, 'n': len(pd.read_csv(f) if f.name.endswith('.csv') else pd.DataFrame())}
                                   for f in uploaded_files for _ in [f.seek(0)]])
            _zf.writestr(f"{_agent_name}_provenance.csv", _prov.to_csv(index=False))

            # Methods docx (Layer 2)
            try:
                from docx import Document as _Doc
                _doc = _Doc()
                _doc.add_heading(f"§4.2 Stability Sweep — {_agent_name}", 1)
                _doc.add_paragraph(f"Generated {_dt.now().strftime('%B %d, %Y · %H:%M')}")
                _doc.add_heading("Data", 2)
                _doc.add_paragraph(
                    f"Pooled n={len(_pool)} from {len(uploaded_files)} CSVs. "
                    f"Conditions: {', '.join(str(c) for c in _conds)}.")
                _doc.add_heading("Methods", 2)
                _doc.add_paragraph(
                    f"All Mapper runs used n_cubes={n_cubes}, overlap={int(overlap*100)}%, "
                    f"DBSCAN(eps={eps}, min_samples={min_samples}). Features = IEP V3 "
                    f"feature matrix (StandardScaler-normalized). PCA fit once on the "
                    f"response embedding column and reused across PCA-using lenses. "
                    f"23 lenses applied per Farzana §4.2 protocol.")
                _vok = _stab[_stab.get('V').notna()] if 'V' in _stab.columns else _stab.head(0)
                if len(_vok) > 0:
                    _doc.add_heading("Findings", 2)
                    _doc.add_paragraph(
                        f"{len(_vok)}/{len(_LENSES)} lenses succeeded. β₀ ranged "
                        f"{int(_vok['beta_0'].min())}-{int(_vok['beta_0'].max())} "
                        f"(mean {_vok['beta_0'].mean():.2f} ± {_vok['beta_0'].std():.2f}). "
                        f"Largest-component fraction averaged {_vok['largest_frac'].mean():.3f}. "
                        f"Q-purity averaged {_vok['Q_purity'].mean():.3f}.")
                _docbuf = _io.BytesIO(); _doc.save(_docbuf); _docbuf.seek(0)
                _zf.writestr(f"{_agent_name}_methods.docx", _docbuf.getvalue())
            except Exception as e:
                _zf.writestr(f"{_agent_name}_methods_FAILED.txt", str(e))

            _zf.close(); _zb.seek(0)

            _stat.text(f"✅ Sweep complete — {len([r for r in _rows if r.get('status')=='OK'])}/{len(_LENSES)} lenses OK")

            # Layer 1: surface table grouped by family
            st.markdown("### 📊 Stability Table")
            _vok2 = _stab[_stab['status']=='OK'] if 'status' in _stab.columns else _stab
            if len(_vok2) > 0:
                m1,m2,m3,m4 = st.columns(4)
                m1.metric("β₀ range", f"{int(_vok2['beta_0'].min())}–{int(_vok2['beta_0'].max())}")
                m2.metric("β₀ SD", f"{_vok2['beta_0'].std():.2f}")
                m3.metric("Mean largest frac", f"{_vok2['largest_frac'].mean():.3f}")
                m4.metric("Mean Q-purity", f"{_vok2['Q_purity'].mean():.3f}")
            for fam in ['AFF','INT','ACT','Cross-IEP','Geometric']:
                _sub = _stab[_stab['family']==fam] if 'family' in _stab.columns else _stab.head(0)
                if len(_sub) > 0:
                    st.markdown(f"**{fam}**")
                    st.dataframe(_sub, use_container_width=True, hide_index=True)

            st.download_button(
                f"📦 Download §4.2 bundle for {_agent_name}",
                _zb.getvalue(),
                file_name=f"{_agent_name}_section_4_2_{_dt.now().strftime('%Y%m%d_%H%M%S')}.zip",
                mime="application/zip",
            )
        st.stop()
    # =========================================================================
    # END v16 STABILITY SWEEP
    # =========================================================================

    if df_combined is not None:
        df = df_combined.copy()
        if 'agent' in df.columns:
            df['agent'] = df['agent'].replace({'Sophia': 'ChatGPT', 'sophia': 'ChatGPT'})
    elif uploaded_file is not None:
        if uploaded_file.name.endswith('.json'):
            raw_data = json.loads(uploaded_file.read().decode('utf-8'))
            df = pd.DataFrame(raw_data)
            rename_map = {}
            if 'full_int_pct' in df.columns and 'int_pct' not in df.columns:
                rename_map.update({'full_int_pct': 'int_pct', 'full_aff_pct': 'aff_pct',
                                   'full_act_pct': 'act_pct', 'full_vader_compound': 'vader_compound',
                                   'full_total_words': 'total_words'})
            if 'condition' in df.columns and 'temperature' not in df.columns:
                rename_map['condition'] = 'temperature'
            if rename_map:
                df = df.rename(columns=rename_map)
            if 'has_italics' in df.columns:
                df['has_italics'] = df['has_italics'].astype(int)
            if 'q2_response' in df.columns:
                df['response_text'] = df['q2_response']
                st.session_state["_sidebar_msg_json_text"] = True
            st.session_state["_sidebar_msg_record_count"] = len(df)
        else:
            df = pd.read_csv(uploaded_file)
        if 'agent' in df.columns:
            df['agent'] = df['agent'].replace({'Sophia': 'ChatGPT', 'sophia': 'ChatGPT'})
    else:
        df = None

    if df is not None:
        if "live score" in use_iep:
            if text_col_choice in df.columns:
                with st.spinner(f"🔬 Scoring with IEP Dictionary V3 ({text_col_choice})…"):
                    df = score_dataframe_iep(df, text_col_choice)
                st.success(f"✅ IEP V3 scoring applied to `{text_col_choice}`")
            else:
                st.warning(
                    f"⚠️ Column `{text_col_choice}` not found. "
                    f"Available: {list(df.columns[:10])}. "
                    f"Falling back to pre-scored columns."
                )

        # ---- CONDITION FILTER ----
        if 'temperature' in df.columns:
            all_conditions = sorted(df['temperature'].unique().tolist())
            st.sidebar.markdown("---")
            st.sidebar.markdown("### 🎯 Condition Filter")
            filter_conditions = st.sidebar.multiselect(
                "Include conditions:", all_conditions, default=all_conditions)
            if filter_conditions and len(filter_conditions) < len(all_conditions):
                df = df[df['temperature'].isin(filter_conditions)].reset_index(drop=True)
                st.sidebar.success(f"✅ Filtered to {len(df)} records")

        # ---- AGENT FILTER ----
        if 'agent' in df.columns:
            all_agents = sorted(df['agent'].unique().tolist())
            filter_agents = st.sidebar.multiselect("Include agents:", all_agents, default=all_agents)
            if filter_agents and len(filter_agents) < len(all_agents):
                df = df[df['agent'].isin(filter_agents)].reset_index(drop=True)
                st.sidebar.success(f"✅ Filtered to {len(df)} records")

        # ---- STATS BOXES ----
        col1, col2, col3, col4 = st.columns(4)
        with col1:
            st.markdown(f'<div class="stats-box"><h3>{len(df)}</h3><p>Records</p></div>', unsafe_allow_html=True)
        with col2:
            n_agents = df['agent'].nunique() if 'agent' in df.columns else 0
            st.markdown(f'<div class="stats-box"><h3>{n_agents}</h3><p>Agents</p></div>', unsafe_allow_html=True)
        with col3:
            n_temps = df['temperature'].nunique() if 'temperature' in df.columns else 0
            st.markdown(f'<div class="stats-box"><h3>{n_temps}</h3><p>Temperatures</p></div>', unsafe_allow_html=True)
        with col4:
            feature_dim = "12D" if "IEP" in use_iep or "live" in use_iep else "384D"
            st.markdown(f'<div class="stats-box"><h3>{feature_dim}</h3><p>Features</p></div>', unsafe_allow_html=True)

        st.markdown("---")

        # ---- PREVIEW ----
        with st.expander("📋 Preview Data"):
            preview_cols = ['agent', 'temperature', 'int_pct', 'aff_pct', 'act_pct', 'vader_compound']
            preview_cols = [c for c in preview_cols if c in df.columns]
            st.dataframe(df[preview_cols].head(20))

        # ---- IEP DISTRIBUTION ----
        if all(c in df.columns for c in ['int_pct', 'aff_pct', 'act_pct']):
            with st.expander("📊 IEP V3 Score Distribution"):
                dist_df = df[['int_pct', 'aff_pct', 'act_pct']].describe().T
                dist_df.columns = ["count", "mean", "std", "min", "25%", "50%", "75%", "max"]
                st.dataframe(dist_df.style.format("{:.2f}"))

        if all(c in df.columns for c in ['int_pct', 'aff_pct', 'act_pct']):
            iep_sum_check = df[['int_pct', 'aff_pct', 'act_pct']].sum(axis=1).mean()
            if iep_sum_check < 50:
                st.error(
                    f"🚨 **IEP Score Warning**: Mean IEP sum = {iep_sum_check:.1f}% (expected ~100%). "
                    f"True dataset means: INT={df['int_pct'].mean():.1f}% | "
                    f"AFF={df['aff_pct'].mean():.1f}% | ACT={df['act_pct'].mean():.1f}%. "
                    f"Check that 'IEP V3 live score' ran successfully on the correct text column."
                )
            else:
                st.success(
                    f"✅ IEP scores valid — Mean sum: {iep_sum_check:.1f}% | "
                    f"INT={df['int_pct'].mean():.1f}% | AFF={df['aff_pct'].mean():.1f}% | ACT={df['act_pct'].mean():.1f}%"
                )

        # ---- RUN MAPPER ----
        if st.button("🚀 Run Mapper Analysis", type="primary"):
            with st.spinner("Preparing data..."):
                if "live score" in use_iep or "pre-scored" in use_iep:
                    data, valid_indices = build_iep_features(df)
                    if len(data) > 0:
                        scaler = StandardScaler()
                        data = scaler.fit_transform(data)
                elif "SBERT" in use_iep:
                    try:
                        from sentence_transformers import SentenceTransformer
                        model = SentenceTransformer('all-MiniLM-L6-v2')
                        texts = df['response_text'].fillna('').tolist()
                        valid_indices = [i for i, t in enumerate(texts) if t.strip()]
                        data = model.encode([texts[i] for i in valid_indices])
                        st.success(f"✅ SBERT embeddings: {data.shape}")
                    except ImportError:
                        st.error("❌ sentence-transformers not installed.")
                        st.stop()
                else:
                    data, valid_indices = parse_embeddings(df)

                if len(data) == 0:
                    st.error("❌ No valid data found! Check your data format.")
                    st.stop()
                st.success(f"✅ {len(data)} data points · {data.shape[1]} dimensions")

            with st.spinner("Running Mapper algorithm..."):
                mapper, graph = run_mapper_analysis(
                    data, df, valid_indices,
                    n_cubes, overlap, eps, min_samples, projection_type
                )

            if len(graph['nodes']) == 0:
                st.error("❌ Mapper produced 0 nodes. Adjust DBSCAN eps / min_samples.")
                st.stop()

            with st.spinner("Analyzing graph..."):
                analysis = analyze_graph(graph, df, valid_indices)

            with st.spinner("Generating visualization..."):
                subset = df.iloc[valid_indices].reset_index(drop=True)

                # ── V18 ROUTER ────────────────────────────────────────────────
                # If user picked a V18 mode (PCA1 projection OR one of the new
                # agent-color modes), bypass kepler-mapper's HTML renderer and
                # use the Plotly fallback instead.
                _is_v18_proj  = (projection_type == "PCA1 (1D)")
                _is_v18_color = color_by in ("agent (overlap=orange)",
                                              "agent (pie composition)")
                if _is_v18_proj or _is_v18_color:
                    # Decide which V18 mode to render
                    if _is_v18_proj:
                        _v18_mode = "pca1_strip"
                    elif color_by == "agent (overlap=orange)":
                        _v18_mode = "overlap_orange"
                    else:
                        _v18_mode = "pie_composition"

                    # PCA1 strip plot needs the projection coordinates
                    _v18_projected = None
                    _v18_explained_var = None
                    if _is_v18_proj:
                        try:
                            _pca1_fit = PCA(n_components=1).fit(data)
                            _v18_projected = _pca1_fit.transform(data)
                            _v18_explained_var = float(_pca1_fit.explained_variance_ratio_[0])
                        except Exception as _e:
                            st.error(f"PCA1 fit failed: {_e}")
                            _v18_projected = None

                    _v18_title = f"SYN-IQ V18 · {projection_type} · {color_by}"
                    _fig = render_plotly_agent_view(
                        graph=graph, df=df, valid_indices=valid_indices,
                        mode=_v18_mode, purity_threshold=purity_threshold,
                        projected=_v18_projected,
                        projection_type=projection_type,
                        explained_variance=_v18_explained_var,
                        title=_v18_title,
                    )
                    if _fig is not None:
                        st.session_state.viz_title = _v18_title
                        st.plotly_chart(_fig, use_container_width=True)
                        _var_note = (f" · PC1 explained_variance={_v18_explained_var*100:.1f}%"
                                     if _v18_explained_var is not None else "")
                        st.caption(
                            f"V18 render · mode={_v18_mode} · "
                            f"purity_threshold={purity_threshold:.0%} · "
                            f"agents colored per AGENT_COLORS contract"
                            f"{_var_note}"
                        )
                        # Skip the legacy HTML pipeline below
                        st.session_state.mapper_results_html = None
                        st.stop()

                if color_by in subset.columns:
                    if subset[color_by].dtype == 'object':
                        cats = subset[color_by].unique().tolist()
                        color_values = subset[color_by].map({c: i for i, c in enumerate(cats)}).values
                    else:
                        color_values = subset[color_by].values
                else:
                    color_values = subset['aff_pct'].values if 'aff_pct' in subset.columns else np.zeros(len(subset))

                if 'temperature' in subset.columns and 'agent' in subset.columns:
                    if 'question_id' in subset.columns:
                        tooltips = np.array([
                            f"{r['agent']} | {r['temperature']} | {r.get('question_id', '')}"
                            for _, r in subset.iterrows()
                        ])
                    else:
                        tooltips = np.array([f"{r['agent']} | {r['temperature']}" for _, r in subset.iterrows()])
                else:
                    tooltips = subset.index.astype(str).values

                # ── Auto-generate meaningful title from data ──────────────────
                _agents = sorted(subset['agent'].unique().tolist()) if 'agent' in subset.columns else []
                _temps  = sorted(subset['temperature'].unique().tolist()) if 'temperature' in subset.columns else []
                _qcol   = 'question_id' if 'question_id' in subset.columns else 'question_label' if 'question_label' in subset.columns else None
                _qs     = sorted(subset[_qcol].unique().tolist()) if _qcol else []

                # Agent string
                _agent_str = " + ".join(_agents) if len(_agents) <= 3 else f"{len(_agents)} Agents"

                # Smart condition string — detect gradient types
                _temp_set = set(_temps)
                if _temp_set & {'AFF_1','AFF_2','AFF_3','AFF_4','AFF_5'}:
                    _has_native = 'NATIVE' in _temp_set
                    _temp_str = "Native + AFF Gradient" if _has_native else "AFF Gradient"
                elif _temp_set & {'INT_1','INT_2','INT_3','INT_4','INT_5'}:
                    _has_native = 'NATIVE' in _temp_set
                    _temp_str = "Native + INT Gradient" if _has_native else "INT Gradient"
                elif _temp_set & {'ACT_1','ACT_2','ACT_3','ACT_4','ACT_5'}:
                    _has_native = 'NATIVE' in _temp_set
                    _temp_str = "Native + ACT Gradient" if _has_native else "ACT Gradient"
                elif _temp_set <= {'COLD','NATIVE','HOT','FIRE'}:
                    if len(_temps) == 1:
                        _temp_str = _temps[0].title()
                    else:
                        _temp_str = "Temp Gradient"
                elif len(_temps) <= 4:
                    _temp_str = " / ".join(_temps)
                else:
                    _temp_str = f"{len(_temps)} Conditions"

                # Smart question string — use name if 1, count if many
                if len(_qs) == 1:
                    _q_str = _qs[0].replace('_', ' ').title()
                elif len(_qs) <= 3:
                    _q_str = " / ".join(q.replace('_',' ').title() for q in _qs)
                else:
                    _q_str = f"{len(_qs)} Questions"

                _parts = [p for p in [_agent_str, _temp_str, _q_str] if p]
                viz_title = " — ".join(_parts) if _parts else "SYN-IQ Mapper Analysis"
                st.session_state.viz_title = viz_title

                with tempfile.NamedTemporaryFile(delete=False, suffix='.html') as tmp:
                    mapper.visualize(
                        graph,
                        path_html=tmp.name,
                        title=viz_title,
                        color_values=color_values,
                        color_function_name=color_by,
                        custom_tooltips=tooltips
                    )
                    with open(tmp.name) as f:
                        html_content = f.read()
                    os.unlink(tmp.name)

                # ── Targeted hex replacement — backgrounds ONLY not color values ─
                # Only replace these specific background hex values in CSS context
                html_content = html_content.replace('background: #212121;', 'background: #ffffff;')
                html_content = html_content.replace('background: #111111;', 'background: #ffffff;')
                html_content = html_content.replace('background: #191919;', 'background: #ffffff;')
                html_content = html_content.replace('background: #212121\n', 'background: #ffffff\n')
                html_content = html_content.replace('background-color: #000000\n', 'background-color: #ffffff\n')

                # ── CSS injection into <head> — works before JS runs ──────────
                white_css = """
<style>
  body, html { background: #ffffff !important; }
  #canvas, #display, #print { background: #ffffff !important; }
  svg { background: #ffffff !important; }
  svg rect { fill: #ffffff !important; }
  line.link, .link { stroke: #444444 !important; stroke-opacity: 0.6 !important; }
  #logo, .wrap-logo { display: none !important; }
  path.circle { stroke: #ffffff !important; stroke-width: 2px !important; }
</style>
"""
                html_content = html_content.replace('</head>', white_css + '</head>')

                # ── Inject size_modifier into graph JSON for scaled node sizing ─
                # Target: publication style matching Dr. Nasrin's PLOS ONE figures
                # KeplerMapper: node area = d.size * 50 * d.size_modifier
                # We want ~15px min diameter, ~45px max diameter
                # D3 symbol area = pi*r^2, so r=15 → area≈707, r=45 → area≈6362
                # With d.size * 50 * modifier: modifier = area / (size * 50)
                import json as _json
                import re as _re
                import math as _math
                try:
                    gm = _re.search(r'const graph = (\{.+?\});\s*\n', html_content, _re.DOTALL)
                    if gm:
                        gj = _json.loads(gm.group(1))
                        gn = gj.get('nodes', [])
                        if gn:
                            max_sz = max(n.get('size', 1) for n in gn)
                            min_sz = min(n.get('size', 1) for n in gn)
                            for n in gn:
                                sz = max(n.get('size', 1), 1)
                                # Linear scale: min_modifier=6, max_modifier=18
                                # Gives clearly visible graduation between node sizes
                                if max_sz == min_sz:
                                    n['size_modifier'] = 10.0
                                else:
                                    t = (sz - min_sz) / (max_sz - min_sz)
                                    n['size_modifier'] = round(6.0 + 12.0 * t, 2)
                            new_graph = 'const graph = ' + _json.dumps(gj) + ';\n'
                            html_content = html_content.replace(gm.group(0), new_graph)
                except Exception:
                    pass

                # ── MutationObserver targeting path.circle (KeplerMapper's actual element) ─
                pub_script = """
<script>
(function() {
  function applyStyle() {
    // KeplerMapper draws nodes as path.circle NOT <circle> elements
    var paths = document.querySelectorAll('path.circle');
    if (!paths.length) return;

    // White stroke for clean node separation on white background
    paths.forEach(function(p) {
      p.style.stroke = '#ffffff';
      p.style.strokeWidth = '2px';
      p.style.fillOpacity = '0.92';
    });

    // White SVG background
    document.querySelectorAll('svg').forEach(function(s) {
      s.style.background = '#ffffff';
    });
    document.querySelectorAll('svg rect').forEach(function(r) {
      r.style.fill = '#ffffff';
    });

    // Darker edges visible on white
    document.querySelectorAll('line.link, .link').forEach(function(el) {
      el.style.stroke = '#444444';
      el.style.strokeOpacity = '0.65';
    });

    // Hide logo
    var logo = document.getElementById('logo');
    if (logo) logo.style.display = 'none';
    var wrap = document.querySelector('.wrap-logo');
    if (wrap) wrap.style.display = 'none';
  }

  // Watch for path.circle to appear
  var observer = new MutationObserver(function() {
    if (document.querySelector('path.circle')) {
      setTimeout(applyStyle, 100);
      observer.disconnect();
    }
  });
  observer.observe(document.documentElement, { childList: true, subtree: true });

  // Safety nets
  [500, 1200, 2500, 4000].forEach(function(t) { setTimeout(applyStyle, t); });
})();
</script>
"""

                # Prevent iframe navigation on node click
                nav_fix = """
<script>
document.addEventListener('DOMContentLoaded', function() {
    document.addEventListener('click', function(e) {
        var a = e.target.closest('a');
        if (a) { e.preventDefault(); e.stopPropagation(); }
    }, true);
    document.querySelectorAll('a[href]').forEach(function(l) {
        l.setAttribute('href', 'javascript:void(0)');
        l.style.cursor = 'pointer';
    });
});
</script>
"""
                html_content = html_content.replace('</body>', pub_script + nav_fix + '</body>')

            st.session_state.mapper_results_html = html_content
            st.session_state.mapper_analysis_data = analysis
            st.session_state.stored_df = df.copy()
            st.session_state.stored_graph = graph
            st.session_state.chat_history = []
            st.session_state.topology_context = None
            st.session_state.chat_active = False
            st.session_state.claude_analysis = None
            st.rerun()

        # ---- DISPLAY RESULTS ----
        if st.session_state.mapper_analysis_data is not None:
            analysis = st.session_state.mapper_analysis_data

            st.markdown("## 📊 Results")
            col1, col2 = st.columns(2)
            col1.metric("Nodes", analysis["n_nodes"])
            col2.metric("Edges", analysis["n_edges"])

            # ── Topology Summary Table (Dr. Nasrin) ──────────────────────────
            st.markdown("### 📐 Topology Summary")
            node_sizes = [info["size"] for info in analysis["nodes"].values()]
            topo_summary = pd.DataFrame({
                "Metric": [
                    "Nodes",
                    "Edges",
                    "Connected Components",
                    "Largest Component (nodes)",
                    "Isolated Nodes",
                    "Largest Node (members)",
                    "Mean Node Size (members)",
                ],
                "Value": [
                    analysis["n_nodes"],
                    analysis["n_edges"],
                    analysis.get("n_components", "N/A"),
                    analysis.get("largest_component", "N/A"),
                    analysis.get("isolated_nodes", "N/A"),
                    max(node_sizes) if node_sizes else 0,
                    round(sum(node_sizes) / len(node_sizes), 1) if node_sizes else 0,
                ]
            })
            st.dataframe(topo_summary, use_container_width=True, hide_index=True)

            # ── Excel download — nicely formatted ─────────────────────────────
            try:
                from openpyxl import Workbook
                from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
                from openpyxl.utils import get_column_letter
                import io as _io

                wb = Workbook()
                ws = wb.active
                ws.title = "Topology Summary"

                # Title row
                ws.merge_cells('A1:B1')
                ws['A1'] = f"SYN-IQ Topology Summary — {st.session_state.get('viz_title', '')}"
                ws['A1'].font = Font(bold=True, size=13, color='FFFFFF')
                ws['A1'].fill = PatternFill('solid', start_color='2E75B6')
                ws['A1'].alignment = Alignment(horizontal='center', vertical='center')
                ws.row_dimensions[1].height = 22

                # Header row
                headers = ['Metric', 'Value']
                for col, h in enumerate(headers, 1):
                    cell = ws.cell(row=2, column=col, value=h)
                    cell.font = Font(bold=True, color='FFFFFF')
                    cell.fill = PatternFill('solid', start_color='4472C4')
                    cell.alignment = Alignment(horizontal='center')

                # Data rows
                thin = Side(style='thin', color='CCCCCC')
                border = Border(left=thin, right=thin, top=thin, bottom=thin)
                for i, row in topo_summary.iterrows():
                    r = i + 3
                    fill = PatternFill('solid', start_color='EBF3FB') if i % 2 == 0 else PatternFill('solid', start_color='FFFFFF')
                    for col, val in enumerate([row['Metric'], row['Value']], 1):
                        cell = ws.cell(row=r, column=col, value=val)
                        cell.fill = fill
                        cell.border = border
                        cell.alignment = Alignment(horizontal='left' if col==1 else 'center')

                # Column widths
                ws.column_dimensions['A'].width = 30
                ws.column_dimensions['B'].width = 18

                buf = _io.BytesIO()
                wb.save(buf)
                buf.seek(0)
                fname = f"{st.session_state.get('viz_title','topology').replace(' — ','_').replace(' ','_')}_summary.xlsx"
                st.download_button(
                    "📥 Download Topology Table (Excel)",
                    buf.getvalue(),
                    file_name=fname,
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                )
            except Exception as _xl_err:
                st.download_button(
                    "📥 Download Topology Table (CSV)",
                    topo_summary.to_csv(index=False),
                    file_name="syniq_topology_summary.csv",
                    mime="text/csv",
                )

            st.markdown("### 🌡️ Temperature Separation")
            col1, col2, col3 = st.columns(3)
            col1.metric("COLD-only nodes", len(analysis["cold_only"]))
            col2.metric("HOT-only nodes",  len(analysis["hot_only"]))
            col3.metric("Mixed nodes",     len(analysis["overlap_nodes"]))

            if analysis["cold_only"] and analysis["hot_only"]:
                st.success("✅ **COLD and HOT show TOPOLOGICAL SEPARATION!**")
            elif analysis["overlap_nodes"] and not (analysis["cold_only"] or analysis["hot_only"]):
                st.warning("⚠️ **COLD and HOT are MIXED**")
            else:
                st.info("🔶 **Partial separation**")

            st.markdown("### 🔬 Node Analysis")
            for node_id, node_info in analysis["nodes"].items():
                with st.expander(f"Node {node_id} ({node_info['size']} points)"):
                    c1, c2 = st.columns(2)
                    with c1:
                        st.write("**Temperatures:**", node_info["temps"])
                        st.write("**Agents:**",       node_info["agents"])
                    with c2:
                        st.markdown("**Composition** *(drives topology)*")
                        st.write(f"INT: {node_info['avg_int']:.1f}%  |  AFF: {node_info['avg_aff']:.1f}%  |  ACT: {node_info['avg_act']:.1f}%")
                        comp_sum = node_info['avg_int'] + node_info['avg_aff'] + node_info['avg_act']
                        st.caption(f"∑ = {comp_sum:.1f}% {'✅' if abs(comp_sum-100)<2 else '⚠️ check scoring'}")
                        if node_info.get('density_mean', 0) > 0:
                            st.markdown("**Density** *(IEP coverage)*")
                            st.write(f"Mean {node_info['density_mean']:.3f} | Med {node_info['density_median']:.3f} | SEM \u00b1{node_info['density_sem']:.3f}")

            st.markdown("### 🖼️ Visualization")
            if st.session_state.mapper_results_html:
                st.download_button(
                    "📥 Download Mapper HTML",
                    st.session_state.mapper_results_html,
                    file_name=f"{st.session_state.get('viz_title', 'syniq_mapper').replace(' — ', '_').replace(' / ', '_').replace(' + ', '_').replace(' ', '_')}.html",
                    mime="text/html"
                )
                st.components.v1.html(st.session_state.mapper_results_html, height=800, scrolling=True)

            # ── AI ANALYSIS ──────────────────────────────────────────────────
            st.markdown("---")
            st.markdown("## 🤖 AI Analysis & Conversation")
            if not api_key:
                st.info("Enter your **Anthropic API key** in the sidebar to unlock AI analysis and chat.")
            else:
                use_df = st.session_state.stored_df if st.session_state.stored_df is not None else (df if df is not None else pd.DataFrame())

                if not st.session_state.chat_active:
                    c1, c2 = st.columns(2)
                    with c1:
                        if st.button("🔬 Analyse Topology", type="primary",
                                     help="Claude describes the topology structure — no hypothesis framing"):
                            with st.spinner("Analysing topology…"):
                                try:
                                    ctx = build_topology_context(analysis, use_df)
                                    st.session_state.topology_context = ctx
                                    prompt = build_claude_analysis_prompt(analysis, use_df)
                                    ai = call_claude_api(api_key, [{"role":"user","content":ctx+"\n\n"+prompt}])
                                    st.session_state.chat_history = [
                                        {"role":"user","content":prompt,"display":False},
                                        {"role":"assistant","content":ai,"display":True}]
                                    st.session_state.claude_analysis = ai
                                    st.session_state.chat_active = True
                                    st.rerun()
                                except Exception as e:
                                    st.error(f"Error: {e}")
                    with c2:
                        if st.button("📋 Full Report",
                                     help="Part A: topological regions · Part B: agent × condition × question patterns"):
                            with st.spinner("Writing full report…"):
                                try:
                                    ctx = build_topology_context(analysis, use_df)
                                    st.session_state.topology_context = ctx
                                    prompt = build_dual_report_prompt(analysis, use_df)
                                    ai = call_claude_api(api_key, [{"role":"user","content":prompt}], max_tokens=6000)
                                    st.session_state.chat_history = [
                                        {"role":"user","content":"[Full Report Request]","display":False},
                                        {"role":"assistant","content":ai,"display":True}]
                                    st.session_state.claude_analysis = ai
                                    st.session_state.chat_active = True
                                    st.rerun()
                                except Exception as e:
                                    st.error(f"Error: {e}")

                if st.session_state.chat_active and st.session_state.chat_history:
                    st.markdown('<div class="ctx-pill">🔬 Topology context loaded · Ask anything</div>', unsafe_allow_html=True)
                    for msg in st.session_state.chat_history:
                        if not msg.get("display", True):
                            continue
                        if msg["role"] == "user":
                            st.markdown(f'<div class="chat-user">💬 <b>You:</b> {msg["content"]}</div>', unsafe_allow_html=True)
                        else:
                            st.markdown(f'<div class="chat-bot"><b style="color:#e94560;">🔬 Claude</b><br>{msg["content"]}</div>', unsafe_allow_html=True)

                    st.markdown("---")
                    st.markdown("### 💬 Ask a Follow-Up Question")
                    user_q = st.text_input("Type your question:", key="chat_in",
                                           placeholder="e.g. What does the largest connected component reveal?")
                    cs, cn = st.columns([3, 1])
                    with cs:
                        if st.button("📤 Send", type="primary") and user_q.strip():
                            with st.spinner("Thinking…"):
                                try:
                                    msgs = [{"role":"user","content":st.session_state.topology_context+"\n\nAnswer follow-up questions about this topology."}]
                                    for m in st.session_state.chat_history:
                                        msgs.append({"role":m["role"],"content":m["content"]})
                                    msgs.append({"role":"user","content":user_q})
                                    ai = call_claude_api(api_key, msgs)
                                    st.session_state.chat_history.append({"role":"user","content":user_q,"display":True})
                                    st.session_state.chat_history.append({"role":"assistant","content":ai,"display":True})
                                    st.rerun()
                                except Exception as e:
                                    st.error(f"Error: {e}")
                    with cn:
                        if st.button("🔄 New Analysis"):
                            st.session_state.chat_history = []
                            st.session_state.topology_context = None
                            st.session_state.chat_active = False
                            st.session_state.claude_analysis = None
                            st.rerun()

                    export = "\n\n".join(
                        f"{'YOU' if m['role']=='user' else 'CLAUDE'}: {m['content']}"
                        for m in st.session_state.chat_history if m.get("display", True))

                    # ── Convert to DOCX ───────────────────────────────────────
                    try:
                        from docx import Document as DocxDocument
                        from docx.shared import Pt, RGBColor
                        from docx.enum.text import WD_ALIGN_PARAGRAPH
                        import io as _io

                        doc = DocxDocument()

                        # Title
                        title = doc.add_heading('SYN-IQ Topology Analysis Report', 0)
                        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

                        # Metadata
                        from datetime import datetime as _dt
                        doc.add_paragraph(f'Generated: {_dt.now().strftime("%B %d, %Y %H:%M")}')
                        doc.add_paragraph()

                        # Content — parse YOU/CLAUDE blocks
                        for block in export.split('\n\n'):
                            if not block.strip():
                                continue
                            if block.startswith('YOU:'):
                                p = doc.add_paragraph()
                                run = p.add_run('YOU: ')
                                run.bold = True
                                run.font.color.rgb = RGBColor(0x37, 0x7e, 0xb8)
                                p.add_run(block[4:].strip())
                            elif block.startswith('CLAUDE:'):
                                p = doc.add_paragraph()
                                run = p.add_run('CLAUDE: ')
                                run.bold = True
                                run.font.color.rgb = RGBColor(0xe4, 0x1a, 0x1c)
                                # Handle markdown bold (**text**)
                                content = block[7:].strip()
                                import re as _re2
                                parts = _re2.split(r'\*\*(.+?)\*\*', content)
                                for j, part in enumerate(parts):
                                    r = p.add_run(part)
                                    if j % 2 == 1:
                                        r.bold = True
                            else:
                                doc.add_paragraph(block.strip())

                        buf = _io.BytesIO()
                        doc.save(buf)
                        buf.seek(0)
                        st.download_button(
                            "📥 Download Report (DOCX)",
                            buf.getvalue(),
                            file_name="syniq_topology_report.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
                    except Exception as _docx_err:
                        # Show error + fallback to markdown
                        st.warning(f"⚠️ DOCX export failed: {_docx_err}")
                        st.download_button("📥 Download Report (MD)", export,
                                           file_name="syniq_topology_report.md", mime="text/markdown")

    else:
        if df_combined is None and uploaded_file is None:
            st.info("👆 Upload a SYN-IQ CSV or JSON file to begin analysis.")
            st.markdown("""
### Expected Data Format

**CSV** should include: `agent`, `temperature`, `int_pct`, `aff_pct`, `act_pct`, `vader_compound`

**V_t columns** (optional, enables V_t lenses): `S_t`, `A_t`, `Q_t`, `D_t`, `R_t`
— or any column with raw response text (set the column name in the sidebar).

**JSON**: auto-extracts response text if present.

Select *"IEP V3 — live score from text"* to score raw responses on-the-fly
using the embedded 1,897-word dictionary. Use *"IEP — pre-scored columns"* (default) for pipeline data.
        """)

# =============================================================================
# FOOTER
# =============================================================================
st.markdown("---")
st.markdown("""
<div style="text-align: center; color: #a0a0a0; padding: 1rem;">
    <strong>SYN-IQ Mapper Analyzer V17</strong><br>
    KeplerMapper + IEP Dictionary V3 — 1,897 words (embedded) · Claude AI Analysis · Topology Summary Table · Publication-Quality Visualization<br>
    <em>SYNINT Team — Tennessee 🎹 CUZ Partnership — March 2026</em>
</div>
""", unsafe_allow_html=True)